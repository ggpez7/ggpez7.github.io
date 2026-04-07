#!/usr/bin/env python3
from __future__ import annotations

import copy
import json
import re
import sys
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is required. Install it with: pip install python-docx")
    sys.exit(1)
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.oxml.ns import qn


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
ROOT = Path(__file__).resolve().parent


UNIVERSAL_SPEC = ROOT / "Universal Spec .docx"
DATA_REQUEST_DIR = ROOT / "DataRequest Samples"
PREVIOUS_REVIEW_DIR = ROOT / "Previous Quarter Template"
OUTPUT_DIR = ROOT / "prototype_output"
DEBUG_OUTPUT_DIR = OUTPUT_DIR / "_debug"


SECTION_HEADINGS = [
    "Business Activities",
    "Financial Update",
    "Business Update",
    "Risk & Exit",
]

SECTION_ALIASES = {
    "Business Activities": [
        "business activities",
        "business information",
        "business introduction",
        "company overview",
        "公司介绍",
    ],
    "Financial Update": [
        "financial update",
        "financial updates",
        "financial updates ",
        "financial update",
        "financial updates (in usd million)",
        "financial update",
        "财务情况",
    ],
    "Business Update": [
        "business update",
        "business highlights",
        "recent trends",
        "运营情况",
        "业务进展",
    ],
    "Risk & Exit": [
        "risk & exit",
        "risk and exit",
        "risks & exit",
        "risks and exit",
    ],
}


def is_question_or_prompt(text: str) -> bool:
    """Detect question/prompt paragraphs generically instead of hardcoding."""
    stripped = norm_space(text)
    if not stripped:
        return False
    if re.search(r"[?？]$", stripped):
        return True
    prompt_patterns = [
        r"^please\b",
        r"^what is\b",
        r"^what are\b",
        r"^how does\b",
        r"^how do\b",
        r"^can you\b",
        r"^could you\b",
        r"^describe\b",
        r"^explain\b",
        r"^elaborate\b",
        r"^provide\b",
        r"^project\b.*\bbalance\b",
        r"^if there is a (?:significant|major)\b.*\bplease\b",
        r"\bplease provide\b.*\bexplanation\b",
        r"^see notes to\b",
        r">\s*see notes to\b",
        r"^what'?s the\b",
        r"^when (?:is|will|did)\b",
        r"^where (?:is|are)\b",
        r"^why (?:is|are|did)\b",
        r"^is there\b",
        r"^are there\b",
        r"^do you\b",
        r"^have you\b",
        r"^请",
        r"^【重要】请",
        r"如有较大的.*请解释",
    ]
    lowered = stripped.lower()
    return any(re.search(p, lowered) for p in prompt_patterns)


def is_standalone_label(text: str) -> bool:
    """Detect short standalone labels/headers that aren’t real content."""
    stripped = norm_space(text)
    if not stripped:
        return False
    if stripped.endswith((":", "：")) and len(stripped) < 80:
        return True
    if re.match(r"^\d+[.)\s]", stripped) and len(stripped) < 60:
        return True
    section_label_patterns = [
        r"^other\s+questions",
        r"^operational\s+data",
        r"^业务发展$",
    ]
    lowered = stripped.lower()
    return any(re.search(p, lowered) for p in section_label_patterns)


def norm_space(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


def is_unit_text(text: str) -> bool:
    normalized = norm_space(text).lower()
    if not normalized:
        return False
    return any(
        token in normalized
        for token in [
            "unit",
            "amount in",
            "in usd",
            "in rmb",
            "usd million",
            "rmb mn",
            "单位",
            "百万",
            "万元",
            "人民币",
        ]
    )


def cell_text(tc: ET.Element) -> str:
    parts: list[str] = []
    for p in tc.findall("w:p", NS):
        text = "".join(t.text or "" for t in p.findall(".//w:t", NS))
        if text.strip():
            parts.append(norm_space(text))
    return " ".join(parts).strip()


def paragraph_text(p: ET.Element) -> str:
    return norm_space("".join(t.text or "" for t in p.findall(".//w:t", NS)))


def load_docx_blocks(path: Path) -> list[dict[str, Any]]:
    with zipfile.ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    body = root.find("w:body", NS)
    assert body is not None

    blocks: list[dict[str, Any]] = []
    for child in body:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            text = paragraph_text(child)
            if text:
                blocks.append({"type": "paragraph", "text": text})
        elif tag == "tbl":
            rows = []
            for tr in child.findall("w:tr", NS):
                rows.append([cell_text(tc) for tc in tr.findall("w:tc", NS)])
            blocks.append({"type": "table", "rows": rows})
    return blocks


def docx_to_text(path: Path) -> str:
    blocks = load_docx_blocks(path)
    lines = []
    for block in blocks:
        if block["type"] == "paragraph":
            lines.append(block["text"])
        else:
            for row in block["rows"]:
                lines.append(" | ".join(row))
    return "\n".join(lines)


def detect_language(text: str) -> tuple[str, str]:
    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
    if chinese_chars > 20:
        return "chinese", f"Detected {chinese_chars} Chinese characters in the company-submitted source."
    return "english", "No meaningful Chinese text detected in the company-submitted source."


def text_language(text: str) -> str:
    return "chinese" if re.search(r"[\u4e00-\u9fff]", text) else "english"


def canonical_heading(text: str) -> str | None:
    lowered = norm_space(text).lower()
    lowered = re.sub(r"^[0-9一二三四五六七八九十]+[\-–、.\s]*", "", lowered)
    for canonical, aliases in SECTION_ALIASES.items():
        for alias in aliases:
            if lowered == alias or lowered.startswith(alias):
                return canonical
    if lowered.startswith("operations update") or lowered.startswith("business progress"):
        return "Business Update"
    if "风险" in lowered and "退出" in lowered:
        return "Risk & Exit"
    return None


def fixed_anchor_heading(text: str) -> str | None:
    lowered = norm_space(text).lower()
    lowered = re.sub(r"^[0-9一二三四五六七八九十]+[\-–、.\s]*", "", lowered)
    for canonical in ["Business Activities", "Financial Update", "Risk & Exit"]:
        for alias in SECTION_ALIASES.get(canonical, []):
            if lowered == alias or lowered.startswith(alias):
                return canonical
    if "风险" in lowered and "退出" in lowered:
        return "Risk & Exit"
    return None


def is_middle_section_heading_candidate(paragraph) -> bool:
    text = norm_space(paragraph.text)
    if not text:
        return False
    if text.endswith((".", "。", ";", "；", ":", "：", "?", "？")):
        return False
    if len(text) > 80:
        return False
    if text.startswith(("-", "•", "*")):
        return False
    style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
    if "list" in style_name or "heading" in style_name or "标题" in style_name or style_name == "normal":
        return True
    return False


def extract_previous_sections(blocks: list[dict[str, Any]]) -> dict[str, Any]:
    sections: dict[str, Any] = {}
    current_heading: str | None = None

    for block in blocks:
        if block["type"] == "paragraph":
            heading = canonical_heading(block["text"])
        else:
            heading = None
        if heading:
            current_heading = heading
            sections[current_heading] = {"paragraphs": [], "tables": []}
            continue
        if not current_heading:
            continue
        if block["type"] == "paragraph":
            sections[current_heading]["paragraphs"].append(block["text"])
        elif block["type"] == "table":
            sections[current_heading]["tables"].append(block["rows"])
    return sections


def extract_previous_section_occurrences(blocks: list[dict[str, Any]]) -> list[SectionOccurrence]:
    occurrences: list[SectionOccurrence] = []
    current: SectionOccurrence | None = None
    for block in blocks:
        heading = canonical_heading(block["text"]) if block["type"] == "paragraph" else None
        if heading:
            current = SectionOccurrence(
                canonical=heading,
                language=text_language(block["text"]),
                heading_text=block["text"],
                paragraphs=[],
                tables=[],
            )
            occurrences.append(current)
            continue
        if current is None:
            continue
        if block["type"] == "paragraph":
            current.paragraphs.append(block["text"])
        elif block["type"] == "table":
            current.tables.append(block["rows"])
    return occurrences


def parse_previous_financial_table(rows: list[list[str]]) -> dict[str, Any]:
    header_row = rows[0]
    unit = header_row[0]
    headers = [norm_space(cell) for cell in header_row[1:]]
    data_rows = []
    for row in rows[1:]:
        label = norm_space(row[0])
        values = {headers[i]: norm_space(row[i + 1]) for i in range(len(headers))}
        data_rows.append({"label": label, "values": values})
    return {"unit": unit, "headers": headers, "rows": data_rows}


def parse_previous_financial_table_with_unit(rows: list[list[str]], unit_label: str | None) -> dict[str, Any]:
    parsed = parse_previous_financial_table(rows)
    if unit_label:
        parsed["unit"] = unit_label
    return parsed


@dataclass(frozen=True, order=True)
class Quarter:
    year: int
    quarter: int

    def display(self) -> str:
        return f"Q{self.quarter} {self.year}"

    def previous(self) -> "Quarter":
        if self.quarter == 1:
            return Quarter(self.year - 1, 4)
        return Quarter(self.year, self.quarter - 1)


@dataclass
class SectionOccurrence:
    canonical: str
    language: str
    heading_text: str
    paragraphs: list[str]
    tables: list[list[list[str]]]


@dataclass
class DocSectionOccurrence:
    canonical: str
    language: str
    heading_index: int


@dataclass(frozen=True)
class UnitSpec:
    label: str
    factor_to_base: Decimal


def parse_quarter_label(text: str) -> Quarter | None:
    match = re.search(r"Q([1-4])\s*(20\d{2})", text)
    if match:
        return Quarter(int(match.group(2)), int(match.group(1)))
    match = re.search(r"(20\d{2})\s*Q([1-4])", text)
    if match:
        return Quarter(int(match.group(1)), int(match.group(2)))
    match = re.search(r"(20\d{2})Q([1-4])", text)
    if match:
        return Quarter(int(match.group(1)), int(match.group(2)))
    return None


def extract_quarter_number(text: str) -> int | None:
    match = re.search(r"Q([1-4])", text, flags=re.I)
    if match:
        return int(match.group(1))
    match = re.search(r"([1-4])季度", text)
    if match:
        return int(match.group(1))
    return None


def parse_month_header(text: str) -> tuple[int, int] | None:
    normalized = norm_space(text)
    match = re.search(r"(20\d{2})[/-]?(0[1-9]|1[0-2])", normalized)
    if match:
        return int(match.group(1)), int(match.group(2))
    return None


def quarter_sort_key(text: str) -> tuple[int, int]:
    quarter = parse_quarter_label(text)
    if quarter is None:
        return (-1, -1)
    return (quarter.year, quarter.quarter)


def parse_decimal(text: str) -> Decimal | None:
    cleaned = text.replace(",", "").strip()
    if not cleaned or cleaned in {"/", "n.a.", "n/a"}:
        return None
    negative = cleaned.startswith("(") and cleaned.endswith(")")
    cleaned = cleaned.strip("()")
    cleaned = cleaned.replace("Positive cf", "").replace("Better", "").strip()
    cleaned = re.sub(r"[^0-9.+-]", "", cleaned)
    if not cleaned:
        return None
    try:
        value = Decimal(cleaned)
    except Exception:
        return None
    return -value if negative else value


def format_decimal(value: Decimal | None, places: int = 2) -> str | None:
    if value is None:
        return None
    quant = Decimal("1." + ("0" * places))
    rounded = value.quantize(quant, rounding=ROUND_HALF_UP)
    return f"{rounded:.{places}f}"


def parse_percent(text: str) -> Decimal | None:
    cleaned = norm_space(text)
    if not cleaned or cleaned == "/":
        return None
    negative = cleaned.startswith("(") and "%" in cleaned
    match = re.search(r"([+-]?\d+(?:\.\d+)?)\s*%", cleaned)
    if not match:
        return None
    value = Decimal(match.group(1))
    if negative and value > 0:
        value = -value
    return value


def format_percent(value: Decimal | None) -> str | None:
    if value is None:
        return None
    rounded = value.quantize(Decimal("1.00"), rounding=ROUND_HALF_UP)
    return f"{rounded:.2f}%"


def parse_current_blocks(blocks: list[dict[str, Any]]) -> dict[str, Any]:
    result: dict[str, Any] = {
        "title": "",
        "finance_table": None,
        "balance_table": None,
        "operation_table": None,
        "business_update_paragraphs": [],
    }

    tables_seen = 0
    for idx, block in enumerate(blocks):
        if idx == 0 and block["type"] == "paragraph":
            result["title"] = block["text"]

        if block["type"] == "table":
            tables_seen += 1
            if result["finance_table"] is None:
                result["finance_table"] = block["rows"]
            elif result["balance_table"] is None:
                result["balance_table"] = block["rows"]
            elif result["operation_table"] is None:
                result["operation_table"] = block["rows"]

        if block["type"] == "paragraph" and block["text"] == "Other Questions:":
            trailing = blocks[idx + 1 :]
            for item in trailing:
                if item["type"] != "paragraph":
                    continue
                text = item["text"]
                if is_question_or_prompt(text):
                    continue
                if is_standalone_label(text):
                    continue
                result["business_update_paragraphs"].append(text)
            break

    if not result["business_update_paragraphs"]:
        seen_table = False
        for block in blocks:
            if block["type"] == "table":
                seen_table = True
                continue
            if not seen_table or block["type"] != "paragraph":
                continue
            text = block["text"]
            if is_question_or_prompt(text):
                continue
            if is_standalone_label(text):
                continue
            if len(text) < 2:
                continue
            result["business_update_paragraphs"].append(text)
    return result


def rows_to_dict(rows: list[list[str]]) -> dict[str, dict[str, str]]:
    headers = rows[0][1:]
    mapping: dict[str, dict[str, str]] = {}
    for row in rows[1:]:
        label = norm_space(row[0])
        mapping[label] = {headers[i]: norm_space(row[i + 1]) for i in range(len(headers))}
    return mapping


def extract_outlet_counts(operation_rows: list[list[str]]) -> tuple[str | None, str | None]:
    op_map = rows_to_dict(operation_rows)
    outlet_row = None
    for label in op_map:
        if "outlet" in label.lower() and "accumul" in label.lower():
            outlet_row = op_map[label]
            break
    if outlet_row is None:
        outlet_row = op_map.get("Accumulated Number of outlets", {})
    ytd_value = None
    for header, value in outlet_row.items():
        if "ytd" in header.lower() or "accum" in header.lower() or "total" in header.lower():
            ytd_value = value
            break
    if not ytd_value:
        values = list(outlet_row.values())
        ytd_value = values[-1] if values else None
    if not ytd_value:
        return None, None
    match = re.search(r"(\d+)\s*\+\s*(\d+)\s*\(JV\)", ytd_value)
    if not match:
        plain = re.search(r"(\d+)", ytd_value)
        return (plain.group(1), None) if plain else (None, None)
    return match.group(1), match.group(2)



def normalize_finance_source(rows: list[list[str]]) -> dict[str, Any]:
    header = [norm_space(cell) for cell in rows[0]]
    row_map = rows_to_dict(rows)
    return {
        "header": header,
        "rows": row_map,
    }


def calc_percent(current: Decimal | None, previous: Decimal | None) -> Decimal | None:
    if current is None or previous is None or previous == 0:
        return None
    denominator = abs(previous)
    if denominator == 0:
        return None
    return ((current - previous) / denominator) * Decimal("100")


def rewrite_third_person(text: str, company_name: str) -> str:
    text = norm_space(text)
    text = re.sub(r"\bwe\b", company_name, text, flags=re.I)
    text = re.sub(r"\bour\b", f"{company_name}'s", text, flags=re.I)
    text = re.sub(r"\bus\b", company_name, text, flags=re.I)
    text = re.sub(r"\bQ([1-4])\.", r"Q\1", text)
    text = re.sub(r"\bthe company\b", company_name, text, flags=re.I)
    return norm_space(text)


def short_sentence(text: str, company_name: str, limit: int = 140) -> str:
    text = rewrite_third_person(text, company_name)
    text = re.sub(r":\s*", ": ", text)
    text = re.sub(r"\s*-\s*", " - ", text)
    if len(text) <= limit:
        return text.rstrip(".") + "."
    clipped = text[:limit].rsplit(" ", 1)[0].rstrip(",;:-")
    return clipped + "."


def score_paragraph_informativeness(text: str) -> float:
    """Score a paragraph by how much useful business information it contains.
    Higher score = more informative (numbers, metrics, dollar amounts, percentages, specifics)."""
    score = 0.0
    score += len(re.findall(r"\$[\d,.]+[kmb]?", text, re.I)) * 3.0
    score += len(re.findall(r"\d+(?:\.\d+)?%", text)) * 2.5
    score += len(re.findall(r"\b\d{1,3}(?:,\d{3})+\b", text)) * 2.0
    score += len(re.findall(r"\b\d+(?:\.\d+)?\s*(?:x|times|million|billion|k)\b", text, re.I)) * 2.0
    score += len(re.findall(r"\bQ[1-4]\s*20\d{2}\b", text, re.I)) * 1.5
    score += len(re.findall(r"\b(?:revenue|profit|ebitda|margin|growth|cost|loss|cash flow|burn rate|funding|financing|expansion|launched|achieved|target|milestone)\b", text, re.I)) * 1.0
    score += len(re.findall(r"\b(?:YoY|QoQ|year.over.year|quarter.over.quarter|month.over.month)\b", text, re.I)) * 1.5
    length_bonus = min(len(text) / 200.0, 1.0)
    score += length_bonus
    if len(text) < 30:
        score *= 0.3
    return score


def split_into_sentences(text: str) -> list[str]:
    """Split a long paragraph into individual sentences for separate bullets."""
    # Don't split short text
    if len(text) < 120:
        return [text]
    # Split on sentence boundaries, but not on abbreviations like "U.S." or "e.g."
    # or on decimal numbers like "2.67m"
    parts = re.split(r'(?<=[.!])\s+(?=[A-Z\d])', text)
    result = []
    for part in parts:
        part = norm_space(part)
        if len(part) >= 20:
            result.append(part)
    return result if result else [text]


def build_business_update_bullets(paragraphs: list[str], company_name: str, max_bullets: int = 9) -> list[str]:
    """Extract the most informative paragraphs as bullet points, scored by content richness."""
    cleaned = clean_update_paragraphs(paragraphs)
    # Split multi-sentence paragraphs into separate candidates
    expanded: list[str] = []
    for text in cleaned:
        text = norm_space(text)
        if len(text) < 20:
            continue
        if text.endswith((":", "：")):
            continue
        expanded.extend(split_into_sentences(text))
    scored: list[tuple[float, str]] = []
    for text in expanded:
        score = score_paragraph_informativeness(text)
        scored.append((score, text))

    scored.sort(key=lambda x: x[0], reverse=True)

    bullets: list[str] = []
    seen_keys: set[str] = set()
    for _score, text in scored:
        bullet = short_sentence(text, company_name, limit=200)
        key = re.sub(r"[^a-z0-9]", "", bullet.lower())
        if key in seen_keys:
            continue
        seen_keys.add(key)
        bullets.append(bullet)
        if len(bullets) >= max_bullets:
            break

    return bullets


def strip_response_prefix(text: str) -> str:
    """Strip Q&A response prefixes like 'CompanyName: ...' or 'Yes, CompanyName ...'."""
    # Strip "CompanyName: " or "CompanyName, " at start
    text = re.sub(r"^[A-Z][A-Za-z\s]{1,30}:\s*", "", text)
    # Strip "Yes, " or "Yes. " at start
    text = re.sub(r"^(?:Yes|No)[,.\s]+\s*", "", text, flags=re.I)
    return norm_space(text)


def clean_update_paragraphs(paragraphs: list[str]) -> list[str]:
    cleaned: list[str] = []
    for text in paragraphs:
        text = norm_space(text)
        if not text:
            continue
        if is_question_or_prompt(text):
            continue
        if is_standalone_label(text):
            continue
        text = strip_response_prefix(text)
        if not text or len(text) < 10:
            continue
        cleaned.append(text)
    return cleaned


def bullet_topic(text: str) -> str:
    lowered = norm_space(text).lower()
    if any(token in lowered for token in ["融资", "股改", "上市", "ipo", "equity", "fund", "loan", "facility", "capital", "cash flow", "现金流"]):
        return "finance"
    if any(token in lowered for token in ["项目", "交付", "签约", "回款", "合同", "project", "delivery", "contract", "customer", "operation", "运营", "milestone"]):
        return "operations"
    if any(token in lowered for token in ["计划", "扩张", "outlet", "growth", "expansion", "rollout", "trend", "progress", "channel", "product", "certification"]):
        return "growth"
    return "general"


def section_heading_preference(heading_text: str) -> list[str]:
    lowered = norm_space(heading_text).lower()
    if any(token in lowered for token in ["运营", "operation", "review", "业务回顾", "recent"]):
        return ["operations", "general", "growth", "finance"]
    if any(token in lowered for token in ["进展", "progress", "highlight", "trend", "业务进展", "business"]):
        return ["growth", "finance", "general", "operations"]
    return ["general", "operations", "growth", "finance"]


def allocate_middle_section_bullets(headings: list[str], bullets: list[str]) -> list[list[str]]:
    if not headings:
        return []
    if len(headings) == 1:
        return [bullets]

    remaining = [(bullet, bullet_topic(bullet)) for bullet in bullets]
    allocations: list[list[str]] = []
    for idx, heading in enumerate(headings):
        if idx == len(headings) - 1:
            allocations.append([bullet for bullet, _topic in remaining])
            break
        preferred = section_heading_preference(heading)
        chosen: list[str] = []
        still_remaining = []
        used_topics = set()
        for bullet, topic in remaining:
            if topic in preferred[:2] and topic not in used_topics:
                chosen.append(bullet)
                used_topics.add(topic)
            else:
                still_remaining.append((bullet, topic))
        if not chosen and remaining:
            chosen.append(remaining[0][0])
            still_remaining = remaining[1:]
        allocations.append(chosen)
        remaining = still_remaining
    while len(allocations) < len(headings):
        allocations.append([])
    return allocations


def build_chinese_business_update_bullets(paragraphs: list[str]) -> list[str]:
    lines = clean_update_paragraphs(paragraphs)
    bullets: list[str] = []
    for text in lines:
        if len(text) < 4:
            continue
        if text.endswith("：") or text.endswith(":"):
            continue
        bullets.append(text.rstrip("；;。") + "。")
    deduped: list[str] = []
    seen = set()
    for bullet in bullets:
        key = bullet.lower()
        if key not in seen:
            seen.add(key)
            deduped.append(bullet)
    return deduped[:8]


def previous_review_has_logo(path: Path) -> bool:
    with zipfile.ZipFile(path) as zf:
        return any(name.startswith("word/media/") for name in zf.namelist())


def normalize_company_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", " ", name).strip()
    parts = [part for part in cleaned.split() if part]
    if not parts:
        return "PortfolioCompany"
    return "".join(parts)


def quarter_filename_label(quarter: Quarter) -> str:
    return f"{quarter.year}Q{quarter.quarter}"


def clean_main_output_folder(final_docx_path: Path) -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    for path in OUTPUT_DIR.iterdir():
        if path == DEBUG_OUTPUT_DIR:
            continue
        if path.is_file() and path.suffix.lower() != ".docx":
            path.unlink()


def clean_main_output_folder_for_batch() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    for path in OUTPUT_DIR.iterdir():
        if path == DEBUG_OUTPUT_DIR:
            continue
        if path.is_file():
            path.unlink()


def list_real_docx(folder: Path) -> list[Path]:
    return sorted(
        path
        for path in folder.glob("*.docx")
        if not path.name.startswith("~$")
    )


def previous_financial_lookup(parsed_prev: dict[str, Any]) -> dict[str, dict[str, str]]:
    return {row["label"]: row["values"] for row in parsed_prev["rows"]}


def latest_quarter_from_current(finance_source: dict[str, Any]) -> Quarter:
    for header in finance_source["header"]:
        quarter = parse_quarter_label(header)
        if quarter:
            return quarter
    raise ValueError("Could not identify current quarter from data request headers.")


def normalize_metric_label(text: str) -> str:
    text = text.lower()
    text = re.sub(r"\(=.*?\)", "", text)
    text = re.sub(r"\([^)]*\)", " ", text)
    text = text.replace("&", " and ")
    text = re.sub(r"[^a-z0-9\u4e00-\u9fff]+", " ", text)
    return norm_space(text)


def metric_tokens(text: str) -> set[str]:
    stopwords = {
        "of",
        "and",
        "the",
        "per",
        "company",
        "owned",
        "number",
        "accumulated",
    }
    normalized = normalize_metric_label(text)
    if re.search(r"[\u4e00-\u9fff]", normalized) and " " not in normalized:
        return {normalized}
    return {token for token in normalized.split() if token not in stopwords}


def find_current_quarter_total_header(finance_source: dict[str, Any], target_quarter: Quarter) -> str | None:
    for header in finance_source["header"]:
        if parse_quarter_label(header) == target_quarter:
            return header
    exact_subtotal = []
    summary_candidates = []
    for header in finance_source["header"]:
        lowered = header.lower()
        if any(token in lowered for token in ["ytd", "accum", "exp", "qoq", "yoy"]):
            continue
        if parse_quarter_label(header) == target_quarter:
            exact_subtotal.append(header)
        if "总计" in header or "subtotal" in lowered:
            summary_candidates.append(header)
    if exact_subtotal:
        return exact_subtotal[0]
    if summary_candidates:
        return summary_candidates[0]
    return None


def find_header_value(values: dict[str, str], patterns: list[str]) -> str | None:
    for header, value in values.items():
        lowered = header.lower()
        if any(pattern in lowered for pattern in patterns):
            return value
    return None


def find_month_headers_for_quarter(finance_source: dict[str, Any], target_quarter: Quarter) -> list[str]:
    headers = []
    for header in finance_source["header"]:
        month_info = parse_month_header(header)
        if not month_info:
            continue
        year, month = month_info
        quarter = ((month - 1) // 3) + 1
        if year == target_quarter.year and quarter == target_quarter.quarter:
            headers.append(header)
    return headers


def sum_month_values(values: dict[str, str], month_headers: list[str]) -> Decimal | None:
    if not month_headers:
        return None
    month_values = [parse_decimal(values.get(header, "")) for header in month_headers]
    if any(value is None for value in month_values):
        return None
    return sum(month_values)  # type: ignore[arg-type]


def find_fy_or_ytd_header(finance_source: dict[str, Any], target_quarter: Quarter) -> str | None:
    candidates = []
    year_text = str(target_quarter.year)
    for header in finance_source["header"]:
        lowered = header.lower()
        if year_text not in header:
            continue
        if any(token in lowered for token in ["exp", "qoq", "yoy"]):
            continue
        if parse_quarter_label(header) is not None:
            continue
        if parse_month_header(header) is not None:
            continue
        if any(token in lowered for token in ["ytd", "accum"]) or "总计" in header or "全年" in header:
            candidates.append(header)
    for header in candidates:
        if "总计" in header or "全年" in header:
            return header
    return candidates[0] if candidates else None


def normalize_previous_quarter_headers(raw_headers: list[str], previous_review_quarter: Quarter | None) -> list[tuple[str, str]]:
    quarter_headers = [header for header in raw_headers if parse_quarter_label(header)]
    if previous_review_quarter is None:
        return [(header, parse_quarter_label(header).display()) for header in quarter_headers if parse_quarter_label(header)]

    expected = []
    cursor = previous_review_quarter
    for _ in quarter_headers:
        expected.append(cursor)
        cursor = cursor.previous()

    normalized: list[tuple[str, str]] = []
    used_expected: set[Quarter] = set()
    for idx, header in enumerate(quarter_headers):
        parsed = parse_quarter_label(header)
        if parsed in expected and parsed not in used_expected:
            normalized.append((header, parsed.display()))
            used_expected.add(parsed)
            continue

        expected_q = expected[idx]
        parsed_qnum = extract_quarter_number(header)
        if expected_q not in used_expected and (parsed is None or parsed_qnum == expected_q.quarter or parsed not in expected):
            normalized.append((header, expected_q.display()))
            used_expected.add(expected_q)
            continue

        if parsed is not None and parsed not in used_expected:
            normalized.append((header, parsed.display()))
            used_expected.add(parsed)
            continue

        normalized.append((header, header))

    return normalized


def detect_unit_spec(text: str) -> UnitSpec | None:
    normalized = norm_space(text).lower()
    compact = normalized.replace(" ", "")
    candidates: list[tuple[str, Decimal]] = [
        ("usdk", Decimal("1000")),
        ("kusd", Decimal("1000")),
        ("k usd", Decimal("1000")),
        ("usd million", Decimal("1000000")),
        ("millions usd", Decimal("1000000")),
        ("million usd", Decimal("1000000")),
        ("in usd million", Decimal("1000000")),
        ("in millions usd", Decimal("1000000")),
        ("amount in usd million", Decimal("1000000")),
        ("unit: usd", Decimal("1")),
        ("financial data (unit: usd)", Decimal("1")),
        ("单位：万元人民币", Decimal("10000")),
        ("万元人民币", Decimal("10000")),
        ("万元 人民币", Decimal("10000")),
        ("单位：百万元人民币", Decimal("1000000")),
        ("百万元人民币", Decimal("1000000")),
        ("单位：百万人民币", Decimal("1000000")),
        ("百万人民币", Decimal("1000000")),
        ("rmb in millions", Decimal("1000000")),
        ("amount in rmb mn", Decimal("1000000")),
        ("rmb mn", Decimal("1000000")),
        ("unit: rmb in millions", Decimal("1000000")),
    ]
    for needle, factor in candidates:
        if needle in normalized or needle.replace(" ", "") in compact:
            return UnitSpec(label=text, factor_to_base=factor)
    if "unit: usd" in normalized or "(unit: usd)" in normalized:
        return UnitSpec(label=text, factor_to_base=Decimal("1"))
    return None


def detect_current_unit_spec(blocks: list[dict[str, Any]]) -> UnitSpec | None:
    for block in blocks[:6]:
        if block["type"] == "paragraph":
            spec = detect_unit_spec(block["text"])
            if spec:
                return spec
    for block in blocks:
        if block["type"] == "paragraph":
            spec = detect_unit_spec(block["text"])
            if spec:
                return spec
    return None


def detect_financial_section_unit(occurrence: SectionOccurrence) -> tuple[str | None, str]:
    outside_unit = None
    for paragraph in occurrence.paragraphs:
        if is_unit_text(paragraph):
            outside_unit = paragraph
            break
    inside_unit = None
    if occurrence.tables and occurrence.tables[0] and occurrence.tables[0][0]:
        inside_unit = norm_space(occurrence.tables[0][0][0])
        if not is_unit_text(inside_unit):
            inside_unit = None
    if outside_unit:
        return outside_unit, "outside"
    if inside_unit:
        return inside_unit, "inside"
    return inside_unit or outside_unit, "none"


def convert_value_between_units(value: Decimal | None, source_unit: UnitSpec | None, output_unit: UnitSpec | None) -> Decimal | None:
    if value is None:
        return None
    if source_unit is None or output_unit is None:
        return value
    return (value * source_unit.factor_to_base) / output_unit.factor_to_base


def determine_target_quarter(current_path: Path, current_title: str, current_finance: dict[str, Any]) -> Quarter:
    monthly_headers = [h for h in current_finance["header"] if parse_month_header(h)]
    if monthly_headers:
        last_header = monthly_headers[-1]
        month_info = parse_month_header(last_header)
        if month_info:
            year, month = month_info
            quarter = ((month - 1) // 3) + 1
            return Quarter(year, quarter)
    for source in [current_title, current_path.stem]:
        quarter = parse_quarter_label(source)
        if quarter:
            return quarter
    return latest_quarter_from_current(current_finance)


def determine_target_quarter_with_debug(current_path: Path, current_title: str, current_finance: dict[str, Any]) -> tuple[Quarter, str]:
    monthly_headers = [h for h in current_finance["header"] if parse_month_header(h)]
    if monthly_headers:
        last_header = monthly_headers[-1]
        month_info = parse_month_header(last_header)
        if month_info:
            year, month = month_info
            quarter = Quarter(year, ((month - 1) // 3) + 1)
            return quarter, f"month-based inference from headers {monthly_headers}"
    for source_name, source in [("title", current_title), ("filename", current_path.stem)]:
        quarter = parse_quarter_label(source)
        if quarter:
            return quarter, f"fallback quarter-label inference from {source_name}: {source}"
    quarter = latest_quarter_from_current(current_finance)
    return quarter, f"fallback latest_quarter_from_current using finance headers {current_finance['header']}"


def choose_finance_row_label(
    previous_label: str,
    finance_rows: dict[str, dict[str, str]],
    current_total_header: str,
    previous_current_value: str | None,
    flags: list[dict[str, Any]],
    used_source_labels: dict[str, str] | None = None,
) -> str | None:
    prev_tokens = metric_tokens(previous_label)
    used_source_labels = used_source_labels or {}

    exact_matches = []
    for candidate, values in finance_rows.items():
        if parse_decimal(values.get(current_total_header, "")) is None:
            continue
        if candidate in used_source_labels and used_source_labels[candidate] != previous_label:
            continue
        if normalize_metric_label(candidate) == normalize_metric_label(previous_label):
            exact_matches.append(candidate)
    if len(exact_matches) == 1:
        return exact_matches[0]
    if len(exact_matches) > 1:
        flags.append(
            {
                "id": f"mapping-{normalize_metric_label(previous_label).replace(' ', '-')}",
                "section": "Financial Update",
                "severity": "warning",
                "message": f"Multiple exact metric matches found for '{previous_label}': {', '.join(exact_matches)}.",
                "source": "current_data_request",
            }
        )
        return exact_matches[0]

    best_label: str | None = None
    best_score: tuple[Decimal, int] | None = None
    ambiguous = False

    for candidate, values in finance_rows.items():
        if candidate in used_source_labels and used_source_labels[candidate] != previous_label:
            continue
        candidate_value = parse_decimal(values.get(current_total_header, ""))
        if candidate_value is None:
            continue
        cand_tokens = metric_tokens(candidate)
        token_score = len(prev_tokens & cand_tokens)

        previous_value = parse_decimal(previous_current_value or "")
        continuity_penalty = Decimal("999999")
        if previous_value is not None:
            candidate_millions = candidate_value / Decimal("1000")
            continuity_penalty = abs(candidate_millions - previous_value)

        score = (-continuity_penalty, token_score)
        if best_score is None or score > best_score:
            best_label = candidate
            best_score = score
            ambiguous = False
        elif score == best_score:
            ambiguous = True

    if best_label is None or best_score is None:
        if any(
            normalize_metric_label(candidate) == normalize_metric_label(previous_label)
            for candidate in finance_rows
        ):
            flags.append(
                {
                    "id": f"mapping-{normalize_metric_label(previous_label).replace(' ', '-')}",
                    "section": "Financial Update",
                    "severity": "warning",
                    "message": f"Metric '{previous_label}' had a potential source-row collision and was left unmapped.",
                    "source": "current_data_request",
                }
            )
            return None
        flags.append(
            {
                "id": f"mapping-{normalize_metric_label(previous_label).replace(' ', '-')}",
                "section": "Financial Update",
                "severity": "warning",
                "message": f"Could not map previous-review metric '{previous_label}' to a current data request row.",
                "source": "current_data_request",
            }
        )
        return None

    if ambiguous:
        flags.append(
            {
                "id": f"mapping-{normalize_metric_label(previous_label).replace(' ', '-')}",
                "section": "Financial Update",
                "severity": "warning",
                "message": f"Metric mapping for '{previous_label}' is ambiguous; selected '{best_label}' by token overlap.",
                "source": "current_data_request",
            }
        )
    return best_label


def build_financial_update(
    prev_financial: dict[str, Any],
    current_finance: dict[str, Any],
    current_operation_rows: list[list[str]] | None,
    flags: list[dict[str, Any]],
    target_quarter: Quarter | None = None,
    source_unit_spec: UnitSpec | None = None,
    output_unit_spec: UnitSpec | None = None,
    previous_review_quarter: Quarter | None = None,
) -> dict[str, Any]:
    previous_lookup = previous_financial_lookup(prev_financial)
    finance_rows = current_finance["rows"]
    target_quarter = target_quarter or latest_quarter_from_current(current_finance)

    normalized_previous_headers = normalize_previous_quarter_headers(prev_financial["headers"], previous_review_quarter)
    normalized_previous_quarter_labels = [normalized for _, normalized in normalized_previous_headers]

    rolling_headers = [target_quarter.display()]
    quarter_cursor = target_quarter.previous()
    while len(rolling_headers) < 5:
        rolling_headers.append(quarter_cursor.display())
        quarter_cursor = quarter_cursor.previous()

    final_total_header = f"{target_quarter.year} FY" if target_quarter.quarter == 4 else f"{target_quarter.year} YTD"
    estimate_header = f"{target_quarter.year + 1}E" if target_quarter.quarter == 4 else f"{target_quarter.year}E"
    output_headers = rolling_headers + ["QoQ", "YoY", final_total_header, estimate_header]
    current_total_header = find_current_quarter_total_header(current_finance, target_quarter)
    current_month_headers = find_month_headers_for_quarter(current_finance, target_quarter)
    fy_or_ytd_header = find_fy_or_ytd_header(current_finance, target_quarter)
    if current_total_header is None:
        current_total_header = ""
    operation_rows = current_operation_rows or [["", ""]]
    company_outlets, jv_outlets = extract_outlet_counts(operation_rows)
    company_outlets_total, jv_outlets_total = extract_outlet_counts(operation_rows)

    rows_output: list[dict[str, Any]] = []
    used_source_labels: dict[str, str] = {}
    for prev_row in prev_financial["rows"]:
        label = prev_row["label"]
        if label == "#of Company-Owned Outlets":
            source_info = {"type": "derived", "label": "company_outlets"}
        elif label == "#of JV Outlets":
            source_info = {"type": "derived", "label": "jv_outlets"}
        else:
            previous_current_value = None
            if previous_review_quarter is not None:
                for raw, normalized in normalized_previous_headers:
                    if normalized == previous_review_quarter.display():
                        previous_current_value = prev_row["values"].get(raw)
                        break
            mapped_label = choose_finance_row_label(
                label,
                finance_rows,
                current_total_header,
                previous_current_value,
                flags,
                used_source_labels,
            )
            source_info = {"type": "finance_row", "label": mapped_label} if mapped_label else None
            if mapped_label:
                used_source_labels[mapped_label] = label
        row_flags: list[str] = []
        values: dict[str, str | None] = {}
        source_trace: dict[str, str] = {}

        source_row = finance_rows.get(source_info["label"]) if source_info and source_info["type"] == "finance_row" else None
        prev_quarters = {normalized: prev_row["values"].get(raw) for raw, normalized in normalized_previous_headers}

        current_quarter_value: str | None = None
        if source_info:
            if source_info["type"] == "finance_row" and source_row:
                decimal_value = sum_month_values(source_row, current_month_headers)
                if decimal_value is None and current_total_header:
                    raw_current = source_row.get(current_total_header)
                    decimal_value = parse_decimal(raw_current)
                if decimal_value is not None:
                    converted_value = convert_value_between_units(decimal_value, source_unit_spec, output_unit_spec)
                    current_quarter_value = format_decimal(converted_value)
                    source_trace[target_quarter.display()] = "current_data_request"
                else:
                    row_flags.append("Missing current-quarter value in the company data request.")
            elif source_info["type"] == "finance_row" and source_row is None:
                row_flags.append("No current data request row could be mapped to this metric.")
            elif source_info["label"] == "company_outlets":
                current_quarter_value = company_outlets
                if company_outlets is None:
                    row_flags.append("Could not parse company-owned outlet count from operation data.")
                else:
                    source_trace[target_quarter.display()] = "current_data_request"
            elif source_info["label"] == "jv_outlets":
                current_quarter_value = jv_outlets
                if jv_outlets is None:
                    row_flags.append("Could not parse JV outlet count from operation data.")
                else:
                    source_trace[target_quarter.display()] = "current_data_request"

        for header in rolling_headers:
            if header == target_quarter.display():
                values[header] = current_quarter_value
                continue
            values[header] = prev_quarters.get(header)
            if prev_quarters.get(header):
                source_trace[header] = "previous_review"
            else:
                row_flags.append(f"Historical quarter '{header}' was not found in the previous review table.")

        qoq_value = None
        yoy_value = None

        if source_info and source_info["type"] == "finance_row" and source_row:
            row_current_dec = parse_decimal(values[target_quarter.display()] or "")
            if row_current_dec is not None and label in {"Revenue", "Gross profit", "EBITDA"}:
                row_current_dec = Decimal(values[target_quarter.display()])  # same unit in output

            previous_q_dec = parse_decimal(values[rolling_headers[1]] or "")
            yoy_dec = parse_decimal(values[rolling_headers[4]] or "")
            calc_qoq = calc_percent(row_current_dec, previous_q_dec)
            calc_yoy = calc_percent(row_current_dec, yoy_dec)
            qoq_value = format_percent(calc_qoq)
            yoy_value = format_percent(calc_yoy)
            if qoq_value is None:
                row_flags.append("QoQ could not be determined from the rolled quarter values.")
            if yoy_value is None:
                row_flags.append("YoY could not be determined from the rolled quarter values.")

            current_fy_source = parse_decimal(source_row.get(fy_or_ytd_header or "", ""))
            estimate_source = parse_decimal(find_header_value(source_row, ["exp", "预计", "全年预计"]) or "")
            if current_fy_source is not None:
                values[final_total_header] = format_decimal(convert_value_between_units(current_fy_source, source_unit_spec, output_unit_spec))
                source_trace[final_total_header] = "current_data_request"
            else:
                q_values = [parse_decimal(values[h] or "") for h in rolling_headers[:4]]
                if all(v is not None for v in q_values):
                    values[final_total_header] = format_decimal(sum(q_values))  # type: ignore[arg-type]
                    source_trace[final_total_header] = "derived_from_quarters"
                else:
                    values[final_total_header] = None
                    row_flags.append(f"{final_total_header} is missing and could not be derived deterministically.")
            if estimate_source is not None:
                values[estimate_header] = format_decimal(convert_value_between_units(estimate_source, source_unit_spec, output_unit_spec))
                source_trace[estimate_header] = "current_data_request"
            else:
                values[estimate_header] = None
                row_flags.append(f"{estimate_header} is missing in the current data request.")

        else:
            row_current_dec = parse_decimal(values[target_quarter.display()] or "")
            previous_q_dec = parse_decimal(values[rolling_headers[1]] or "")
            yoy_dec = parse_decimal(values[rolling_headers[4]] or "")
            qoq_value = format_percent(calc_percent(row_current_dec, previous_q_dec))
            yoy_value = format_percent(calc_percent(row_current_dec, yoy_dec))
            if qoq_value is None:
                row_flags.append("QoQ could not be determined from the rolled quarter values.")
            if yoy_value is None:
                row_flags.append("YoY could not be determined from the rolled quarter values.")

            if source_info["label"] == "company_outlets":
                values[final_total_header] = company_outlets_total
                if company_outlets_total is not None:
                    source_trace[final_total_header] = "current_data_request"
                else:
                    row_flags.append(f"{final_total_header} is missing for company-owned outlets.")
            elif source_info["label"] == "jv_outlets":
                values[final_total_header] = jv_outlets_total
                if jv_outlets_total is not None:
                    source_trace[final_total_header] = "current_data_request"
                else:
                    row_flags.append(f"{final_total_header} is missing for JV outlets.")
            values[estimate_header] = None
            row_flags.append(f"{estimate_header} is missing in the current data request.")

        values["QoQ"] = qoq_value
        values["YoY"] = yoy_value

        rows_output.append(
            {
                "label": label,
                "mapped_source_label": source_info["label"] if source_info else None,
                "values": {header: values.get(header) for header in output_headers},
                "source_trace": source_trace,
                "flags": row_flags,
            }
        )

    return {
        "status": "updated_with_flags",
        "unit": prev_financial["unit"] or "in millions USD",
        "columns": output_headers,
        "rows": rows_output,
    }


def build_markdown(
    business_activities: str,
    financial_update: dict[str, Any],
    business_update_bullets: list[str],
    risk_exit_bullets: list[str],
) -> str:
    lines = [
        "# Portfolio Review Draft",
        "",
        "## 1. Business Activities",
        business_activities,
        "",
        "## 2. Financial Update",
        f"({financial_update['unit']})",
        "",
    ]

    headers = [financial_update["unit"]] + financial_update["columns"]
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in financial_update["rows"]:
        row_values = [row["label"]] + [row["values"].get(header) or "" for header in financial_update["columns"]]
        lines.append("| " + " | ".join(row_values) + " |")

    lines.extend(["", "## 3. Business Update"])
    for bullet in business_update_bullets:
        lines.append(f"- {bullet}")

    lines.extend(["", "## 4. Risk & Exit"])
    for bullet in risk_exit_bullets:
        lines.append(f"- {bullet}")

    return "\n".join(lines) + "\n"


def split_evenly(items: list[str], count: int) -> list[list[str]]:
    if count <= 0:
        return []
    if not items:
        return [[] for _ in range(count)]
    base, remainder = divmod(len(items), count)
    chunks: list[list[str]] = []
    cursor = 0
    for idx in range(count):
        size = base + (1 if idx < remainder else 0)
        chunks.append(items[cursor : cursor + size])
        cursor += size
    return chunks


def detect_doc_section_occurrences(doc: Document) -> list[DocSectionOccurrence]:
    occurrences: list[DocSectionOccurrence] = []
    after_financial: dict[str, bool] = {"english": False, "chinese": False}
    after_risk: dict[str, bool] = {"english": False, "chinese": False}
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        language = text_language(text)
        canonical = fixed_anchor_heading(text)
        if canonical:
            occurrences.append(
                DocSectionOccurrence(
                    canonical=canonical,
                    language=language,
                    heading_index=idx,
                )
            )
            if canonical == "Financial Update":
                after_financial[language] = True
                after_risk[language] = False
            elif canonical == "Risk & Exit":
                after_risk[language] = True
            continue

        canonical = canonical_heading(text)
        if canonical == "Business Update":
            occurrences.append(
                DocSectionOccurrence(
                    canonical=canonical,
                    language=language,
                    heading_index=idx,
                )
            )
            continue

        if after_financial[language] and not after_risk[language] and is_middle_section_heading_candidate(paragraph):
            occurrences.append(
                DocSectionOccurrence(
                    canonical="Business Update",
                    language=language,
                    heading_index=idx,
                )
            )
    return occurrences


def resolve_section_plan_to_doc(doc: Document, section_plan: list[SectionOccurrence]) -> list[DocSectionOccurrence]:
    occurrences: list[DocSectionOccurrence] = []
    paragraph_index = 0
    for planned in section_plan:
        target = norm_space(planned.heading_text)
        for idx in range(paragraph_index, len(doc.paragraphs)):
            if norm_space(doc.paragraphs[idx].text) == target:
                occurrences.append(
                    DocSectionOccurrence(
                        canonical=planned.canonical,
                        language=planned.language,
                        heading_index=idx,
                    )
                )
                paragraph_index = idx + 1
                break
    return occurrences


def template_language_mode(occurrences: list[DocSectionOccurrence]) -> str:
    languages = {occ.language for occ in occurrences}
    return "bilingual" if "english" in languages and "chinese" in languages else "English-only"


def middle_section_indices_by_language(occurrences: list[DocSectionOccurrence]) -> dict[str, list[int]]:
    result: dict[str, list[int]] = {"english": [], "chinese": []}
    for language in ["english", "chinese"]:
        language_occurrences = [(idx, occ) for idx, occ in enumerate(occurrences) if occ.language == language]
        financial_idx = next((idx for idx, occ in language_occurrences if occ.canonical == "Financial Update"), None)
        risk_idx = next((idx for idx, occ in language_occurrences if occ.canonical == "Risk & Exit" and financial_idx is not None and idx > financial_idx), None)
        if financial_idx is None:
            continue
        upper = risk_idx if risk_idx is not None else len(occurrences)
        for idx, occ in enumerate(occurrences):
            if occ.language != language:
                continue
            if idx <= financial_idx or idx >= upper:
                continue
            if occ.canonical in {"Business Activities", "Financial Update", "Risk & Exit"}:
                continue
            result[language].append(idx)
    return result


def get_body_paragraphs_for_occurrence(doc: Document, occurrences: list[DocSectionOccurrence], occurrence_index: int):
    start = occurrences[occurrence_index].heading_index + 1
    end = occurrences[occurrence_index + 1].heading_index if occurrence_index + 1 < len(occurrences) else len(doc.paragraphs)
    return doc.paragraphs[start:end]


def set_paragraph_text_preserve(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def capture_run_format(paragraph) -> tuple[str | None, float | None, bool | None]:
    for run in paragraph.runs:
        font_name = run.font.name
        size_pt = float(run.font.size.pt) if run.font.size is not None else None
        bold = run.font.bold
        return font_name, size_pt, bold
    return None, None, None


def apply_paragraph_run_format(paragraph, font_name: str | None, size_pt: float | None, bold: bool | None) -> None:
    for run in paragraph.runs:
        if font_name:
            run.font.name = font_name
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:ascii"), font_name)
            rfonts.set(qn("w:hAnsi"), font_name)
            rfonts.set(qn("w:cs"), font_name)
            rfonts.set(qn("w:eastAsia"), font_name)
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        run.font.bold = bold


def clear_paragraph_preserve(paragraph, normal_style=None) -> None:
    set_paragraph_text_preserve(paragraph, "")
    if normal_style is not None:
        try:
            paragraph.style = normal_style
        except Exception:
            pass


def copy_paragraph_layout(target, source) -> None:
    try:
        target.style = source.style
    except Exception:
        pass
    target.alignment = source.alignment
    src_fmt = source.paragraph_format
    dst_fmt = target.paragraph_format
    dst_fmt.left_indent = src_fmt.left_indent
    dst_fmt.right_indent = src_fmt.right_indent
    dst_fmt.first_line_indent = src_fmt.first_line_indent
    dst_fmt.space_before = src_fmt.space_before
    dst_fmt.space_after = src_fmt.space_after
    dst_fmt.line_spacing = src_fmt.line_spacing
    dst_fmt.line_spacing_rule = src_fmt.line_spacing_rule
    dst_fmt.keep_together = src_fmt.keep_together
    dst_fmt.keep_with_next = src_fmt.keep_with_next
    dst_fmt.page_break_before = src_fmt.page_break_before
    dst_fmt.widow_control = src_fmt.widow_control
    # Copy bullet/numbering properties (numPr) from source to target
    src_pPr = source._element.find(qn("w:pPr"))
    if src_pPr is not None:
        src_numPr = src_pPr.find(qn("w:numPr"))
        if src_numPr is not None:
            tgt_pPr = target._element.get_or_add_pPr()
            old_numPr = tgt_pPr.find(qn("w:numPr"))
            if old_numPr is not None:
                tgt_pPr.remove(old_numPr)
            tgt_pPr.append(copy.deepcopy(src_numPr))


def paragraph_is_list_like(paragraph) -> bool:
    style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
    if "list" in style_name or "列表" in style_name:
        return True
    ppr = getattr(paragraph._element, "pPr", None)
    if ppr is not None and getattr(ppr, "numPr", None) is not None:
        return True
    stripped = paragraph.text.strip()
    return stripped.startswith(("-", "•", "*"))


def fill_section_paragraphs(doc: Document, paragraphs, texts: list[str], force_bullets: bool = False) -> None:
    slots = list(paragraphs)
    normal_style = None
    try:
        normal_style = doc.styles["Normal"]
    except Exception:
        pass

    template_style = None
    template_run_format = (None, None, None)
    for paragraph in slots:
        if paragraph.text.strip():
            template_style = paragraph.style
            template_run_format = capture_run_format(paragraph)
            break
    if template_style is None and slots:
        template_style = slots[0].style
        template_run_format = capture_run_format(slots[0])

    working_slots = slots
    bullet_reference = None
    if force_bullets:
        list_like = [paragraph for paragraph in slots if paragraph_is_list_like(paragraph)]
        if list_like:
            working_slots = list_like
            bullet_reference = list_like[0]
        elif slots:
            working_slots = slots
            bullet_reference = slots[0]

    remove_later = [paragraph for paragraph in slots if paragraph not in working_slots]
    for idx, paragraph in enumerate(working_slots):
        if idx < len(texts):
            set_paragraph_text_preserve(paragraph, texts[idx])
            if force_bullets and bullet_reference is not None:
                copy_paragraph_layout(paragraph, bullet_reference)
            apply_paragraph_run_format(paragraph, *template_run_format)
        else:
            remove_later.append(paragraph)

    if len(texts) <= len(working_slots):
        for paragraph in reversed(remove_later):
            element = paragraph._element
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
        return

    anchor = working_slots[-1] if working_slots else None
    for text in texts[len(working_slots) :]:
        if anchor is None:
            break
        new_paragraph = anchor.insert_paragraph_before(text)
        if force_bullets and bullet_reference is not None:
            copy_paragraph_layout(new_paragraph, bullet_reference)
        elif template_style is not None:
            try:
                new_paragraph.style = template_style
            except Exception:
                pass
        apply_paragraph_run_format(new_paragraph, *template_run_format)
        anchor = new_paragraph

    for paragraph in reversed(remove_later):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//*[local-name()='drawing' or local-name()='pict']"))


def paragraph_has_page_break(paragraph) -> bool:
    for element in paragraph._element.iter():
        if element.tag.endswith("}br") and element.get(qn("w:type")) == "page":
            return True
    return False


def is_plain_blank_paragraph(paragraph) -> bool:
    return paragraph.text.strip() == "" and not paragraph_has_drawing(paragraph) and not paragraph_has_page_break(paragraph)


def find_logo_info(doc: Document, template_path: Path) -> tuple[Path | None, Any | None, Any | None, Any | None]:
    logo_paragraph = None
    for paragraph in doc.paragraphs:
        if paragraph_has_drawing(paragraph):
            logo_paragraph = paragraph
            break
    if logo_paragraph is None:
        return None, None, None, None

    alignment = logo_paragraph.alignment
    style = logo_paragraph.style
    width = None
    height = None
    if doc.inline_shapes:
        shape = doc.inline_shapes[0]
        width = shape.width
        height = shape.height

    with zipfile.ZipFile(template_path) as zf:
        media_names = sorted(name for name in zf.namelist() if name.startswith("word/media/"))
        if not media_names:
            return None, alignment, style, None
        media_name = media_names[0]
        suffix = Path(media_name).suffix or ".png"
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        temp_file.write(zf.read(media_name))
        temp_file.flush()
        temp_file.close()
    return Path(temp_file.name), alignment, style, (width, height)


def insert_page_break_and_repeated_logo(doc: Document, template_path: Path, occurrences: list[DocSectionOccurrence] | None = None) -> None:
    occurrences = occurrences or detect_doc_section_occurrences(doc)
    english_occurrences = [occ for occ in occurrences if occ.language == "english"]
    chinese_occurrences = [occ for occ in occurrences if occ.language == "chinese"]
    if not english_occurrences or not chinese_occurrences:
        return

    english_start = doc.paragraphs[english_occurrences[0].heading_index]
    logo_path, alignment, style, size = find_logo_info(doc, template_path)

    # Insert logo first, then page break before it, so final order is:
    # [page break] [logo] [english heading]
    logo_ref = english_start  # reference for page break insertion
    if logo_path is not None:
        logo_paragraph = english_start.insert_paragraph_before("")
        if style is not None:
            try:
                logo_paragraph.style = style
            except Exception:
                pass
        logo_paragraph.alignment = alignment
        run = logo_paragraph.add_run()
        width, height = size if size is not None else (None, None)
        if width is not None:
            run.add_picture(str(logo_path), width=width, height=height)
        else:
            run.add_picture(str(logo_path))
        logo_ref = logo_paragraph  # insert page break before the logo
        try:
            logo_path.unlink()
        except Exception:
            pass

    break_paragraph = logo_ref.insert_paragraph_before("")
    break_paragraph.add_run().add_break(WD_BREAK.PAGE)


def remove_blank_paragraphs_inside_sections(doc: Document, occurrences: list[DocSectionOccurrence] | None = None) -> None:
    occurrences = occurrences or detect_doc_section_occurrences(doc)
    to_remove = []
    for occ_idx, _occurrence in enumerate(occurrences):
        body = list(get_body_paragraphs_for_occurrence(doc, occurrences, occ_idx))
        for paragraph in body:
            if is_plain_blank_paragraph(paragraph):
                to_remove.append(paragraph)
    for paragraph in reversed(to_remove):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def normalize_bullet_sections(doc: Document, occurrences: list[DocSectionOccurrence] | None = None) -> None:
    occurrences = occurrences or detect_doc_section_occurrences(doc)
    for occ_idx, occurrence in enumerate(occurrences):
        if occurrence.canonical not in {"Business Update", "Risk & Exit"}:
            continue
        body = list(get_body_paragraphs_for_occurrence(doc, occurrences, occ_idx))
        content_paragraphs = [p for p in body if p.text.strip()]
        if not content_paragraphs:
            continue
        bullet_reference = next((p for p in content_paragraphs if paragraph_is_list_like(p)), content_paragraphs[0])
        run_format = capture_run_format(bullet_reference)
        for paragraph in content_paragraphs:
            copy_paragraph_layout(paragraph, bullet_reference)
            apply_paragraph_run_format(paragraph, *run_format)


def normalize_exact_single_blank_between_sections(doc: Document, occurrences: list[DocSectionOccurrence] | None = None) -> None:
    occurrences = occurrences or detect_doc_section_occurrences(doc)
    if not occurrences:
        return

    first_english_heading_index = None
    if any(occ.language == "english" for occ in occurrences) and any(occ.language == "chinese" for occ in occurrences):
        for occ in occurrences:
            if occ.language == "english":
                first_english_heading_index = occ.heading_index
                break

    def remove_paragraph(paragraph) -> None:
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    for idx, occurrence in enumerate(occurrences):
        paragraph = doc.paragraphs[occurrence.heading_index]
        previous = paragraph._element.getprevious()
        while previous is not None:
            prev_paragraph = next((p for p in doc.paragraphs if p._element is previous), None)
            if prev_paragraph is None or not is_plain_blank_paragraph(prev_paragraph):
                break
            remove_paragraph(prev_paragraph)
            previous = paragraph._element.getprevious()

        if idx == 0:
            continue
        if first_english_heading_index is not None and occurrence.heading_index == first_english_heading_index:
            continue

        previous = paragraph._element.getprevious()
        prev_paragraph = next((p for p in doc.paragraphs if p._element is previous), None) if previous is not None else None
        if prev_paragraph is None or not is_plain_blank_paragraph(prev_paragraph):
            blank = paragraph.insert_paragraph_before("")
            try:
                blank.style = doc.styles["Normal"]
            except Exception:
                pass


def center_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")


def write_docx(
    output_path: Path,
    template_path: Path,
    section_plan: list[SectionOccurrence],
    business_activities_map: dict[str, str],
    financial_updates: list[dict[str, Any]],
    business_update_map: dict[str, list[str]],
    risk_exit_map: dict[str, list[str]],
) -> None:
    doc = Document(str(template_path))

    def apply_table_run_format(run, font_name: str | None, size_pt: float | None, bold: bool | None) -> None:
        if font_name:
            run.font.name = font_name
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:ascii"), font_name)
            rfonts.set(qn("w:hAnsi"), font_name)
            rfonts.set(qn("w:cs"), font_name)
            rfonts.set(qn("w:eastAsia"), font_name)
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        run.font.bold = bold

    # --- Fill Business Activities (re-resolve before each major step) ---
    occurrences = resolve_section_plan_to_doc(doc, section_plan)
    for occ_idx, occurrence in enumerate(occurrences):
        body = get_body_paragraphs_for_occurrence(doc, occurrences, occ_idx)
        if occurrence.canonical == "Business Activities":
            fill_section_paragraphs(doc, body, [business_activities_map.get(occurrence.language, business_activities_map.get("english", ""))])

    # --- Fill Business Update middle sections (re-resolve after Business Activities may have shifted indices) ---
    occurrences = resolve_section_plan_to_doc(doc, section_plan)
    update_occurrences_by_language = middle_section_indices_by_language(occurrences)

    for language, occ_indices in update_occurrences_by_language.items():
        if not occ_indices:
            continue
        # Re-resolve fresh for each language group to get current indices
        occurrences = resolve_section_plan_to_doc(doc, section_plan)
        update_occ = middle_section_indices_by_language(occurrences).get(language, [])
        if not update_occ:
            continue
        heading_texts = [doc.paragraphs[occurrences[occ_idx].heading_index].text for occ_idx in update_occ]
        chunks = allocate_middle_section_bullets(heading_texts, business_update_map.get(language, []))
        for occ_idx, texts in zip(update_occ, chunks):
            # Re-resolve before each fill to handle shifting from prior fills
            occurrences = resolve_section_plan_to_doc(doc, section_plan)
            body = get_body_paragraphs_for_occurrence(doc, occurrences, occ_idx)
            fill_section_paragraphs(doc, body, texts, force_bullets=True)

    # Risk & Exit: DO NOT rewrite. The template (previous quarter) already has
    # the correct content. Rewriting via XML extraction and fill_section_paragraphs
    # is lossy (merges paragraphs, creates duplicates, loses formatting).

    financial_occurrences = [occ for occ in occurrences if occ.canonical == "Financial Update"]
    for table_idx, table in enumerate(doc.tables):
        if table_idx >= len(financial_updates):
            break
        center_table(table)
        financial_update = financial_updates[table_idx]
        formatting: dict[tuple[int, int], tuple[str | None, float | None, bool | None]] = {}
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                font_name = None
                size_pt = None
                bold = None
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font_name = run.font.name or font_name
                        if run.font.size is not None:
                            size_pt = float(run.font.size.pt)
                        if run.font.bold is not None:
                            bold = run.font.bold
                        break
                    if font_name or size_pt is not None or bold is not None:
                        break
                formatting[(r_idx, c_idx)] = (font_name, size_pt, bold)

        unit_label = financial_update["unit"]
        top_left_original = norm_space(table.cell(0, 0).text)
        inside_unit = is_unit_text(top_left_original)
        outside_unit_paragraph = None
        if table_idx < len(financial_occurrences):
            occurrence = financial_occurrences[table_idx]
            occ_index = next((i for i, occ in enumerate(occurrences) if occ.heading_index == occurrence.heading_index), None)
            if occ_index is not None:
                for paragraph in get_body_paragraphs_for_occurrence(doc, occurrences, occ_index):
                    if paragraph.text.strip() and is_unit_text(paragraph.text):
                        outside_unit_paragraph = paragraph
                        break

        if outside_unit_paragraph is not None:
            unit_run_format = capture_run_format(outside_unit_paragraph)
            set_paragraph_text_preserve(outside_unit_paragraph, unit_label)
            apply_paragraph_run_format(outside_unit_paragraph, *unit_run_format)
            table.cell(0, 0).text = ""
        elif inside_unit:
            table.cell(0, 0).text = unit_label

        for idx, header in enumerate(financial_update["columns"], start=1):
            if idx < len(table.rows[0].cells):
                table.cell(0, idx).text = header
        for row_idx, row_data in enumerate(financial_update["rows"], start=1):
            if row_idx >= len(table.rows):
                break
            table.cell(row_idx, 0).text = row_data["label"]
            for col_idx, header in enumerate(financial_update["columns"], start=1):
                if col_idx >= len(table.rows[row_idx].cells):
                    break
                table.cell(row_idx, col_idx).text = row_data["values"].get(header) or ""

        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                # Remove extra blank paragraphs inside cells (keeps only the first)
                while len(cell.paragraphs) > 1:
                    last = cell.paragraphs[-1]
                    last._element.getparent().remove(last._element)
                font_name, size_pt, bold = formatting.get((r_idx, c_idx), (None, None, None))
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        apply_table_run_format(run, font_name, size_pt, bold)

    # Use detect_doc_section_occurrences for post-edit passes (safer than resolve
    # which depends on exact text matching against a potentially stale section_plan)
    def safe_occurrences() -> list[DocSectionOccurrence]:
        return [occ for occ in detect_doc_section_occurrences(doc)
                if occ.heading_index < len(doc.paragraphs)]

    occurrences = safe_occurrences()
    remove_blank_paragraphs_inside_sections(doc, occurrences)
    insert_page_break_and_repeated_logo(doc, template_path, occurrences)
    occurrences = safe_occurrences()
    remove_blank_paragraphs_inside_sections(doc, occurrences)
    normalize_bullet_sections(doc, occurrences)
    normalize_exact_single_blank_between_sections(doc, occurrences)

    # Ensure all section headings are bold (after all paragraph shifts are done)
    for occ in safe_occurrences():
        heading_paragraph = doc.paragraphs[occ.heading_index]
        for run in heading_paragraph.runs:
            run.font.bold = True

    doc.save(str(output_path))


def extract_company_name(title: str, business_activities: str) -> str:
    match = re.search(r',\s*([^,]+?) \(the "Company"\)', business_activities)
    if match:
        return norm_space(match.group(1))
    cleaned = re.sub(r"\b20\d{2}\b", "", title)
    cleaned = re.sub(r"\bQ[1-4]\b", "", cleaned, flags=re.I)
    cleaned = re.sub(r"[_-]+", " ", cleaned)
    cleaned = norm_space(cleaned)
    if cleaned:
        return re.sub(r"(?<=[a-z])(?=[A-Z])", " ", cleaned)
    return "The Company"


def normalize_name_tokens(text: str) -> list[str]:
    text = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", text)  # split CamelCase generically
    text = re.sub(r"20\d{2}Q[1-4]", " ", text, flags=re.I)
    text = re.sub(r"20\d{2}", " ", text)
    text = re.sub(r"q[1-4]", " ", text, flags=re.I)
    text = re.sub(r"data request|datarequest|previous quarter template|template|samples|review|portfolio|quarter", " ", text, flags=re.I)
    text = re.sub(r"[^A-Za-z\u4e00-\u9fff]+", " ", text)
    tokens = [t.lower() for t in text.split() if len(t) > 1]
    stop = {"docx", "ventech"}
    return [t for t in tokens if t not in stop]


def jaccard_score(a: set[str], b: set[str]) -> float:
    if not a or not b:
        return 0.0
    return len(a & b) / len(a | b)


def infer_current_company_name(path: Path) -> str:
    blocks = load_docx_blocks(path)
    parsed = parse_current_blocks(blocks)
    title = parsed.get("title", "")
    if title:
        cleaned = re.sub(r"\b20\d{2}\b", "", title)
        cleaned = re.sub(r"\bQ[1-4]\b", "", cleaned, flags=re.I)
        cleaned = re.sub(r"第[一二三四1234]季度.*", "", cleaned)
        cleaned = re.sub(r"\(.*?view.*?\)", "", cleaned, flags=re.I)
        cleaned = re.sub(r"业务概览.*", "", cleaned)
        cleaned = norm_space(cleaned)
        if cleaned:
            return re.sub(r"(?<=[a-z])(?=[A-Z])", " ", cleaned)
    return path.stem


def infer_previous_company_name(path: Path) -> str:
    blocks = load_docx_blocks(path)
    sections = extract_previous_sections(blocks)
    business = " ".join(sections.get("Business Activities", {}).get("paragraphs", []))
    if business:
        name = extract_company_name("", business)
        if name != "The Company":
            return name
    cleaned = re.sub(r"^\d+\s+", "", path.stem)
    cleaned = re.sub(r"\b20\d{2}Q[1-4]\b", "", cleaned, flags=re.I)
    cleaned = norm_space(cleaned)
    return cleaned


def extract_identity_text(path: Path, is_previous: bool) -> str:
    blocks = load_docx_blocks(path)
    if is_previous:
        sections = extract_previous_sections(blocks)
        return " ".join(sections.get("Business Activities", {}).get("paragraphs", []))
    parsed = parse_current_blocks(blocks)
    return " ".join(parsed.get("business_update_paragraphs", [])[:3] + [parsed.get("title", "")])


def match_current_to_previous(current_files: list[Path], previous_files: list[Path]) -> list[dict[str, Any]]:
    previous_meta = []
    for prev in previous_files:
        prev_name = infer_previous_company_name(prev)
        prev_tokens = set(normalize_name_tokens(prev.stem + " " + prev_name))
        prev_identity = set(normalize_name_tokens(extract_identity_text(prev, is_previous=True)))
        prev_quarter = parse_quarter_label(prev.stem) or parse_quarter_label(docx_to_text(prev)[:50])
        previous_meta.append(
            {
                "path": prev,
                "company_name": prev_name,
                "name_tokens": prev_tokens,
                "identity_tokens": prev_identity,
                "quarter": prev_quarter,
            }
        )

    reports = []
    for current in current_files:
        current_name = infer_current_company_name(current)
        current_tokens = set(normalize_name_tokens(current.stem + " " + current_name))
        current_identity = set(normalize_name_tokens(extract_identity_text(current, is_previous=False)))
        current_quarter = parse_quarter_label(current.stem) or parse_quarter_label(docx_to_text(current)[:50])

        scored = []
        for prev in previous_meta:
            score = 0.0
            reasons = []

            name_score = jaccard_score(current_tokens, prev["name_tokens"])
            score += name_score * 0.55
            if name_score > 0:
                reasons.append(f"filename/company tokens overlap ({name_score:.2f})")

            current_name_tokens = set(normalize_name_tokens(current_name))
            prev_name_tokens = set(normalize_name_tokens(prev["company_name"]))
            direct_name_overlap = current_name_tokens & prev_name_tokens
            if direct_name_overlap:
                score += 0.20
                reasons.append(f"direct company-name overlap ({', '.join(sorted(direct_name_overlap))})")

            identity_score = jaccard_score(current_identity, prev["identity_tokens"])
            score += identity_score * 0.30
            if identity_score > 0:
                reasons.append(f"content identity overlap ({identity_score:.2f})")

            if current_quarter and prev["quarter"]:
                if current_quarter.year == prev["quarter"].year and current_quarter.quarter == prev["quarter"].quarter + 1:
                    score += 0.15
                    reasons.append("quarter progression aligns (current is next quarter)")
                elif current_quarter.year == prev["quarter"].year and current_quarter.quarter != prev["quarter"].quarter:
                    score += 0.05
                    reasons.append("same year quarter reference present")

            scored.append((score, prev, reasons))

        scored.sort(key=lambda x: x[0], reverse=True)
        best_score, best_prev, best_reasons = scored[0]
        second_score = scored[1][0] if len(scored) > 1 else 0.0
        gap = best_score - second_score

        if best_score >= 0.65 and gap >= 0.15:
            confidence = "high"
        elif best_score >= 0.45 and gap >= 0.08:
            confidence = "medium"
        else:
            confidence = "ambiguous / review needed"

        reports.append(
            {
                "current_file": str(current.relative_to(ROOT)),
                "matched_previous_file": str(best_prev["path"].relative_to(ROOT)),
                "confidence": confidence,
                "score": round(best_score, 3),
                "explanation": "; ".join(best_reasons) if best_reasons else "weak textual similarity only",
            }
        )
    return reports


def run_matching_report() -> None:
    current_files = list_real_docx(DATA_REQUEST_DIR)
    previous_files = list_real_docx(PREVIOUS_REVIEW_DIR)
    reports = match_current_to_previous(current_files, previous_files)
    print("Matching Report")
    print("")
    for item in reports:
        print(f"Current: {item['current_file']}")
        print(f"Matched previous: {item['matched_previous_file']}")
        print(f"Confidence: {item['confidence']}")
        print(f"Why: {item['explanation']}")
        print("")


def generate_review_for_pair(current_path: Path, previous_path: Path) -> dict[str, Any]:
    OUTPUT_DIR.mkdir(exist_ok=True)
    DEBUG_OUTPUT_DIR.mkdir(exist_ok=True)

    spec_text = docx_to_text(UNIVERSAL_SPEC)
    prev_blocks = load_docx_blocks(previous_path)
    current_blocks = load_docx_blocks(current_path)
    previous_sections = extract_previous_sections(prev_blocks)
    previous_occurrences = extract_previous_section_occurrences(prev_blocks)
    previous_financial_occurrences = [occ for occ in previous_occurrences if occ.canonical == "Financial Update" and occ.tables]
    previous_tables = []
    for occurrence in previous_financial_occurrences:
        unit_label, _ = detect_financial_section_unit(occurrence)
        previous_tables.append(parse_previous_financial_table_with_unit(occurrence.tables[0], unit_label))
    current_parsed = parse_current_blocks(current_blocks)
    current_finance = normalize_finance_source(current_parsed["finance_table"])
    current_unit_spec = detect_current_unit_spec(current_blocks)
    target_quarter = determine_target_quarter(current_path, current_parsed["title"], current_finance)
    previous_review_quarter = parse_quarter_label(previous_path.stem)

    review_flags: list[dict[str, Any]] = []
    source_language, language_evidence = detect_language(docx_to_text(current_path))
    company_name = infer_previous_company_name(previous_path)
    final_company_name = normalize_company_filename(company_name)
    final_quarter_label = quarter_filename_label(target_quarter)
    final_docx_path = OUTPUT_DIR / f"{final_company_name}_{final_quarter_label}.docx"
    json_output = DEBUG_OUTPUT_DIR / f"{final_company_name}_{final_quarter_label}.json"
    markdown_output = DEBUG_OUTPUT_DIR / f"{final_company_name}_{final_quarter_label}.md"

    financial_updates = [
        build_financial_update(
            prev_financial,
            current_finance,
            current_parsed.get("operation_table"),
            review_flags,
            target_quarter,
            current_unit_spec,
            detect_unit_spec(prev_financial.get("unit", "")),
            previous_review_quarter,
        )
        for prev_financial in previous_tables
    ]

    business_activity_map: dict[str, str] = {}
    risk_exit_map: dict[str, list[str]] = {}
    previous_business_update_map: dict[str, list[str]] = {}
    for occurrence in previous_occurrences:
        if occurrence.canonical == "Business Activities" and occurrence.paragraphs:
            business_activity_map.setdefault(occurrence.language, " ".join(occurrence.paragraphs).strip())
        elif occurrence.canonical == "Risk & Exit" and occurrence.paragraphs:
            if occurrence.language not in risk_exit_map:
                seen = set()
                deduped = []
                for p in occurrence.paragraphs:
                    key = norm_space(p).lower()
                    if key not in seen:
                        seen.add(key)
                        deduped.append(p)
                risk_exit_map[occurrence.language] = deduped
        elif occurrence.canonical == "Business Update" and occurrence.paragraphs:
            previous_business_update_map.setdefault(occurrence.language, [])
            previous_business_update_map[occurrence.language].extend(occurrence.paragraphs)

    business_update_map: dict[str, list[str]] = {}
    if business_activity_map.get("english"):
        business_update_map["english"] = build_business_update_bullets(current_parsed["business_update_paragraphs"], company_name)
    if business_activity_map.get("chinese"):
        business_update_map["chinese"] = build_chinese_business_update_bullets(current_parsed["business_update_paragraphs"])
    if source_language == "chinese" and business_activity_map.get("english") and not clean_update_paragraphs([p for p in current_parsed["business_update_paragraphs"] if text_language(p) == "english"]):
        business_update_map["english"] = previous_business_update_map.get("english", business_update_map.get("english", []))

    if not business_update_map.get("english") and previous_business_update_map.get("english"):
        business_update_map["english"] = previous_business_update_map["english"]
    if not business_update_map.get("chinese") and previous_business_update_map.get("chinese"):
        business_update_map["chinese"] = previous_business_update_map["chinese"]

    if not previous_review_has_logo(previous_path):
        review_flags.append(
            {
                "id": "logo-missing",
                "section": "Business Activities",
                "severity": "warning",
                "message": "No reusable company logo was found in the previous review.",
                "source": "previous_review",
            }
        )

    for table_update in financial_updates:
        for row in table_update["rows"]:
            for idx, message in enumerate(row["flags"], start=1):
                review_flags.append(
                    {
                        "id": f"{row['label'].lower().replace('#', 'num').replace(' ', '-')}-flag-{idx}",
                        "section": "Financial Update",
                        "severity": "warning" if "missing" in message.lower() or "could not" in message.lower() else "info",
                        "message": f"{row['label']}: {message}",
                        "source": "current_data_request" if "current data request" in message.lower() else "previous_review",
                    }
                )

    template_occurrences = detect_doc_section_occurrences(Document(str(previous_path)))
    language_mode = template_language_mode(template_occurrences)
    primary_financial_update = financial_updates[-1] if financial_updates else {}
    primary_business_activities = business_activity_map.get("english") or next(iter(business_activity_map.values()), "")
    primary_risk = risk_exit_map.get("english") or next(iter(risk_exit_map.values()), [])
    primary_business_update = business_update_map.get("english") or next(iter(business_update_map.values()), [])

    output_payload = {
        "prototype_version": "v1",
        "test_case": {
            "company": company_name,
            "target_quarter": target_quarter.display(),
            "files": {
                "universal_spec": {"path": str(UNIVERSAL_SPEC)},
                "current_data_request": {"path": str(current_path)},
                "previous_review": {"path": str(previous_path)},
            },
        },
        "language_detection": {
            "source_language": source_language,
            "template_language_mode": language_mode,
            "evidence": language_evidence,
        },
        "extracted_inputs": {
            "universal_spec": {"found": True, "raw_text": spec_text},
            "previous_review": {
                "business_activities": {"found": True, "raw_text": primary_business_activities},
                "financial_update": {"found": True, "parsed_table": primary_financial_update},
                "business_update": {"found": True, "raw_bullets": previous_business_update_map.get("english") or previous_business_update_map.get("chinese", [])},
                "risk_exit": {"found": True, "raw_bullets": primary_risk},
            },
            "current_data_request": {
                "financial_source": {"found": True, "raw_table": current_parsed["finance_table"], "parsed_metrics": current_finance},
                "operation_source": {"found": current_parsed.get("operation_table") is not None, "raw_table": current_parsed.get("operation_table")},
                "business_update_source": {"found": True, "source_paragraphs": current_parsed["business_update_paragraphs"]},
            },
        },
        "proposed_outputs": {
            "business_activities": {"status": "carried_forward", "text": primary_business_activities},
            "financial_update": primary_financial_update,
            "business_update": {"status": "drafted_from_current_data", "bullets": primary_business_update, "source_trace": current_parsed["business_update_paragraphs"]},
            "risk_exit": {"status": "carried_forward", "bullets": primary_risk},
        },
        "review_flags": review_flags,
    }

    json_output.write_text(json.dumps(output_payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    markdown_output.write_text(
        build_markdown(primary_business_activities, primary_financial_update, primary_business_update, primary_risk),
        encoding="utf-8",
    )
    write_docx(final_docx_path, previous_path, previous_occurrences, business_activity_map, financial_updates, business_update_map, risk_exit_map)

    return {
        "current_file": str(current_path.relative_to(ROOT)),
        "matched_previous_file": str(previous_path.relative_to(ROOT)),
        "output_file": str(final_docx_path.relative_to(ROOT)),
        "language_mode": language_mode,
        "company_naming_preserved": "yes",
    }


def debug_generation_for_pair(current_path: Path, previous_path: Path) -> None:
    output_path = OUTPUT_DIR / f"{normalize_company_filename(infer_previous_company_name(previous_path))}_{quarter_filename_label(determine_target_quarter(current_path, parse_current_blocks(load_docx_blocks(current_path))['title'], normalize_finance_source(parse_current_blocks(load_docx_blocks(current_path))['finance_table'])))}.docx"
    existed_before = output_path.exists()
    before_mtime = datetime.fromtimestamp(output_path.stat().st_mtime).astimezone().isoformat(timespec="seconds") if existed_before else None

    prev_blocks = load_docx_blocks(previous_path)
    current_blocks = load_docx_blocks(current_path)
    previous_occurrences = extract_previous_section_occurrences(prev_blocks)
    previous_financial_occurrences = [occ for occ in previous_occurrences if occ.canonical == "Financial Update" and occ.tables]
    current_parsed = parse_current_blocks(current_blocks)
    current_finance = normalize_finance_source(current_parsed["finance_table"])
    target_quarter, quarter_debug = determine_target_quarter_with_debug(current_path, current_parsed["title"], current_finance)
    current_unit_spec = detect_current_unit_spec(current_blocks)
    previous_review_quarter = parse_quarter_label(previous_path.stem)

    print("DEBUG_TARGET_CURRENT_FILE", str(current_path.resolve()))
    print("DEBUG_TARGET_PREVIOUS_FILE", str(previous_path.resolve()))
    print("DEBUG_OUTPUT_PATH", str(output_path.resolve()))
    print("DEBUG_OUTPUT_EXISTED_BEFORE", existed_before)
    print("DEBUG_OUTPUT_MTIME_BEFORE", before_mtime or "missing")
    print("DEBUG_CURRENT_SOURCE_UNIT_LABEL", current_unit_spec.label if current_unit_spec else "None")
    print("DEBUG_CURRENT_SOURCE_UNIT_FACTOR_TO_BASE", str(current_unit_spec.factor_to_base) if current_unit_spec else "None")
    print("DEBUG_TARGET_QUARTER", target_quarter.display())
    print("DEBUG_TARGET_QUARTER_PATH", quarter_debug)
    print("DEBUG_CURRENT_FINANCE_HEADERS", json.dumps(current_finance["header"], ensure_ascii=False))

    for idx, occurrence in enumerate(previous_financial_occurrences):
        unit_label, placement = detect_financial_section_unit(occurrence)
        prev_financial = parse_previous_financial_table_with_unit(occurrence.tables[0], unit_label)
        output_unit_spec = detect_unit_spec(prev_financial.get("unit", ""))
        conversion_factor = None
        if current_unit_spec and output_unit_spec:
            conversion_factor = current_unit_spec.factor_to_base / output_unit_spec.factor_to_base
        historical_headers = [header for header in prev_financial["headers"] if parse_quarter_label(header)]
        normalized_historical = normalize_previous_quarter_headers(prev_financial["headers"], previous_review_quarter)
        flags: list[dict[str, Any]] = []
        financial_update = build_financial_update(
            prev_financial,
            current_finance,
            current_parsed.get("operation_table"),
            flags,
            target_quarter,
            current_unit_spec,
            output_unit_spec,
            previous_review_quarter,
        )

        print(f"DEBUG_TABLE_{idx}_UNIT_PLACEMENT", placement)
        print(f"DEBUG_TABLE_{idx}_TEMPLATE_UNIT_LABEL", unit_label or "None")
        print(f"DEBUG_TABLE_{idx}_TEMPLATE_OUTPUT_UNIT_FACTOR_TO_BASE", str(output_unit_spec.factor_to_base) if output_unit_spec else "None")
        print(f"DEBUG_TABLE_{idx}_CONVERSION_FACTOR_APPLIED", str(conversion_factor) if conversion_factor is not None else "None")
        print(f"DEBUG_TABLE_{idx}_PARSED_HISTORICAL_HEADERS_RAW", json.dumps(historical_headers, ensure_ascii=False))
        print(f"DEBUG_TABLE_{idx}_NORMALIZED_HISTORICAL_HEADERS", json.dumps(normalized_historical, ensure_ascii=False))
        print(f"DEBUG_TABLE_{idx}_ROLLED_HEADERS", json.dumps(financial_update['columns'], ensure_ascii=False))
        for row in financial_update["rows"]:
            print(
                f"DEBUG_TABLE_{idx}_ROW_{row['label']}",
                json.dumps({"mapped_source_label": row.get("mapped_source_label"), "values": row["values"]}, ensure_ascii=False),
            )
        if flags:
            print(f"DEBUG_TABLE_{idx}_FLAGS", json.dumps(flags, ensure_ascii=False))
        else:
            print(f"DEBUG_TABLE_{idx}_FLAGS", "[]")

    result = generate_review_for_pair(current_path, previous_path)
    regenerated_output = ROOT / result["output_file"]
    existed_after = regenerated_output.exists()
    after_mtime = datetime.fromtimestamp(regenerated_output.stat().st_mtime).astimezone().isoformat(timespec="seconds") if existed_after else None
    print("DEBUG_OUTPUT_PATH_AFTER_SAVE", str(regenerated_output.resolve()))
    print("DEBUG_OUTPUT_OVERWRITE_BEHAVIOR", "overwritten existing file" if existed_before and regenerated_output == output_path else "created new file" if not existed_before else "wrote different path")
    print("DEBUG_OUTPUT_MTIME_AFTER", after_mtime or "missing")

    saved_doc = Document(str(regenerated_output))
    for idx, table in enumerate(saved_doc.tables):
        print(f"DEBUG_SAVED_TABLE_{idx}_ROW_0", json.dumps([cell.text for cell in table.rows[0].cells], ensure_ascii=False))
        for row_idx, row in enumerate(table.rows[1:], start=1):
            print(f"DEBUG_SAVED_TABLE_{idx}_ROW_{row_idx}", json.dumps([cell.text for cell in row.cells], ensure_ascii=False))


def run_batch_generation() -> None:
    current_files = list_real_docx(DATA_REQUEST_DIR)
    previous_files = list_real_docx(PREVIOUS_REVIEW_DIR)
    reports = match_current_to_previous(current_files, previous_files)
    clean_main_output_folder_for_batch()

    print("| Current data request | Matched previous review | Output file generated | Language mode used | Company naming preserved |")
    print("| --- | --- | --- | --- | --- |")
    for item in reports:
        current_path = ROOT / item["current_file"]
        previous_path = ROOT / item["matched_previous_file"]
        result = generate_review_for_pair(current_path, previous_path)
        print(
            f"| {Path(result['current_file']).name} | {Path(result['matched_previous_file']).name} | {Path(result['output_file']).name} | {result['language_mode']} | {result['company_naming_preserved']} |"
        )


def main() -> None:
    current_files = list_real_docx(DATA_REQUEST_DIR)
    previous_files = list_real_docx(PREVIOUS_REVIEW_DIR)
    if not current_files:
        print(f"ERROR: No .docx files found in {DATA_REQUEST_DIR}")
        sys.exit(1)
    if not previous_files:
        print(f"ERROR: No .docx files found in {PREVIOUS_REVIEW_DIR}")
        sys.exit(1)
    if not UNIVERSAL_SPEC.exists():
        print(f"WARNING: Universal Spec not found at {UNIVERSAL_SPEC}")
    reports = match_current_to_previous(current_files, previous_files)
    OUTPUT_DIR.mkdir(exist_ok=True)
    for item in reports:
        current_path = ROOT / item["current_file"]
        previous_path = ROOT / item["matched_previous_file"]
        print(f"Processing: {current_path.name} <-> {previous_path.name} (confidence: {item['confidence']})")
        result = generate_review_for_pair(current_path, previous_path)
        print(f"  -> {result['output_file']}")


if __name__ == "__main__":
    if "--match-report" in sys.argv:
        run_matching_report()
    elif "--generate-matched" in sys.argv:
        run_batch_generation()
    elif "--debug-company" in sys.argv:
        idx = sys.argv.index("--debug-company")
        keyword = sys.argv[idx + 1]
        reports = match_current_to_previous(list_real_docx(DATA_REQUEST_DIR), list_real_docx(PREVIOUS_REVIEW_DIR))
        matched = None
        for item in reports:
            if keyword.lower() in item["current_file"].lower() or keyword.lower() in item["matched_previous_file"].lower():
                matched = item
                break
        if matched is None:
            raise SystemExit(f"No matched pair found for keyword: {keyword}")
        debug_generation_for_pair(ROOT / matched["current_file"], ROOT / matched["matched_previous_file"])
    else:
        main()
