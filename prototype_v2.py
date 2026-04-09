#!/usr/bin/env python3
from __future__ import annotations

import copy
import json
import re
import sys
import tempfile
import unicodedata
import zipfile
from dataclasses import dataclass
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is required. Install it with: pip install python-docx")
    sys.exit(1)
try:
    import pypdf
except ImportError:
    pypdf = None
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Pt, Emu, Twips
from docx.oxml.ns import qn


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
ROOT = Path(__file__).resolve().parent


UNIVERSAL_SPEC = ROOT / "Universal Spec .docx"
DATA_REQUEST_DIR = ROOT / "DataRequest Samples"
PREVIOUS_REVIEW_DIR = ROOT / "Previous Quarter Template"
OUTPUT_DIR = ROOT / "prototype_output"
DEBUG_OUTPUT_DIR = OUTPUT_DIR / "_debug"


SECTION_HEADINGS = [
    "Business Activities",
    "Financial Update",
    "Business Update",
    "Risk & Exit",
]

SECTION_ALIASES = {
    "Business Activities": [
        "business activities",
        "business information",
        "business introduction",
        "company overview",
        "公司介绍",
    ],
    "Financial Update": [
        "financial update",
        "financial updates",
        "financial updates ",
        "financial update",
        "financial updates (in usd million)",
        "financial update",
        "财务情况",
    ],
    "Business Update": [
        "business update",
        "business highlights",
        "recent trends",
        "operational status",
        "operations update",
        "operating status",
        "business progress",
        "运营情况",
        "业务进展",
    ],
    "Risk & Exit": [
        "risk & exit",
        "risk and exit",
        "risks & exit",
        "risks and exit",
    ],
}


def is_question_or_prompt(text: str) -> bool:
    """Detect question/prompt paragraphs generically instead of hardcoding."""
    stripped = norm_space(text)
    if not stripped:
        return False
    if re.search(r"[?？]$", stripped):
        return True
    prompt_patterns = [
        r"^please\b",
        r"^what is\b",
        r"^what are\b",
        r"^how does\b",
        r"^how do\b",
        r"^how much\b",
        r"^how many\b",
        r"^can you\b",
        r"^could you\b",
        r"^would you\b",
        r"^describe\b",
        r"^explain\b",
        r"^elaborate\b",
        r"^provide\b",
        r"^list\b.*\b(?:key|main|major|top)\b",
        r"^note\b.*\bany\b",
        r"^project\b.*\bbalance\b",
        r"^if there is a (?:significant|major)\b.*\bplease\b",
        r"^if there (?:is|are)\b.*\bplease\b",
        r"\bplease provide\b.*\bexplanation\b",
        r"\bplease\b.*\b(?:describe|explain|elaborate|provide|list|detail|share|clarify|specify|confirm)\b",
        r"^see notes to\b",
        r">\s*see notes to\b",
        r"^what'?s the\b",
        r"^when (?:is|will|did)\b",
        r"^where (?:is|are)\b",
        r"^why (?:is|are|did)\b",
        r"^is there\b",
        r"^are there\b",
        r"^do you\b",
        r"^have you\b",
        r"^does the\b",
        r"^did the\b",
        r"^will the\b",
        r"^has the\b",
        # Chinese question/prompt patterns (without question marks)
        r"^请",
        r"^【重要】请",
        r"如有较大的.*请解释",
        r"^请问",
        r"^是否",
        r"^能否",
        r"^有没有",
        r"^有无",
        r"请说明",
        r"请提供",
        r"请描述",
        r"请列出",
        r"请解释",
        r"请详细",
        r"请补充",
        r"请确认",
        r"^如.*请详述",
        r"^如.*请解释",
    ]
    # Normalize curly quotes to straight for matching
    lowered = stripped.lower().replace("\u2018", "'").replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"')
    if any(re.search(p, lowered) for p in prompt_patterns):
        return True
    # Detect data-request prompts: short topic + parenthetical list of sub-topics
    # e.g. "新一轮融资的计划和进展（预计时间线，融资额度，估值，潜在投资方）。"
    # Pattern: Chinese text ending with (list of Chinese comma-separated items)。
    if re.search(r"[（(][^）)]*[，,][^）)]*[，,][^）)]*[）)]", stripped):
        # Has a parenthetical with 3+ comma-separated items → likely a prompt
        # But only flag if there are no numbers (actual data would have numbers)
        if not re.search(r"\d", stripped.replace("Q1", "").replace("Q2", "").replace("Q3", "").replace("Q4", "")):
            return True
    # Detect topic labels that just name an area without providing data
    # e.g. "现阶段重组进展，卡睿业务的进展。" — just listing topics with abstract words
    if re.search(r"[\u4e00-\u9fff]", stripped) and len(stripped) < 40:
        if not re.search(r"\d", stripped):
            # Ends with abstract topic words (进展/情况/计划/方案) without concrete details
            if re.search(r"(?:的进展|的情况|的计划|的方案|的安排)[。.]?$", stripped):
                return True
    if re.search(r"[\u4e00-\u9fff]", stripped):
        if re.search(r"(?:如何|是什么|有哪些|时间线|结果如何|进展如何|最新进展是什么|情况如何)[。.]?$", stripped):
            return True
        if re.search(r"(?:足够支持多久|能支持多久|支持多久的运营)[。.]?$", stripped):
            return True
    return False


def is_standalone_label(text: str) -> bool:
    """Detect short standalone labels/headers that aren’t real content."""
    stripped = norm_space(text)
    if not stripped:
        return False
    if stripped.endswith((":", "：")) and len(stripped) < 80:
        return True
    if re.match(r"^\d+[.)\s]", stripped) and len(stripped) < 60:
        return True
    section_label_patterns = [
        r"^other\s+questions",
        r"^operational\s+data",
        r"^业务发展$",
    ]
    lowered = stripped.lower()
    return any(re.search(p, lowered) for p in section_label_patterns)


def norm_space(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


def is_unit_text(text: str) -> bool:
    normalized = norm_space(text).lower()
    if not normalized:
        return False
    return any(
        token in normalized
        for token in [
            "unit",
            "amount in",
            "in usd",
            "in rmb",
            "usd million",
            "rmb mn",
            "单位",
            "百万",
            "万元",
            "人民币",
        ]
    )


def cell_text(tc: ET.Element) -> str:
    parts: list[str] = []
    for p in tc.findall("w:p", NS):
        text = "".join(t.text or "" for t in p.findall(".//w:t", NS))
        if text.strip():
            parts.append(norm_space(text))
    return " ".join(parts).strip()


def paragraph_text(p: ET.Element) -> str:
    return norm_space("".join(t.text or "" for t in p.findall(".//w:t", NS)))


def load_docx_blocks(path: Path) -> list[dict[str, Any]]:
    with zipfile.ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    body = root.find("w:body", NS)
    assert body is not None

    blocks: list[dict[str, Any]] = []
    for child in body:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            text = paragraph_text(child)
            if text:
                blocks.append({"type": "paragraph", "text": text})
        elif tag == "tbl":
            rows = []
            for tr in child.findall("w:tr", NS):
                rows.append([cell_text(tc) for tc in tr.findall("w:tc", NS)])
            blocks.append({"type": "table", "rows": rows})
    return blocks


PDF_TABLE_VALUE_RE = re.compile(r"(?:−|-)?\d+(?:,\d{3})*(?:\.\d+)?%?|N/?A|/")


def extract_pdf_lines(path: Path) -> list[str]:
    if pypdf is None:
        raise RuntimeError("PDF support requires pypdf to be installed.")
    reader = pypdf.PdfReader(str(path))
    lines: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        for raw in text.splitlines():
            cleaned = norm_space(raw)
            if cleaned:
                lines.append(cleaned)
    return lines


def is_numbered_section_line(text: str) -> bool:
    return bool(re.match(r"^\d+\s*[.)、．]?\s*", norm_space(text)))


def extract_pdf_header_tokens(lines: list[str]) -> list[str]:
    combined = " ".join(lines)
    patterns = [
        r"\d{4}/\d{1,2}",
        r"\d{1,2}\s*月",
        r"\d{4}Q[1-4]\s*(?:合计|总计)?",
        r"\d{4}\s*(?:合计|总计)",
        r"\d{4}\s*年\s*预计",
        r"Q[1-4]\s*20\d{2}",
        r"20\d{2}\s*YTD",
        r"20\d{2}E",
        r"QoQ",
        r"YoY",
    ]
    matches: list[tuple[int, int, str]] = []
    for pattern in patterns:
        for match in re.finditer(pattern, combined, flags=re.I):
            matches.append((match.start(), match.end(), norm_space(match.group(0))))
    matches.sort(key=lambda item: item[0])
    headers: list[str] = []
    seen: set[tuple[int, int]] = set()
    for start, end, token in matches:
        if (start, end) in seen:
            continue
        seen.add((start, end))
        if not headers or headers[-1] != token:
            headers.append(token)
    return headers


def parse_pdf_table_lines(lines: list[str]) -> list[list[str]] | None:
    if not lines:
        return None
    header_idx = next((idx for idx, line in enumerate(lines) if "日期" in line), None)
    if header_idx is None:
        return None
    data_start = None
    for idx in range(header_idx + 1, len(lines)):
        line = lines[idx]
        if PDF_TABLE_VALUE_RE.search(line) and not line.startswith("截至"):
            data_start = idx
            break
    if data_start is None:
        return None
    headers = extract_pdf_header_tokens(lines[header_idx:data_start])
    if not headers:
        return None
    rows: list[list[str]] = [[""] + headers]
    for line in lines[data_start:]:
        if is_numbered_section_line(line) or line.startswith("截至"):
            break
        match = PDF_TABLE_VALUE_RE.search(line)
        if match is None:
            continue
        label = norm_space(line[: match.start()])
        values = [norm_space(token) for token in PDF_TABLE_VALUE_RE.findall(line[match.start() :])]
        if not label:
            continue
        while len(values) < len(headers):
            values.append("")
        rows.append([label] + values[: len(headers)])
    return rows if len(rows) > 1 else None


def parse_pdf_balance_table(lines: list[str]) -> list[list[str]] | None:
    start = next((idx for idx, line in enumerate(lines) if line.startswith("截至")), None)
    if start is None:
        return None
    rows = [["", lines[start]]]
    idx = start + 1
    while idx < len(lines):
        line = lines[idx]
        if is_numbered_section_line(line):
            break
        if "银行账户现金余额" in line:
            value_parts = []
            remainder = norm_space(line.replace("银行账户现金余额", "", 1))
            if remainder:
                value_parts.append(remainder)
            j = idx + 1
            while j < len(lines) and not is_numbered_section_line(lines[j]) and ("美元" in lines[j] or "人民币" in lines[j]):
                value_parts.append(lines[j])
                j += 1
            rows.append(["银行账户现金余额", " ".join(value_parts)])
            idx = j
            continue
        if "总员工人数" in line:
            value_parts = []
            remainder = norm_space(line.replace("总员工人数", "", 1))
            if remainder:
                value_parts.append(remainder)
            j = idx + 1
            while j < len(lines) and not is_numbered_section_line(lines[j]) and ("兼职" in lines[j] or "全职" in lines[j]):
                value_parts.append(lines[j])
                j += 1
            rows.append(["总员工人数", " ".join(value_parts)])
            idx = j
            continue
        idx += 1
    return rows if len(rows) > 1 else None


def load_pdf_blocks(path: Path) -> list[dict[str, Any]]:
    lines = extract_pdf_lines(path)
    if not lines:
        return []

    blocks: list[dict[str, Any]] = [{"type": "paragraph", "text": lines[0]}]

    section_positions = [idx for idx, line in enumerate(lines) if idx > 0 and is_numbered_section_line(line)]
    section_positions.append(len(lines))

    finance_section = None
    operation_section = None
    update_start = None
    for idx in range(len(section_positions) - 1):
        start = section_positions[idx]
        end = section_positions[idx + 1]
        heading = lines[start]
        if finance_section is None and ("财务数据" in heading or "finance" in heading.lower()):
            finance_section = lines[start:end]
        elif operation_section is None and ("运营数据" in heading or "经营数据" in heading):
            operation_section = lines[start:end]
        elif update_start is None and ("业务发展" in heading or "业务进展" in heading or "business" in heading.lower()):
            update_start = start

    if finance_section:
        finance_table = parse_pdf_table_lines(finance_section)
        if finance_table:
            blocks.append({"type": "table", "rows": finance_table})
        balance_table = parse_pdf_balance_table(finance_section)
        if balance_table:
            blocks.append({"type": "table", "rows": balance_table})

    if operation_section:
        operation_table = parse_pdf_table_lines(operation_section)
        if operation_table:
            blocks.append({"type": "table", "rows": operation_table})

    if update_start is not None:
        for line in lines[update_start + 1 :]:
            if is_numbered_section_line(line):
                continue
            blocks.append({"type": "paragraph", "text": line})

    return blocks


def load_blocks(path: Path) -> list[dict[str, Any]]:
    if path.suffix.lower() == ".pdf":
        return load_pdf_blocks(path)
    return load_docx_blocks(path)


def docx_to_text(path: Path) -> str:
    blocks = load_blocks(path)
    lines = []
    for block in blocks:
        if block["type"] == "paragraph":
            lines.append(block["text"])
        else:
            for row in block["rows"]:
                lines.append(" | ".join(row))
    return "\n".join(lines)


def detect_language(text: str) -> tuple[str, str]:
    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
    if chinese_chars > 20:
        return "chinese", f"Detected {chinese_chars} Chinese characters in the company-submitted source."
    return "english", "No meaningful Chinese text detected in the company-submitted source."


def text_language(text: str) -> str:
    return "chinese" if re.search(r"[\u4e00-\u9fff]", text) else "english"


def canonical_heading(text: str) -> str | None:
    lowered = norm_space(text).lower()
    lowered = re.sub(r"^[0-9一二三四五六七八九十]+[\-–、.\s]*", "", lowered)
    for canonical, aliases in SECTION_ALIASES.items():
        for alias in aliases:
            if lowered == alias or lowered.startswith(alias):
                return canonical
    if lowered.startswith("operations update") or lowered.startswith("business progress"):
        return "Business Update"
    if "风险" in lowered and "退出" in lowered:
        return "Risk & Exit"
    return None


def fixed_anchor_heading(text: str) -> str | None:
    lowered = norm_space(text).lower()
    lowered = re.sub(r"^[0-9一二三四五六七八九十]+[\-–、.\s]*", "", lowered)
    for canonical in ["Business Activities", "Financial Update", "Risk & Exit"]:
        for alias in SECTION_ALIASES.get(canonical, []):
            if lowered == alias or lowered.startswith(alias):
                return canonical
    if "风险" in lowered and "退出" in lowered:
        return "Risk & Exit"
    return None


def is_middle_section_heading_candidate(paragraph) -> bool:
    text = norm_space(paragraph.text)
    if not text:
        return False
    if text.endswith((".", "。", ";", "；", ":", "：", "?", "？", ")", "）")):
        return False
    if len(text) > 80:
        return False
    if text.startswith(("-", "•", "*", "(", "（")):
        return False
    # Unit text should never be a section heading
    if is_unit_text(text):
        return False
    style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
    if "list" in style_name or "heading" in style_name or "标题" in style_name or style_name == "normal":
        return True
    return False


def extract_previous_sections(blocks: list[dict[str, Any]]) -> dict[str, Any]:
    sections: dict[str, Any] = {}
    current_heading: str | None = None

    for block in blocks:
        if block["type"] == "paragraph":
            heading = canonical_heading(block["text"])
        else:
            heading = None
        if heading:
            current_heading = heading
            sections[current_heading] = {"paragraphs": [], "tables": []}
            continue
        if not current_heading:
            continue
        if block["type"] == "paragraph":
            sections[current_heading]["paragraphs"].append(block["text"])
        elif block["type"] == "table":
            sections[current_heading]["tables"].append(block["rows"])
    return sections


def extract_previous_section_occurrences(blocks: list[dict[str, Any]]) -> list[SectionOccurrence]:
    occurrences: list[SectionOccurrence] = []
    current: SectionOccurrence | None = None
    for block in blocks:
        heading = canonical_heading(block["text"]) if block["type"] == "paragraph" else None
        if heading:
            current = SectionOccurrence(
                canonical=heading,
                language=text_language(block["text"]),
                heading_text=block["text"],
                paragraphs=[],
                tables=[],
            )
            occurrences.append(current)
            continue
        if current is None:
            continue
        if block["type"] == "paragraph":
            current.paragraphs.append(block["text"])
        elif block["type"] == "table":
            current.tables.append(block["rows"])
    return occurrences


def parse_previous_financial_table(rows: list[list[str]]) -> dict[str, Any]:
    header_row = rows[0]
    unit = header_row[0]
    headers = [norm_space(cell) for cell in header_row[1:]]
    data_rows = []
    for row in rows[1:]:
        label = norm_space(row[0])
        values = {headers[i]: norm_space(row[i + 1]) for i in range(len(headers))}
        data_rows.append({"label": label, "values": values})
    return {"unit": unit, "headers": headers, "rows": data_rows}


def parse_previous_financial_table_with_unit(rows: list[list[str]], unit_label: str | None) -> dict[str, Any]:
    parsed = parse_previous_financial_table(rows)
    if unit_label:
        parsed["unit"] = unit_label
    return parsed


@dataclass(frozen=True, order=True)
class Quarter:
    year: int
    quarter: int

    def display(self) -> str:
        return f"Q{self.quarter} {self.year}"

    def previous(self) -> "Quarter":
        if self.quarter == 1:
            return Quarter(self.year - 1, 4)
        return Quarter(self.year, self.quarter - 1)


@dataclass
class SectionOccurrence:
    canonical: str
    language: str
    heading_text: str
    paragraphs: list[str]
    tables: list[list[list[str]]]


@dataclass
class DocSectionOccurrence:
    canonical: str
    language: str
    heading_index: int


@dataclass(frozen=True)
class UnitSpec:
    label: str
    factor_to_base: Decimal
    currency: str | None = None


@dataclass(frozen=True)
class ExchangeRates:
    rmb_per_usd: Decimal | None = None


def parse_quarter_label(text: str) -> Quarter | None:
    match = re.search(r"Q([1-4])\s*(20\d{2})", text)
    if match:
        return Quarter(int(match.group(2)), int(match.group(1)))
    match = re.search(r"(20\d{2})\s*Q([1-4])", text)
    if match:
        return Quarter(int(match.group(1)), int(match.group(2)))
    match = re.search(r"(20\d{2})Q([1-4])", text)
    if match:
        return Quarter(int(match.group(1)), int(match.group(2)))
    # Chinese quarter formats: "2025年第4季度", "2025年Q4", "2025第四季度"
    cn_digits = {"一": 1, "二": 2, "三": 3, "四": 4}
    match = re.search(r"(20\d{2})\s*年?\s*第?\s*([1-4一二三四])\s*季度", text)
    if match:
        q = cn_digits.get(match.group(2), None) or int(match.group(2))
        return Quarter(int(match.group(1)), q)
    return None


def extract_quarter_number(text: str) -> int | None:
    match = re.search(r"Q([1-4])", text, flags=re.I)
    if match:
        return int(match.group(1))
    match = re.search(r"([1-4])季度", text)
    if match:
        return int(match.group(1))
    return None


def parse_month_header(text: str) -> tuple[int, int] | None:
    normalized = norm_space(text)
    match = re.search(r"(20\d{2})[/-]?(0[1-9]|1[0-2])", normalized)
    if match:
        return int(match.group(1)), int(match.group(2))
    # Chinese month formats: "2025年1月", "2025年01月", "1月" (without year)
    match = re.search(r"(20\d{2})\s*年\s*(1[0-2]|0?[1-9])\s*月", normalized)
    if match:
        return int(match.group(1)), int(match.group(2))
    return None


def quarter_sort_key(text: str) -> tuple[int, int]:
    quarter = parse_quarter_label(text)
    if quarter is None:
        return (-1, -1)
    return (quarter.year, quarter.quarter)


def parse_decimal(text: str) -> Decimal | None:
    cleaned = text.replace(",", "").strip()
    if not cleaned or cleaned in {"/", "n.a.", "n/a"}:
        return None
    negative = cleaned.startswith("(") and cleaned.endswith(")")
    cleaned = cleaned.strip("()")
    cleaned = cleaned.replace("Positive cf", "").replace("Better", "").strip()
    # Detect K/M/B suffix multiplier before stripping letters
    suffix_multiplier = Decimal("1")
    suffix_match = re.search(r"(\d)\s*([KkMmBb])\s*$", cleaned)
    if suffix_match:
        suffix_char = suffix_match.group(2).upper()
        if suffix_char == "K":
            suffix_multiplier = Decimal("1000")
        elif suffix_char == "M":
            suffix_multiplier = Decimal("1000000")
        elif suffix_char == "B":
            suffix_multiplier = Decimal("1000000000")
    cleaned = re.sub(r"[^0-9.+-]", "", cleaned)
    if not cleaned:
        return None
    try:
        value = Decimal(cleaned) * suffix_multiplier
    except Exception:
        return None
    return -value if negative else value


def format_decimal(value: Decimal | None, places: int = 2) -> str | None:
    if value is None:
        return None
    quant = Decimal("1." + ("0" * places))
    rounded = value.quantize(quant, rounding=ROUND_HALF_UP)
    return f"{rounded:.{places}f}"


def parse_percent(text: str) -> Decimal | None:
    cleaned = norm_space(text)
    if not cleaned or cleaned == "/":
        return None
    negative = cleaned.startswith("(") and "%" in cleaned
    match = re.search(r"([+-]?\d+(?:\.\d+)?)\s*%", cleaned)
    if not match:
        return None
    value = Decimal(match.group(1))
    if negative and value > 0:
        value = -value
    return value


def format_percent(value: Decimal | None) -> str | None:
    if value is None:
        return None
    rounded = value.quantize(Decimal("1.00"), rounding=ROUND_HALF_UP)
    return f"{rounded:.2f}%"


def parse_current_blocks(blocks: list[dict[str, Any]]) -> dict[str, Any]:
    result: dict[str, Any] = {
        "title": "",
        "finance_table": None,
        "balance_table": None,
        "operation_table": None,
        "extra_tables": [],
        "business_update_paragraphs": [],
    }

    tables_seen = 0
    for idx, block in enumerate(blocks):
        if idx == 0 and block["type"] == "paragraph":
            result["title"] = block["text"]

        if block["type"] == "table":
            tables_seen += 1
            if result["finance_table"] is None:
                result["finance_table"] = block["rows"]
            elif result["balance_table"] is None:
                result["balance_table"] = block["rows"]
            else:
                result["extra_tables"].append(block["rows"])
                if result["operation_table"] is None:
                    result["operation_table"] = block["rows"]

        if block["type"] == "paragraph" and block["text"] == "Other Questions:":
            trailing = blocks[idx + 1 :]
            for item in trailing:
                if item["type"] != "paragraph":
                    continue
                text = item["text"]
                if is_question_or_prompt(text):
                    continue
                if is_standalone_label(text):
                    continue
                result["business_update_paragraphs"].append(text)
            break

    if not result["business_update_paragraphs"]:
        seen_table = False
        for block in blocks:
            if block["type"] == "table":
                seen_table = True
                continue
            if not seen_table or block["type"] != "paragraph":
                continue
            text = block["text"]
            if is_question_or_prompt(text):
                continue
            if is_standalone_label(text):
                continue
            if len(text) < 2:
                continue
            result["business_update_paragraphs"].append(text)
    return result


def rows_to_dict(rows: list[list[str]]) -> dict[str, dict[str, str]]:
    headers = rows[0][1:]
    mapping: dict[str, dict[str, str]] = {}
    for row in rows[1:]:
        label = norm_space(row[0])
        mapping[label] = {headers[i]: norm_space(row[i + 1]) for i in range(len(headers))}
    return mapping


def extract_outlet_counts(operation_rows: list[list[str]]) -> tuple[str | None, str | None]:
    op_map = rows_to_dict(operation_rows)
    outlet_row = None
    for label in op_map:
        if "outlet" in label.lower() and "accumul" in label.lower():
            outlet_row = op_map[label]
            break
    if outlet_row is None:
        outlet_row = op_map.get("Accumulated Number of outlets", {})
    ytd_value = None
    for header, value in outlet_row.items():
        if "ytd" in header.lower() or "accum" in header.lower() or "total" in header.lower():
            ytd_value = value
            break
    if not ytd_value:
        values = list(outlet_row.values())
        ytd_value = values[-1] if values else None
    if not ytd_value:
        return None, None
    match = re.search(r"(\d+)\s*\+\s*(\d+)\s*\(JV\)", ytd_value)
    if not match:
        plain = re.search(r"(\d+)", ytd_value)
        return (plain.group(1), None) if plain else (None, None)
    return match.group(1), match.group(2)



def normalize_finance_source(rows: list[list[str]]) -> dict[str, Any]:
    header = [norm_space(cell) for cell in rows[0]]
    row_map = rows_to_dict(rows)
    return {
        "header": header,
        "rows": row_map,
    }


def is_empty_numeric_cell(text: str) -> bool:
    cleaned = norm_space(text)
    return cleaned in {"", "/", "-", "—", "–"}


def finance_row_current_period_is_blank(
    source_row: dict[str, str],
    current_month_headers: list[str],
    current_total_header: str,
    fy_or_ytd_header: str | None,
) -> bool:
    relevant_headers = list(current_month_headers)
    if current_total_header:
        relevant_headers.append(current_total_header)
    if fy_or_ytd_header:
        relevant_headers.append(fy_or_ytd_header)
    if not relevant_headers:
        return False
    for header in relevant_headers:
        if not is_empty_numeric_cell(source_row.get(header, "")):
            return False
    return True


def extract_cash_balance_line(balance_rows: list[list[str]] | None, language: str) -> str | None:
    if not balance_rows or len(balance_rows) < 2:
        return None
    for row in balance_rows[1:]:
        if len(row) < 2:
            continue
        label = norm_space(row[0])
        value = norm_space(row[1])
        if not label or not value:
            continue
        label_lower = label.lower()
        if "银行账户现金余额" not in label and "bank account cash balance" not in label_lower and "cash balance" not in label_lower:
            continue

        rmb_match = re.search(r"(?:人民币|RMB)\s*[:：]?\s*([0-9,]+(?:\.\d+)?)", value, flags=re.I)
        usd_match = re.search(r"(?:美元|USD)\s*[:：]?\s*([0-9,]+(?:\.\d+)?)", value, flags=re.I)
        rmb_value = rmb_match.group(1) if rmb_match else None
        usd_value = usd_match.group(1) if usd_match else None

        if language == "chinese":
            parts = []
            if rmb_value is not None:
                parts.append(f"人民币：{rmb_value}")
            if usd_value is not None:
                parts.append(f"美元：{usd_value}")
            return f"银行账户现金余额：{' '.join(parts)}" if parts else f"银行账户现金余额：{value}"

        parts = []
        if rmb_value is not None:
            parts.append(f"RMB {rmb_value}")
        if usd_value is not None:
            parts.append(f"USD {usd_value}")
        return f"Bank Account Cash Balance: {', '.join(parts)}" if parts else f"Bank Account Cash Balance: {value}"
    return None


def calc_percent(current: Decimal | None, previous: Decimal | None) -> Decimal | None:
    if current is None or previous is None or previous == 0:
        return None
    denominator = abs(previous)
    if denominator == 0:
        return None
    return ((current - previous) / denominator) * Decimal("100")


def rewrite_third_person(text: str, company_name: str) -> str:
    text = norm_space(text)
    text = re.sub(r"\bwe\b", company_name, text, flags=re.I)
    text = re.sub(r"\bour\b", f"{company_name}'s", text, flags=re.I)
    text = re.sub(r"\bus\b", company_name, text, flags=re.I)
    text = re.sub(r"\bQ([1-4])\.", r"Q\1", text)
    text = re.sub(r"\bthe company\b", company_name, text, flags=re.I)
    return norm_space(text)


def short_sentence(text: str, company_name: str, limit: int = 140) -> str:
    text = rewrite_third_person(text, company_name)
    text = re.sub(r":\s*", ": ", text)
    text = re.sub(r"\s+-\s+", " - ", text)
    if len(text) <= limit:
        return text.rstrip(".") + "."
    clipped = text[:limit].rsplit(" ", 1)[0].rstrip(",;:-")
    return clipped + "."


def score_paragraph_informativeness(text: str) -> float:
    """Score a paragraph by how much useful business information it contains.
    Higher score = more informative (numbers, metrics, dollar amounts, percentages, specifics)."""
    score = 0.0
    score += len(re.findall(r"\$[\d,.]+[kmb]?", text, re.I)) * 3.0
    score += len(re.findall(r"\d+(?:\.\d+)?%", text)) * 2.5
    score += len(re.findall(r"\b\d{1,3}(?:,\d{3})+\b", text)) * 2.0
    score += len(re.findall(r"\b\d+(?:\.\d+)?\s*(?:x|times|million|billion|k)\b", text, re.I)) * 2.0
    score += len(re.findall(r"\bQ[1-4]\s*20\d{2}\b", text, re.I)) * 1.5
    score += len(re.findall(r"\b(?:revenue|profit|ebitda|margin|growth|cost|loss|cash flow|burn rate|funding|financing|expansion|launched|achieved|target|milestone)\b", text, re.I)) * 1.0
    score += len(re.findall(r"\b(?:YoY|QoQ|year.over.year|quarter.over.quarter|month.over.month)\b", text, re.I)) * 1.5
    length_bonus = min(len(text) / 200.0, 1.0)
    score += length_bonus
    if len(text) < 30:
        score *= 0.3
    return score


def split_into_sentences(text: str) -> list[str]:
    """Split a long paragraph into individual sentences for separate bullets."""
    # Don't split short text
    if len(text) < 120:
        return [text]
    # Split on sentence boundaries, but not on abbreviations like "U.S." or "e.g."
    # or on decimal numbers like "2.67m"
    parts = re.split(r'(?<=[.!])\s+(?=[A-Z\d])', text)
    result = []
    for part in parts:
        part = norm_space(part)
        if len(part) >= 20:
            result.append(part)
    return result if result else [text]


CONTINUATION_PATTERNS = re.compile(
    r"^(?:among them|in addition|additionally|furthermore|moreover|also|"
    r"in particular|specifically|for (?:example|instance)|this (?:includes?|means?)|"
    r"these |the key |the planned |the new )\b",
    re.I,
)


def merge_continuations(items: list[str]) -> list[str]:
    """Merge continuation sentences with their preceding item."""
    if not items:
        return items
    merged: list[str] = [items[0]]
    for item in items[1:]:
        if CONTINUATION_PATTERNS.search(item) and merged:
            merged[-1] = merged[-1].rstrip(".") + ". " + item
        else:
            merged.append(item)
    return merged


def build_business_update_bullets(paragraphs: list[str], company_name: str, max_bullets: int = 9) -> list[str]:
    """Extract the most informative paragraphs as bullet points, scored by content richness."""
    cleaned = clean_update_paragraphs(paragraphs)
    # Merge continuation sentences before splitting
    cleaned = merge_continuations(cleaned)
    # Split multi-sentence paragraphs into separate candidates
    expanded: list[str] = []
    for text in cleaned:
        text = norm_space(text)
        if len(text) < 20:
            continue
        if text.endswith((":", "：")):
            continue
        expanded.extend(split_into_sentences(text))
    # Merge again after splitting (in case split created new continuations)
    expanded = merge_continuations(expanded)
    scored: list[tuple[float, str]] = []
    for text in expanded:
        score = score_paragraph_informativeness(text)
        scored.append((score, text))

    scored.sort(key=lambda x: x[0], reverse=True)

    bullets: list[str] = []
    seen_keys: set[str] = set()
    for _score, text in scored:
        bullet = short_sentence(text, company_name, limit=200)
        key = re.sub(r"[^a-z0-9]", "", bullet.lower())
        if key in seen_keys:
            continue
        seen_keys.add(key)
        bullets.append(bullet)
        if len(bullets) >= max_bullets:
            break

    return bullets


def format_decimal_compact(value: Decimal) -> str:
    text = f"{value.quantize(Decimal('0.01'))}"
    if "." in text:
        text = text.rstrip("0").rstrip(".")
    return text


def convert_chinese_amounts_to_rmb_millions(text: str) -> str:
    def repl_yi(match: re.Match[str]) -> str:
        value = Decimal(match.group(1))
        million = value * Decimal("100")
        return f"RMB {format_decimal_compact(million)} million"

    def repl_wan(match: re.Match[str]) -> str:
        value = Decimal(match.group(1))
        million = value / Decimal("100")
        return f"RMB {format_decimal_compact(million)} million"

    text = re.sub(r"(\d+(?:\.\d+)?)\s*亿元", repl_yi, text)
    text = re.sub(r"(\d+(?:\.\d+)?)\s*万(?:元人民币|人民币)?", repl_wan, text)
    text = re.sub(r"(\d+(?:\.\d+)?)\s*万", repl_wan, text)
    return text


def english_company_name(company_name: str) -> str:
    match = re.match(r"\s*([A-Za-z][A-Za-z0-9& .'-]*)", company_name)
    if match:
        return norm_space(match.group(1))
    return company_name


def strip_leading_list_marker(text: str) -> str:
    text = norm_space(text)
    text = re.sub(r"^[0-9一二三四五六七八九十]+\s*[)）.\-、:：]\s*", "", text)
    return norm_space(text)


def normalize_translation_memory_key(text: str) -> str:
    text = strip_leading_list_marker(text)
    text = re.sub(r"(20\d{2})年\s*Q([1-4])", "<Q>", text, flags=re.I)
    text = re.sub(r"Q([1-4])\s*(20\d{2})", "<Q>", text, flags=re.I)
    text = re.sub(r"第[一二三四1234]季度", "<Q>", text)
    text = re.sub(r"\d+(?:\.\d+)?\s*亿元", "<AMOUNT>", text)
    text = re.sub(r"\d+(?:\.\d+)?\s*万(?:元人民币|人民币)?", "<AMOUNT>", text)
    text = re.sub(r"[\s，,。；;：:（）()]+", "", text)
    return text.lower()


def extract_quarter_display_from_chinese(text: str) -> str | None:
    match = re.search(r"(20\d{2})年\s*Q([1-4])", text, re.I)
    if match:
        return f"{match.group(1)} Q{match.group(2)}"
    match = re.search(r"(20\d{2})年第?([一二三四1234])季度", text)
    if match:
        mapping = {"一": "1", "二": "2", "三": "3", "四": "4"}
        q = mapping.get(match.group(2), match.group(2))
        return f"{match.group(1)} Q{q}"
    return None


def extract_rmb_million_display(text: str) -> str | None:
    converted = convert_chinese_amounts_to_rmb_millions(text)
    match = re.search(r"RMB ([\d.]+) million", converted)
    if match:
        return match.group(1)
    return None


def build_business_update_translation_memory(previous_business_update_map: dict[str, list[str]]) -> dict[str, str]:
    memory: dict[str, str] = {}
    zh_lines = previous_business_update_map.get("chinese", [])
    en_lines = previous_business_update_map.get("english", [])
    for zh, en in zip(zh_lines, en_lines):
        zh_clean = strip_leading_list_marker(zh)
        en_clean = strip_leading_list_marker(en)
        if not zh_clean or not en_clean:
            continue
        memory[normalize_translation_memory_key(zh_clean)] = en_clean
    return memory


def translation_memory_match(text: str, translation_memory: dict[str, str]) -> str | None:
    key = normalize_translation_memory_key(text)
    if key in translation_memory:
        return translation_memory[key]
    for memory_key, memory_value in translation_memory.items():
        if not memory_key:
            continue
        overlap = len(set(key) & set(memory_key)) / max(len(set(memory_key)), 1)
        if overlap >= 0.7:
            return memory_value
    return None


def should_use_translation_memory(current_chinese: str, previous_english: str) -> bool:
    zh_numbers = re.findall(r"\d+(?:\.\d+)?", current_chinese)
    en_numbers = re.findall(r"\d+(?:\.\d+)?", previous_english)
    if len(zh_numbers) > len(en_numbers) + 2:
        return False
    if any(token in current_chinese for token in ["人员变动", "销售量", "万台", "兼职", "全职"]) and not any(
        token in previous_english.lower() for token in ["personnel", "employee", "staff", "sales volume", "units", "full-time", "part-time"]
    ):
        return False
    return True


def adapt_previous_english_translation(previous_english: str, current_chinese: str, company_name: str) -> str:
    text = previous_english
    quarter_display = extract_quarter_display_from_chinese(current_chinese)
    if quarter_display:
        text = re.sub(r"\b20\d{2}\s*Q[1-4]\b", quarter_display, text)
        text = re.sub(r"\bQ[1-4]\s*20\d{2}\b", quarter_display, text)
        text = re.sub(r"\bQ[1-4]\b", quarter_display, text, count=1)
    amount_display = extract_rmb_million_display(current_chinese)
    if amount_display:
        text = re.sub(r"RMB\s+[\d.]+\s+million", f"RMB {amount_display} million", text)
    if any(token in current_chinese for token in ["平稳", "稳定"]):
        text = re.sub(
            r"which increased compared to the last quarter",
            "with collections remaining stable",
            text,
            flags=re.I,
        )
    return short_sentence(text, english_company_name(company_name), limit=220)


# Generic Chinese-to-English business term dictionary.
# Ordered longest-first within each group so longer phrases match before their
# substrings (e.g. "应收账款回款" before "应收账款", "净利润" before "净利").
# Groups: multi-char phrases first, then two-char terms.
ZH_EN_TERM_MAP: list[tuple[str, str]] = [
    # ── Finance / accounting (long phrases first) ──────────────────────
    ("应收账款回款", "accounts receivable collections"),
    ("应付账款", "accounts payable"),
    ("应收账款", "accounts receivable"),
    ("经营性现金流", "operating cash flow"),
    ("正向经营现金流", "positive operating cash flow"),
    ("净现金消耗", "net cash burn"),
    ("现金流入", "cash inflow"),
    ("现金支出", "cash outflow"),
    ("现金余额", "cash balance"),
    ("现金流", "cash flow"),
    ("毛利率", "gross margin"),
    ("净利率", "net margin"),
    ("毛利润", "gross profit"),
    ("净利润", "net profit"),
    ("营业收入", "revenue"),
    ("销售收入", "sales revenue"),
    ("主营业务收入", "core business revenue"),
    ("资产负债表", "balance sheet"),
    ("利润表", "income statement"),
    ("现金流量表", "cash flow statement"),
    ("财务报表", "financial statements"),
    ("银行贷款", "bank loan"),
    ("信贷额度", "credit facility"),
    ("贷款额度", "loan facility"),
    ("股权融资", "equity financing"),
    ("债务融资", "debt financing"),
    ("下一轮融资", "the next financing round"),
    ("融资额度", "financing amount"),
    ("融资计划", "financing plan"),
    ("融资进展", "financing progress"),
    ("资金需求", "funding needs"),
    ("资金链", "funding chain"),
    ("烧钱率", "burn rate"),
    ("年化收入", "annualized revenue"),
    ("同比增长", "year-over-year growth"),
    ("环比增长", "quarter-over-quarter growth"),
    ("同比下降", "year-over-year decline"),
    ("环比下降", "quarter-over-quarter decline"),
    ("营业成本", "cost of revenue"),
    ("运营成本", "operating costs"),
    ("运营费用", "operating expenses"),
    ("管理费用", "administrative expenses"),
    ("销售费用", "selling expenses"),
    ("研发费用", "R&D expenses"),
    ("折旧摊销", "depreciation and amortization"),
    # ── Finance / accounting (short terms) ─────────────────────────────
    ("收入", "revenue"),
    ("毛利", "gross profit"),
    ("净利", "net profit"),
    ("利润", "profit"),
    ("亏损", "loss"),
    ("成本", "cost"),
    ("费用", "expenses"),
    ("融资", "financing"),
    ("贷款", "loan"),
    ("投资", "investment"),
    ("估值", "valuation"),
    ("股权", "equity"),
    ("负债", "liabilities"),
    ("资产", "assets"),
    ("现金", "cash"),
    # ── Operations / business ──────────────────────────────────────────
    ("供应链", "supply chain"),
    ("供应商", "supplier"),
    ("经销商", "distributor"),
    ("合作伙伴", "partner"),
    ("合作方", "partner"),
    ("战略合作", "strategic partnership"),
    ("业务进展", "business progress"),
    ("业务发展", "business development"),
    ("业务拓展", "business expansion"),
    ("业务回顾", "business review"),
    ("市场拓展", "market expansion"),
    ("市场份额", "market share"),
    ("市场需求", "market demand"),
    ("客户数量", "customer count"),
    ("新增客户", "new customers acquired"),
    ("客户留存", "customer retention"),
    ("运营数据", "operational data"),
    ("运营效率", "operational efficiency"),
    ("产能利用率", "capacity utilization"),
    ("生产线", "production line"),
    ("量产", "mass production"),
    ("批产", "batch production"),
    ("试产", "trial production"),
    ("产能", "production capacity"),
    ("产量", "output volume"),
    ("销量", "sales volume"),
    ("出货量", "shipment volume"),
    ("库存", "inventory"),
    ("备货", "inventory build-up"),
    ("清关入库", "customs clearance and warehousing"),
    ("清关", "customs clearance"),
    ("入库", "warehousing"),
    ("上架销售", "listed for sale"),
    ("上架", "listed"),
    ("采购", "procurement"),
    ("订单", "order"),
    ("合同", "contract"),
    ("签约", "contract signed"),
    ("交付", "delivery"),
    ("发货", "shipment"),
    ("回款", "payment collection"),
    ("结算", "settlement"),
    ("项目", "project"),
    ("工厂", "factory"),
    ("门店", "store"),
    ("网点", "outlet"),
    # ── Growth / strategy ──────────────────────────────────────────────
    ("里程碑", "milestone"),
    ("时间线", "timeline"),
    ("路线图", "roadmap"),
    ("战略规划", "strategic plan"),
    ("扩张计划", "expansion plan"),
    ("增长目标", "growth target"),
    ("销售目标", "sales target"),
    ("下一步计划", "next steps"),
    ("未来计划", "future plans"),
    ("短期计划", "near-term plan"),
    ("中长期计划", "medium-to-long-term plan"),
    # ── People / organization ──────────────────────────────────────────
    ("全职员工", "full-time employees"),
    ("兼职员工", "part-time employees"),
    ("总员工人数", "total headcount"),
    ("人员变动", "personnel changes"),
    ("团队建设", "team building"),
    ("管理层", "management team"),
    ("核心团队", "core team"),
    # ── IPO / corporate ────────────────────────────────────────────────
    ("上市计划", "IPO plan"),
    ("上市进展", "IPO progress"),
    ("上市", "listing/IPO"),
    ("股改", "shareholding restructuring"),
    ("尽职调查", "due diligence"),
    ("审计", "audit"),
    ("合规", "compliance"),
    ("注册资料", "registration documents"),
    ("核名", "name approval"),
    # ── Risk ───────────────────────────────────────────────────────────
    ("风险", "risk"),
    ("退出", "exit"),
    ("退出路径", "exit path"),
    # ── Trend / status words ───────────────────────────────────────────
    ("同比", "YoY"),
    ("环比", "QoQ"),
    ("持续增长", "continued to grow"),
    ("大幅增长", "grew significantly"),
    ("稳步增长", "grew steadily"),
    ("保持稳定", "remained stable"),
    ("有所下降", "declined slightly"),
    ("大幅下降", "declined significantly"),
    ("超出预期", "exceeded expectations"),
    ("低于预期", "below expectations"),
    ("不及预期", "below expectations"),
    ("符合预期", "in line with expectations"),
    ("预计", "is expected to"),
    ("计划", "plans to"),
    ("目标", "target"),
    ("完成", "completed"),
    ("实现", "achieved"),
    ("推进", "advanced"),
    ("启动", "launched"),
    ("落地", "implemented"),
    ("待完成", "pending completion"),
    # ── Time ───────────────────────────────────────────────────────────
    ("一季度", "Q1"),
    ("二季度", "Q2"),
    ("三季度", "Q3"),
    ("四季度", "Q4"),
    ("第一季度", "Q1"),
    ("第二季度", "Q2"),
    ("第三季度", "Q3"),
    ("第四季度", "Q4"),
    ("上半年", "H1"),
    ("下半年", "H2"),
    ("年初", "the beginning of the year"),
    ("年底", "year-end"),
    ("月底", "month-end"),
    ("季度末", "quarter-end"),
    # ── Geography ──────────────────────────────────────────────────────
    ("台湾", "Taiwan"),
    ("大陆", "Mainland China"),
    ("海外", "overseas"),
    ("境外", "offshore"),
    ("境内", "domestic"),
    # ── General connectors / action words ──────────────────────────────
    ("业务", "business"),
    ("进展", "progress"),
    ("情况", "status"),
    ("平稳", "stable"),
    ("稳定", "stable"),
    ("增加", "increased"),
    ("减少", "decreased"),
    ("提升", "improved"),
    ("优化", "optimized"),
    ("调整", "adjusted"),
    ("评估", "evaluated"),
    ("观察", "monitored"),
    ("公司", "the company"),
]

# Pre-sort so longer Chinese phrases are matched first (avoids partial replacement).
ZH_EN_TERM_MAP.sort(key=lambda pair: len(pair[0]), reverse=True)


def _chinese_fraction(text: str) -> float:
    """Return the fraction of characters that are Chinese."""
    if not text:
        return 0.0
    chinese = len(re.findall(r"[\u4e00-\u9fff]", text))
    return chinese / len(text)


def generic_translate_chinese_sentence(text: str, company_name: str) -> str:
    """Translate a Chinese sentence to English using the generic term dictionary.

    The approach:
    1. Convert Chinese number expressions (亿/万) to "RMB X million".
    2. Normalize dates (YYYY年M月D日 → Month D, YYYY).
    3. Apply the generic term dictionary longest-first.
    4. Polish spacing and punctuation.
    5. If significant Chinese remains, prepend [MT] to flag for human review.
    """
    english_name = english_company_name(company_name)
    translated = convert_chinese_amounts_to_rmb_millions(text)

    # Normalize Chinese date expressions to English
    _MONTHS = ["", "January", "February", "March", "April", "May", "June",
               "July", "August", "September", "October", "November", "December"]

    def _replace_date_ymd(m: re.Match[str]) -> str:
        y, mo, d = m.group(1), int(m.group(2)), m.group(3)
        return f"{_MONTHS[mo]} {d}, {y}"

    def _replace_date_ym(m: re.Match[str]) -> str:
        y, mo = m.group(1), int(m.group(2))
        return f"{_MONTHS[mo]} {y}"

    # Match full dates first (年月日), then year-month only (年月)
    translated = re.sub(r"(20\d{2})年\s*(1[0-2]|0?[1-9])月\s*(\d{1,2})日", _replace_date_ymd, translated)
    translated = re.sub(r"(20\d{2})年\s*(1[0-2]|0?[1-9])月", _replace_date_ym, translated)

    # Normalize "YYYY年QN" to "QN YYYY"
    translated = re.sub(r"(20\d{2})年\s*Q([1-4])", r"Q\2 \1", translated, flags=re.I)
    # Normalize "第N季度" already handled in the term map

    # Normalize Chinese percentage expressions: "增长了30%" → "grew by 30%"
    translated = re.sub(r"增长了?\s*(\d+(?:\.\d+)?%)", r"grew by \1", translated)
    translated = re.sub(r"下降了?\s*(\d+(?:\.\d+)?%)", r"declined by \1", translated)

    # Apply the generic term dictionary
    for zh, en in ZH_EN_TERM_MAP:
        translated = translated.replace(zh, f" {en} ")

    # Replace Chinese punctuation
    translated = translated.replace("，", ", ")
    translated = translated.replace("。", ". ")
    translated = translated.replace("；", "; ")
    translated = translated.replace("：", ": ")
    translated = translated.replace("（", " (")
    translated = translated.replace("）", ") ")
    translated = translated.replace("、", ", ")

    # Clean up spacing
    translated = re.sub(r"\s+", " ", translated)
    translated = translated.replace(" ,", ",")
    translated = translated.replace(" .", ".")
    translated = translated.replace(" ;", ";")
    translated = translated.replace("( ", "(")
    translated = translated.replace(" )", ")")
    translated = translated.strip(" .")

    if not translated:
        return f"{english_name}: {text}."

    # Capitalize first letter
    translated = translated[0].upper() + translated[1:] if translated else translated
    if not translated.endswith("."):
        translated += "."

    # Flag if significant Chinese characters remain (> 20% of text)
    if _chinese_fraction(translated) > 0.20:
        translated = f"[MT] {translated}"

    return norm_space(translated)


def polish_english_business_bullet(text: str, company_name: str) -> str:
    """Clean up spacing, punctuation, and formatting artefacts in a translated bullet."""
    text = norm_space(text)
    company = english_company_name(company_name)
    # Replace any remaining Chinese punctuation
    text = re.sub(r"([A-Za-z])（", r"\1 (", text)
    text = re.sub(r"）", ")", text)
    text = re.sub(r"，", ", ", text)
    # Fix missing spaces between camelCase-like joins from term substitution
    text = re.sub(r"([a-z])([A-Z])", r"\1 \2", text)
    text = re.sub(r"(\d)([A-Za-z])", r"\1 \2", text)
    text = re.sub(r"([A-Za-z])(\d)", r"\1 \2", text)
    # Fix "Q 4" → "Q4"
    text = re.sub(r"\bQ\s+([1-4])\b", r"Q\1", text)
    # Remove redundant "update:" prefix
    text = text.replace(f"{company} update:", f"{company}")
    # Clean up spacing
    text = re.sub(r"\s+", " ", text)
    text = text.replace(" ,", ",")
    text = text.replace(" .", ".")
    text = text.replace(" ;", ";")
    text = text.strip()
    if not text.endswith("."):
        text += "."
    return text


def translate_chinese_update_bullet_to_english(text: str, company_name: str, translation_memory: dict[str, str] | None = None) -> str:
    """Translate a Chinese business-update bullet to English.

    Strategy (in priority order):
    1. Translation memory — reuse the previous quarter's known Chinese→English pair.
    2. Structured pattern extraction — handle common financial/operational patterns
       generically (amounts, trends, dates) without company-specific rules.
    3. Generic term-dictionary translation — apply ZH_EN_TERM_MAP and flag
       incomplete results with [MT].
    """
    text = strip_leading_list_marker(norm_space(text).rstrip("。"))
    english_name = english_company_name(company_name)

    # ── 1. Translation memory ──────────────────────────────────────────
    if translation_memory:
        memory_hit = translation_memory_match(text, translation_memory)
        if memory_hit and should_use_translation_memory(text, memory_hit):
            return adapt_previous_english_translation(memory_hit, text, english_name)

    # ── 2. Structured pattern extraction ───────────────────────────────
    converted = convert_chinese_amounts_to_rmb_millions(text)

    # "YYYY年QN业务回顾" → "QN YYYY business review."
    quarter_match = re.search(r"(20\d{2})年\s*Q([1-4])", converted, re.I)
    if quarter_match and "业务回顾" in converted:
        return f"Q{quarter_match.group(2)} {quarter_match.group(1)} business review."

    # Accounts receivable pattern: extract amount + trend generically
    if "应收账款" in converted:
        amount_match = re.search(r"RMB ([\d.]+) million", converted)
        amount = f" approximately RMB {amount_match.group(1)} million" if amount_match else ""
        if any(kw in converted for kw in ["平稳", "稳定"]):
            trend = ", with collections remaining stable"
        elif any(kw in converted for kw in ["上涨", "增加", "增长"]):
            trend = ", with collections increasing versus the previous quarter"
        elif any(kw in converted for kw in ["下降", "减少"]):
            trend = ", with collections declining versus the previous quarter"
        else:
            trend = ""
        return f"Accounts receivable collections were{amount}{trend}."

    # Cash insufficient for operations
    if "现金" in converted and any(kw in converted for kw in ["无法支持", "不足以支持", "难以维持"]):
        return f"{english_name}'s current cash position is no longer sufficient to support operations."

    # Generic financing pattern
    if "融资" in converted and len(converted) < 80:
        amount_match = re.search(r"RMB ([\d.]+) million", converted)
        amount = f" of RMB {amount_match.group(1)} million" if amount_match else ""
        return f"{english_name} is advancing its financing plan{amount}."

    # IPO / listing pattern
    if ("上市" in converted or "股改" in converted) and len(converted) < 100:
        return f"{english_name} continued preparations for its listing/IPO."

    # ── 3. Generic term-dictionary translation ─────────────────────────
    return polish_english_business_bullet(generic_translate_chinese_sentence(text, company_name), company_name)


def build_english_business_update_from_chinese(paragraphs: list[str], company_name: str, max_bullets: int = 9, translation_memory: dict[str, str] | None = None) -> list[str]:
    chinese_bullets = build_chinese_business_update_bullets(paragraphs)
    bullets: list[str] = []
    seen_keys: set[str] = set()
    for text in chinese_bullets:
        bullet = translate_chinese_update_bullet_to_english(text, company_name, translation_memory)
        bullet = polish_english_business_bullet(bullet, company_name)
        bullet = short_sentence(bullet, company_name, limit=600)
        key = re.sub(r"[^a-z0-9]", "", bullet.lower())
        if key in seen_keys:
            continue
        seen_keys.add(key)
        bullets.append(bullet)
        if len(bullets) >= max_bullets:
            break
    return bullets


def ensure_table_dimensions(table, required_rows: int, required_cols: int) -> bool:
    expanded = False
    if not table.rows:
        return False
    current_cols = len(table.rows[0].cells) if table.rows else 0
    grid_widths = get_table_grid_widths(table)
    last_width = grid_widths[-1] if grid_widths else 900

    while current_cols < required_cols:
        source_props = []
        for row in table.rows:
            source_cell = row.cells[-1]
            tc_pr = copy.deepcopy(source_cell._tc.tcPr) if source_cell._tc.tcPr is not None else None
            p_pr = None
            if source_cell.paragraphs:
                existing_ppr = source_cell.paragraphs[0]._element.find(qn("w:pPr"))
                if existing_ppr is not None:
                    p_pr = copy.deepcopy(existing_ppr)
            source_props.append((tc_pr, p_pr))
        table.add_column(Twips(last_width))
        if not source_props:
            current_cols += 1
            expanded = True
            continue
        for row_idx, row in enumerate(table.rows):
            new_cell = row.cells[-1]
            tc = new_cell._tc
            existing_tcpr = tc.tcPr
            if existing_tcpr is not None:
                tc.remove(existing_tcpr)
            source_tcpr, source_ppr = source_props[row_idx] if row_idx < len(source_props) else source_props[-1]
            if source_tcpr is not None:
                tc.insert(0, copy.deepcopy(source_tcpr))
            if new_cell.paragraphs:
                p = new_cell.paragraphs[0]
                existing_ppr = p._element.find(qn("w:pPr"))
                if existing_ppr is not None:
                    p._element.remove(existing_ppr)
                if source_ppr is not None:
                    p._element.insert(0, copy.deepcopy(source_ppr))
        current_cols += 1
        expanded = True

    while len(table.rows) < required_rows:
        source_row = table.rows[-1]
        source_props = []
        for source_cell in source_row.cells:
            tc_pr = copy.deepcopy(source_cell._tc.tcPr) if source_cell._tc.tcPr is not None else None
            p_pr = None
            if source_cell.paragraphs:
                existing_ppr = source_cell.paragraphs[0]._element.find(qn("w:pPr"))
                if existing_ppr is not None:
                    p_pr = copy.deepcopy(existing_ppr)
            source_props.append((tc_pr, p_pr))
        new_row = table.add_row()
        for cell_idx, new_cell in enumerate(new_row.cells[: len(source_props)]):
            tc = new_cell._tc
            existing_tcpr = tc.tcPr
            if existing_tcpr is not None:
                tc.remove(existing_tcpr)
            source_tcpr, source_ppr = source_props[cell_idx]
            if source_tcpr is not None:
                tc.insert(0, copy.deepcopy(source_tcpr))
            if new_cell.paragraphs:
                p = new_cell.paragraphs[0]
                existing_ppr = p._element.find(qn("w:pPr"))
                if existing_ppr is not None:
                    p._element.remove(existing_ppr)
                if source_ppr is not None:
                    p._element.insert(0, copy.deepcopy(source_ppr))
        expanded = True
    return expanded


def strip_response_prefix(text: str) -> str:
    """Strip Q&A response prefixes like 'CompanyName: ...' or 'Yes, CompanyName ...'."""
    # Strip "CompanyName: " or "CompanyName, " at start
    text = re.sub(r"^[A-Z][A-Za-z\s]{1,30}:\s*", "", text)
    # Strip "Yes, " or "Yes. " at start
    text = re.sub(r"^(?:Yes|No)[,.\s]+\s*", "", text, flags=re.I)
    return norm_space(text)


def clean_update_paragraphs(paragraphs: list[str]) -> list[str]:
    cleaned: list[str] = []
    for text in paragraphs:
        text = norm_space(text)
        if not text:
            continue
        if is_question_or_prompt(text):
            continue
        if is_standalone_label(text):
            continue
        text = strip_response_prefix(text)
        if not text or len(text) < 10:
            continue
        cleaned.append(text)
    return cleaned


def bullet_topic(text: str) -> str:
    lowered = norm_space(text).lower()
    if any(token in lowered for token in ["融资", "股改", "上市", "ipo", "equity", "fund", "loan", "facility", "capital", "cash flow", "现金流"]):
        return "finance"
    if any(token in lowered for token in ["项目", "交付", "签约", "回款", "合同", "project", "delivery", "contract", "customer", "operation", "运营", "milestone"]):
        return "operations"
    if any(token in lowered for token in ["计划", "扩张", "outlet", "growth", "expansion", "rollout", "trend", "progress", "channel", "product", "certification"]):
        return "growth"
    return "general"


def section_heading_preference(heading_text: str) -> list[str]:
    lowered = norm_space(heading_text).lower()
    if any(token in lowered for token in ["运营", "operation", "review", "业务回顾", "recent"]):
        return ["operations", "general", "growth", "finance"]
    if any(token in lowered for token in ["进展", "progress", "highlight", "trend", "业务进展", "business"]):
        return ["growth", "finance", "general", "operations"]
    return ["general", "operations", "growth", "finance"]


def allocate_middle_section_bullets(headings: list[str], bullets: list[str]) -> list[list[str]]:
    if not headings:
        return []
    if len(headings) == 1:
        return [bullets]

    remaining = [(bullet, bullet_topic(bullet)) for bullet in bullets]
    allocations: list[list[str]] = []
    for idx, heading in enumerate(headings):
        if idx == len(headings) - 1:
            allocations.append([bullet for bullet, _topic in remaining])
            break
        preferred = section_heading_preference(heading)
        chosen: list[str] = []
        still_remaining = []
        used_topics = set()
        for bullet, topic in remaining:
            if topic in preferred[:2] and topic not in used_topics:
                chosen.append(bullet)
                used_topics.add(topic)
            else:
                still_remaining.append((bullet, topic))
        if not chosen and remaining:
            chosen.append(remaining[0][0])
            still_remaining = remaining[1:]
        allocations.append(chosen)
        remaining = still_remaining
    while len(allocations) < len(headings):
        allocations.append([])
    return allocations


def build_chinese_business_update_bullets(paragraphs: list[str]) -> list[str]:
    lines = clean_update_paragraphs(paragraphs)
    bullets: list[str] = []
    for text in lines:
        if len(text) < 4:
            continue
        if text.endswith("：") or text.endswith(":"):
            continue
        bullets.append(text.rstrip("；;。") + "。")
    deduped: list[str] = []
    seen = set()
    for bullet in bullets:
        key = bullet.lower()
        if key not in seen:
            seen.add(key)
            deduped.append(bullet)
    return deduped[:8]


def previous_review_has_logo(path: Path) -> bool:
    with zipfile.ZipFile(path) as zf:
        return any(name.startswith("word/media/") for name in zf.namelist())


def normalize_company_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", " ", name).strip()
    parts = [part for part in cleaned.split() if part]
    if not parts:
        return "PortfolioCompany"
    return "".join(parts)


def quarter_filename_label(quarter: Quarter) -> str:
    return f"{quarter.year}Q{quarter.quarter}"


def clean_main_output_folder(final_docx_path: Path) -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    for path in OUTPUT_DIR.iterdir():
        if path == DEBUG_OUTPUT_DIR:
            continue
        if path.is_file() and path.suffix.lower() != ".docx":
            path.unlink()


def clean_main_output_folder_for_batch() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    for path in OUTPUT_DIR.iterdir():
        if path == DEBUG_OUTPUT_DIR:
            continue
        if path.is_file():
            path.unlink()


def list_real_docx(folder: Path) -> list[Path]:
    return sorted(
        path
        for path in folder.glob("*.docx")
        if not path.name.startswith("~$")
    )


def list_current_inputs(folder: Path) -> list[Path]:
    return sorted(
        path
        for path in folder.iterdir()
        if path.is_file()
        and path.suffix.lower() in {".docx", ".pdf"}
        and not path.name.startswith("~$")
    )


def previous_financial_lookup(parsed_prev: dict[str, Any]) -> dict[str, dict[str, str]]:
    return {row["label"]: row["values"] for row in parsed_prev["rows"]}


def latest_quarter_from_current(finance_source: dict[str, Any]) -> Quarter:
    for header in finance_source["header"]:
        quarter = parse_quarter_label(header)
        if quarter:
            return quarter
    raise ValueError("Could not identify current quarter from data request headers.")


def normalize_metric_label(text: str) -> str:
    text = text.lower()
    text = re.sub(r"\(=.*?\)", "", text)
    text = re.sub(r"\([^)]*\)", " ", text)
    text = text.replace("&", " and ")
    text = re.sub(r"[^a-z0-9\u4e00-\u9fff]+", " ", text)
    return norm_space(text)


def metric_tokens(text: str) -> set[str]:
    stopwords = {
        "of",
        "and",
        "the",
        "per",
        "company",
        "owned",
        "number",
        "accumulated",
    }
    normalized = normalize_metric_label(text)
    # For Chinese text, tokenize at the character level so partial matches work
    # e.g. "收入" tokens = {"收", "入", "收入"} which overlaps with "营业收入"
    chinese_chars = re.findall(r"[\u4e00-\u9fff]", normalized)
    if chinese_chars:
        tokens = set(chinese_chars)
        # Also add the full Chinese string and any contiguous Chinese substrings of length 2+
        chinese_only = "".join(chinese_chars)
        tokens.add(chinese_only)
        for length in range(2, len(chinese_chars)):
            for start in range(len(chinese_chars) - length + 1):
                tokens.add(chinese_only[start:start + length])
        # Also include any English tokens
        english_parts = re.sub(r"[\u4e00-\u9fff]+", " ", normalized)
        tokens |= {t for t in english_parts.split() if t not in stopwords}
        return tokens
    return {token for token in normalized.split() if token not in stopwords}


def find_current_quarter_total_header(finance_source: dict[str, Any], target_quarter: Quarter) -> str | None:
    for header in finance_source["header"]:
        if parse_quarter_label(header) == target_quarter:
            return header
    exact_subtotal = []
    yearly_total = []
    for header in finance_source["header"]:
        lowered = header.lower()
        if any(token in lowered for token in ["ytd", "accum", "exp", "qoq", "yoy", "预计", "预估"]):
            continue
        if parse_quarter_label(header) == target_quarter:
            exact_subtotal.append(header)
        # Match quarter number without year (e.g., "Q4", "第四季度")
        q_num = extract_quarter_number(header)
        if q_num == target_quarter.quarter and parse_quarter_label(header) is None:
            if not any(skip in lowered for skip in ["exp", "预计"]):
                exact_subtotal.append(header)
        # Match "总计/subtotal" with target quarter identifier
        if "总计" in header or "小计" in header or "subtotal" in lowered:
            parsed_q = parse_quarter_label(header)
            if parsed_q == target_quarter:
                exact_subtotal.append(header)
            elif parsed_q is None:
                q_in_header = extract_quarter_number(header)
                if q_in_header == target_quarter.quarter:
                    exact_subtotal.append(header)
                elif q_in_header is None:
                    # Generic yearly total (e.g. "2025总计") — use as last resort
                    if str(target_quarter.year) in header:
                        yearly_total.append(header)
    if exact_subtotal:
        return exact_subtotal[0]
    # Fall back to yearly total if no quarter-specific header found
    # (the actual Q4 value will come from month summation, but this header
    # is needed for choose_finance_row_label to filter rows that have data)
    if yearly_total:
        return yearly_total[0]
    # Last resort: any non-metadata column with numeric data
    skip_patterns = {"qoq", "yoy", "exp", "预计", "预估"}
    for header in finance_source["header"]:
        lowered = header.lower()
        if any(s in lowered for s in skip_patterns):
            continue
        if parse_month_header(header) is not None:
            continue
        # Check if at least one row has a parseable value
        for row_data in finance_source["rows"].values():
            if parse_decimal(row_data.get(header, "")) is not None:
                return header
    return None


def find_header_value(values: dict[str, str], patterns: list[str]) -> str | None:
    for header, value in values.items():
        lowered = header.lower()
        if any(pattern in lowered for pattern in patterns):
            return value
    return None


def find_month_headers_for_quarter(finance_source: dict[str, Any], target_quarter: Quarter) -> list[str]:
    headers = []
    for header in finance_source["header"]:
        month_info = parse_month_header(header)
        if not month_info:
            continue
        year, month = month_info
        quarter = ((month - 1) // 3) + 1
        if year == target_quarter.year and quarter == target_quarter.quarter:
            headers.append(header)
    return headers


def sum_month_values(values: dict[str, str], month_headers: list[str]) -> Decimal | None:
    if not month_headers:
        return None
    month_values = [parse_decimal(values.get(header, "")) for header in month_headers]
    if any(value is None for value in month_values):
        return None
    return sum(month_values)  # type: ignore[arg-type]


def find_fy_or_ytd_header(finance_source: dict[str, Any], target_quarter: Quarter) -> str | None:
    candidates = []
    year_text = str(target_quarter.year)
    for header in finance_source["header"]:
        lowered = header.lower()
        if year_text not in header:
            continue
        if any(token in lowered for token in ["exp", "qoq", "yoy"]):
            continue
        if parse_quarter_label(header) is not None:
            continue
        if parse_month_header(header) is not None:
            continue
        if any(token in lowered for token in ["ytd", "accum"]) or "总计" in header or "全年" in header:
            candidates.append(header)
    for header in candidates:
        if "总计" in header or "全年" in header:
            return header
    return candidates[0] if candidates else None


def normalize_previous_quarter_headers(raw_headers: list[str], previous_review_quarter: Quarter | None) -> list[tuple[str, str]]:
    quarter_headers = [header for header in raw_headers if parse_quarter_label(header)]
    if previous_review_quarter is None:
        return [(header, parse_quarter_label(header).display()) for header in quarter_headers if parse_quarter_label(header)]

    expected = []
    cursor = previous_review_quarter
    for _ in quarter_headers:
        expected.append(cursor)
        cursor = cursor.previous()

    normalized: list[tuple[str, str]] = []
    used_expected: set[Quarter] = set()
    for idx, header in enumerate(quarter_headers):
        parsed = parse_quarter_label(header)
        if parsed in expected and parsed not in used_expected:
            normalized.append((header, parsed.display()))
            used_expected.add(parsed)
            continue

        expected_q = expected[idx]
        parsed_qnum = extract_quarter_number(header)
        if expected_q not in used_expected and (parsed is None or parsed_qnum == expected_q.quarter or parsed not in expected):
            normalized.append((header, expected_q.display()))
            used_expected.add(expected_q)
            continue

        if parsed is not None and parsed not in used_expected:
            normalized.append((header, parsed.display()))
            used_expected.add(parsed)
            continue

        normalized.append((header, header))

    return normalized


def detect_unit_spec(text: str) -> UnitSpec | None:
    normalized = norm_space(text).lower()
    compact = normalized.replace(" ", "")
    # IMPORTANT: More specific (longer) patterns MUST come before shorter ones
    # e.g. "百万元人民币" before "万元人民币" since the latter is a substring of the former
    candidates: list[tuple[str, Decimal, str | None]] = [
        ("usdk", Decimal("1000"), "USD"),
        ("kusd", Decimal("1000"), "USD"),
        ("k usd", Decimal("1000"), "USD"),
        ("amount in usd million", Decimal("1000000"), "USD"),
        ("in millions usd", Decimal("1000000"), "USD"),
        ("in usd million", Decimal("1000000"), "USD"),
        ("usd million", Decimal("1000000"), "USD"),
        ("millions usd", Decimal("1000000"), "USD"),
        ("million usd", Decimal("1000000"), "USD"),
        ("financial data (unit: usd)", Decimal("1"), "USD"),
        ("unit: usd", Decimal("1"), "USD"),
        # Chinese: 百万 (million) patterns BEFORE 万 (ten-thousand) patterns
        ("单位：百万元人民币", Decimal("1000000"), "RMB"),
        ("百万元人民币", Decimal("1000000"), "RMB"),
        ("单位：百万人民币", Decimal("1000000"), "RMB"),
        ("百万人民币", Decimal("1000000"), "RMB"),
        ("百万元", Decimal("1000000"), "RMB"),
        ("百万", Decimal("1000000"), "RMB"),
        ("单位：万元人民币", Decimal("10000"), "RMB"),
        ("万元人民币", Decimal("10000"), "RMB"),
        ("万元 人民币", Decimal("10000"), "RMB"),
        ("万元", Decimal("10000"), "RMB"),
        ("单位：人民币", Decimal("1"), "RMB"),
        ("元人民币", Decimal("1"), "RMB"),
        ("人民币", Decimal("1"), "RMB"),
        ("unit: rmb in millions", Decimal("1000000"), "RMB"),
        ("rmb in millions", Decimal("1000000"), "RMB"),
        ("amount in rmb mn", Decimal("1000000"), "RMB"),
        ("rmb mn", Decimal("1000000"), "RMB"),
        ("单位：美元", Decimal("1"), "USD"),
        ("美元", Decimal("1"), "USD"),
    ]
    for needle, factor, currency in candidates:
        if needle in normalized or needle.replace(" ", "") in compact:
            return UnitSpec(label=text, factor_to_base=factor, currency=currency)
    if "unit: usd" in normalized or "(unit: usd)" in normalized:
        return UnitSpec(label=text, factor_to_base=Decimal("1"), currency="USD")
    return None


def extract_exchange_rates(blocks: list[dict[str, Any]]) -> ExchangeRates:
    texts: list[str] = []
    for block in blocks:
        if block["type"] == "paragraph":
            texts.append(block["text"])
        elif block["type"] == "table":
            for row in block["rows"]:
                texts.extend(row)

    patterns = [
        r"1\s*(?:美元|usd|us\$)\s*[=＝:：]\s*([0-9]+(?:\.[0-9]+)?)\s*(?:元\s*人民币|人民币|rmb|cny)",
        r"(?:人民币|rmb|cny)\s*([0-9]+(?:\.[0-9]+)?)\s*(?:元)?\s*[=＝:：]\s*1\s*(?:美元|usd|us\$)",
    ]
    for text in texts:
        normalized = norm_space(text)
        for pattern in patterns:
            match = re.search(pattern, normalized, flags=re.I)
            if match:
                return ExchangeRates(rmb_per_usd=Decimal(match.group(1)))
    return ExchangeRates()


def detect_current_unit_spec(blocks: list[dict[str, Any]]) -> UnitSpec | None:
    for block in blocks[:6]:
        if block["type"] == "paragraph":
            spec = detect_unit_spec(block["text"])
            if spec:
                return spec
    for block in blocks:
        if block["type"] == "paragraph":
            spec = detect_unit_spec(block["text"])
            if spec:
                return spec
    return None


def detect_financial_section_unit(occurrence: SectionOccurrence) -> tuple[str | None, str]:
    outside_unit = None
    for paragraph in occurrence.paragraphs:
        if is_unit_text(paragraph):
            outside_unit = paragraph
            break
    inside_unit = None
    if occurrence.tables and occurrence.tables[0] and occurrence.tables[0][0]:
        inside_unit = norm_space(occurrence.tables[0][0][0])
        if not is_unit_text(inside_unit):
            inside_unit = None
    if outside_unit:
        return outside_unit, "outside"
    if inside_unit:
        return inside_unit, "inside"
    return inside_unit or outside_unit, "none"


def convert_value_between_units(
    value: Decimal | None,
    source_unit: UnitSpec | None,
    output_unit: UnitSpec | None,
    exchange_rates: ExchangeRates | None = None,
) -> Decimal | None:
    if value is None:
        return None
    if source_unit is None or output_unit is None:
        return value
    source_amount = value * source_unit.factor_to_base
    source_currency = source_unit.currency
    output_currency = output_unit.currency

    if source_currency and output_currency and source_currency != output_currency:
        rmb_per_usd = exchange_rates.rmb_per_usd if exchange_rates else None
        if rmb_per_usd:
            if source_currency == "RMB" and output_currency == "USD":
                source_amount = source_amount / rmb_per_usd
            elif source_currency == "USD" and output_currency == "RMB":
                source_amount = source_amount * rmb_per_usd
            else:
                return value
        else:
            return value

    return source_amount / output_unit.factor_to_base


def determine_target_quarter(current_path: Path, current_title: str, current_finance: dict[str, Any]) -> Quarter:
    monthly_headers = [h for h in current_finance["header"] if parse_month_header(h)]
    if monthly_headers:
        last_header = monthly_headers[-1]
        month_info = parse_month_header(last_header)
        if month_info:
            year, month = month_info
            quarter = ((month - 1) // 3) + 1
            return Quarter(year, quarter)
    for source in [current_title, current_path.stem]:
        quarter = parse_quarter_label(source)
        if quarter:
            return quarter
    return latest_quarter_from_current(current_finance)


def determine_target_quarter_with_debug(current_path: Path, current_title: str, current_finance: dict[str, Any]) -> tuple[Quarter, str]:
    monthly_headers = [h for h in current_finance["header"] if parse_month_header(h)]
    if monthly_headers:
        last_header = monthly_headers[-1]
        month_info = parse_month_header(last_header)
        if month_info:
            year, month = month_info
            quarter = Quarter(year, ((month - 1) // 3) + 1)
            return quarter, f"month-based inference from headers {monthly_headers}"
    for source_name, source in [("title", current_title), ("filename", current_path.stem)]:
        quarter = parse_quarter_label(source)
        if quarter:
            return quarter, f"fallback quarter-label inference from {source_name}: {source}"
    quarter = latest_quarter_from_current(current_finance)
    return quarter, f"fallback latest_quarter_from_current using finance headers {current_finance['header']}"


# Cross-language metric aliases: map English labels to their Chinese equivalents and vice versa
METRIC_ALIASES: dict[str, list[str]] = {
    "revenue": ["收入", "收入金额", "营业收入", "主营业务收入", "销售收入"],
    "gross profit": ["毛利", "毛利润", "毛利额"],
    "net profit": ["净利", "净利润", "净利额", "纯利"],
    "net income": ["净利", "净利润", "净收入"],
    "ebitda": ["ebitda", "息税折旧摊销前利润"],
    "cash inflow": ["现金流入", "经营性现金流入"],
    "cash outflow": ["现金支出", "经营性现金支出"],
    "burn rate": ["烧钱率", "月均净现金消耗"],
    "收入": ["revenue", "net revenue", "收入金额"],
    "毛利": ["gross profit", "毛利润"],
    "净利": ["net profit", "net income", "净利润"],
}


def choose_finance_row_label(
    previous_label: str,
    finance_rows: dict[str, dict[str, str]],
    current_total_header: str,
    previous_current_value: str | None,
    flags: list[dict[str, Any]],
    used_source_labels: dict[str, str] | None = None,
    source_unit_spec: "UnitSpec | None" = None,
    output_unit_spec: "UnitSpec | None" = None,
    exchange_rates: "ExchangeRates | None" = None,
) -> str | None:
    prev_tokens = metric_tokens(previous_label)
    used_source_labels = used_source_labels or {}

    exact_matches = []
    for candidate, values in finance_rows.items():
        if candidate in used_source_labels and used_source_labels[candidate] != previous_label:
            continue
        norm_cand = normalize_metric_label(candidate)
        norm_prev = normalize_metric_label(previous_label)
        if norm_cand == norm_prev:
            exact_matches.append(candidate)
    if len(exact_matches) == 1:
        return exact_matches[0]
    if len(exact_matches) > 1:
        flags.append(
            {
                "id": f"mapping-{normalize_metric_label(previous_label).replace(' ', '-')}",
                "section": "Financial Update",
                "severity": "warning",
                "message": f"Multiple exact metric matches found for '{previous_label}': {', '.join(exact_matches)}.",
                "source": "current_data_request",
            }
        )
        return exact_matches[0]

    # Try cross-language alias matching before fuzzy
    norm_prev = normalize_metric_label(previous_label)
    aliases = METRIC_ALIASES.get(norm_prev, [])
    alias_matches = []
    for candidate, values in finance_rows.items():
        if candidate in used_source_labels and used_source_labels[candidate] != previous_label:
            continue
        norm_cand = normalize_metric_label(candidate)
        for alias in aliases:
            if norm_cand == normalize_metric_label(alias) or normalize_metric_label(alias) in norm_cand:
                alias_matches.append(candidate)
                break
    if len(alias_matches) == 1:
        return alias_matches[0]
    if len(alias_matches) > 1:
        return alias_matches[0]

    best_label: str | None = None
    best_score: tuple[int, Decimal] | None = None
    ambiguous = False

    # Also check for Chinese substring containment: if one label contains the other
    # (e.g. "收入" in "收入金额"), give bonus token score
    def containment_bonus(prev_label: str, candidate_label: str) -> int:
        prev_cn = "".join(re.findall(r"[\u4e00-\u9fff]", prev_label))
        cand_cn = "".join(re.findall(r"[\u4e00-\u9fff]", candidate_label))
        if prev_cn and cand_cn:
            if prev_cn in cand_cn or cand_cn in prev_cn:
                return 5
        return 0

    for candidate, values in finance_rows.items():
        if candidate in used_source_labels and used_source_labels[candidate] != previous_label:
            continue
        candidate_value = parse_decimal(values.get(current_total_header, ""))
        if candidate_value is None:
            continue
        cand_tokens = metric_tokens(candidate)
        token_score = len(prev_tokens & cand_tokens) + containment_bonus(previous_label, candidate)

        previous_value = parse_decimal(previous_current_value or "")
        continuity_penalty = Decimal("999999")
        if previous_value is not None:
            # Convert candidate value to output units using actual unit specs
            if source_unit_spec is not None and output_unit_spec is not None:
                candidate_converted = convert_value_between_units(candidate_value, source_unit_spec, output_unit_spec, exchange_rates) or candidate_value
            else:
                candidate_converted = candidate_value / Decimal("1000")
            continuity_penalty = abs(candidate_converted - previous_value)

        # Token overlap is PRIMARY (higher = better); continuity penalty is tiebreaker (lower = better)
        score = (token_score, -continuity_penalty)
        if best_score is None or score > best_score:
            best_label = candidate
            best_score = score
            ambiguous = False
        elif score == best_score:
            ambiguous = True

    if best_label is None or best_score is None:
        if any(
            normalize_metric_label(candidate) == normalize_metric_label(previous_label)
            for candidate in finance_rows
        ):
            flags.append(
                {
                    "id": f"mapping-{normalize_metric_label(previous_label).replace(' ', '-')}",
                    "section": "Financial Update",
                    "severity": "warning",
                    "message": f"Metric '{previous_label}' had a potential source-row collision and was left unmapped.",
                    "source": "current_data_request",
                }
            )
            return None
        flags.append(
            {
                "id": f"mapping-{normalize_metric_label(previous_label).replace(' ', '-')}",
                "section": "Financial Update",
                "severity": "warning",
                "message": f"Could not map previous-review metric '{previous_label}' to a current data request row.",
                "source": "current_data_request",
            }
        )
        return None

    if ambiguous:
        flags.append(
            {
                "id": f"mapping-{normalize_metric_label(previous_label).replace(' ', '-')}",
                "section": "Financial Update",
                "severity": "warning",
                "message": f"Metric mapping for '{previous_label}' is ambiguous; selected '{best_label}' by token overlap.",
                "source": "current_data_request",
            }
        )
    return best_label


def build_financial_update(
    prev_financial: dict[str, Any],
    current_finance: dict[str, Any],
    current_operation_rows: list[list[str]] | None,
    flags: list[dict[str, Any]],
    target_quarter: Quarter | None = None,
    source_unit_spec: UnitSpec | None = None,
    output_unit_spec: UnitSpec | None = None,
    previous_review_quarter: Quarter | None = None,
    exchange_rates: ExchangeRates | None = None,
) -> dict[str, Any]:
    previous_lookup = previous_financial_lookup(prev_financial)
    finance_rows = current_finance["rows"]
    target_quarter = target_quarter or latest_quarter_from_current(current_finance)

    normalized_previous_headers = normalize_previous_quarter_headers(prev_financial["headers"], previous_review_quarter)
    normalized_previous_quarter_labels = [normalized for _, normalized in normalized_previous_headers]

    rolling_headers = [target_quarter.display()]
    quarter_cursor = target_quarter.previous()
    while len(rolling_headers) < 5:
        rolling_headers.append(quarter_cursor.display())
        quarter_cursor = quarter_cursor.previous()

    final_total_header = f"{target_quarter.year} FY" if target_quarter.quarter == 4 else f"{target_quarter.year} YTD"
    estimate_header = f"{target_quarter.year + 1}E" if target_quarter.quarter == 4 else f"{target_quarter.year}E"
    output_headers = rolling_headers + ["QoQ", "YoY", final_total_header, estimate_header]
    current_total_header = find_current_quarter_total_header(current_finance, target_quarter)
    current_month_headers = find_month_headers_for_quarter(current_finance, target_quarter)
    fy_or_ytd_header = find_fy_or_ytd_header(current_finance, target_quarter)
    if current_total_header is None:
        current_total_header = ""
    operation_rows = current_operation_rows or [["", ""]]
    company_outlets, jv_outlets = extract_outlet_counts(operation_rows)
    company_outlets_total, jv_outlets_total = extract_outlet_counts(operation_rows)

    rows_output: list[dict[str, Any]] = []
    used_source_labels: dict[str, str] = {}
    for prev_row in prev_financial["rows"]:
        label = prev_row["label"]
        if label == "#of Company-Owned Outlets":
            source_info = {"type": "derived", "label": "company_outlets"}
        elif label == "#of JV Outlets":
            source_info = {"type": "derived", "label": "jv_outlets"}
        else:
            previous_current_value = None
            if previous_review_quarter is not None:
                for raw, normalized in normalized_previous_headers:
                    if normalized == previous_review_quarter.display():
                        previous_current_value = prev_row["values"].get(raw)
                        break
            mapped_label = choose_finance_row_label(
                label,
                finance_rows,
                current_total_header,
                previous_current_value,
                flags,
                used_source_labels,
                source_unit_spec,
                output_unit_spec,
                exchange_rates,
            )
            source_info = {"type": "finance_row", "label": mapped_label} if mapped_label else None
            if mapped_label:
                used_source_labels[mapped_label] = label
        row_flags: list[str] = []
        values: dict[str, str | None] = {}
        source_trace: dict[str, str] = {}
        backfilled_current_year_history = False

        source_row = finance_rows.get(source_info["label"]) if source_info and source_info["type"] == "finance_row" else None
        prev_quarters = {normalized: prev_row["values"].get(raw) for raw, normalized in normalized_previous_headers}

        current_quarter_value: str | None = None
        if source_info:
            if source_info["type"] == "finance_row" and source_row:
                decimal_value = sum_month_values(source_row, current_month_headers)
                if decimal_value is None and current_total_header:
                    raw_current = source_row.get(current_total_header)
                    decimal_value = parse_decimal(raw_current)
                if decimal_value is None and finance_row_current_period_is_blank(
                    source_row,
                    current_month_headers,
                    current_total_header,
                    fy_or_ytd_header,
                ):
                    decimal_value = Decimal("0")
                if decimal_value is not None:
                    converted_value = convert_value_between_units(decimal_value, source_unit_spec, output_unit_spec, exchange_rates)
                    current_quarter_value = format_decimal(converted_value)
                    source_trace[target_quarter.display()] = "current_data_request"
                else:
                    row_flags.append("Missing current-quarter value in the company data request.")
            elif source_info["type"] == "finance_row" and source_row is None:
                row_flags.append("No current data request row could be mapped to this metric.")
            elif source_info["label"] == "company_outlets":
                current_quarter_value = company_outlets
                if company_outlets is None:
                    row_flags.append("Could not parse company-owned outlet count from operation data.")
                else:
                    source_trace[target_quarter.display()] = "current_data_request"
            elif source_info["label"] == "jv_outlets":
                current_quarter_value = jv_outlets
                if jv_outlets is None:
                    row_flags.append("Could not parse JV outlet count from operation data.")
                else:
                    source_trace[target_quarter.display()] = "current_data_request"

        for header in rolling_headers:
            if header == target_quarter.display():
                values[header] = current_quarter_value
                continue
            values[header] = prev_quarters.get(header) or "0"
            if prev_quarters.get(header):
                source_trace[header] = "previous_review"
            elif parse_quarter_label(header) and parse_quarter_label(header).year == target_quarter.year:
                backfilled_current_year_history = True

        qoq_value = None
        yoy_value = None

        if source_info and source_info["type"] == "finance_row" and source_row:
            row_current_dec = parse_decimal(values[target_quarter.display()] or "")
            if row_current_dec is not None and label in {"Revenue", "Gross profit", "EBITDA"}:
                row_current_dec = Decimal(values[target_quarter.display()])  # same unit in output

            previous_q_dec = parse_decimal(values[rolling_headers[1]] or "")
            yoy_dec = parse_decimal(values[rolling_headers[4]] or "")
            calc_qoq = calc_percent(row_current_dec, previous_q_dec)
            calc_yoy = calc_percent(row_current_dec, yoy_dec)
            qoq_value = format_percent(calc_qoq)
            yoy_value = format_percent(calc_yoy)
            if qoq_value is None:
                row_flags.append("QoQ could not be determined from the rolled quarter values.")
            if yoy_value is None:
                row_flags.append("YoY could not be determined from the rolled quarter values.")

            current_fy_source = parse_decimal(source_row.get(fy_or_ytd_header or "", ""))
            estimate_source = parse_decimal(find_header_value(source_row, ["exp", "预计", "全年预计"]) or "")
            if current_fy_source is not None and not backfilled_current_year_history:
                values[final_total_header] = format_decimal(convert_value_between_units(current_fy_source, source_unit_spec, output_unit_spec, exchange_rates))
                source_trace[final_total_header] = "current_data_request"
            else:
                q_values = [parse_decimal(values[h] or "") for h in rolling_headers if parse_quarter_label(h) and parse_quarter_label(h).year == target_quarter.year]
                if all(v is not None for v in q_values):
                    values[final_total_header] = format_decimal(sum(q_values))  # type: ignore[arg-type]
                    source_trace[final_total_header] = "derived_from_quarters"
                else:
                    values[final_total_header] = "0.00"
                    row_flags.append(f"{final_total_header} is missing and was defaulted to 0.00.")
            if estimate_source is not None:
                values[estimate_header] = format_decimal(convert_value_between_units(estimate_source, source_unit_spec, output_unit_spec, exchange_rates))
                source_trace[estimate_header] = "current_data_request"
            else:
                values[estimate_header] = None
                row_flags.append(f"{estimate_header} is missing in the current data request.")

        else:
            row_current_dec = parse_decimal(values[target_quarter.display()] or "")
            previous_q_dec = parse_decimal(values[rolling_headers[1]] or "")
            yoy_dec = parse_decimal(values[rolling_headers[4]] or "")
            qoq_value = format_percent(calc_percent(row_current_dec, previous_q_dec))
            yoy_value = format_percent(calc_percent(row_current_dec, yoy_dec))
            if qoq_value is None:
                row_flags.append("QoQ could not be determined from the rolled quarter values.")
            if yoy_value is None:
                row_flags.append("YoY could not be determined from the rolled quarter values.")

            if source_info and source_info["label"] == "company_outlets":
                values[final_total_header] = company_outlets_total
                if company_outlets_total is not None:
                    source_trace[final_total_header] = "current_data_request"
                else:
                    values[final_total_header] = "0"
                    row_flags.append(f"{final_total_header} is missing for company-owned outlets and was defaulted to 0.")
            elif source_info and source_info["label"] == "jv_outlets":
                values[final_total_header] = jv_outlets_total
                if jv_outlets_total is not None:
                    source_trace[final_total_header] = "current_data_request"
                else:
                    values[final_total_header] = "0"
                    row_flags.append(f"{final_total_header} is missing for JV outlets and was defaulted to 0.")
            values[estimate_header] = None
            row_flags.append(f"{estimate_header} is missing in the current data request.")

        values["QoQ"] = qoq_value
        values["YoY"] = yoy_value

        rows_output.append(
            {
                "label": label,
                "mapped_source_label": source_info["label"] if source_info else None,
                "values": {header: values.get(header) for header in output_headers},
                "source_trace": source_trace,
                "flags": row_flags,
            }
        )

    return {
        "status": "updated_with_flags",
        "unit": prev_financial["unit"] or "in millions USD",
        "columns": output_headers,
        "rows": rows_output,
    }


def build_markdown(
    business_activities: str,
    financial_update: dict[str, Any],
    business_update_bullets: list[str],
    risk_exit_bullets: list[str],
) -> str:
    lines = [
        "# Portfolio Review Draft",
        "",
        "## 1. Business Activities",
        business_activities,
        "",
        "## 2. Financial Update",
        f"({financial_update['unit']})",
        "",
    ]

    headers = [financial_update["unit"]] + financial_update["columns"]
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in financial_update["rows"]:
        row_values = [row["label"]] + [row["values"].get(header) or "" for header in financial_update["columns"]]
        lines.append("| " + " | ".join(row_values) + " |")

    lines.extend(["", "## 3. Business Update"])
    for bullet in business_update_bullets:
        lines.append(f"- {bullet}")

    lines.extend(["", "## 4. Risk & Exit"])
    for bullet in risk_exit_bullets:
        lines.append(f"- {bullet}")

    return "\n".join(lines) + "\n"


def split_evenly(items: list[str], count: int) -> list[list[str]]:
    if count <= 0:
        return []
    if not items:
        return [[] for _ in range(count)]
    base, remainder = divmod(len(items), count)
    chunks: list[list[str]] = []
    cursor = 0
    for idx in range(count):
        size = base + (1 if idx < remainder else 0)
        chunks.append(items[cursor : cursor + size])
        cursor += size
    return chunks


def detect_doc_section_occurrences(doc: Document) -> list[DocSectionOccurrence]:
    occurrences: list[DocSectionOccurrence] = []
    after_financial: dict[str, bool] = {"english": False, "chinese": False}
    after_risk: dict[str, bool] = {"english": False, "chinese": False}
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        language = text_language(text)
        canonical = fixed_anchor_heading(text)
        if canonical:
            occurrences.append(
                DocSectionOccurrence(
                    canonical=canonical,
                    language=language,
                    heading_index=idx,
                )
            )
            if canonical == "Financial Update":
                after_financial[language] = True
                after_risk[language] = False
            elif canonical == "Risk & Exit":
                after_risk[language] = True
            continue

        canonical = canonical_heading(text)
        if canonical == "Business Update":
            occurrences.append(
                DocSectionOccurrence(
                    canonical=canonical,
                    language=language,
                    heading_index=idx,
                )
            )
            continue

        if after_financial[language] and not after_risk[language] and is_middle_section_heading_candidate(paragraph):
            occurrences.append(
                DocSectionOccurrence(
                    canonical="Business Update",
                    language=language,
                    heading_index=idx,
                )
            )
    return occurrences


def resolve_section_plan_to_doc(doc: Document, section_plan: list[SectionOccurrence]) -> list[DocSectionOccurrence]:
    occurrences: list[DocSectionOccurrence] = []
    paragraph_index = 0
    for planned in section_plan:
        target = norm_space(planned.heading_text)
        for idx in range(paragraph_index, len(doc.paragraphs)):
            if norm_space(doc.paragraphs[idx].text) == target:
                occurrences.append(
                    DocSectionOccurrence(
                        canonical=planned.canonical,
                        language=planned.language,
                        heading_index=idx,
                    )
                )
                paragraph_index = idx + 1
                break
    return occurrences


def find_resolved_occurrence_index(
    doc: Document,
    occurrences: list[DocSectionOccurrence],
    planned: SectionOccurrence,
) -> int | None:
    target = norm_space(planned.heading_text)
    for idx, occurrence in enumerate(occurrences):
        if occurrence.canonical != planned.canonical or occurrence.language != planned.language:
            continue
        current = norm_space(doc.paragraphs[occurrence.heading_index].text)
        if current == target:
            return idx
    fallback_matches = [
        idx
        for idx, occurrence in enumerate(occurrences)
        if occurrence.canonical == planned.canonical and occurrence.language == planned.language
    ]
    if len(fallback_matches) == 1:
        return fallback_matches[0]
    return None


def template_language_mode(occurrences: list[DocSectionOccurrence]) -> str:
    languages = {occ.language for occ in occurrences}
    if "english" in languages and "chinese" in languages:
        return "bilingual"
    if "english" in languages:
        return "English-only"
    return "Chinese-only"


def middle_section_indices_by_language(occurrences: list[DocSectionOccurrence]) -> dict[str, list[int]]:
    result: dict[str, list[int]] = {"english": [], "chinese": []}
    for language in ["english", "chinese"]:
        language_occurrences = [(idx, occ) for idx, occ in enumerate(occurrences) if occ.language == language]
        financial_idx = next((idx for idx, occ in language_occurrences if occ.canonical == "Financial Update"), None)
        risk_idx = next((idx for idx, occ in language_occurrences if occ.canonical == "Risk & Exit" and financial_idx is not None and idx > financial_idx), None)
        if financial_idx is None:
            continue
        upper = risk_idx if risk_idx is not None else len(occurrences)
        for idx, occ in enumerate(occurrences):
            if occ.language != language:
                continue
            if idx <= financial_idx or idx >= upper:
                continue
            if occ.canonical in {"Business Activities", "Financial Update", "Risk & Exit"}:
                continue
            result[language].append(idx)
    return result


def get_body_paragraphs_for_occurrence(doc: Document, occurrences: list[DocSectionOccurrence], occurrence_index: int):
    start = occurrences[occurrence_index].heading_index + 1
    end = occurrences[occurrence_index + 1].heading_index if occurrence_index + 1 < len(occurrences) else len(doc.paragraphs)
    return doc.paragraphs[start:end]


def set_paragraph_text_preserve(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text_preserve(cell, text: str) -> None:
    """Set cell text while preserving paragraph and run formatting.
    Unlike cell.text = value, this doesn't destroy the cell's formatting."""
    # Remove ALL paragraphs except the first
    while len(cell.paragraphs) > 1:
        last_p = cell.paragraphs[-1]
        last_p._element.getparent().remove(last_p._element)
    p = cell.paragraphs[0]
    # Remove ALL child elements except pPr (paragraph properties)
    for child in list(p._element):
        if child.tag != qn("w:pPr"):
            p._element.remove(child)
    # Add a single clean run with the text
    run_elem = OxmlElement("w:r")
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")
    run_elem.append(t_elem)
    p._element.append(run_elem)
    # Force compact spacing by creating/replacing <w:spacing> in <w:pPr>
    pPr = p._element.get_or_add_pPr()
    old_spacing = pPr.find(qn("w:spacing"))
    if old_spacing is not None:
        pPr.remove(old_spacing)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def table_display_text(text: str, language: str) -> str:
    text = text.replace("\xa0", " ")
    if re.match(r"^-\d", text):
        text = "\u2212" + text[1:]
    if text.endswith("%"):
        text = text[:-1] + "\u2060%"
    if language != "english":
        return text
    normalized = text
    if is_unit_text(normalized) and " " in normalized:
        return normalized.replace(" ", "\xa0")
    if normalized.startswith("Q") and re.match(r"^Q[1-4] 20\d{2}$", normalized):
        return normalized.replace(" ", "\xa0")
    protected = {
        "Gross Profit",
        "Net Profit",
        "Cash Flow",
        "Business Update",
        "Risk & Exit",
        "Financial Update",
        "Financial Updates",
        "2025 FY",
        "2026E",
    }
    if normalized in protected:
        return normalized.replace(" ", "\xa0")
    if " " in normalized and re.fullmatch(r"[A-Za-z#()&/\- ]+", normalized):
        return normalized.replace(" ", "\xa0")
    return normalized


def set_cell_no_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is None:
        no_wrap = OxmlElement("w:noWrap")
        tc_pr.append(no_wrap)


def clear_cell_no_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is not None:
        tc_pr.remove(no_wrap)


def get_table_layout_type(table) -> str | None:
    tbl_layout = table._tbl.tblPr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        return None
    return tbl_layout.get(qn("w:type"))


def get_table_grid_widths(table) -> list[int]:
    grid = table._tbl.tblGrid
    if grid is None:
        return []
    widths: list[int] = []
    for col in grid.iterchildren():
        width = col.get(qn("w:w"))
        widths.append(int(width) if width else 0)
    return widths


def estimate_table_text_width(text: str, language: str) -> int:
    total = 180
    for char in text:
        if char == "\u2060":
            continue
        if char in {" ", "\xa0"}:
            total += 35
        elif char.isdigit():
            total += 80
        elif char in {".", ",", ":", ";", "!", "|"}:
            total += 40
        elif char in {"-", "\u2212", "%", "/", "&", "#", "(", ")"}:
            total += 55
        elif ord(char) > 127 or unicodedata.east_asian_width(char) in {"W", "F"}:
            total += 120
        elif char.isupper():
            total += 95
        else:
            total += 85
    if language != "english":
        total += 40
    return total


def financial_column_min_width(header_text: str, language: str, col_idx: int) -> int:
    normalized = header_text.replace("\xa0", " ").strip()
    if col_idx == 0:
        return 1320 if language == "english" else 1200
    if re.fullmatch(r"Q[1-4] 20\d{2}", normalized):
        return 780 if language == "english" else 940
    if normalized in {"QoQ", "YoY"}:
        return 800 if language == "english" else 1030
    if normalized.endswith("FY") or normalized.endswith("YTD"):
        return 860 if language == "english" else 980
    if re.fullmatch(r"20\d{2}E", normalized):
        return 660 if language == "english" else 900
    return 0


def page_usable_width_twips(doc: Document) -> int | None:
    try:
        section = doc.sections[0]
    except IndexError:
        return None
    usable_emu = section.page_width - section.left_margin - section.right_margin
    if usable_emu is None:
        return None
    return max(0, int(round(usable_emu / 635)))


def rebalance_table_widths(widths: list[int], required: list[int], minimums: list[int], max_total: int | None = None) -> list[int]:
    if not widths or len(widths) != len(required) or len(widths) != len(minimums):
        return widths
    adjusted = [max(widths[i], minimums[i], required[i]) for i in range(len(widths))]
    if max_total is not None and sum(adjusted) > max_total:
        excess = sum(adjusted) - max_total
        shrinkable = [max(0, adjusted[i] - minimums[i]) for i in range(len(widths))]
        total_shrinkable = sum(shrinkable)
        if total_shrinkable > 0:
            for i, available in enumerate(shrinkable):
                if excess <= 0:
                    break
                if available <= 0:
                    continue
                shrink = min(available, round((sum(adjusted) - max_total) * available / total_shrinkable))
                adjusted[i] -= shrink
                excess -= shrink
            if excess > 0:
                for i, available in enumerate(shrinkable):
                    if excess <= 0:
                        break
                    extra = min(max(0, adjusted[i] - minimums[i]), excess)
                    adjusted[i] -= extra
                    excess -= extra
    return adjusted


def set_table_grid_widths(table, widths: list[int], sync_cell_widths: bool = True) -> None:
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(1, tbl_grid)
    grid_cols = list(tbl_grid.iterchildren())
    while len(grid_cols) < len(widths):
        new_col = OxmlElement("w:gridCol")
        tbl_grid.append(new_col)
        grid_cols = list(tbl_grid.iterchildren())
    for idx, width in enumerate(widths):
        grid_cols[idx].set(qn("w:w"), str(int(width)))
    for extra in grid_cols[len(widths):]:
        tbl_grid.remove(extra)
    if sync_cell_widths:
        for row in table.rows:
            for idx, cell in enumerate(row.cells[: len(widths)]):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:type"), "dxa")
                tc_w.set(qn("w:w"), str(int(widths[idx])))


def fit_financial_table_columns(doc: Document, table, language: str) -> None:
    widths = get_table_grid_widths(table)
    layout_type = get_table_layout_type(table)
    if widths:
        required: list[int] = []
        minimums: list[int] = []
        num_cols = min(len(widths), max(len(row.cells) for row in table.rows))
        for col_idx in range(num_cols):
            max_width = widths[col_idx]
            header_text = ""
            if table.rows and col_idx < len(table.rows[0].cells):
                header_text = table.rows[0].cells[col_idx].text.replace("\n", " ").strip()
            for row in table.rows:
                if col_idx >= len(row.cells):
                    continue
                cell_text = row.cells[col_idx].text.replace("\n", " ")
                if not cell_text:
                    continue
                max_width = max(max_width, estimate_table_text_width(cell_text, language))
            minimum_width = financial_column_min_width(header_text, language, col_idx)
            max_width = max(max_width, minimum_width)
            required.append(max_width)
            minimums.append(minimum_width)
        adjusted = rebalance_table_widths(
            widths[: len(required)],
            required,
            minimums,
            page_usable_width_twips(doc),
        )
        set_table_grid_widths(table, adjusted, sync_cell_widths=(layout_type == "fixed"))
    for row in table.rows:
        for cell in row.cells:
            if layout_type == "fixed":
                set_cell_no_wrap(cell)
            else:
                clear_cell_no_wrap(cell)


def capture_run_format(paragraph) -> tuple[str | None, str | None, float | None, bool | None]:
    """Returns (font_name, east_asia_font, size_pt, bold)."""
    for run in paragraph.runs:
        font_name = run.font.name
        east_asia_font = None
        rpr = run._element.find(qn("w:rPr"))
        if rpr is not None:
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                east_asia_font = rfonts.get(qn("w:eastAsia"))
        size_pt = float(run.font.size.pt) if run.font.size is not None else None
        bold = run.font.bold
        return font_name, east_asia_font, size_pt, bold
    return None, None, None, None


def apply_paragraph_run_format(paragraph, font_name: str | None, east_asia_font: str | None, size_pt: float | None, bold: bool | None) -> None:
    for run in paragraph.runs:
        if font_name:
            run.font.name = font_name
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:ascii"), font_name)
            rfonts.set(qn("w:hAnsi"), font_name)
            rfonts.set(qn("w:cs"), font_name)
            # Use the East Asian font if captured, otherwise fall back to font_name
            rfonts.set(qn("w:eastAsia"), east_asia_font or font_name)
        elif east_asia_font:
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:eastAsia"), east_asia_font)
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        run.font.bold = bold


def clear_paragraph_preserve(paragraph, normal_style=None) -> None:
    set_paragraph_text_preserve(paragraph, "")
    if normal_style is not None:
        try:
            paragraph.style = normal_style
        except Exception:
            pass


def copy_paragraph_layout(target, source) -> None:
    try:
        target.style = source.style
    except Exception:
        pass
    target.alignment = source.alignment
    src_fmt = source.paragraph_format
    dst_fmt = target.paragraph_format
    dst_fmt.left_indent = src_fmt.left_indent
    dst_fmt.right_indent = src_fmt.right_indent
    dst_fmt.first_line_indent = src_fmt.first_line_indent
    dst_fmt.space_before = src_fmt.space_before
    dst_fmt.space_after = src_fmt.space_after
    dst_fmt.line_spacing = src_fmt.line_spacing
    dst_fmt.line_spacing_rule = src_fmt.line_spacing_rule
    dst_fmt.keep_together = src_fmt.keep_together
    dst_fmt.keep_with_next = src_fmt.keep_with_next
    dst_fmt.page_break_before = src_fmt.page_break_before
    dst_fmt.widow_control = src_fmt.widow_control
    # Copy bullet/numbering properties (numPr) from source to target
    src_pPr = source._element.find(qn("w:pPr"))
    if src_pPr is not None:
        src_numPr = src_pPr.find(qn("w:numPr"))
        if src_numPr is not None:
            tgt_pPr = target._element.get_or_add_pPr()
            old_numPr = tgt_pPr.find(qn("w:numPr"))
            if old_numPr is not None:
                tgt_pPr.remove(old_numPr)
            tgt_pPr.append(copy.deepcopy(src_numPr))


def paragraph_is_list_like(paragraph) -> bool:
    style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
    if "list" in style_name or "列表" in style_name:
        return True
    ppr = getattr(paragraph._element, "pPr", None)
    if ppr is not None and getattr(ppr, "numPr", None) is not None:
        return True
    stripped = paragraph.text.strip()
    return stripped.startswith(("-", "•", "*"))


def fill_section_paragraphs(doc: Document, paragraphs, texts: list[str], force_bullets: bool = False) -> None:
    slots = list(paragraphs)
    normal_style = None
    try:
        normal_style = doc.styles["Normal"]
    except Exception:
        pass

    template_style = None
    template_run_format = (None, None, None, None)
    for paragraph in slots:
        if paragraph.text.strip():
            template_style = paragraph.style
            template_run_format = capture_run_format(paragraph)
            break
    if template_style is None and slots:
        template_style = slots[0].style
        template_run_format = capture_run_format(slots[0])

    working_slots = slots
    bullet_reference = None
    if force_bullets:
        list_like = [paragraph for paragraph in slots if paragraph_is_list_like(paragraph)]
        if list_like:
            working_slots = list_like
            bullet_reference = list_like[0]
        elif slots:
            working_slots = slots
            bullet_reference = slots[0]

    remove_later = [paragraph for paragraph in slots if paragraph not in working_slots]
    for idx, paragraph in enumerate(working_slots):
        if idx < len(texts):
            set_paragraph_text_preserve(paragraph, texts[idx])
            if force_bullets and bullet_reference is not None:
                copy_paragraph_layout(paragraph, bullet_reference)
            apply_paragraph_run_format(paragraph, *template_run_format)
        else:
            remove_later.append(paragraph)

    if len(texts) <= len(working_slots):
        for paragraph in reversed(remove_later):
            element = paragraph._element
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
        return

    anchor = working_slots[-1] if working_slots else None
    for text in texts[len(working_slots) :]:
        if anchor is None:
            break
        new_paragraph = anchor.insert_paragraph_before(text)
        if force_bullets and bullet_reference is not None:
            copy_paragraph_layout(new_paragraph, bullet_reference)
        elif template_style is not None:
            try:
                new_paragraph.style = template_style
            except Exception:
                pass
        apply_paragraph_run_format(new_paragraph, *template_run_format)
        anchor = new_paragraph

    for paragraph in reversed(remove_later):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//*[local-name()='drawing' or local-name()='pict']"))


def paragraph_has_page_break(paragraph) -> bool:
    for element in paragraph._element.iter():
        if element.tag.endswith("}br") and element.get(qn("w:type")) == "page":
            return True
    return False


def is_plain_blank_paragraph(paragraph) -> bool:
    return paragraph.text.strip() == "" and not paragraph_has_drawing(paragraph) and not paragraph_has_page_break(paragraph)


def find_logo_info(doc: Document, template_path: Path) -> tuple[Path | None, Any | None, Any | None, Any | None]:
    logo_paragraph = None
    for paragraph in doc.paragraphs:
        if paragraph_has_drawing(paragraph):
            logo_paragraph = paragraph
            break
    if logo_paragraph is None:
        return None, None, None, None

    alignment = logo_paragraph.alignment
    style = logo_paragraph.style
    width = None
    height = None
    if doc.inline_shapes:
        shape = doc.inline_shapes[0]
        width = shape.width
        height = shape.height

    with zipfile.ZipFile(template_path) as zf:
        media_names = sorted(name for name in zf.namelist() if name.startswith("word/media/"))
        if not media_names:
            return None, alignment, style, None
        media_name = media_names[0]
        suffix = Path(media_name).suffix or ".png"
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        temp_file.write(zf.read(media_name))
        temp_file.flush()
        temp_file.close()
    return Path(temp_file.name), alignment, style, (width, height)


def insert_page_break_and_repeated_logo(doc: Document, template_path: Path, occurrences: list[DocSectionOccurrence] | None = None) -> None:
    occurrences = occurrences or detect_doc_section_occurrences(doc)
    english_occurrences = [occ for occ in occurrences if occ.language == "english"]
    chinese_occurrences = [occ for occ in occurrences if occ.language == "chinese"]
    if not english_occurrences or not chinese_occurrences:
        return

    english_start = doc.paragraphs[english_occurrences[0].heading_index]
    logo_path, alignment, style, size = find_logo_info(doc, template_path)

    # Insert logo first, then page break before it, so final order is:
    # [page break] [logo] [english heading]
    logo_ref = english_start  # reference for page break insertion
    if logo_path is not None:
        logo_paragraph = english_start.insert_paragraph_before("")
        if style is not None:
            try:
                logo_paragraph.style = style
            except Exception:
                pass
        logo_paragraph.alignment = alignment
        run = logo_paragraph.add_run()
        width, height = size if size is not None else (None, None)
        if width is not None:
            run.add_picture(str(logo_path), width=width, height=height)
        else:
            run.add_picture(str(logo_path))
        logo_ref = logo_paragraph  # insert page break before the logo
        try:
            logo_path.unlink()
        except Exception:
            pass

    break_paragraph = logo_ref.insert_paragraph_before("")
    break_paragraph.add_run().add_break(WD_BREAK.PAGE)


def remove_blank_paragraphs_inside_sections(doc: Document, occurrences: list[DocSectionOccurrence] | None = None) -> None:
    occurrences = occurrences or detect_doc_section_occurrences(doc)
    to_remove = []
    for occ_idx, _occurrence in enumerate(occurrences):
        body = list(get_body_paragraphs_for_occurrence(doc, occurrences, occ_idx))
        for paragraph in body:
            if is_plain_blank_paragraph(paragraph):
                to_remove.append(paragraph)
    for paragraph in reversed(to_remove):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def normalize_bullet_sections(doc: Document, occurrences: list[DocSectionOccurrence] | None = None) -> None:
    occurrences = occurrences or detect_doc_section_occurrences(doc)
    for occ_idx, occurrence in enumerate(occurrences):
        if occurrence.canonical not in {"Business Update", "Risk & Exit"}:
            continue
        body = list(get_body_paragraphs_for_occurrence(doc, occurrences, occ_idx))
        content_paragraphs = [p for p in body if p.text.strip()]
        if not content_paragraphs:
            continue
        bullet_reference = next((p for p in content_paragraphs if paragraph_is_list_like(p)), content_paragraphs[0])
        font_name, east_asia_font, size_pt, _bold = capture_run_format(bullet_reference)
        for paragraph in content_paragraphs:
            copy_paragraph_layout(paragraph, bullet_reference)
            # Always set bold=False for content paragraphs; headings get bold separately
            apply_paragraph_run_format(paragraph, font_name, east_asia_font, size_pt, False)


def normalize_exact_single_blank_between_sections(doc: Document, occurrences: list[DocSectionOccurrence] | None = None) -> None:
    """Insert exactly one blank paragraph before each section heading (except the very first)."""
    occurrences = occurrences or detect_doc_section_occurrences(doc)
    if not occurrences:
        return

    # Collect heading paragraph elements by matching text (indices may be stale)
    heading_texts = set()
    for occ in occurrences:
        if occ.heading_index < len(doc.paragraphs):
            heading_texts.add(norm_space(doc.paragraphs[occ.heading_index].text))

    heading_paragraphs = []
    for p in doc.paragraphs:
        if norm_space(p.text) in heading_texts:
            heading_paragraphs.append(p)

    for idx, paragraph in enumerate(heading_paragraphs):
        # Skip the very first heading in the document (no blank needed above it)
        if idx == 0:
            continue

        # Check if there's already a blank paragraph immediately before this heading
        previous = paragraph._element.getprevious()
        if previous is not None:
            prev_p = next((p for p in doc.paragraphs if p._element is previous), None)
            if prev_p is not None and is_plain_blank_paragraph(prev_p):
                continue  # already has a blank

        # Insert a blank paragraph before this heading
        blank = paragraph.insert_paragraph_before("")
        try:
            blank.style = doc.styles["Normal"]
        except Exception:
            pass


def center_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")


def financial_table_backbone_score(prev_financial: dict[str, Any], previous_review_quarter: Quarter | None) -> tuple[int, int]:
    normalized_headers = normalize_previous_quarter_headers(prev_financial["headers"], previous_review_quarter)
    labels = [normalized for _raw, normalized in normalized_headers]
    expected: list[str] = []
    if previous_review_quarter is not None:
        cursor = previous_review_quarter
        for _ in range(5):
            expected.append(cursor.display())
            cursor = cursor.previous()
    coverage = len([label for label in labels if label in expected]) if expected else len(labels)
    nonempty = 0
    for row in prev_financial["rows"]:
        for raw, _normalized in normalized_headers:
            if norm_space(row["values"].get(raw, "")):
                nonempty += 1
    return coverage, nonempty


def localize_financial_update(source_update: dict[str, Any], target_prev_financial: dict[str, Any]) -> dict[str, Any]:
    localized_rows: list[dict[str, Any]] = []
    for idx, source_row in enumerate(source_update["rows"]):
        target_label = target_prev_financial["rows"][idx]["label"] if idx < len(target_prev_financial["rows"]) else source_row["label"]
        localized_rows.append(
            {
                "label": target_label,
                "values": dict(source_row["values"]),
                "source_trace": dict(source_row.get("source_trace", {})),
                "flags": list(source_row.get("flags", [])),
            }
        )
    return {
        "status": source_update.get("status"),
        "unit": target_prev_financial.get("unit") or source_update.get("unit"),
        "columns": list(source_update["columns"]),
        "rows": localized_rows,
    }


def section_table_entries(section_plan: list[SectionOccurrence], canonical_filter: set[str] | None = None) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    global_table_index = 0
    for occurrence in section_plan:
        for local_table_index, _rows in enumerate(occurrence.tables):
            if canonical_filter is None or occurrence.canonical in canonical_filter:
                entries.append(
                    {
                        "occurrence": occurrence,
                        "local_table_index": local_table_index,
                        "global_table_index": global_table_index,
                    }
                )
            global_table_index += 1
    return entries


def group_middle_table_occurrences(previous_occurrences: list[SectionOccurrence]) -> list[list[SectionOccurrence]]:
    groups: dict[int, list[SectionOccurrence]] = {}
    counters: dict[str, int] = {"chinese": 0, "english": 0}
    for occurrence in previous_occurrences:
        if occurrence.canonical != "Business Update" or not occurrence.tables:
            continue
        group_index = counters.get(occurrence.language, 0)
        counters[occurrence.language] = group_index + 1
        groups.setdefault(group_index, []).append(occurrence)
    return [groups[idx] for idx in sorted(groups)]


def build_section_table_update(
    prev_table: dict[str, Any],
    current_table: dict[str, Any],
    flags: list[dict[str, Any]],
    target_quarter: Quarter | None = None,
    source_unit_spec: UnitSpec | None = None,
    output_unit_spec: UnitSpec | None = None,
    previous_review_quarter: Quarter | None = None,
    exchange_rates: ExchangeRates | None = None,
) -> dict[str, Any]:
    target_quarter = target_quarter or latest_quarter_from_current(current_table)
    normalized_previous_headers = normalize_previous_quarter_headers(prev_table["headers"], previous_review_quarter)
    previous_quarter_labels = [normalized for _raw, normalized in normalized_previous_headers if normalized]
    quarter_count = len(previous_quarter_labels)
    rolling_headers: list[str] = []
    quarter_cursor = target_quarter
    while len(rolling_headers) < quarter_count:
        rolling_headers.append(quarter_cursor.display())
        quarter_cursor = quarter_cursor.previous()

    trailing_headers = [header for header in prev_table["headers"] if not parse_quarter_label(header)]
    output_headers = rolling_headers + trailing_headers
    current_total_header = find_current_quarter_total_header(current_table, target_quarter) or ""
    current_month_headers = find_month_headers_for_quarter(current_table, target_quarter)
    fy_or_ytd_header = find_fy_or_ytd_header(current_table, target_quarter)

    rows_output: list[dict[str, Any]] = []
    used_source_labels: dict[str, str] = {}
    for prev_row in prev_table["rows"]:
        label = prev_row["label"]
        previous_current_value = None
        if previous_review_quarter is not None:
            for raw, normalized in normalized_previous_headers:
                if normalized == previous_review_quarter.display():
                    previous_current_value = prev_row["values"].get(raw)
                    break
        mapped_label = choose_finance_row_label(
            label,
            current_table["rows"],
            current_total_header,
            previous_current_value,
            flags,
            used_source_labels,
            source_unit_spec,
            output_unit_spec,
            exchange_rates,
        )
        if mapped_label:
            used_source_labels[mapped_label] = label
        source_row = current_table["rows"].get(mapped_label) if mapped_label else None

        values: dict[str, str | None] = {}
        source_trace: dict[str, str] = {}
        row_flags: list[str] = []
        prev_quarters = {normalized: prev_row["values"].get(raw) for raw, normalized in normalized_previous_headers}

        current_quarter_value: str | None = None
        if source_row:
            decimal_value = sum_month_values(source_row, current_month_headers)
            if decimal_value is None and current_total_header:
                decimal_value = parse_decimal(source_row.get(current_total_header, ""))
            if decimal_value is None and finance_row_current_period_is_blank(source_row, current_month_headers, current_total_header, fy_or_ytd_header):
                decimal_value = Decimal("0")
            if decimal_value is not None:
                converted_value = convert_value_between_units(decimal_value, source_unit_spec, output_unit_spec, exchange_rates)
                current_quarter_value = format_decimal(converted_value)
                source_trace[target_quarter.display()] = "current_data_request"
            else:
                row_flags.append("Missing current-quarter value in the company data request.")
        else:
            row_flags.append("No current data request row could be mapped to this metric.")

        for header in rolling_headers:
            if header == target_quarter.display():
                values[header] = current_quarter_value
            else:
                values[header] = prev_quarters.get(header)
                if prev_quarters.get(header):
                    source_trace[header] = "previous_review"

        for header in trailing_headers:
            normalized_header = header.replace("\xa0", " ").strip()
            if normalized_header in {"QoQ", "YoY"}:
                row_current_dec = parse_decimal(values.get(target_quarter.display()) or "")
                comparison_header = rolling_headers[1] if normalized_header == "QoQ" and len(rolling_headers) > 1 else rolling_headers[4] if normalized_header == "YoY" and len(rolling_headers) > 4 else None
                comparison_value = parse_decimal(values.get(comparison_header, "") or "") if comparison_header else None
                values[header] = format_percent(calc_percent(row_current_dec, comparison_value))
            elif normalized_header.endswith("FY") or normalized_header.endswith("YTD"):
                source_value = parse_decimal(source_row.get(fy_or_ytd_header or "", "")) if source_row and fy_or_ytd_header else None
                if source_value is not None:
                    values[header] = format_decimal(convert_value_between_units(source_value, source_unit_spec, output_unit_spec, exchange_rates))
                    source_trace[header] = "current_data_request"
                else:
                    quarter_values = [parse_decimal(values.get(qh) or "") for qh in rolling_headers if qh.startswith(f"Q") and qh.endswith(str(target_quarter.year))]
                    if quarter_values and all(v is not None for v in quarter_values):
                        values[header] = format_decimal(sum(quarter_values))  # type: ignore[arg-type]
                        source_trace[header] = "derived_from_quarters"
                    else:
                        values[header] = prev_row["values"].get(header) or None
            elif re.fullmatch(r"20\d{2}E", normalized_header):
                estimate_source = None
                if source_row:
                    estimate_source = parse_decimal(source_row.get(header, ""))
                    if estimate_source is None:
                        estimate_source = parse_decimal(find_header_value(source_row, ["exp", "预计", "全年预计"]) or "")
                if estimate_source is not None:
                    values[header] = format_decimal(convert_value_between_units(estimate_source, source_unit_spec, output_unit_spec, exchange_rates))
                    source_trace[header] = "current_data_request"
                else:
                    values[header] = prev_row["values"].get(header) or None
            else:
                if source_row:
                    raw_source_value = source_row.get(header, "")
                    decimal_source = parse_decimal(raw_source_value)
                    if decimal_source is not None:
                        values[header] = format_decimal(convert_value_between_units(decimal_source, source_unit_spec, output_unit_spec, exchange_rates))
                        source_trace[header] = "current_data_request"
                    else:
                        values[header] = norm_space(raw_source_value) or prev_row["values"].get(header) or None
                else:
                    values[header] = prev_row["values"].get(header) or None

        rows_output.append(
            {
                "label": label,
                "mapped_source_label": mapped_label,
                "values": {header: values.get(header) for header in output_headers},
                "source_trace": source_trace,
                "flags": row_flags,
            }
        )

    return {
        "status": "updated_with_flags",
        "unit": prev_table.get("unit") or "",
        "columns": output_headers,
        "rows": rows_output,
    }


def enforce_bilingual_page_break(doc: Document, occurrences: list[DocSectionOccurrence]) -> None:
    english_occurrences = [occ for occ in occurrences if occ.language == "english"]
    chinese_occurrences = [occ for occ in occurrences if occ.language == "chinese"]
    if not english_occurrences or not chinese_occurrences:
        return

    first_english_heading = english_occurrences[0].heading_index
    start_idx = first_english_heading
    while start_idx > 0:
        previous = doc.paragraphs[start_idx - 1]
        if is_plain_blank_paragraph(previous) or paragraph_has_drawing(previous):
            start_idx -= 1
            continue
        break

    # Remove blank separator paragraphs immediately before the English block so the
    # Chinese page does not end with a large empty gap.
    to_remove = []
    idx = start_idx - 1
    while idx >= 0 and is_plain_blank_paragraph(doc.paragraphs[idx]):
        to_remove.append(doc.paragraphs[idx])
        idx -= 1
    for paragraph in reversed(to_remove):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    target = doc.paragraphs[start_idx - len(to_remove)]
    previous = target._element.getprevious()
    previous_para = next((p for p in doc.paragraphs if p._element is previous), None) if previous is not None else None
    if previous_para is not None and paragraph_has_page_break(previous_para):
        return

    break_paragraph = target.insert_paragraph_before("")
    break_paragraph.add_run().add_break(WD_BREAK.PAGE)


def write_docx(
    output_path: Path,
    template_path: Path,
    section_plan: list[SectionOccurrence],
    business_activities_map: dict[str, str],
    financial_updates: list[dict[str, Any]],
    middle_table_updates: dict[tuple[str, str, int], dict[str, Any]] | None,
    business_update_map: dict[str, list[str]],
    risk_exit_map: dict[str, list[str]],
    financial_note_map: dict[str, str] | None = None,
) -> None:
    doc = Document(str(template_path))
    financial_note_map = financial_note_map or {}
    middle_table_updates = middle_table_updates or {}

    def apply_table_run_format(run, font_name: str | None, east_asia_font: str | None, size_pt: float | None, bold: bool | None) -> None:
        if font_name:
            run.font.name = font_name
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:ascii"), font_name)
            rfonts.set(qn("w:hAnsi"), font_name)
            rfonts.set(qn("w:cs"), font_name)
            rfonts.set(qn("w:eastAsia"), east_asia_font or font_name)
        elif east_asia_font:
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:eastAsia"), east_asia_font)
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        run.font.bold = bold

    # --- Fill Business Activities (re-resolve before each fill so later headings don't drift) ---
    for planned in section_plan:
        if planned.canonical != "Business Activities":
            continue
        occurrences = resolve_section_plan_to_doc(doc, section_plan)
        occ_idx = find_resolved_occurrence_index(doc, occurrences, planned)
        if occ_idx is None:
            continue
        occurrence = occurrences[occ_idx]
        body = get_body_paragraphs_for_occurrence(doc, occurrences, occ_idx)
        fill_section_paragraphs(doc, body, [business_activities_map.get(occurrence.language, business_activities_map.get("english", ""))])

    # --- Fill Business Update middle sections (re-resolve after Business Activities may have shifted indices) ---
    occurrences = resolve_section_plan_to_doc(doc, section_plan)
    update_occurrences_by_language = middle_section_indices_by_language(occurrences)

    for language, occ_indices in update_occurrences_by_language.items():
        if not occ_indices:
            continue
        # Re-resolve fresh for each language group to get current indices
        occurrences = resolve_section_plan_to_doc(doc, section_plan)
        update_occ = middle_section_indices_by_language(occurrences).get(language, [])
        if not update_occ:
            continue
        heading_texts = [doc.paragraphs[occurrences[occ_idx].heading_index].text for occ_idx in update_occ]
        chunks = allocate_middle_section_bullets(heading_texts, business_update_map.get(language, []))
        for occ_idx, texts in zip(update_occ, chunks):
            # Re-resolve before each fill to handle shifting from prior fills
            occurrences = resolve_section_plan_to_doc(doc, section_plan)
            body = get_body_paragraphs_for_occurrence(doc, occurrences, occ_idx)
            fill_section_paragraphs(doc, body, texts, force_bullets=True)

    # Risk & Exit: DO NOT rewrite. The template (previous quarter) already has
    # the correct content. Rewriting via XML extraction and fill_section_paragraphs
    # is lossy (merges paragraphs, creates duplicates, loses formatting).

    occurrences = resolve_section_plan_to_doc(doc, section_plan)
    financial_occurrences = [occ for occ in occurrences if occ.canonical == "Financial Update"]
    for table_idx, table in enumerate(doc.tables):
        if table_idx >= len(financial_updates):
            break
        center_table(table)
        financial_update = financial_updates[table_idx]
        ensure_table_dimensions(table, len(financial_update["rows"]) + 1, len(financial_update["columns"]) + 1)
        formatting: dict[tuple[int, int], tuple[str | None, str | None, float | None, bool | None]] = {}
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                font_name = None
                east_asia_font = None
                size_pt = None
                bold = None
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font_name = run.font.name or font_name
                        rpr = run._element.find(qn("w:rPr"))
                        if rpr is not None:
                            rfonts = rpr.find(qn("w:rFonts"))
                            if rfonts is not None:
                                east_asia_font = rfonts.get(qn("w:eastAsia")) or east_asia_font
                        if run.font.size is not None:
                            size_pt = float(run.font.size.pt)
                        if run.font.bold is not None:
                            bold = run.font.bold
                        break
                    if font_name or size_pt is not None or bold is not None:
                        break
                formatting[(r_idx, c_idx)] = (font_name, east_asia_font, size_pt, bold)

        unit_label = financial_update["unit"]
        top_left_original = norm_space(table.cell(0, 0).text)
        inside_unit = is_unit_text(top_left_original)
        outside_unit_paragraph = None
        if table_idx < len(financial_occurrences):
            occurrence = financial_occurrences[table_idx]
            occ_index = next((i for i, occ in enumerate(occurrences) if occ.heading_index == occurrence.heading_index), None)
            if occ_index is not None:
                for paragraph in get_body_paragraphs_for_occurrence(doc, occurrences, occ_index):
                    if paragraph.text.strip() and is_unit_text(paragraph.text):
                        outside_unit_paragraph = paragraph
                        break

        if outside_unit_paragraph is not None:
            # Only rewrite if the text actually changed (rewriting can lose font formatting)
            if norm_space(outside_unit_paragraph.text) != norm_space(unit_label):
                unit_run_format = capture_run_format(outside_unit_paragraph)
                set_paragraph_text_preserve(outside_unit_paragraph, unit_label)
                apply_paragraph_run_format(outside_unit_paragraph, *unit_run_format)
            set_cell_text_preserve(table.cell(0, 0), "")
        elif inside_unit:
            set_cell_text_preserve(table.cell(0, 0), unit_label)

        if table_idx < len(financial_occurrences):
            occurrence = financial_occurrences[table_idx]
            occ_index = next((i for i, occ in enumerate(occurrences) if occ.heading_index == occurrence.heading_index), None)
            if occ_index is not None:
                note_text = financial_note_map.get(occurrence.language)
                if note_text:
                    for paragraph in get_body_paragraphs_for_occurrence(doc, occurrences, occ_index):
                        existing = norm_space(paragraph.text)
                        if not existing or is_unit_text(existing):
                            continue
                        if (
                            "银行账户现金余额" in existing
                            or "cash balance" in existing.lower()
                            or "bank account cash balance" in existing.lower()
                        ):
                            para_format = capture_run_format(paragraph)
                            set_paragraph_text_preserve(paragraph, note_text)
                            apply_paragraph_run_format(paragraph, para_format[0], para_format[1], para_format[2], False)
                            break

        for idx, header in enumerate(financial_update["columns"], start=1):
            if idx < len(table.rows[0].cells):
                cell = table.cell(0, idx)
                set_cell_text_preserve(cell, table_display_text(header, occurrence.language if table_idx < len(financial_occurrences) else "english"))
                set_cell_no_wrap(cell)
        for row_idx, row_data in enumerate(financial_update["rows"], start=1):
            if row_idx >= len(table.rows):
                break
            cell_language = occurrence.language if table_idx < len(financial_occurrences) else "english"
            label_cell = table.cell(row_idx, 0)
            set_cell_text_preserve(label_cell, table_display_text(row_data["label"], cell_language))
            set_cell_no_wrap(label_cell)
            for col_idx, header in enumerate(financial_update["columns"], start=1):
                if col_idx >= len(table.rows[row_idx].cells):
                    break
                value_cell = table.cell(row_idx, col_idx)
                set_cell_text_preserve(value_cell, table_display_text(row_data["values"].get(header) or "", cell_language))
                if header in {"QoQ", "YoY", "2025 FY", "2026E"} or header.endswith("FY") or header.endswith("YTD") or header.endswith("E"):
                    set_cell_no_wrap(value_cell)

        # Build numeric-cell fallback format per row so added QoQ/YoY/FY/Estimate cells
        # inherit the same visual style as existing numeric cells, not label-cell styling.
        row_defaults: dict[int, tuple] = {}
        for r_idx in range(len(table.rows)):
            preferred_indices = list(range(1, len(table.rows[r_idx].cells))) + [0]
            for c_idx in preferred_indices:
                fmt = formatting.get((r_idx, c_idx), (None, None, None, None))
                if any(v is not None for v in fmt):
                    row_defaults[r_idx] = fmt
                    break

        for r_idx, row in enumerate(table.rows):
            fallback = row_defaults.get(r_idx, (None, None, None, None))
            for c_idx, cell in enumerate(row.cells):
                fmt = formatting.get((r_idx, c_idx), (None, None, None, None))
                # If this cell had no formatting, use the row default
                if not any(v is not None for v in fmt):
                    fmt = fallback
                font_name, east_asia_font, size_pt, bold = fmt
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        apply_table_run_format(run, font_name, east_asia_font, size_pt, bold)
        fit_financial_table_columns(doc, table, occurrence.language if table_idx < len(financial_occurrences) else "english")

    additional_table_entries = section_table_entries(section_plan, {"Business Update"})
    for entry in additional_table_entries:
        key = (entry["occurrence"].language, entry["occurrence"].heading_text, entry["local_table_index"])
        table_update = middle_table_updates.get(key)
        if table_update is None:
            continue
        table_idx = entry["global_table_index"]
        if table_idx >= len(doc.tables):
            continue
        table = doc.tables[table_idx]
        center_table(table)
        formatting: dict[tuple[int, int], tuple[str | None, str | None, float | None, bool | None]] = {}
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                font_name = None
                east_asia_font = None
                size_pt = None
                bold = None
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font_name = run.font.name or font_name
                        rpr = run._element.find(qn("w:rPr"))
                        if rpr is not None:
                            rfonts = rpr.find(qn("w:rFonts"))
                            if rfonts is not None:
                                east_asia_font = rfonts.get(qn("w:eastAsia")) or east_asia_font
                        if run.font.size is not None:
                            size_pt = float(run.font.size.pt)
                        if run.font.bold is not None:
                            bold = run.font.bold
                        break
                    if font_name or size_pt is not None or bold is not None:
                        break
                formatting[(r_idx, c_idx)] = (font_name, east_asia_font, size_pt, bold)

        occurrence = entry["occurrence"]
        unit_label = table_update.get("unit") or ""
        top_left_original = norm_space(table.cell(0, 0).text)
        inside_unit = is_unit_text(top_left_original)
        if inside_unit:
            set_cell_text_preserve(table.cell(0, 0), unit_label if unit_label else "")

        for idx, header in enumerate(table_update["columns"], start=1):
            if idx < len(table.rows[0].cells):
                cell = table.cell(0, idx)
                set_cell_text_preserve(cell, table_display_text(header, occurrence.language))
        for row_idx, row_data in enumerate(table_update["rows"], start=1):
            if row_idx >= len(table.rows):
                break
            label_cell = table.cell(row_idx, 0)
            set_cell_text_preserve(label_cell, table_display_text(row_data["label"], occurrence.language))
            for col_idx, header in enumerate(table_update["columns"], start=1):
                if col_idx >= len(table.rows[row_idx].cells):
                    break
                value_cell = table.cell(row_idx, col_idx)
                set_cell_text_preserve(value_cell, table_display_text(row_data["values"].get(header) or "", occurrence.language))

        row_defaults: dict[int, tuple] = {}
        for r_idx in range(len(table.rows)):
            for c_idx in range(len(table.rows[r_idx].cells)):
                fmt = formatting.get((r_idx, c_idx), (None, None, None, None))
                if any(v is not None for v in fmt):
                    row_defaults[r_idx] = fmt
                    break

        for r_idx, row in enumerate(table.rows):
            fallback = row_defaults.get(r_idx, (None, None, None, None))
            for c_idx, cell in enumerate(row.cells):
                fmt = formatting.get((r_idx, c_idx), (None, None, None, None))
                if not any(v is not None for v in fmt):
                    fmt = fallback
                font_name, east_asia_font, size_pt, bold = fmt
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        apply_table_run_format(run, font_name, east_asia_font, size_pt, bold)
        fit_financial_table_columns(doc, table, occurrence.language)

    def safe_occurrences() -> list[DocSectionOccurrence]:
        resolved = [
            occ
            for occ in resolve_section_plan_to_doc(doc, section_plan)
            if occ.heading_index < len(doc.paragraphs)
        ]
        if resolved:
            return resolved
        return [
            occ
            for occ in detect_doc_section_occurrences(doc)
            if occ.heading_index < len(doc.paragraphs)
        ]

    occurrences = safe_occurrences()
    remove_blank_paragraphs_inside_sections(doc, occurrences)
    occurrences = safe_occurrences()
    remove_blank_paragraphs_inside_sections(doc, occurrences)
    normalize_bullet_sections(doc, occurrences)

    # Ensure all section headings are bold (after all paragraph shifts are done)
    for occ in safe_occurrences():
        heading_paragraph = doc.paragraphs[occ.heading_index]
        for run in heading_paragraph.runs:
            run.font.bold = True

    # Remove ALL blank paragraphs first, then re-insert exactly one before each heading
    to_remove = []
    for idx, paragraph in enumerate(doc.paragraphs):
        if is_plain_blank_paragraph(paragraph):
            to_remove.append(paragraph)
    for paragraph in reversed(to_remove):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    # Now insert exactly one blank paragraph before each section heading
    normalize_exact_single_blank_between_sections(doc, safe_occurrences())
    enforce_bilingual_page_break(doc, safe_occurrences())

    doc.save(str(output_path))


def extract_company_name(title: str, business_activities: str) -> str:
    match = re.search(r',\s*([^,]+?) \(the "Company"\)', business_activities)
    if match:
        return norm_space(match.group(1))
    cleaned = re.sub(r"\b20\d{2}\b", "", title)
    cleaned = re.sub(r"\bQ[1-4]\b", "", cleaned, flags=re.I)
    cleaned = re.sub(r"[_-]+", " ", cleaned)
    cleaned = norm_space(cleaned)
    if cleaned:
        return re.sub(r"(?<=[a-z])(?=[A-Z])", " ", cleaned)
    return "The Company"


def normalize_name_tokens(text: str) -> list[str]:
    text = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", text)  # split CamelCase generically
    text = re.sub(r"20\d{2}Q[1-4]", " ", text, flags=re.I)
    text = re.sub(r"20\d{2}", " ", text)
    text = re.sub(r"q[1-4]", " ", text, flags=re.I)
    text = re.sub(r"data request|datarequest|previous quarter template|template|samples|review|portfolio|quarter", " ", text, flags=re.I)
    text = re.sub(r"[^A-Za-z\u4e00-\u9fff]+", " ", text)
    tokens = [t.lower() for t in text.split() if len(t) > 1]
    stop = {"docx", "ventech"}
    return [t for t in tokens if t not in stop]


def jaccard_score(a: set[str], b: set[str]) -> float:
    if not a or not b:
        return 0.0
    return len(a & b) / len(a | b)


def infer_current_company_name(path: Path) -> str:
    blocks = load_blocks(path)
    parsed = parse_current_blocks(blocks)
    title = parsed.get("title", "")
    if title:
        cleaned = re.sub(r"\b20\d{2}\b", "", title)
        cleaned = re.sub(r"\bQ[1-4]\b", "", cleaned, flags=re.I)
        cleaned = re.sub(r"第[一二三四1234]季度.*", "", cleaned)
        cleaned = re.sub(r"\(.*?view.*?\)", "", cleaned, flags=re.I)
        cleaned = re.sub(r"业务概览.*", "", cleaned)
        cleaned = norm_space(cleaned)
        if cleaned:
            return re.sub(r"(?<=[a-z])(?=[A-Z])", " ", cleaned)
    return path.stem


def infer_previous_company_name(path: Path) -> str:
    blocks = load_docx_blocks(path)
    sections = extract_previous_sections(blocks)
    business = " ".join(sections.get("Business Activities", {}).get("paragraphs", []))
    if business:
        name = extract_company_name("", business)
        if name != "The Company":
            return name
    cleaned = re.sub(r"^\d+\s+", "", path.stem)
    cleaned = re.sub(r"\b20\d{2}Q[1-4]\b", "", cleaned, flags=re.I)
    cleaned = norm_space(cleaned)
    return cleaned


def extract_identity_text(path: Path, is_previous: bool) -> str:
    blocks = load_blocks(path) if not is_previous else load_docx_blocks(path)
    if is_previous:
        sections = extract_previous_sections(blocks)
        return " ".join(sections.get("Business Activities", {}).get("paragraphs", []))
    parsed = parse_current_blocks(blocks)
    return " ".join(parsed.get("business_update_paragraphs", [])[:3] + [parsed.get("title", "")])


def match_current_to_previous(current_files: list[Path], previous_files: list[Path]) -> list[dict[str, Any]]:
    previous_meta = []
    for prev in previous_files:
        prev_name = infer_previous_company_name(prev)
        prev_tokens = set(normalize_name_tokens(prev.stem + " " + prev_name))
        prev_identity = set(normalize_name_tokens(extract_identity_text(prev, is_previous=True)))
        prev_quarter = parse_quarter_label(prev.stem) or parse_quarter_label(docx_to_text(prev)[:50])
        previous_meta.append(
            {
                "path": prev,
                "company_name": prev_name,
                "name_tokens": prev_tokens,
                "identity_tokens": prev_identity,
                "quarter": prev_quarter,
            }
        )

    reports = []
    for current in current_files:
        current_name = infer_current_company_name(current)
        current_tokens = set(normalize_name_tokens(current.stem + " " + current_name))
        current_identity = set(normalize_name_tokens(extract_identity_text(current, is_previous=False)))
        current_quarter = parse_quarter_label(current.stem) or parse_quarter_label(docx_to_text(current)[:50])

        scored = []
        for prev in previous_meta:
            score = 0.0
            reasons = []

            name_score = jaccard_score(current_tokens, prev["name_tokens"])
            score += name_score * 0.55
            if name_score > 0:
                reasons.append(f"filename/company tokens overlap ({name_score:.2f})")

            current_name_tokens = set(normalize_name_tokens(current_name))
            prev_name_tokens = set(normalize_name_tokens(prev["company_name"]))
            direct_name_overlap = current_name_tokens & prev_name_tokens
            if direct_name_overlap:
                score += 0.20
                reasons.append(f"direct company-name overlap ({', '.join(sorted(direct_name_overlap))})")

            identity_score = jaccard_score(current_identity, prev["identity_tokens"])
            score += identity_score * 0.30
            if identity_score > 0:
                reasons.append(f"content identity overlap ({identity_score:.2f})")

            if current_quarter and prev["quarter"]:
                if current_quarter.year == prev["quarter"].year and current_quarter.quarter == prev["quarter"].quarter + 1:
                    score += 0.15
                    reasons.append("quarter progression aligns (current is next quarter)")
                elif current_quarter.year == prev["quarter"].year and current_quarter.quarter != prev["quarter"].quarter:
                    score += 0.05
                    reasons.append("same year quarter reference present")

            scored.append((score, prev, reasons))

        scored.sort(key=lambda x: x[0], reverse=True)
        best_score, best_prev, best_reasons = scored[0]
        second_score = scored[1][0] if len(scored) > 1 else 0.0
        gap = best_score - second_score

        if best_score >= 0.65 and gap >= 0.15:
            confidence = "high"
        elif best_score >= 0.45 and gap >= 0.08:
            confidence = "medium"
        else:
            confidence = "ambiguous / review needed"

        reports.append(
            {
                "current_file": str(current.relative_to(ROOT)),
                "matched_previous_file": str(best_prev["path"].relative_to(ROOT)),
                "confidence": confidence,
                "score": round(best_score, 3),
                "explanation": "; ".join(best_reasons) if best_reasons else "weak textual similarity only",
            }
        )
    return reports


def run_matching_report() -> None:
    current_files = list_current_inputs(DATA_REQUEST_DIR)
    previous_files = list_real_docx(PREVIOUS_REVIEW_DIR)
    reports = match_current_to_previous(current_files, previous_files)
    for item in reports:
        pass


def generate_review_for_pair(current_path: Path, previous_path: Path) -> dict[str, Any]:
    OUTPUT_DIR.mkdir(exist_ok=True)
    DEBUG_OUTPUT_DIR.mkdir(exist_ok=True)

    spec_text = docx_to_text(UNIVERSAL_SPEC)
    prev_blocks = load_docx_blocks(previous_path)
    current_blocks = load_blocks(current_path)
    template_occurrences = detect_doc_section_occurrences(Document(str(previous_path)))
    template_languages = {occ.language for occ in template_occurrences}
    previous_sections = extract_previous_sections(prev_blocks)
    previous_occurrences = extract_previous_section_occurrences(prev_blocks)
    previous_financial_occurrences = [occ for occ in previous_occurrences if occ.canonical == "Financial Update" and occ.tables]
    previous_tables = []
    for occurrence in previous_financial_occurrences:
        unit_label, _ = detect_financial_section_unit(occurrence)
        previous_tables.append(parse_previous_financial_table_with_unit(occurrence.tables[0], unit_label))
    current_parsed = parse_current_blocks(current_blocks)
    current_finance = normalize_finance_source(current_parsed["finance_table"])
    current_unit_spec = detect_current_unit_spec(current_blocks)
    current_exchange_rates = extract_exchange_rates(current_blocks)
    target_quarter = determine_target_quarter(current_path, current_parsed["title"], current_finance)
    previous_review_quarter = parse_quarter_label(previous_path.stem)
    financial_note_map = {
        language: note
        for language, note in {
            "chinese": extract_cash_balance_line(current_parsed.get("balance_table"), "chinese"),
            "english": extract_cash_balance_line(current_parsed.get("balance_table"), "english"),
        }.items()
        if note
    }

    review_flags: list[dict[str, Any]] = []
    source_language, language_evidence = detect_language(docx_to_text(current_path))
    company_name = infer_previous_company_name(previous_path)
    final_company_name = normalize_company_filename(company_name)
    final_quarter_label = quarter_filename_label(target_quarter)
    final_docx_path = OUTPUT_DIR / f"{final_company_name}_{final_quarter_label}.docx"
    json_output = DEBUG_OUTPUT_DIR / f"{final_company_name}_{final_quarter_label}.json"
    markdown_output = DEBUG_OUTPUT_DIR / f"{final_company_name}_{final_quarter_label}.md"

    strongest_financial_idx = 0
    if previous_tables:
        strongest_financial_idx = max(
            range(len(previous_tables)),
            key=lambda idx: financial_table_backbone_score(previous_tables[idx], previous_review_quarter),
        )

    source_financial_update = None
    if previous_tables:
        source_prev_financial = previous_tables[strongest_financial_idx]
        source_financial_update = build_financial_update(
            source_prev_financial,
            current_finance,
            current_parsed.get("operation_table"),
            review_flags,
            target_quarter,
            current_unit_spec,
            detect_unit_spec(source_prev_financial.get("unit", "")),
            previous_review_quarter,
            current_exchange_rates,
        )

    financial_updates = []
    for idx, prev_financial in enumerate(previous_tables):
        if source_financial_update is not None and len(prev_financial.get("rows", [])) == len(source_financial_update.get("rows", [])):
            financial_updates.append(localize_financial_update(source_financial_update, prev_financial))
        else:
            financial_updates.append(
                build_financial_update(
                    prev_financial,
                    current_finance,
                    current_parsed.get("operation_table"),
                    review_flags,
                    target_quarter,
                    current_unit_spec,
                    detect_unit_spec(prev_financial.get("unit", "")),
                    previous_review_quarter,
                    current_exchange_rates,
                )
            )

    middle_table_updates: dict[tuple[str, str, int], dict[str, Any]] = {}
    current_extra_tables: list[list[list[str]]] = list(current_parsed.get("extra_tables") or [])
    middle_table_groups = group_middle_table_occurrences(previous_occurrences)
    for group_idx, group in enumerate(middle_table_groups):
        if group_idx >= len(current_extra_tables):
            continue
        current_source_table = normalize_finance_source(current_extra_tables[group_idx])
        parsed_group_tables = []
        for occurrence in group:
            unit_label, _placement = detect_financial_section_unit(occurrence)
            parsed_group_tables.append((occurrence, parse_previous_financial_table_with_unit(occurrence.tables[0], unit_label)))
        strongest_idx = max(
            range(len(parsed_group_tables)),
            key=lambda idx: financial_table_backbone_score(parsed_group_tables[idx][1], previous_review_quarter),
        )
        source_occurrence, source_prev_table = parsed_group_tables[strongest_idx]
        source_middle_update = build_section_table_update(
            source_prev_table,
            current_source_table,
            review_flags,
            target_quarter,
            current_unit_spec,
            detect_unit_spec(source_prev_table.get("unit", "")),
            previous_review_quarter,
            current_exchange_rates,
        )
        for occurrence, prev_table in parsed_group_tables:
            if len(prev_table.get("rows", [])) == len(source_middle_update.get("rows", [])):
                middle_table_updates[(occurrence.language, occurrence.heading_text, 0)] = localize_financial_update(source_middle_update, prev_table)
            else:
                middle_table_updates[(occurrence.language, occurrence.heading_text, 0)] = build_section_table_update(
                    prev_table,
                    current_source_table,
                    review_flags,
                    target_quarter,
                    current_unit_spec,
                    detect_unit_spec(prev_table.get("unit", "")),
                    previous_review_quarter,
                    current_exchange_rates,
                )

    business_activity_map: dict[str, str] = {}
    risk_exit_map: dict[str, list[str]] = {}
    previous_business_update_map: dict[str, list[str]] = {}
    for occurrence in previous_occurrences:
        if occurrence.canonical == "Business Activities" and occurrence.paragraphs:
            business_activity_map.setdefault(occurrence.language, " ".join(occurrence.paragraphs).strip())
        elif occurrence.canonical == "Risk & Exit" and occurrence.paragraphs:
            if occurrence.language not in risk_exit_map:
                seen = set()
                deduped = []
                for p in occurrence.paragraphs:
                    key = norm_space(p).lower()
                    if key not in seen:
                        seen.add(key)
                        deduped.append(p)
                risk_exit_map[occurrence.language] = deduped
        elif occurrence.canonical == "Business Update" and occurrence.paragraphs:
            previous_business_update_map.setdefault(occurrence.language, [])
            previous_business_update_map[occurrence.language].extend(occurrence.paragraphs)

    business_update_translation_memory = build_business_update_translation_memory(previous_business_update_map)

    business_update_map: dict[str, list[str]] = {}
    current_clean_english_updates = clean_update_paragraphs([p for p in current_parsed["business_update_paragraphs"] if text_language(p) == "english"])
    current_clean_chinese_updates = clean_update_paragraphs([p for p in current_parsed["business_update_paragraphs"] if text_language(p) == "chinese"])
    if "english" in template_languages and business_activity_map.get("english"):
        if current_clean_english_updates:
            business_update_map["english"] = build_business_update_bullets(current_clean_english_updates, company_name)
        elif current_clean_chinese_updates:
            business_update_map["english"] = build_english_business_update_from_chinese(current_clean_chinese_updates, company_name, translation_memory=business_update_translation_memory)
    if "chinese" in template_languages and business_activity_map.get("chinese"):
        business_update_map["chinese"] = build_chinese_business_update_bullets(current_parsed["business_update_paragraphs"])

    if not business_update_map.get("english") and previous_business_update_map.get("english"):
        business_update_map["english"] = previous_business_update_map["english"]
    if not business_update_map.get("chinese") and previous_business_update_map.get("chinese"):
        business_update_map["chinese"] = previous_business_update_map["chinese"]

    if not previous_review_has_logo(previous_path):
        review_flags.append(
            {
                "id": "logo-missing",
                "section": "Business Activities",
                "severity": "warning",
                "message": "No reusable company logo was found in the previous review.",
                "source": "previous_review",
            }
        )

    for table_update in financial_updates:
        for row in table_update["rows"]:
            for idx, message in enumerate(row["flags"], start=1):
                review_flags.append(
                    {
                        "id": f"{row['label'].lower().replace('#', 'num').replace(' ', '-')}-flag-{idx}",
                        "section": "Financial Update",
                        "severity": "warning" if "missing" in message.lower() or "could not" in message.lower() else "info",
                        "message": f"{row['label']}: {message}",
                        "source": "current_data_request" if "current data request" in message.lower() else "previous_review",
                    }
                )

    language_mode = template_language_mode(template_occurrences)
    primary_financial_update = financial_updates[-1] if financial_updates else {}
    primary_business_activities = business_activity_map.get("english") or next(iter(business_activity_map.values()), "")
    primary_risk = risk_exit_map.get("english") or next(iter(risk_exit_map.values()), [])
    primary_business_update = business_update_map.get("english") or next(iter(business_update_map.values()), [])

    output_payload = {
        "prototype_version": "v2",
        "test_case": {
            "company": company_name,
            "target_quarter": target_quarter.display(),
            "files": {
                "universal_spec": {"path": str(UNIVERSAL_SPEC)},
                "current_data_request": {"path": str(current_path)},
                "previous_review": {"path": str(previous_path)},
            },
        },
        "language_detection": {
            "source_language": source_language,
            "template_language_mode": language_mode,
            "evidence": language_evidence,
        },
        "extracted_inputs": {
            "universal_spec": {"found": True, "raw_text": spec_text},
            "previous_review": {
                "business_activities": {"found": True, "raw_text": primary_business_activities},
                "financial_update": {"found": True, "parsed_table": primary_financial_update},
                "business_update": {"found": True, "raw_bullets": previous_business_update_map.get("english") or previous_business_update_map.get("chinese", [])},
                "risk_exit": {"found": True, "raw_bullets": primary_risk},
            },
            "current_data_request": {
                "financial_source": {"found": True, "raw_table": current_parsed["finance_table"], "parsed_metrics": current_finance},
                "operation_source": {"found": current_parsed.get("operation_table") is not None, "raw_table": current_parsed.get("operation_table")},
                "business_update_source": {"found": True, "source_paragraphs": current_parsed["business_update_paragraphs"]},
            },
        },
        "proposed_outputs": {
            "business_activities": {"status": "carried_forward", "text": primary_business_activities},
            "financial_update": primary_financial_update,
            "business_update": {"status": "drafted_from_current_data", "bullets": primary_business_update, "source_trace": current_parsed["business_update_paragraphs"]},
            "risk_exit": {"status": "carried_forward", "bullets": primary_risk},
        },
        "review_flags": review_flags,
    }

    json_output.write_text(json.dumps(output_payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    markdown_output.write_text(
        build_markdown(primary_business_activities, primary_financial_update, primary_business_update, primary_risk),
        encoding="utf-8",
    )
    write_docx(
        final_docx_path,
        previous_path,
        previous_occurrences,
        business_activity_map,
        financial_updates,
        middle_table_updates,
        business_update_map,
        risk_exit_map,
        financial_note_map,
    )

    return {
        "current_file": str(current_path.relative_to(ROOT)),
        "matched_previous_file": str(previous_path.relative_to(ROOT)),
        "output_file": str(final_docx_path.relative_to(ROOT)),
        "language_mode": language_mode,
        "company_naming_preserved": "yes",
    }


def debug_generation_for_pair(current_path: Path, previous_path: Path) -> None:
    output_path = OUTPUT_DIR / f"{normalize_company_filename(infer_previous_company_name(previous_path))}_{quarter_filename_label(determine_target_quarter(current_path, parse_current_blocks(load_blocks(current_path))['title'], normalize_finance_source(parse_current_blocks(load_blocks(current_path))['finance_table'])))}.docx"
    existed_before = output_path.exists()
    before_mtime = datetime.fromtimestamp(output_path.stat().st_mtime).astimezone().isoformat(timespec="seconds") if existed_before else None

    prev_blocks = load_docx_blocks(previous_path)
    current_blocks = load_blocks(current_path)
    previous_occurrences = extract_previous_section_occurrences(prev_blocks)
    previous_financial_occurrences = [occ for occ in previous_occurrences if occ.canonical == "Financial Update" and occ.tables]
    current_parsed = parse_current_blocks(current_blocks)
    current_finance = normalize_finance_source(current_parsed["finance_table"])
    target_quarter, quarter_debug = determine_target_quarter_with_debug(current_path, current_parsed["title"], current_finance)
    current_unit_spec = detect_current_unit_spec(current_blocks)
    current_exchange_rates = extract_exchange_rates(current_blocks)
    previous_review_quarter = parse_quarter_label(previous_path.stem)


    for idx, occurrence in enumerate(previous_financial_occurrences):
        unit_label, placement = detect_financial_section_unit(occurrence)
        prev_financial = parse_previous_financial_table_with_unit(occurrence.tables[0], unit_label)
        output_unit_spec = detect_unit_spec(prev_financial.get("unit", ""))
        conversion_factor = None
        if current_unit_spec and output_unit_spec:
            conversion_factor = current_unit_spec.factor_to_base / output_unit_spec.factor_to_base
        historical_headers = [header for header in prev_financial["headers"] if parse_quarter_label(header)]
        normalized_historical = normalize_previous_quarter_headers(prev_financial["headers"], previous_review_quarter)
        flags: list[dict[str, Any]] = []
        financial_update = build_financial_update(
            prev_financial,
            current_finance,
            current_parsed.get("operation_table"),
            flags,
            target_quarter,
            current_unit_spec,
            output_unit_spec,
            previous_review_quarter,
            current_exchange_rates,
        )

    result = generate_review_for_pair(current_path, previous_path)
    regenerated_output = ROOT / result["output_file"]
    existed_after = regenerated_output.exists()
    after_mtime = datetime.fromtimestamp(regenerated_output.stat().st_mtime).astimezone().isoformat(timespec="seconds") if existed_after else None

def run_batch_generation() -> None:
    current_files = list_current_inputs(DATA_REQUEST_DIR)
    previous_files = list_real_docx(PREVIOUS_REVIEW_DIR)
    reports = match_current_to_previous(current_files, previous_files)
    clean_main_output_folder_for_batch()

    for item in reports:
        current_path = ROOT / item["current_file"]
        previous_path = ROOT / item["matched_previous_file"]
        result = generate_review_for_pair(current_path, previous_path)


def main() -> None:
    current_files = list_current_inputs(DATA_REQUEST_DIR)
    previous_files = list_real_docx(PREVIOUS_REVIEW_DIR)
    if not current_files:
        print(f"ERROR: No .docx or .pdf files found in {DATA_REQUEST_DIR}")
        sys.exit(1)
    if not previous_files:
        print(f"ERROR: No .docx files found in {PREVIOUS_REVIEW_DIR}")
        sys.exit(1)
    if not UNIVERSAL_SPEC.exists():
        print(f"WARNING: Universal Spec not found at {UNIVERSAL_SPEC}")
    reports = match_current_to_previous(current_files, previous_files)
    OUTPUT_DIR.mkdir(exist_ok=True)
    for item in reports:
        current_path = ROOT / item["current_file"]
        previous_path = ROOT / item["matched_previous_file"]
        print(f"Processing: {current_path.name} <-> {previous_path.name} (confidence: {item['confidence']})")
        result = generate_review_for_pair(current_path, previous_path)
        print(f"  -> {result['output_file']}")


if __name__ == "__main__":
    if "--match-report" in sys.argv:
        run_matching_report()
    elif "--generate-matched" in sys.argv:
        run_batch_generation()
    elif "--debug-company" in sys.argv:
        idx = sys.argv.index("--debug-company")
        keyword = sys.argv[idx + 1]
        reports = match_current_to_previous(list_current_inputs(DATA_REQUEST_DIR), list_real_docx(PREVIOUS_REVIEW_DIR))
        matched = None
        for item in reports:
            if keyword.lower() in item["current_file"].lower() or keyword.lower() in item["matched_previous_file"].lower():
                matched = item
                break
        if matched is None:
            raise SystemExit(f"No matched pair found for keyword: {keyword}")
        debug_generation_for_pair(ROOT / matched["current_file"], ROOT / matched["matched_previous_file"])
    else:
        main()
