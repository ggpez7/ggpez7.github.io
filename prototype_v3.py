#!/usr/bin/env python3
from __future__ import annotations

import copy
import contextlib
import io
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import unicodedata
import zipfile
from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is required. Install it with: pip install python-docx")
    sys.exit(1)
try:
    import pypdf
except ImportError:
    pypdf = None
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Pt, Emu, Twips
from docx.oxml.ns import qn


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
ROOT = Path(__file__).resolve().parent


UNIVERSAL_SPEC = ROOT / "Universal Spec .docx"
DATA_REQUEST_DIR = ROOT / "DataRequest Samples"
PREVIOUS_REVIEW_DIR = ROOT / "Previous Quarter Template"
OUTPUT_DIR = ROOT / "prototype_output_v3"
DEBUG_OUTPUT_DIR = OUTPUT_DIR / "_debug"
USE_OLLAMA_TRANSLATION = os.getenv("PR_USE_OLLAMA_TRANSLATION", "1").strip().lower() not in {"0", "false", "no", "off"}
OLLAMA_MODEL = os.getenv("PR_OLLAMA_MODEL", "").strip()
DEFAULT_OLLAMA_MODELS = [
    "qwen3.5:4b",
    "qwen2.5:7b",
    "qwen2.5:3b",
]
OLLAMA_TIMEOUT_SECONDS = int(os.getenv("PR_OLLAMA_TIMEOUT_SECONDS", "120"))

OLLAMA_TRANSLATION_STATUS: dict[str, Any] = {
    "checked": False,
    "available": False,
    "reason": "",
    "model": "",
    "warned": False,
}
USE_ARGOS_TRANSLATION = os.getenv("PR_USE_ARGOS_TRANSLATION", "1").strip().lower() not in {"0", "false", "no", "off"}
ARGOS_AUTO_INSTALL = os.getenv("PR_ARGOS_AUTO_INSTALL", "0").strip().lower() in {"1", "true", "yes", "on"}
ARGOS_TRANSLATION_STATUS: dict[str, Any] = {
    "checked": False,
    "available": False,
    "reason": "",
    "warned": False,
}
ARGOS_RESOURCES: dict[str, Any] = {
    "translation": None,
}
ARGOS_RUNTIME_ROOT = ROOT / ".runtime" / "argos"
PREFER_WORD_PDF_CONVERSION = os.getenv("PR_PREFER_WORD_PDF_CONVERSION", "1").strip().lower() not in {"0", "false", "no", "off"}
WORD_PDF_CONVERSION_TIMEOUT_SECONDS = int(os.getenv("PR_WORD_PDF_CONVERSION_TIMEOUT_SECONDS", "180"))
FORCE_BILINGUAL_PAGE_BREAK = os.getenv("PR_FORCE_BILINGUAL_PAGE_BREAK", "1").strip().lower() in {"1", "true", "yes", "on"}


SECTION_HEADINGS = [
    "Business Activities",
    "Financial Update",
    "Business Update",
    "Risk & Exit",
]

SECTION_ALIASES = {
    "Business Activities": [
        "business activities",
        "business activity",
        "business information",
        "company information",
        "business introduction",
        "business overview",
        "company intro",
        "company introduction",
        "company overview",
        "overview",
        "公司介绍",
        "公司简介",
        "公司概况",
        "公司介绍和概况",
        "业务概况",
        "业务概述",
    ],
    "Financial Update": [
        "financial update",
        "financial updates",
        "financial updates (in usd million)",
        "financial data",
        "financial status",
        "financial situation",
        "财务情况",
        "财务数据",
        "财务状况",
        "财务信息",
        "财务概况",
    ],
    "Business Update": [
        "business update",
        "business highlights",
        "business operations",
        "recent trends",
        "recent highlights",
        "business review",
        "operational status",
        "operations update",
        "operating status",
        "business progress",
        "运营情况",
        "业务进展",
        "业务更新",
        "业务亮点",
        "业务回顾",
    ],
    "Risk & Exit": [
        "risk & exit",
        "risk&exit",
        "risk and exit",
        "risks & exit",
        "risks&exit",
        "risks and exit",
        "风险和退出",
        "风险与退出",
        "风险及退出",
        "风险 & 退出目标",
        "风险&退出目标",
        "风险与退出目标",
    ],
}


def is_question_or_prompt(text: str) -> bool:
    """Detect question/prompt paragraphs generically instead of hardcoding."""
    stripped = norm_space(text)
    if not stripped:
        return False
    if re.search(r"[?？]$", stripped):
        return True
    prompt_patterns = [
        r"^please\b",
        r"^what is\b",
        r"^what are\b",
        r"^how does\b",
        r"^how do\b",
        r"^how much\b",
        r"^how many\b",
        r"^can you\b",
        r"^could you\b",
        r"^would you\b",
        r"^describe\b",
        r"^explain\b",
        r"^elaborate\b",
        r"^provide\b",
        r"^list\b.*\b(?:key|main|major|top)\b",
        r"^note\b.*\bany\b",
        r"^project\b.*\bbalance\b",
        r"^if there is a (?:significant|major)\b.*\bplease\b",
        r"^if there (?:is|are)\b.*\bplease\b",
        r"\bplease provide\b.*\bexplanation\b",
        r"\bplease\b.*\b(?:describe|explain|elaborate|provide|list|detail|share|clarify|specify|confirm)\b",
        r"^see notes to\b",
        r">\s*see notes to\b",
        r"^what'?s the\b",
        r"^when (?:is|will|did)\b",
        r"^where (?:is|are)\b",
        r"^why (?:is|are|did)\b",
        r"^is there\b",
        r"^are there\b",
        r"^do you\b",
        r"^have you\b",
        r"^does the\b",
        r"^did the\b",
        r"^will the\b",
        r"^has the\b",
        # Chinese question/prompt patterns (without question marks)
        r"^请",
        r"^【重要】请",
        r"如有较大的.*请解释",
        r"^请问",
        r"^是否",
        r"^能否",
        r"^有没有",
        r"^有无",
        r"请说明",
        r"请提供",
        r"请描述",
        r"请列出",
        r"请解释",
        r"请详细",
        r"请补充",
        r"请确认",
        r"^如.*请详述",
        r"^如.*请解释",
    ]
    # Normalize curly quotes to straight for matching
    lowered = stripped.lower().replace("\u2018", "'").replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"')
    if any(re.search(p, lowered) for p in prompt_patterns):
        return True
    # Detect data-request prompts: short topic + parenthetical list of sub-topics
    # e.g. "新一轮融资的计划和进展（预计时间线，融资额度，估值，潜在投资方）。"
    # Pattern: Chinese text ending with (list of Chinese comma-separated items)。
    if re.search(r"[（(][^）)]*[，,][^）)]*[，,][^）)]*[）)]", stripped):
        # Has a parenthetical with 3+ comma-separated items → likely a prompt
        # But only flag if there are no numbers (actual data would have numbers)
        if not re.search(r"\d", stripped.replace("Q1", "").replace("Q2", "").replace("Q3", "").replace("Q4", "")):
            return True
    # Detect topic labels that just name an area without providing data
    # e.g. "现阶段重组进展，卡睿业务的进展。" — just listing topics with abstract words
    if re.search(r"[\u4e00-\u9fff]", stripped) and len(stripped) < 40:
        if not re.search(r"\d", stripped):
            # Ends with abstract topic words (进展/情况/计划/方案) without concrete details
            if re.search(r"(?:的进展|的情况|的计划|的方案|的安排)[。.]?$", stripped):
                return True
    if re.search(r"[\u4e00-\u9fff]", stripped):
        if re.search(r"(?:如何|是什么|有哪些|时间线|结果如何|进展如何|最新进展是什么|情况如何)[。.]?$", stripped):
            return True
        if re.search(r"(?:足够支持多久|能支持多久|支持多久的运营)[。.]?$", stripped):
            return True
    return False


def is_standalone_label(text: str) -> bool:
    """Detect short standalone labels/headers that aren’t real content."""
    stripped = norm_space(text)
    if not stripped:
        return False
    if stripped.endswith((":", "：")) and len(stripped) < 80:
        return True
    if re.match(r"^\d+[.)\s]", stripped) and len(stripped) < 60:
        return True
    section_label_patterns = [
        r"^other\s+questions",
        r"^operational\s+data",
        r"^业务发展$",
    ]
    lowered = stripped.lower()
    return any(re.search(p, lowered) for p in section_label_patterns)


def norm_space(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


def is_unit_text(text: str) -> bool:
    normalized = norm_space(text).lower()
    if not normalized:
        return False
    return any(
        token in normalized
        for token in [
            "unit",
            "amount in",
            "in usd",
            "in rmb",
            "thousand usd",
            "usd million",
            "million usd",
            "k of us$",
            "k of usd",
            "rmb mn",
            "单位",
            "百万",
            "万元",
            "人民币",
        ]
    )


def detect_explicit_unit_spec(text: str) -> UnitSpec | None:
    normalized = norm_space(text)
    lowered = normalized.lower()
    explicit_markers = [
        "单位",
        "unit",
        "amount in",
        "financial data",
        "财务数据",
        "指标变动",
    ]
    short_standalone = len(normalized) <= 32 and not re.search(r"\d{2,}", normalized)
    if any(marker in lowered for marker in explicit_markers) or short_standalone:
        return detect_unit_spec(normalized)
    return None


def cell_text(tc: ET.Element) -> str:
    parts: list[str] = []
    for p in tc.findall("w:p", NS):
        text = "".join(t.text or "" for t in p.findall(".//w:t", NS))
        if text.strip():
            parts.append(norm_space(text))
    return " ".join(parts).strip()


def paragraph_text(p: ET.Element) -> str:
    return norm_space("".join(t.text or "" for t in p.findall(".//w:t", NS)))


def load_docx_blocks(path: Path) -> list[dict[str, Any]]:
    with zipfile.ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    body = root.find("w:body", NS)
    assert body is not None

    blocks: list[dict[str, Any]] = []
    for child in body:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            text = paragraph_text(child)
            if text:
                blocks.append({"type": "paragraph", "text": text})
        elif tag == "tbl":
            rows = []
            for tr in child.findall("w:tr", NS):
                rows.append([cell_text(tc) for tc in tr.findall("w:tc", NS)])
            blocks.append({"type": "table", "rows": rows})
    return blocks


PDF_TABLE_VALUE_RE = re.compile(r"(?:−|-)?\d+(?:,\d{3})*(?:\.\d+)?%?|N/?A|/")


def extract_pdf_lines(path: Path) -> list[str]:
    if pypdf is None:
        raise RuntimeError("PDF support requires pypdf to be installed.")
    reader = pypdf.PdfReader(str(path))
    lines: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        for raw in text.splitlines():
            cleaned = norm_space(raw)
            if cleaned:
                lines.append(cleaned)
    return lines


def is_numbered_section_line(text: str) -> bool:
    stripped = norm_space(text)
    if re.match(r"^20\d{2}\b", stripped):
        return False
    return bool(re.match(r"^(?:[1-9]\d?)(?:\s*[.)、．]\s*|\s+(?=[A-Za-z\u4e00-\u9fff]))", stripped))


def is_top_level_numbered_section_line(text: str) -> bool:
    stripped = norm_space(text)
    if re.match(r"^20\d{2}\b", stripped):
        return False
    return bool(re.match(r"^(?:[1-9]\d?)(?:\.\s*|\s+)[A-Za-z\u4e00-\u9fff]", stripped))


def is_decimal_section_heading_line(text: str) -> bool:
    return bool(re.match(r"^\d+(?:\.\d+)+\s+\S", norm_space(text)))


def is_numbered_list_item_line(text: str) -> bool:
    stripped = norm_space(text)
    return bool(re.match(r"^(?:[（(]?\d+[)）.、-]|[一二三四五六七八九十]+[、.])", stripped))


def should_merge_pdf_update_continuation(previous: str, current: str) -> bool:
    prev = norm_space(previous)
    curr = norm_space(current)
    if not prev or not curr:
        return False
    if is_top_level_numbered_section_line(curr) or is_decimal_section_heading_line(curr) or is_numbered_list_item_line(curr):
        return False
    if curr in {"无", "无变化", "暂无"}:
        return False
    if prev.endswith(("，", ",", "；", ";", "：", ":")):
        return True
    if not prev.endswith(("。", ".", "！", "!", "？", "?")) and len(curr) <= 28 and not re.search(
        r"(有限公司|科技有限公司|股份有限公司|有限责任公司|公司)", curr
    ):
        return True
    return False


def merge_pdf_update_lines(lines: list[str]) -> list[str]:
    merged: list[str] = []
    for raw in lines:
        text = norm_space(raw)
        if not text:
            continue
        if merged and should_merge_pdf_update_continuation(merged[-1], text):
            merged[-1] = norm_space(f"{merged[-1]}{text}")
        else:
            merged.append(text)
    return merged


def extract_pdf_header_tokens(lines: list[str]) -> list[str]:
    combined = " ".join(lines)
    patterns = [
        r"\d{6}",
        r"\d{4}/\d{1,2}",
        r"\d{1,2}\s*月",
        r"\d{4}Q[1-4]\s*(?:合计|总计)?",
        r"\d{4}\s*(?:合计|总计)",
        r"\d{4}\s*年\s*预计",
        r"Q[1-4]\s*20\d{2}",
        r"20\d{2}\s*YTD",
        r"20\d{2}E",
        r"QoQ",
        r"YoY",
    ]
    matches: list[tuple[int, int, str]] = []
    for pattern in patterns:
        for match in re.finditer(pattern, combined, flags=re.I):
            matches.append((match.start(), match.end(), norm_space(match.group(0))))
    matches.sort(key=lambda item: item[0])
    headers: list[str] = []
    seen: set[tuple[int, int]] = set()
    for start, end, token in matches:
        if (start, end) in seen:
            continue
        seen.add((start, end))
        if not headers or headers[-1] != token:
            headers.append(token)
    return headers


def extract_pdfplumber_tables(path: Path) -> list[list[list[str]]]:
    if pdfplumber is None:
        return []
    tables: list[list[list[str]]] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables() or []:
                cleaned_rows: list[list[str]] = []
                for row in table:
                    if row is None:
                        continue
                    cleaned_row: list[str] = []
                    for cell in row:
                        text = norm_space(str(cell or ""))
                        if re.search(r"\d,\s*\d", text):
                            text = re.sub(r",\s+", ",", text)
                        text = text.replace("\n", " ")
                        text = norm_space(text)
                        cleaned_row.append(text)
                    if any(cell for cell in cleaned_row):
                        cleaned_rows.append(cleaned_row)
                if cleaned_rows:
                    tables.append(cleaned_rows)
    return tables


def looks_like_pdf_finance_table(rows: list[list[str]]) -> bool:
    if len(rows) < 2:
        return False
    first_row = " ".join(rows[0]).lower()
    month_hits = len(re.findall(r"20\d{2}/\d{1,2}|20\d{4}|2025q[1-4]|qoq|yoy", first_row, flags=re.I))
    labels = " ".join(" ".join(row[:1]) for row in rows[1:6]).lower()
    metric_hits = sum(
        token in labels
        for token in ["revenue", "gross profit", "net profit", "net income", "收入", "毛利", "净利", "现金流入", "现金支出"]
    )
    return month_hits >= 2 and metric_hits >= 2


def looks_like_pdf_balance_table(rows: list[list[str]]) -> bool:
    flattened = " ".join(" ".join(row) for row in rows).lower()
    return any(
        token in flattened
        for token in ["cash balance", "shareholder's equity", "bank", "银行账户现金余额", "员工人数", "number of employees", "截至"]
    )


def looks_like_pdf_operation_table(rows: list[list[str]]) -> bool:
    if len(rows) < 2:
        return False
    first_row = " ".join(rows[0]).lower()
    return any(token in first_row for token in ["签约", "合作", "operation", "2026"]) and any(
        re.search(r"20\d{2}/\d{1,2}", cell) for cell in rows[0]
    )


def normalize_pdfplumber_finance_table(rows: list[list[str]]) -> list[list[str]] | None:
    if not rows:
        return None
    header = rows[0]
    normalized_header: list[str] = [""]
    seen_headers: set[str] = set()
    for idx, cell in enumerate(header[1:], start=1):
        token = norm_space(cell)
        if not token:
            continue
        if token.lower() == "subtotal":
            continue
        if token in seen_headers:
            continue
        seen_headers.add(token)
        normalized_header.append(token)
    normalized_rows: list[list[str]] = [normalized_header]
    target_len = len(normalized_header) - 1
    for row in rows[1:]:
        label = norm_space(row[0] if row else "")
        if not label:
            continue
        values = [norm_space(cell) for cell in row[1:] if norm_space(cell)]
        if not values:
            continue
        while len(values) < target_len:
            values.append("")
        normalized_rows.append([label] + values[:target_len])
    return normalized_rows if len(normalized_rows) > 1 else None


def normalize_pdfplumber_balance_table(rows: list[list[str]]) -> list[list[str]] | None:
    if not rows:
        return None
    normalized: list[list[str]] = [["", norm_space(" ".join(rows[0]))]]
    pending_label = ""
    pending_value_parts: list[str] = []
    for row in rows[1:]:
        cells = [norm_space(cell) for cell in row if norm_space(cell)]
        if not cells:
            continue
        if pending_label and cells[0] in {"人民币：70,614.33", "美元：179,899.07"}:
            pending_value_parts.extend(cells)
            continue
        if pending_label:
            normalized.append([pending_label, " ".join(pending_value_parts)])
            pending_label = ""
            pending_value_parts = []
        if len(cells) == 1:
            pending_label = cells[0]
            pending_value_parts = []
        else:
            pending_label = cells[0]
            pending_value_parts = cells[1:]
    if pending_label:
        normalized.append([pending_label, " ".join(pending_value_parts)])
    return normalized if len(normalized) > 1 else None


def parse_pdf_metric_headers(lines: list[str]) -> list[str]:
    headers: list[str] = []
    seen: set[str] = set()

    def add(token: str) -> None:
        normalized = norm_space(token)
        if not normalized or normalized in seen:
            return
        seen.add(normalized)
        headers.append(normalized)

    for line in lines:
        stripped = norm_space(line)
        if not stripped:
            continue
        stripped = stripped.replace("数据及指标", " ")
        stripped = stripped.replace("項目", " ")
        stripped = stripped.replace("项目", " ")
        for match in re.finditer(r"\d{6}|\d{4}/\d{1,2}|\d{4}Q[1-4]|Q[1-4]\s*20\d{2}|20\d{2}\s*YTD|20\d{2}E", stripped, flags=re.I):
            add(match.group(0))
        for token in ["本季数", "上季度", "上年同季", "本年累计", "本年累计数", "累计", "QOQ", "YOY"]:
            if token.lower() in stripped.lower():
                add(token)

    return headers


def collect_pdf_table_rows(lines: list[str]) -> list[tuple[str, list[str]]]:
    collected: list[tuple[str, list[str]]] = []
    pending_label = ""
    for line in lines:
        stripped = norm_space(line)
        if (
            not stripped
            or is_numbered_section_line(stripped)
            or stripped.startswith("截至")
            or stripped.lower().startswith("date as of")
        ):
            break
        match = PDF_TABLE_VALUE_RE.search(stripped)
        if match is None:
            pending_label = norm_space(f"{pending_label} {stripped}")
            continue
        label = norm_space(f"{pending_label} {stripped[: match.start()]}")
        pending_label = ""
        values = [norm_space(token) for token in PDF_TABLE_VALUE_RE.findall(stripped[match.start() :])]
        if not label or not values:
            continue
        collected.append((label, values))
    return collected


def parse_pdf_table_lines(lines: list[str]) -> list[list[str]] | None:
    if not lines:
        return None
    header_idx = next(
        (
            idx
            for idx, line in enumerate(lines)
            if "日期" in line
            or "数据及指标" in line
            or re.search(r"^date\b", line, flags=re.I)
            or ("data" in line.lower() and "indicator" in line.lower())
        ),
        None,
    )
    if header_idx is None:
        return None
    data_start = None
    for idx in range(header_idx + 1, len(lines)):
        line = lines[idx]
        if re.search(r"(?:−|-)?\d", line) and not line.startswith("截至"):
            data_start = idx
            break
    if data_start is None:
        return None
    header_lines = lines[header_idx:data_start]
    headers = parse_pdf_metric_headers(header_lines) or extract_pdf_header_tokens(header_lines)
    if not headers:
        return None
    rows: list[list[str]] = [[""] + headers]
    for label, values in collect_pdf_table_rows(lines[data_start:]):
        while len(values) < len(headers):
            values.append("")
        rows.append([label] + values[: len(headers)])
    return rows if len(rows) > 1 else None


def parse_pdf_balance_table(lines: list[str]) -> list[list[str]] | None:
    start = next(
        (
            idx
            for idx, line in enumerate(lines)
            if line.startswith("截至") or line.lower().startswith("date as of")
        ),
        None,
    )
    if start is None:
        return None
    rows = [["", lines[start]]]
    idx = start + 1
    while idx < len(lines):
        line = lines[idx]
        if is_numbered_section_line(line):
            break
        if "银行账户现金余额" in line:
            value_parts = []
            remainder = norm_space(line.replace("银行账户现金余额", "", 1))
            if remainder:
                value_parts.append(remainder)
            j = idx + 1
            while j < len(lines) and not is_numbered_section_line(lines[j]) and ("美元" in lines[j] or "人民币" in lines[j]):
                value_parts.append(lines[j])
                j += 1
            rows.append(["银行账户现金余额", " ".join(value_parts)])
            idx = j
            continue
        if line.lower().startswith("cash balance"):
            rows.append(["Cash balance", norm_space(line[len("cash balance") :]) or line])
            idx += 1
            continue
        if "Shareholder's equity" in line or "Shareholders' equity" in line:
            rows.append(["Shareholder's equity", norm_space(re.sub(r"^Shareholders?'s equity", "", line, flags=re.I)) or line])
            idx += 1
            continue
        if line.lower().startswith("number of employees"):
            value_parts = []
            remainder = norm_space(re.sub(r"^Number of employees", "", line, flags=re.I))
            if remainder:
                value_parts.append(remainder)
            j = idx + 1
            while j < len(lines) and not is_numbered_section_line(lines[j]) and any(
                token in lines[j].lower() for token in ["full time", "part time"]
            ):
                value_parts.append(lines[j])
                j += 1
            rows.append(["Number of employees", " ".join(value_parts)])
            idx = j
            continue
        if "总员工人数" in line:
            value_parts = []
            remainder = norm_space(line.replace("总员工人数", "", 1))
            if remainder:
                value_parts.append(remainder)
            j = idx + 1
            while j < len(lines) and not is_numbered_section_line(lines[j]) and ("兼职" in lines[j] or "全职" in lines[j]):
                value_parts.append(lines[j])
                j += 1
            rows.append(["总员工人数", " ".join(value_parts)])
            idx = j
            continue
        idx += 1
    return rows if len(rows) > 1 else None


def find_microsoft_word_app() -> Path | None:
    candidates = [
        Path("/Applications/Microsoft Word.app"),
        Path.home() / "Applications" / "Microsoft Word.app",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def converted_pdf_docx_path(path: Path) -> Path:
    converted_dir = path.parent / "_converted_pdf_docx"
    converted_dir.mkdir(parents=True, exist_ok=True)
    return converted_dir / f"_converted_{path.stem}_word.docx"


def convert_pdf_via_word(path: Path) -> Path | None:
    if not PREFER_WORD_PDF_CONVERSION or path.suffix.lower() != ".pdf":
        return None
    if shutil.which("osascript") is None:
        return None
    if find_microsoft_word_app() is None:
        return None

    output_path = converted_pdf_docx_path(path)
    try:
        if output_path.exists() and output_path.stat().st_mtime >= path.stat().st_mtime:
            return output_path
    except OSError:
        pass

    with contextlib.suppress(FileNotFoundError):
        output_path.unlink()

    script_lines = [
        f'set srcPath to POSIX file "{path}" as alias',
        f'set outPath to POSIX file "{output_path}"',
        'tell application "Microsoft Word"',
        'open srcPath confirm conversions false',
        'set docRef to active document',
        'save as docRef file name outPath file format format document default',
        'close docRef saving no',
        'end tell',
    ]
    command = ["osascript"]
    for line in script_lines:
        command.extend(["-e", line])

    try:
        subprocess.run(
            command,
            check=True,
            capture_output=True,
            text=True,
            timeout=WORD_PDF_CONVERSION_TIMEOUT_SECONDS,
        )
    except (subprocess.SubprocessError, OSError):
        return None

    return output_path if output_path.exists() else None


def score_finance_table_quality(rows: list[list[str]] | None) -> int:
    if not rows or not rows[0]:
        return 0
    headers = [norm_space(cell).lower() for cell in rows[0][1:]]
    labels = [norm_space(row[0]).lower() for row in rows[1:] if row and norm_space(row[0])]
    joined_headers = " ".join(headers)
    joined_labels = " ".join(labels[:12])
    score = 0

    if re.search(r"\d{4}/\d{1,2}|\d{6}|\d{4}q[1-4]|q[1-4]\s*20\d{2}", joined_headers, flags=re.I):
        score += 3
    if any(token in joined_headers for token in ["本季", "上季度", "上年同季", "qoq", "yoy", "fy", "total", "全年预期", "2026e"]):
        score += 3
    if any(
        token in joined_labels
        for token in [
            "收入",
            "revenue",
            "net profit",
            "net income",
            "净利",
            "gross profit",
            "毛利",
            "现金流入",
            "现金支出",
            "ebitda",
        ]
    ):
        score += 5

    # Penalize KPI/headcount/patent tables masquerading as finance tables.
    if any(
        token in (joined_headers + " " + joined_labels)
        for token in ["员工", "研发人员", "专利", "patent", "headcount", "employee", "customer count"]
    ) and not any(token in joined_labels for token in ["收入", "revenue", "净利", "net profit", "毛利", "gross profit"]):
        score -= 8

    return max(score, 0)


def score_balance_table_quality(rows: list[list[str]] | None) -> int:
    if not rows:
        return 0
    labels = " ".join(norm_space(row[0]).lower() for row in rows[1:] if row and norm_space(row[0]))
    score = 0
    if any(token in labels for token in ["账面资金余额", "cash balance", "shareholder", "equity", "总员工人数", "number of employees"]):
        score += 4
    return score


def score_operation_table_quality(rows: list[list[str]] | None) -> int:
    if not rows:
        return 0
    labels = " ".join(norm_space(row[0]).lower() for row in rows[1:] if row and norm_space(row[0]))
    score = 0
    if any(token in labels for token in ["产量", "销量", "门店", "outlet", "shipments", "delivery", "运营数据", "经营数据"]):
        score += 3
    return score


def score_extracted_blocks(blocks: list[dict[str, Any]]) -> int:
    if not blocks:
        return -1
    try:
        parsed = parse_current_blocks(blocks)
    except Exception:
        return -1
    score = 0
    score += score_finance_table_quality(parsed.get("finance_table")) * 5
    score += score_balance_table_quality(parsed.get("balance_table")) * 3
    score += score_operation_table_quality(parsed.get("operation_table")) * 2
    score += min(len(parsed.get("business_update_raw_paragraphs", [])), 12)
    score += min(sum(1 for block in blocks if block.get("type") == "table"), 5) * 2
    return score


def load_pdf_blocks(path: Path) -> list[dict[str, Any]]:
    lines = extract_pdf_lines(path)
    if not lines:
        return []
    plumber_tables = extract_pdfplumber_tables(path)

    blocks: list[dict[str, Any]] = [{"type": "paragraph", "text": lines[0]}]

    section_positions = [idx for idx, line in enumerate(lines) if idx > 0 and is_numbered_section_line(line)]
    section_positions.append(len(lines))
    top_level_positions = [idx for idx, line in enumerate(lines) if idx > 0 and is_top_level_numbered_section_line(line)]
    top_level_positions.append(len(lines))

    finance_section = None
    operation_section = None
    update_start = None
    update_end = len(lines)
    for idx in range(len(section_positions) - 1):
        start = section_positions[idx]
        end = section_positions[idx + 1]
        heading = lines[start]
        if finance_section is None and ("财务数据" in heading or "finance" in heading.lower()):
            finance_section = lines[start:end]
        elif operation_section is None and ("运营数据" in heading or "经营数据" in heading):
            operation_section = lines[start:end]
    for idx in range(len(section_positions) - 1):
        start = section_positions[idx]
        end = section_positions[idx + 1]
        heading = lines[start]
        if update_start is None and ("业务发展" in heading or "业务进展" in heading or "business" in heading.lower()):
            update_start = start
            update_end = end
            break

    if finance_section:
        finance_heading = finance_section[0]
        explicit_spec = detect_explicit_unit_spec(finance_heading)
        if explicit_spec is not None:
            blocks.append({"type": "paragraph", "text": finance_heading})
        finance_table = parse_pdf_table_lines(finance_section)
        if not finance_table:
            finance_table = next(
                (
                    normalize_pdfplumber_finance_table(table)
                    for table in plumber_tables
                    if looks_like_pdf_finance_table(table)
                ),
                None,
            )
        if finance_table:
            blocks.append({"type": "table", "rows": finance_table})
        balance_table = parse_pdf_balance_table(finance_section)
        if not balance_table:
            balance_table = next(
                (
                    normalize_pdfplumber_balance_table(table)
                    for table in plumber_tables
                    if looks_like_pdf_balance_table(table)
                ),
                None,
            )
        if balance_table:
            blocks.append({"type": "table", "rows": balance_table})

    if operation_section:
        operation_table = parse_pdf_table_lines(operation_section)
        if not operation_table:
            operation_table = next(
                (table for table in plumber_tables if looks_like_pdf_operation_table(table)),
                None,
            )
        if operation_table:
            blocks.append({"type": "table", "rows": operation_table})

    if update_start is not None:
        update_lines = merge_pdf_update_lines(lines[update_start + 1 : update_end])
        for line in update_lines:
            if is_top_level_numbered_section_line(line) or is_decimal_section_heading_line(line):
                continue
            blocks.append({"type": "paragraph", "text": line})

    return blocks


def load_blocks(path: Path) -> list[dict[str, Any]]:
    if path.suffix.lower() == ".pdf":
        raw_blocks = load_pdf_blocks(path)
        converted_path = convert_pdf_via_word(path)
        if converted_path is None:
            return raw_blocks
        try:
            converted_blocks = load_docx_blocks(converted_path)
        except Exception:
            return raw_blocks
        raw_parsed = parse_current_blocks(raw_blocks) if raw_blocks else {}
        converted_parsed = parse_current_blocks(converted_blocks) if converted_blocks else {}
        raw_finance_quality = score_finance_table_quality(raw_parsed.get("finance_table"))
        converted_finance_quality = score_finance_table_quality(converted_parsed.get("finance_table"))
        if raw_finance_quality >= 8 and raw_finance_quality > converted_finance_quality:
            return raw_blocks
        if score_extracted_blocks(converted_blocks) >= score_extracted_blocks(raw_blocks):
            return converted_blocks
        return raw_blocks
    return load_docx_blocks(path)


def blocks_to_text(blocks: list[dict[str, Any]]) -> str:
    lines = []
    for block in blocks:
        if block["type"] == "paragraph":
            lines.append(block["text"])
        else:
            for row in block["rows"]:
                lines.append(" | ".join(row))
    return "\n".join(lines)


def docx_to_text(path: Path) -> str:
    blocks = load_blocks(path)
    return blocks_to_text(blocks)


def detect_language(text: str) -> tuple[str, str]:
    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
    if chinese_chars > 20:
        return "chinese", f"Detected {chinese_chars} Chinese characters in the company-submitted source."
    return "english", "No meaningful Chinese text detected in the company-submitted source."


def text_language(text: str) -> str:
    return "chinese" if re.search(r"[\u4e00-\u9fff]", text) else "english"


def normalize_heading_probe_text(text: str) -> str:
    lowered = norm_space(text).lower()
    # Strip common numbering prefixes first (e.g. "3-", "2.", "三、").
    lowered = re.sub(r"^[0-9一二三四五六七八九十]+(?:(?:[\-–、.．）\)]\s*)|\s+)", "", lowered)

    # Strip leading quarter/date prefixes so headings like
    # "2025Q3业务进展" / "2025年Q4业务进展" resolve to the same aliases.
    quarter_prefix_patterns = [
        r"^20\d{2}\s*[-_/]?\s*q[1-4][\-–、.\s:：]*",
        r"^q[1-4]\s*[-_/]?\s*20\d{2}[\-–、.\s:：]*",
        r"^20\d{2}\s*年\s*q?[1-4][\-–、.\s:：]*",
        r"^20\d{2}\s*年?\s*第?\s*[一二三四1234]\s*季度[\-–、.\s:：]*",
    ]
    changed = True
    while changed:
        changed = False
        for pattern in quarter_prefix_patterns:
            updated = re.sub(pattern, "", lowered)
            if updated != lowered:
                lowered = updated
                changed = True
    return lowered


def canonical_heading(text: str) -> str | None:
    lowered = normalize_heading_probe_text(text)
    for canonical, aliases in SECTION_ALIASES.items():
        for alias in aliases:
            if lowered == alias or lowered.startswith(alias):
                return canonical
    if lowered.startswith("operations update") or lowered.startswith("business progress"):
        return "Business Update"
    if "风险" in lowered and "退出" in lowered:
        return "Risk & Exit"
    return None


def fixed_anchor_heading(text: str) -> str | None:
    lowered = normalize_heading_probe_text(text)
    for canonical in ["Business Activities", "Financial Update", "Risk & Exit"]:
        for alias in SECTION_ALIASES.get(canonical, []):
            if lowered == alias or lowered.startswith(alias):
                return canonical
    if "风险" in lowered and "退出" in lowered:
        return "Risk & Exit"
    return None


def is_middle_section_heading_candidate(paragraph) -> bool:
    text = norm_space(paragraph.text)
    if not text:
        return False
    if text.endswith((".", "。", ";", "；", ":", "：", "?", "？", ")", "）")):
        return False
    if len(text) > 80:
        return False
    if text.startswith(("-", "•", "*", "(", "（")):
        return False
    # Unit text should never be a section heading
    if is_unit_text(text):
        return False
    style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
    if "list" in style_name or "heading" in style_name or "标题" in style_name or style_name == "normal":
        return True
    return False


def is_middle_section_heading_text_candidate(text: str) -> bool:
    stripped = norm_space(text)
    if not stripped:
        return False
    if len(stripped) > 80:
        return False
    if stripped.startswith(("-", "•", "*", "(", "（")):
        return False
    if stripped.endswith((".", "。", ";", "；", ":", "：", "?", "？", ")", "）")):
        return False
    if is_unit_text(stripped):
        return False
    if is_question_or_prompt(stripped):
        return False
    if re.search(r"\d{2,}", stripped):
        return False
    if re.search(r"[。.!?？；;:：]", stripped):
        return False
    return True


def extract_previous_sections(blocks: list[dict[str, Any]]) -> dict[str, Any]:
    sections: dict[str, Any] = {}
    current_heading: str | None = None

    for block in blocks:
        if block["type"] == "paragraph":
            heading = canonical_heading(block["text"])
        else:
            heading = None
        if heading:
            current_heading = heading
            sections[current_heading] = {"paragraphs": [], "tables": []}
            continue
        if not current_heading:
            continue
        if block["type"] == "paragraph":
            sections[current_heading]["paragraphs"].append(block["text"])
        elif block["type"] == "table":
            sections[current_heading]["tables"].append(block["rows"])
    return sections


def extract_previous_section_occurrences(blocks: list[dict[str, Any]]) -> list[SectionOccurrence]:
    occurrences: list[SectionOccurrence] = []
    current: SectionOccurrence | None = None
    after_financial: dict[str, bool] = {"english": False, "chinese": False}
    after_risk: dict[str, bool] = {"english": False, "chinese": False}
    for block in blocks:
        heading = None
        language = "english"
        if block["type"] == "paragraph":
            text = block["text"]
            language = text_language(text)
            heading = fixed_anchor_heading(text) or canonical_heading(text)
            if heading is None and after_financial.get(language, False) and not after_risk.get(language, False):
                if is_middle_section_heading_text_candidate(text):
                    heading = "Business Update"
        if heading:
            current = SectionOccurrence(
                canonical=heading,
                language=language,
                heading_text=block["text"],
                paragraphs=[],
                tables=[],
            )
            occurrences.append(current)
            if heading == "Financial Update":
                after_financial[language] = True
                after_risk[language] = False
            elif heading == "Risk & Exit":
                after_risk[language] = True
            continue
        if current is None:
            continue
        if block["type"] == "paragraph":
            current.paragraphs.append(block["text"])
        elif block["type"] == "table":
            current.tables.append(block["rows"])
    return occurrences


def parse_previous_financial_table(rows: list[list[str]]) -> dict[str, Any]:
    header_row = rows[0]
    unit = header_row[0]
    headers = [norm_space(cell) for cell in header_row[1:]]
    data_rows = []
    for row in rows[1:]:
        label = norm_space(row[0])
        values = {headers[i]: norm_space(row[i + 1]) for i in range(len(headers))}
        data_rows.append({"label": label, "values": values})
    return {"unit": unit, "headers": headers, "rows": data_rows}


def parse_previous_financial_table_with_unit(rows: list[list[str]], unit_label: str | None) -> dict[str, Any]:
    parsed = parse_previous_financial_table(rows)
    if unit_label:
        parsed["unit"] = unit_label
    return parsed


@dataclass(frozen=True, order=True)
class Quarter:
    year: int
    quarter: int

    def display(self) -> str:
        return f"Q{self.quarter} {self.year}"

    def previous(self) -> "Quarter":
        if self.quarter == 1:
            return Quarter(self.year - 1, 4)
        return Quarter(self.year, self.quarter - 1)


@dataclass
class SectionOccurrence:
    canonical: str
    language: str
    heading_text: str
    paragraphs: list[str]
    tables: list[list[list[str]]]


@dataclass
class DocSectionOccurrence:
    canonical: str
    language: str
    heading_index: int


@dataclass
class MiddleSectionCandidates:
    raw_paragraphs: list[str]
    source_groups: list[dict[str, Any]]
    chinese_lines: list[str]
    english_lines: list[str]


@dataclass
class MiddleSectionOutputs:
    flat_by_language: dict[str, list[str]]
    grouped_by_language: dict[str, list[list[str]]]
    translation_engine: str
    translation_engine_note: str


@dataclass
class MiddleSectionSelection:
    selected_groups: list[dict[str, Any]]
    chinese_lines: list[str]
    english_lines: list[str]


@dataclass(frozen=True)
class UnitSpec:
    label: str
    factor_to_base: Decimal
    currency: str | None = None


@dataclass(frozen=True)
class ExchangeRates:
    rmb_per_usd: Decimal | None = None


def parse_quarter_label(text: str) -> Quarter | None:
    match = re.search(r"Q([1-4])\s*(20\d{2})", text)
    if match:
        return Quarter(int(match.group(2)), int(match.group(1)))
    match = re.search(r"(20\d{2})\s*Q([1-4])", text)
    if match:
        return Quarter(int(match.group(1)), int(match.group(2)))
    match = re.search(r"(20\d{2})Q([1-4])", text)
    if match:
        return Quarter(int(match.group(1)), int(match.group(2)))
    # Chinese quarter formats: "2025年第4季度", "2025年Q4", "2025第四季度"
    cn_digits = {"一": 1, "二": 2, "三": 3, "四": 4}
    match = re.search(r"(20\d{2})\s*年?\s*第?\s*([1-4一二三四])\s*季度", text)
    if match:
        q = cn_digits.get(match.group(2), None) or int(match.group(2))
        return Quarter(int(match.group(1)), q)
    return None


def extract_quarter_number(text: str) -> int | None:
    match = re.search(r"Q([1-4])", text, flags=re.I)
    if match:
        return int(match.group(1))
    match = re.search(r"([1-4])季度", text)
    if match:
        return int(match.group(1))
    return None


def parse_month_header(text: str) -> tuple[int, int] | None:
    normalized = norm_space(text)
    match = re.search(r"(20\d{2})[/-]?(0[1-9]|1[0-2])", normalized)
    if match:
        return int(match.group(1)), int(match.group(2))
    # Chinese month formats: "2025年1月", "2025年01月", "1月" (without year)
    match = re.search(r"(20\d{2})\s*年\s*(1[0-2]|0?[1-9])\s*月", normalized)
    if match:
        return int(match.group(1)), int(match.group(2))
    return None


def quarter_sort_key(text: str) -> tuple[int, int]:
    quarter = parse_quarter_label(text)
    if quarter is None:
        return (-1, -1)
    return (quarter.year, quarter.quarter)


def parse_decimal(text: str) -> Decimal | None:
    cleaned = text.replace(",", "").strip()
    if not cleaned or cleaned in {"/", "n.a.", "n/a"}:
        return None
    negative = cleaned.startswith("(") and cleaned.endswith(")")
    cleaned = cleaned.strip("()")
    cleaned = cleaned.replace("Positive cf", "").replace("Better", "").strip()
    # Detect K/M/B suffix multiplier before stripping letters
    suffix_multiplier = Decimal("1")
    suffix_match = re.search(r"(\d)\s*([KkMmBb])\s*$", cleaned)
    if suffix_match:
        suffix_char = suffix_match.group(2).upper()
        if suffix_char == "K":
            suffix_multiplier = Decimal("1000")
        elif suffix_char == "M":
            suffix_multiplier = Decimal("1000000")
        elif suffix_char == "B":
            suffix_multiplier = Decimal("1000000000")
    cleaned = re.sub(r"[^0-9.+-]", "", cleaned)
    if not cleaned:
        return None
    try:
        value = Decimal(cleaned) * suffix_multiplier
    except Exception:
        return None
    return -value if negative else value


def finance_parse_decimal(text: str | None) -> Decimal | None:
    cleaned = norm_space(text or "")
    if cleaned == "/":
        return Decimal("0")
    return parse_decimal(cleaned)


def format_decimal(value: Decimal | None, places: int = 2) -> str | None:
    if value is None:
        return None
    quant = Decimal("1." + ("0" * places))
    rounded = value.quantize(quant, rounding=ROUND_HALF_UP)
    return f"{rounded:.{places}f}"


def parse_percent(text: str) -> Decimal | None:
    cleaned = norm_space(text)
    if not cleaned or cleaned == "/":
        return None
    negative = cleaned.startswith("(") and "%" in cleaned
    match = re.search(r"([+-]?\d+(?:\.\d+)?)\s*%", cleaned)
    if not match:
        return None
    value = Decimal(match.group(1))
    if negative and value > 0:
        value = -value
    return value


def format_percent(value: Decimal | None) -> str | None:
    if value is None:
        return None
    rounded = value.quantize(Decimal("1.00"), rounding=ROUND_HALF_UP)
    return f"{rounded:.2f}%"


def display_percent_text(current_text: str | None, previous_text: str | None) -> str | None:
    current_clean = norm_space(current_text or "")
    previous_clean = norm_space(previous_text or "")
    current = finance_parse_decimal(current_clean)
    previous = finance_parse_decimal(previous_clean)
    if current == 0 and previous == 0:
        return "0.00%"
    return format_percent(calc_percent(current, previous))


def default_missing_financial_value(label: str | None = None) -> str:
    normalized = normalize_metric_label(label or "")
    if normalized in {"#of company-owned outlets", "#of jv outlets"}:
        return "0"
    return "0.00"


def percent_cell_font_size(text: str | None) -> float:
    normalized = norm_space((text or "").replace("\u2060", ""))
    if not normalized:
        return 8.0
    # Keep normal percentages readable, but compress unusually long outliers
    # like 55975.45% or -650.00% so they stay on one line in narrow templates.
    if len(normalized) >= 10:
        return 7.0
    if len(normalized) >= 8:
        return 7.5
    return 8.0


def get_cell_width_twips(cell) -> int | None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        return None
    try:
        return int(tc_w.get(qn("w:w")))
    except (TypeError, ValueError):
        return None


def should_drop_font_one_point(cell, text: str, language: str) -> bool:
    normalized = norm_space((text or "").replace("\u2060", ""))
    if not normalized:
        return False
    cell_width = get_cell_width_twips(cell)
    if not cell_width:
        return False
    # Approximate whether Word is likely to wrap based on the current cell width
    # and the text-width estimator already used for table sizing.
    return estimate_table_text_width(normalized, language) > max(0, cell_width - 70)


def table_requires_font_drop(table, language: str) -> bool:
    for row in table.rows:
        for cell in row.cells:
            if should_drop_font_one_point(cell, cell.text, language):
                return True
    return False


def center_cell_paragraphs(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def current_quarter_from_financial_updates(financial_updates: list[dict[str, Any]]) -> Quarter | None:
    for update in financial_updates:
        for column in update.get("columns", []):
            parsed = parse_quarter_label(column)
            if parsed is not None:
                return parsed
    return None


def normalize_business_update_heading_text(text: str, target_quarter: Quarter | None) -> str:
    stripped = norm_space(text)
    if not stripped or target_quarter is None:
        return stripped
    cn_label = f"{target_quarter.year}Q{target_quarter.quarter}"
    zh_quarter_word = {1: "第一", 2: "第二", 3: "第三", 4: "第四"}[target_quarter.quarter]
    en_label = target_quarter.display()
    en_label_alt = f"{target_quarter.year} Q{target_quarter.quarter}"

    if re.search(r"[\u4e00-\u9fff]", stripped):
        stripped = re.sub(
            r"20\d{2}\s*年\s*第[一二三四1-4]\s*季度",
            f"{target_quarter.year}年{zh_quarter_word}季度",
            stripped,
            flags=re.I,
        )
        stripped = re.sub(r"20\d{2}\s*Q\s*[1-4]", cn_label, stripped, flags=re.I)
        stripped = re.sub(r"Q\s*[1-4]\s*20\d{2}", cn_label, stripped, flags=re.I)
        return stripped

    if re.match(r"^\s*20\d{2}\s*Q\s*[1-4]\b", stripped, flags=re.I):
        stripped = re.sub(r"^\s*20\d{2}\s*Q\s*[1-4]\b", en_label_alt, stripped, flags=re.I)
    else:
        stripped = re.sub(r"20\d{2}\s*Q\s*[1-4]", en_label, stripped, flags=re.I)
        stripped = re.sub(r"Q\s*[1-4]\s*20\d{2}", en_label, stripped, flags=re.I)
    return stripped


def normalize_middle_section_heading_quarters(doc: Document, occurrences: list[DocSectionOccurrence], target_quarter: Quarter | None) -> None:
    for indices in middle_section_indices_by_language(occurrences).values():
        for occ_idx in indices:
            occurrence = occurrences[occ_idx]
            heading_paragraph = doc.paragraphs[occurrence.heading_index]
            original = norm_space(heading_paragraph.text)
            normalized = normalize_business_update_heading_text(original, target_quarter)
            if normalized and normalized != original:
                para_format = capture_run_format(heading_paragraph)
                set_paragraph_text_preserve(heading_paragraph, normalized)
                apply_paragraph_run_format(heading_paragraph, *para_format)


def is_middle_heading_like_paragraph(text: str) -> bool:
    stripped = norm_space(text)
    if not stripped or len(stripped) > 90:
        return False
    lowered = stripped.lower()
    return any(
        token in stripped or token in lowered
        for token in ["业务回顾", "业务进展", "运营情况", "Business Review", "Business Update", "Business Progress", "Business highlight"]
    )


def normalize_middle_heading_like_paragraphs(doc: Document, target_quarter: Quarter | None) -> None:
    if target_quarter is None:
        return
    for paragraph in doc.paragraphs:
        original = norm_space(paragraph.text)
        if not is_middle_heading_like_paragraph(original):
            continue
        normalized = normalize_business_update_heading_text(original, target_quarter)
        if normalized and normalized != original:
            para_format = capture_run_format(paragraph)
            set_paragraph_text_preserve(paragraph, normalized)
            apply_paragraph_run_format(paragraph, *para_format)


def is_ignorable_middle_heading_text(text: str) -> bool:
    stripped = norm_space(text)
    if not stripped:
        return True
    if is_business_update_root_heading_text(stripped):
        return False
    if is_admin_or_meta_chinese_line(stripped) or is_admin_or_meta_english_line(stripped):
        return True
    return stripped in {"无变化", "暂无变化", "No change", "No material change"}


def remove_ignorable_middle_headings(doc: Document, occurrences: list[DocSectionOccurrence]) -> None:
    grouped = middle_section_indices_by_language(occurrences)
    all_indices = sorted([idx for indices in grouped.values() for idx in indices], reverse=True)
    for occ_idx in all_indices:
        occurrence = occurrences[occ_idx]
        heading_paragraph = doc.paragraphs[occurrence.heading_index]
        if is_ignorable_middle_heading_text(heading_paragraph.text):
            remove_paragraph(heading_paragraph)


def is_business_update_root_label(text: str) -> bool:
    lowered = norm_space(text).lower()
    return lowered in {
        "业务发展",
        "业务进展",
        "运营情况",
        "business update",
        "business progress",
        "business highlights",
        "recent trends",
    }


def is_current_business_update_heading(text: str) -> bool:
    stripped = norm_space(text)
    if not stripped:
        return False
    lowered = stripped.lower()
    if is_business_update_root_label(stripped):
        return False
    if re.search(r"\d+(?:\.\d+)?\s*(?:万|亿|million|mn|m)\b", stripped, flags=re.I):
        return False
    if re.match(r"^[0-9一二三四五六七八九十]+[)）.、-]", stripped):
        remainder = re.sub(r"^[0-9一二三四五六七八九十]+[)）.、-]\s*", "", stripped)
        if (("：" in remainder or ":" in remainder) and len(remainder) > 12) or is_company_amount_line_item(remainder):
            return False
        return True
    if re.search(r"[?？]$", stripped) and len(stripped) <= 120:
        return True
    if "业务回顾" in stripped or "business review" in lowered:
        return True
    if stripped.endswith(("情况", "安排", "规划", "进展")) and len(stripped) <= 40 and "，" not in stripped and "," not in stripped:
        return True
    return False


def normalize_current_business_update_heading(text: str) -> str:
    stripped = norm_space(text).strip()
    return stripped.rstrip("：:？?。；;")


def clean_grouped_update_paragraphs(paragraphs: list[str]) -> list[str]:
    cleaned: list[str] = []
    for text in paragraphs:
        text = norm_space(text)
        if not text:
            continue
        if is_standalone_label(text):
            continue
        text = strip_response_prefix(text)
        if not text or len(text) < 2:
            continue
        cleaned.append(text)
    return cleaned


def group_current_business_update_paragraphs(paragraphs: list[str]) -> list[dict[str, Any]]:
    groups: list[dict[str, Any]] = []
    current_heading: str | None = None
    current_items: list[str] = []

    def flush() -> None:
        nonlocal current_heading, current_items
        if current_heading is None and not current_items:
            return
        cleaned_items = clean_grouped_update_paragraphs(current_items)
        if cleaned_items:
            groups.append({"heading": current_heading, "paragraphs": cleaned_items})
        current_heading = None
        current_items = []

    for raw_text in paragraphs:
        text = norm_space(raw_text)
        if not text or is_business_update_root_label(text):
            continue
        if is_current_business_update_heading(text):
            flush()
            current_heading = normalize_current_business_update_heading(text)
            continue
        current_items.append(text)

    flush()
    if groups:
        return groups
    cleaned = clean_update_paragraphs(paragraphs)
    return [{"heading": None, "paragraphs": cleaned}] if cleaned else []


def parse_current_blocks(blocks: list[dict[str, Any]]) -> dict[str, Any]:
    result: dict[str, Any] = {
        "title": "",
        "finance_table": None,
        "balance_table": None,
        "operation_table": None,
        "extra_tables": [],
        "business_update_raw_paragraphs": [],
        "business_update_paragraphs": [],
        "business_update_groups": [],
    }

    tables_seen = 0
    for idx, block in enumerate(blocks):
        if idx == 0 and block["type"] == "paragraph":
            result["title"] = block["text"]

        if block["type"] == "table":
            tables_seen += 1
            if result["finance_table"] is None:
                result["finance_table"] = block["rows"]
            elif result["balance_table"] is None:
                result["balance_table"] = block["rows"]
            else:
                result["extra_tables"].append(block["rows"])
                if result["operation_table"] is None:
                    result["operation_table"] = block["rows"]

        if block["type"] == "paragraph" and block["text"] == "Other Questions:":
            trailing = blocks[idx + 1 :]
            for item in trailing:
                if item["type"] != "paragraph":
                    continue
                result["business_update_raw_paragraphs"].append(item["text"])
            break

    if not result["business_update_raw_paragraphs"]:
        seen_table = False
        raw_after_tables: list[str] = []
        for block in blocks:
            if block["type"] == "table":
                seen_table = True
                continue
            if not seen_table or block["type"] != "paragraph":
                continue
            raw_after_tables.append(block["text"])

        root_idx = next((idx for idx, text in enumerate(raw_after_tables) if is_business_update_root_label(text)), None)
        result["business_update_raw_paragraphs"] = raw_after_tables[root_idx + 1 :] if root_idx is not None else raw_after_tables

    result["business_update_groups"] = group_current_business_update_paragraphs(result["business_update_raw_paragraphs"])
    result["business_update_paragraphs"] = clean_update_paragraphs(result["business_update_raw_paragraphs"])
    return result


def rows_to_dict(rows: list[list[str]]) -> dict[str, dict[str, str]]:
    headers = rows[0][1:]
    mapping: dict[str, dict[str, str]] = {}
    for row in rows[1:]:
        if not row:
            continue
        label = norm_space(row[0])
        if not label:
            continue
        mapping[label] = {
            headers[i]: norm_space(row[i + 1]) if i + 1 < len(row) else ""
            for i in range(len(headers))
        }
    return mapping


def extract_outlet_counts(operation_rows: list[list[str]]) -> tuple[str | None, str | None]:
    op_map = rows_to_dict(operation_rows)
    outlet_row = None
    for label in op_map:
        if "outlet" in label.lower() and "accumul" in label.lower():
            outlet_row = op_map[label]
            break
    if outlet_row is None:
        outlet_row = op_map.get("Accumulated Number of outlets", {})
    ytd_value = None
    for header, value in outlet_row.items():
        if "ytd" in header.lower() or "accum" in header.lower() or "total" in header.lower():
            ytd_value = value
            break
    if not ytd_value:
        values = list(outlet_row.values())
        ytd_value = values[-1] if values else None
    if not ytd_value:
        return None, None
    match = re.search(r"(\d+)\s*\+\s*(\d+)\s*\(JV\)", ytd_value)
    if not match:
        plain = re.search(r"(\d+)", ytd_value)
        return (plain.group(1), None) if plain else (None, None)
    return match.group(1), match.group(2)



def normalize_finance_source(rows: list[list[str]]) -> dict[str, Any]:
    header = [norm_space(cell) for cell in rows[0]]
    row_map = rows_to_dict(rows)
    return {
        "header": header,
        "rows": row_map,
    }


def is_empty_numeric_cell(text: str) -> bool:
    cleaned = norm_space(text)
    return cleaned in {"", "/", "-", "—", "–"}


def finance_row_current_period_is_blank(
    source_row: dict[str, str],
    current_month_headers: list[str],
    current_total_header: str,
    fy_or_ytd_header: str | None,
) -> bool:
    relevant_headers = list(current_month_headers)
    if current_total_header:
        relevant_headers.append(current_total_header)
    if fy_or_ytd_header:
        relevant_headers.append(fy_or_ytd_header)
    if not relevant_headers:
        return False
    for header in relevant_headers:
        if not is_empty_numeric_cell(source_row.get(header, "")):
            return False
    return True


def extract_cash_balance_line(balance_rows: list[list[str]] | None, language: str) -> str | None:
    if not balance_rows or len(balance_rows) < 2:
        return None
    for row in balance_rows[1:]:
        if len(row) < 2:
            continue
        label = norm_space(row[0])
        value = norm_space(row[1])
        if not label or not value:
            continue
        label_lower = label.lower()
        if "银行账户现金余额" not in label and "bank account cash balance" not in label_lower and "cash balance" not in label_lower:
            continue

        rmb_match = re.search(r"(?:人民币|RMB)\s*[:：]?\s*([0-9,]+(?:\.\d+)?)", value, flags=re.I)
        usd_match = re.search(r"(?:美元|USD)\s*[:：]?\s*([0-9,]+(?:\.\d+)?)", value, flags=re.I)
        rmb_value = rmb_match.group(1) if rmb_match else None
        usd_value = usd_match.group(1) if usd_match else None

        if language == "chinese":
            parts = []
            if rmb_value is not None:
                parts.append(f"人民币：{rmb_value}")
            if usd_value is not None:
                parts.append(f"美元：{usd_value}")
            return f"银行账户现金余额：{' '.join(parts)}" if parts else f"银行账户现金余额：{value}"

        parts = []
        if rmb_value is not None:
            parts.append(f"RMB {rmb_value}")
        if usd_value is not None:
            parts.append(f"USD {usd_value}")
        return f"Bank Account Cash Balance: {', '.join(parts)}" if parts else f"Bank Account Cash Balance: {value}"
    return None


def calc_percent(current: Decimal | None, previous: Decimal | None) -> Decimal | None:
    if current is None or previous is None:
        return None
    if previous == 0:
        if current == 0:
            return Decimal("0")
        return None
    denominator = abs(previous)
    if denominator == 0:
        return None
    return ((current - previous) / denominator) * Decimal("100")


def rewrite_third_person(text: str, company_name: str) -> str:
    company_label = english_company_name(company_name)
    if not company_label:
        company_label = "The company"
    text = norm_space(text)
    text = re.sub(r"\bwe\b", company_label, text, flags=re.I)
    text = re.sub(r"\bour\b", f"{company_label}'s", text, flags=re.I)
    text = re.sub(r"\bus\b", company_label, text, flags=re.I)
    text = re.sub(r"\bQ([1-4])\.", r"Q\1", text)
    text = re.sub(r"\bthe company\b", company_label, text, flags=re.I)
    text = re.sub(r"[\u4e00-\u9fff]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return norm_space(text)


def short_sentence(text: str, company_name: str, limit: int = 140) -> str:
    text = rewrite_third_person(text, company_name)
    text = re.sub(r":\s*", ": ", text)
    text = re.sub(r"\s+-\s+", " - ", text)
    if len(text) <= limit:
        return text.rstrip(".") + "."
    clipped = text[:limit].rsplit(" ", 1)[0].rstrip(",;:-")
    return clipped + "."


def score_paragraph_informativeness(text: str) -> float:
    """Score a paragraph by how much useful business information it contains.
    Higher score = more informative (numbers, metrics, dollar amounts, percentages, specifics)."""
    score = 0.0
    score += len(re.findall(r"\$[\d,.]+[kmb]?", text, re.I)) * 3.0
    score += len(re.findall(r"\d+(?:\.\d+)?%", text)) * 2.5
    score += len(re.findall(r"\b\d{1,3}(?:,\d{3})+\b", text)) * 2.0
    score += len(re.findall(r"\b\d+(?:\.\d+)?\s*(?:x|times|million|billion|k)\b", text, re.I)) * 2.0
    score += len(re.findall(r"\bQ[1-4]\s*20\d{2}\b", text, re.I)) * 1.5
    score += len(re.findall(r"\b(?:revenue|profit|ebitda|margin|growth|cost|loss|cash flow|burn rate|funding|financing|expansion|launched|achieved|target|milestone)\b", text, re.I)) * 1.0
    score += len(re.findall(r"\b(?:YoY|QoQ|year.over.year|quarter.over.quarter|month.over.month)\b", text, re.I)) * 1.5
    length_bonus = min(len(text) / 200.0, 1.0)
    score += length_bonus
    if len(text) < 30:
        score *= 0.3
    return score


def split_into_sentences(text: str) -> list[str]:
    """Split a long paragraph into individual sentences for separate bullets."""
    # Don't split short text
    if len(text) < 120:
        return [text]
    # Split on sentence boundaries, but not on abbreviations like "U.S." or "e.g."
    # or on decimal numbers like "2.67m"
    parts = re.split(r'(?<=[.!])\s+(?=[A-Z\d])', text)
    result = []
    for part in parts:
        part = norm_space(part)
        if len(part) >= 20:
            result.append(part)
    return result if result else [text]


CONTINUATION_PATTERNS = re.compile(
    r"^(?:among them|in addition|additionally|furthermore|moreover|also|"
    r"in particular|specifically|for (?:example|instance)|this (?:includes?|means?)|"
    r"these |the key |the planned |the new )\b",
    re.I,
)


def merge_continuations(items: list[str]) -> list[str]:
    """Merge continuation sentences with their preceding item."""
    if not items:
        return items
    merged: list[str] = [items[0]]
    for item in items[1:]:
        if CONTINUATION_PATTERNS.search(item) and merged:
            merged[-1] = merged[-1].rstrip(".") + ". " + item
        else:
            merged.append(item)
    return merged


def build_business_update_bullets(paragraphs: list[str], company_name: str, max_bullets: int = 9) -> list[str]:
    """Extract the most informative paragraphs as bullet points, scored by content richness."""
    cleaned = [strip_leading_list_marker(text) for text in clean_update_paragraphs(paragraphs)]
    # Merge continuation sentences before splitting
    cleaned = merge_continuations(cleaned)
    # Split multi-sentence paragraphs into separate candidates
    expanded: list[str] = []
    for text in cleaned:
        text = norm_space(text)
        if len(text) < 20:
            continue
        if text.endswith((":", "：")):
            continue
        expanded.extend(split_into_sentences(text))
    # Merge again after splitting (in case split created new continuations)
    expanded = merge_continuations(expanded)
    scored: list[tuple[float, str]] = []
    for text in expanded:
        score = score_paragraph_informativeness(text)
        scored.append((score, text))

    scored.sort(key=lambda x: x[0], reverse=True)

    bullets: list[str] = []
    seen_keys: set[str] = set()
    for _score, text in scored:
        bullet = short_sentence(text, company_name, limit=200)
        key = re.sub(r"[^a-z0-9]", "", bullet.lower())
        if key in seen_keys:
            continue
        seen_keys.add(key)
        bullets.append(bullet)
        if len(bullets) >= max_bullets:
            break

    return bullets


def format_decimal_compact(value: Decimal) -> str:
    text = f"{value.quantize(Decimal('0.01'))}"
    if "." in text:
        text = text.rstrip("0").rstrip(".")
    return text


def convert_chinese_amounts_to_rmb_millions(text: str) -> str:
    def repl_yi(match: re.Match[str]) -> str:
        value = Decimal(match.group(1))
        million = value * Decimal("100")
        return f"RMB {format_decimal_compact(million)} million"

    def repl_wan(match: re.Match[str]) -> str:
        value = Decimal(match.group(1))
        million = value / Decimal("100")
        return f"RMB {format_decimal_compact(million)} million"

    text = re.sub(r"(\d+(?:\.\d+)?)\s*亿元", repl_yi, text)
    text = re.sub(r"(\d+(?:\.\d+)?)\s*万(?:元人民币|人民币)?", repl_wan, text)
    text = re.sub(r"(\d+(?:\.\d+)?)\s*万", repl_wan, text)
    return text


def extract_trailing_chinese_amount_display(text: str) -> tuple[str, str] | None:
    match = re.search(r"^(.*?)(\d+(?:\.\d+)?)\s*(亿|万)(?:元人民币|人民币)?[。.]?$", norm_space(text))
    if not match:
        return None
    prefix = norm_space(match.group(1)).rstrip("，,；;：:")
    amount_value = Decimal(match.group(2))
    amount_unit = match.group(3)
    amount_million = amount_value * Decimal("100") if amount_unit == "亿" else amount_value / Decimal("100")
    return prefix, f"RMB {format_decimal_compact(amount_million)} million"


def translate_project_amount_bullet_with_template(text: str, company_name: str) -> str | None:
    extracted = extract_trailing_chinese_amount_display(text)
    if not extracted:
        return None
    project_text, amount_display = extracted
    if len(project_text) < 4:
        return None

    translated_project: str | None = None
    ready, _reason = ensure_argos_ready()
    if ready:
        translation = ARGOS_RESOURCES.get("translation")
        if translation is not None:
            try:
                translated_project = norm_space(translation.translate(project_text)).strip(" \"'")
            except Exception:
                translated_project = None

    if not translated_project:
        translated_project = convert_chinese_amounts_to_rmb_millions(project_text)

    translated_project = polish_english_business_bullet(translated_project, company_name)
    translated_project = translated_project.rstrip(".")
    translated_project = re.sub(r"\s+", " ", translated_project).strip(" ,.;:-")
    if contains_chinese_chars(translated_project):
        translated_project = re.sub(r"[\u4e00-\u9fff]+", " ", translated_project)
        translated_project = re.sub(r"\s+", " ", translated_project).strip(" ,.;:-")
    if len(re.findall(r"[A-Za-z]{2,}", translated_project)) < 3:
        return None
    return f"{translated_project}: contract amount of {amount_display}."


def english_company_name(company_name: str) -> str:
    cleaned = re.sub(r"[\u4e00-\u9fff]+", " ", company_name or "")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    match = re.search(r"[A-Za-z][A-Za-z0-9& .'-]*", cleaned)
    if match:
        return norm_space(match.group(0))
    return ""


def strip_leading_list_marker(text: str) -> str:
    text = norm_space(text)
    text = re.sub(r"^[（(]\s*[0-9一二三四五六七八九十]+\s*[）)]\s*", "", text)
    text = re.sub(r"^[0-9一二三四五六七八九十]+\s*[)）.\-、:：]\s*", "", text)
    return norm_space(text)


def normalize_translation_memory_key(text: str) -> str:
    text = strip_leading_list_marker(text)
    text = re.sub(r"(20\d{2})年\s*Q([1-4])", "<Q>", text, flags=re.I)
    text = re.sub(r"Q([1-4])\s*(20\d{2})", "<Q>", text, flags=re.I)
    text = re.sub(r"第[一二三四1234]季度", "<Q>", text)
    text = re.sub(r"\d+(?:\.\d+)?\s*亿元", "<AMOUNT>", text)
    text = re.sub(r"\d+(?:\.\d+)?\s*万(?:元人民币|人民币)?", "<AMOUNT>", text)
    text = re.sub(r"[\s，,。；;：:（）()]+", "", text)
    return text.lower()


def extract_quarter_display_from_chinese(text: str) -> str | None:
    match = re.search(r"(20\d{2})年\s*Q([1-4])", text, re.I)
    if match:
        return f"{match.group(1)} Q{match.group(2)}"
    match = re.search(r"(20\d{2})年第?([一二三四1234])季度", text)
    if match:
        mapping = {"一": "1", "二": "2", "三": "3", "四": "4"}
        q = mapping.get(match.group(2), match.group(2))
        return f"{match.group(1)} Q{q}"
    return None


def extract_rmb_million_display(text: str) -> str | None:
    converted = convert_chinese_amounts_to_rmb_millions(text)
    match = re.search(r"RMB ([\d.]+) million", converted)
    if match:
        return match.group(1)
    return None


def build_business_update_translation_memory(previous_business_update_map: dict[str, list[str]]) -> dict[str, str]:
    memory: dict[str, str] = {}
    zh_lines = previous_business_update_map.get("chinese", [])
    en_lines = previous_business_update_map.get("english", [])
    for zh, en in zip(zh_lines, en_lines):
        zh_clean = strip_leading_list_marker(zh)
        en_clean = strip_leading_list_marker(en)
        if not zh_clean or not en_clean:
            continue
        memory[normalize_translation_memory_key(zh_clean)] = en_clean
    return memory


def translation_memory_match(text: str, translation_memory: dict[str, str]) -> str | None:
    key = normalize_translation_memory_key(text)
    if key in translation_memory:
        return translation_memory[key]
    for memory_key, memory_value in translation_memory.items():
        if not memory_key:
            continue
        overlap = len(set(key) & set(memory_key)) / max(len(set(memory_key)), 1)
        if overlap >= 0.7:
            return memory_value
    return None


def should_use_translation_memory(current_chinese: str, previous_english: str) -> bool:
    zh_numbers = re.findall(r"\d+(?:\.\d+)?", current_chinese)
    en_numbers = re.findall(r"\d+(?:\.\d+)?", previous_english)
    if len(zh_numbers) > len(en_numbers) + 2:
        return False
    if any(token in current_chinese for token in ["人员变动", "销售量", "万台", "兼职", "全职"]) and not any(
        token in previous_english.lower() for token in ["personnel", "employee", "staff", "sales volume", "units", "full-time", "part-time"]
    ):
        return False
    return True


def adapt_previous_english_translation(previous_english: str, current_chinese: str, company_name: str) -> str:
    text = previous_english
    quarter_display = extract_quarter_display_from_chinese(current_chinese)
    if quarter_display:
        text = re.sub(r"\b20\d{2}\s*Q[1-4]\b", quarter_display, text)
        text = re.sub(r"\bQ[1-4]\s*20\d{2}\b", quarter_display, text)
        text = re.sub(r"\bQ[1-4]\b", quarter_display, text, count=1)
    amount_display = extract_rmb_million_display(current_chinese)
    if amount_display:
        text = re.sub(r"RMB\s+[\d.]+\s+million", f"RMB {amount_display} million", text)
    if any(token in current_chinese for token in ["平稳", "稳定"]):
        text = re.sub(
            r"which increased compared to the last quarter",
            "with collections remaining stable",
            text,
            flags=re.I,
        )
    return short_sentence(text, english_company_name(company_name), limit=220)


COMMON_ZH_EN_TERMS = [
    ("签署", "signed"),
    ("签约", "signed"),
    ("合作", "partnership"),
    ("交付", "delivery"),
    ("验收", "acceptance"),
    ("测试", "testing"),
    ("试飞", "test flight"),
    ("量产", "mass production"),
    ("融资", "financing"),
    ("募资", "fundraising"),
    ("尽调", "due diligence"),
    ("上市", "IPO preparation"),
    ("回款", "collections"),
    ("人员", "personnel"),
    ("团队", "team"),
    ("现金", "cash"),
    ("收入", "revenue"),
    ("毛利", "gross profit"),
    ("净利", "net profit"),
    ("客户", "customer"),
    ("渠道", "channel"),
    ("产品", "product"),
    ("认证", "certification"),
    ("合规", "compliance"),
    ("扩张", "expansion"),
    ("项目", "project"),
    ("合同", "contract"),
]


def polish_english_business_bullet(text: str, company_name: str) -> str:
    text = norm_space(text)
    text = re.sub(r"[\u4e00-\u9fff]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"([A-Za-z])（", r"\1 (", text)
    text = re.sub(r"）", ")", text)
    text = re.sub(r"，", ", ", text)
    text = re.sub(r"([a-z])([A-Z])", r"\1 \2", text)
    text = re.sub(r"(\d)([A-Za-z])", r"\1 \2", text)
    text = re.sub(r"([A-Za-z])(\d)", r"\1 \2", text)
    text = re.sub(r"\bQ\s+([1-4])\b", r"Q\1", text)
    text = re.sub(r"\bA\+\s+round of finance\b", "A+ financing round", text, flags=re.I)
    text = re.sub(r"\bTarget income\b", "Target revenue", text, flags=re.I)
    text = re.sub(r"\bpositive net gains\b", "positive net profit", text, flags=re.I)
    text = re.sub(r"\bfinancial and legal adjustments\b", "financial and legal due diligence", text, flags=re.I)
    text = re.sub(r"\bAt the negotiation stage\b", "The company is currently in the negotiation stage", text, flags=re.I)
    text = re.sub(r"\bbusiness changes\b", "industrial and commercial registration changes", text, flags=re.I)
    text = re.sub(
        r"\bcompleted and industrial and commercial registration changes\b",
        "has been completed, and the industrial and commercial registration change has been completed",
        text,
        flags=re.I,
    )
    text = re.sub(r"\bUnsanitary and unsanitary vehicle\b", "sanitation vehicle", text, flags=re.I)
    text = re.sub(r"\bThe new round is in hand\b", "The new round is being closed", text, flags=re.I)
    text = re.sub(r"\b(\d+)\s+clients completed sales\b", r"sales were completed with \1 customers", text, flags=re.I)
    text = re.sub(r"\$\s*([\d.]+(?:\s*-\s*[\d.]+)?)\s*million\b", r"RMB \1 million", text)
    text = re.sub(r"^\s*We are\b", "The company is", text, flags=re.I)
    text = re.sub(r"^\s*We're\b", "The company is", text, flags=re.I)
    text = re.sub(r"^\s*We have\b", "The company has", text, flags=re.I)
    text = re.sub(r"^\s*We had\b", "The company had", text, flags=re.I)
    text = re.sub(r"^\s*We expect\b", "The company expects", text, flags=re.I)
    text = re.sub(r"^\s*We target\b", "The company targets", text, flags=re.I)
    text = re.sub(r"^\s*We plan\b", "The company plans", text, flags=re.I)
    text = re.sub(r"^\s*We did not\b", "The company did not", text, flags=re.I)
    text = re.sub(r"^\s*We do not\b", "The company does not", text, flags=re.I)
    text = re.sub(r"^\s*We\b", "The company", text, flags=re.I)
    text = re.sub(r"^\s*Our\b", "The company's", text, flags=re.I)
    text = re.sub(r"\bwe are\b", "the company is", text, flags=re.I)
    text = re.sub(r"\bwe have\b", "the company has", text, flags=re.I)
    text = re.sub(r"\bwe had\b", "the company had", text, flags=re.I)
    text = re.sub(r"\bwe expect\b", "the company expects", text, flags=re.I)
    text = re.sub(r"\bwe target\b", "the company targets", text, flags=re.I)
    text = re.sub(r"\bwe plan\b", "the company plans", text, flags=re.I)
    text = re.sub(r"\bwe did not\b", "the company did not", text, flags=re.I)
    text = re.sub(r"\bwe do not\b", "the company does not", text, flags=re.I)
    text = re.sub(r"\bwe\b", "the company", text, flags=re.I)
    text = re.sub(r"\bour\b", "the company's", text, flags=re.I)
    text = re.sub(r"\bours\b", "the company's", text, flags=re.I)
    text = re.sub(r"\ballow us to\b", "allow the company to", text, flags=re.I)
    text = re.sub(r"\ballows us to\b", "allows the company to", text, flags=re.I)
    text = re.sub(r"\bfor us\b", "for the company", text, flags=re.I)
    text = text.replace("  ", " ")
    text = text.replace(" ,", ",")
    text = text.replace(" .", ".")
    text = re.sub(r"\b([A-Za-z]+)\s+\1\b", r"\1", text)
    text = text.strip()
    if text and not text.endswith("."):
        text += "."
    return text


def contains_chinese_chars(text: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", text))


def is_clean_english_bullet(text: str) -> bool:
    if contains_chinese_chars(text):
        return False
    if re.search(r"[，。；、“”]", text):
        return False
    if re.search(r"(?:,\s*){3,}", text):
        return False
    words = re.findall(r"[A-Za-z]{2,}", text)
    if len(words) < 6:
        return False
    unique_ratio = len(set(w.lower() for w in words)) / max(len(words), 1)
    if unique_ratio < 0.45:
        return False
    return True


def extract_arabic_numbers(text: str) -> list[str]:
    return re.findall(r"\d+(?:\.\d+)?", text)


def validate_model_translation(source_text: str, translated_text: str) -> bool:
    if not translated_text or contains_chinese_chars(translated_text):
        return False
    if len(re.findall(r"[A-Za-z]{2,}", translated_text)) < 5:
        return False
    source_numbers = extract_arabic_numbers(source_text)
    if source_numbers and not extract_arabic_numbers(translated_text):
        return False
    return True


def is_usable_english_translation(text: str) -> bool:
    if not text:
        return False
    english_words = re.findall(r"[A-Za-z]{2,}", text)
    if len(english_words) < 4:
        return False
    if contains_chinese_chars(text):
        stripped = re.sub(r"[\u4e00-\u9fff]+", " ", text)
        stripped_words = re.findall(r"[A-Za-z]{2,}", stripped)
        return len(stripped_words) >= 6
    return True


def ensure_argos_ready() -> tuple[bool, str]:
    if not USE_ARGOS_TRANSLATION:
        return False, "Argos translation disabled by PR_USE_ARGOS_TRANSLATION."
    if ARGOS_TRANSLATION_STATUS["checked"]:
        return ARGOS_TRANSLATION_STATUS["available"], ARGOS_TRANSLATION_STATUS["reason"]

    # Keep Argos runtime/cache/config under workspace to avoid ~/.local permission issues.
    data_home = ARGOS_RUNTIME_ROOT / "data"
    config_home = ARGOS_RUNTIME_ROOT / "config"
    cache_home = ARGOS_RUNTIME_ROOT / "cache"
    data_home.mkdir(parents=True, exist_ok=True)
    config_home.mkdir(parents=True, exist_ok=True)
    cache_home.mkdir(parents=True, exist_ok=True)
    os.environ["XDG_DATA_HOME"] = str(data_home)
    os.environ["XDG_CONFIG_HOME"] = str(config_home)
    os.environ["XDG_CACHE_HOME"] = str(cache_home)

    ARGOS_TRANSLATION_STATUS["checked"] = True
    try:
        import argostranslate.package  # type: ignore
        import argostranslate.sbd  # type: ignore
        import argostranslate.settings  # type: ignore
        import argostranslate.translate  # type: ignore
    except Exception as exc:
        ARGOS_TRANSLATION_STATUS["available"] = False
        ARGOS_TRANSLATION_STATUS["reason"] = f"Argos dependencies not installed: {exc}"
        return False, ARGOS_TRANSLATION_STATUS["reason"]

    # Force Stanza to reuse the packaged local resources instead of trying to
    # refresh resources.json from GitHub on every translation run.
    if not getattr(argostranslate.sbd.StanzaSentencizer, "_codex_offline_patch", False):
        def _offline_lazy_pipeline(self):  # type: ignore[no-untyped-def]
            if self.stanza_pipeline is None:
                import stanza  # type: ignore
                from stanza.pipeline.core import DownloadMethod  # type: ignore

                stanza_dir = self.pkg.package_path / "stanza"
                pipeline_kwargs: dict[str, Any] = {
                    "lang": self.stanza_lang_code,
                    "dir": str(stanza_dir),
                    "processors": "tokenize",
                    "use_gpu": argostranslate.settings.device == "cuda",
                    "logging_level": "WARNING",
                    "download_method": DownloadMethod.REUSE_RESOURCES,
                }
                resources_path = stanza_dir / "resources.json"
                if resources_path.exists():
                    pipeline_kwargs["resources_filepath"] = str(resources_path)
                self.stanza_pipeline = stanza.Pipeline(**pipeline_kwargs)
            return self.stanza_pipeline

        argostranslate.sbd.StanzaSentencizer.lazy_pipeline = _offline_lazy_pipeline  # type: ignore[assignment]
        argostranslate.sbd.StanzaSentencizer._codex_offline_patch = True  # type: ignore[attr-defined]

    try:
        installed_languages = argostranslate.translate.get_installed_languages()
        from_lang = next((lang for lang in installed_languages if lang.code == "zh"), None)
        to_lang = next((lang for lang in installed_languages if lang.code == "en"), None)

        translation = None
        if from_lang is not None and to_lang is not None:
            try:
                translation = from_lang.get_translation(to_lang)
            except Exception:
                translation = None

        if from_lang is None or to_lang is None or translation is None:
            if not ARGOS_AUTO_INSTALL:
                raise RuntimeError(
                    "No local Argos zh->en package installed. "
                    "Set PR_ARGOS_AUTO_INSTALL=1 to auto-download, or preinstall Argos zh->en package."
                )
            with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
                argostranslate.package.update_package_index()
                available_packages = argostranslate.package.get_available_packages()
                package_to_install = next(
                    (
                        pkg
                        for pkg in available_packages
                        if pkg.from_code == "zh" and pkg.to_code == "en"
                    ),
                    None,
                )
                if package_to_install is None:
                    raise RuntimeError("No Argos zh->en package found in package index.")
                package_path = package_to_install.download()
                argostranslate.package.install_from_path(package_path)
            installed_languages = argostranslate.translate.get_installed_languages()
            from_lang = next((lang for lang in installed_languages if lang.code == "zh"), None)
            to_lang = next((lang for lang in installed_languages if lang.code == "en"), None)
            translation = None
            if from_lang is not None and to_lang is not None:
                try:
                    translation = from_lang.get_translation(to_lang)
                except Exception:
                    translation = None

        if from_lang is None or to_lang is None:
            raise RuntimeError("Argos zh/en languages not available after installation.")
        if translation is None:
            raise RuntimeError("Argos zh->en translation package is not installed.")
    except Exception as exc:
        ARGOS_TRANSLATION_STATUS["available"] = False
        ARGOS_TRANSLATION_STATUS["reason"] = f"Argos model load failed: {exc}"
        return False, ARGOS_TRANSLATION_STATUS["reason"]

    ARGOS_RESOURCES["translation"] = translation
    ARGOS_TRANSLATION_STATUS["available"] = True
    ARGOS_TRANSLATION_STATUS["reason"] = ""
    return True, ""


def translate_with_argos(source_text: str, company_name: str) -> tuple[str | None, str]:
    ready, reason = ensure_argos_ready()
    if not ready:
        return None, reason

    translation = ARGOS_RESOURCES.get("translation")
    if translation is None:
        return None, "Argos translation resource was not initialized."

    try:
        candidate = translation.translate(source_text)
    except Exception as exc:
        return None, f"Argos translation execution failed: {exc}"

    candidate = norm_space(candidate).strip(" \"'")
    if not candidate:
        return None, "Argos returned empty output."
    candidate = polish_english_business_bullet(candidate, company_name)
    if not is_usable_english_translation(candidate):
        return None, "Argos output was not usable English."
    return candidate, ""


def ollama_model_candidates() -> list[str]:
    candidates: list[str] = []
    for model in [OLLAMA_MODEL, *DEFAULT_OLLAMA_MODELS]:
        normalized = model.strip()
        if normalized and normalized not in candidates:
            candidates.append(normalized)
    return candidates


def selected_ollama_model() -> str:
    return str(OLLAMA_TRANSLATION_STATUS.get("model", "") or "")


def ensure_ollama_ready() -> tuple[bool, str]:
    if not USE_OLLAMA_TRANSLATION:
        return False, "Ollama translation disabled by PR_USE_OLLAMA_TRANSLATION."
    if OLLAMA_TRANSLATION_STATUS["checked"]:
        return OLLAMA_TRANSLATION_STATUS["available"], OLLAMA_TRANSLATION_STATUS["reason"]

    OLLAMA_TRANSLATION_STATUS["checked"] = True
    ollama_path = shutil.which("ollama")
    if not ollama_path:
        OLLAMA_TRANSLATION_STATUS["available"] = False
        OLLAMA_TRANSLATION_STATUS["reason"] = "Ollama CLI not found in PATH."
        OLLAMA_TRANSLATION_STATUS["model"] = ""
        return False, OLLAMA_TRANSLATION_STATUS["reason"]

    try:
        proc = subprocess.run(
            ["ollama", "list"],
            capture_output=True,
            text=True,
            timeout=20,
            check=False,
        )
    except Exception as exc:
        OLLAMA_TRANSLATION_STATUS["available"] = False
        OLLAMA_TRANSLATION_STATUS["reason"] = f"Ollama check failed: {exc}"
        OLLAMA_TRANSLATION_STATUS["model"] = ""
        return False, OLLAMA_TRANSLATION_STATUS["reason"]

    output = f"{proc.stdout}\n{proc.stderr}".strip().lower()
    if proc.returncode != 0:
        OLLAMA_TRANSLATION_STATUS["available"] = False
        OLLAMA_TRANSLATION_STATUS["reason"] = "Ollama is not running or unreachable."
        OLLAMA_TRANSLATION_STATUS["model"] = ""
        return False, OLLAMA_TRANSLATION_STATUS["reason"]

    available_models = {line.split()[0].strip().lower() for line in proc.stdout.splitlines() if line.strip()}
    selected_model = next((model for model in ollama_model_candidates() if model.lower() in available_models), "")
    if not selected_model:
        OLLAMA_TRANSLATION_STATUS["available"] = False
        candidates = ", ".join(ollama_model_candidates())
        OLLAMA_TRANSLATION_STATUS["reason"] = f"No preferred Ollama model is pulled. Tried: {candidates}."
        OLLAMA_TRANSLATION_STATUS["model"] = ""
        return False, OLLAMA_TRANSLATION_STATUS["reason"]

    OLLAMA_TRANSLATION_STATUS["available"] = True
    OLLAMA_TRANSLATION_STATUS["reason"] = ""
    OLLAMA_TRANSLATION_STATUS["model"] = selected_model
    return True, ""


def translate_with_ollama(source_text: str, company_name: str) -> tuple[str | None, str]:
    ready, reason = ensure_ollama_ready()
    if not ready:
        return None, reason
    model_name = selected_ollama_model()
    if not model_name:
        return None, "Ollama reported ready, but no selected model was recorded."

    company = english_company_name(company_name)
    prompt = (
        "Translate the following Chinese business update bullet into concise professional English for a VC portfolio review.\n"
        "Rules:\n"
        "1) Keep all numbers, percentages, dates, and currency values.\n"
        "2) Keep proper nouns and project names readable in English (transliterate if needed).\n"
        "3) Use third-person investor-report tone.\n"
        "4) Output exactly one sentence, no bullet symbol, no extra commentary.\n"
        f"Company: {company}\n"
        f"Chinese bullet: {source_text}\n"
        "English sentence:"
    )

    try:
        proc = subprocess.run(
            ["ollama", "run", model_name],
            input=prompt,
            capture_output=True,
            text=True,
            timeout=OLLAMA_TIMEOUT_SECONDS,
            check=False,
        )
    except Exception as exc:
        return None, f"Ollama translation execution failed: {exc}"

    if proc.returncode != 0:
        stderr = norm_space(proc.stderr) or "Ollama returned non-zero status."
        return None, stderr

    candidate = norm_space(proc.stdout)
    candidate = re.sub(r"^\s*[-•*]\s*", "", candidate)
    candidate = re.sub(r"^English sentence:\s*", "", candidate, flags=re.I)
    candidate = candidate.strip(" \"'")
    if not candidate:
        return None, "Ollama returned empty output."
    if not candidate.endswith("."):
        candidate += "."
    if not validate_model_translation(source_text, candidate):
        return None, "Ollama output failed translation quality checks."
    return candidate, ""


def infer_update_topic(text: str) -> str:
    if any(token in text for token in ["融资", "募资", "估值", "投资方", "基金", "债务", "贷款", "尽调"]):
        return "financing progress"
    if any(token in text for token in ["人员", "团队", "员工", "组织", "编制"]):
        return "team and organizational updates"
    if any(token in text for token in ["现金", "回款", "利润", "收入", "毛利", "净利", "成本"]):
        return "financial and cash flow updates"
    if any(token in text for token in ["认证", "合规", "监管", "审批"]):
        return "regulatory and certification updates"
    if any(token in text for token in ["产品", "渠道", "发布", "上线", "推广"]):
        return "product and channel updates"
    return "operational progress"


def extract_numeric_signals_for_translation(text: str, max_items: int = 3) -> list[str]:
    converted = convert_chinese_amounts_to_rmb_millions(text)
    signals: list[str] = []
    for match in re.findall(r"RMB\s+[\d.]+\s+million", converted):
        signals.append(match)
    for match in re.findall(r"\d+(?:\.\d+)?%", converted):
        signals.append(match)
    for match in re.findall(r"\d+(?:\.\d+)?\s*(?:台|人|家|项|辆|月|年|天)", text):
        signals.append(match)
    for match in re.findall(r"\d+(?:\.\d+)?\s*(?:万|亿)", text):
        signals.append(match)
    deduped: list[str] = []
    seen = set()
    for signal in signals:
        key = signal.strip().lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(signal.strip())
        if len(deduped) >= max_items:
            break
    return deduped


def extract_common_term_hits(text: str, max_items: int = 3) -> list[str]:
    hits: list[str] = []
    for zh, en in COMMON_ZH_EN_TERMS:
        if zh in text and en not in hits:
            hits.append(en)
        if len(hits) >= max_items:
            break
    return hits


def compose_universal_english_business_bullet(text: str, company_name: str) -> str:
    company = english_company_name(company_name)
    clean = strip_leading_list_marker(norm_space(text).rstrip("。"))
    quarter_display = extract_quarter_display_from_chinese(clean)
    topic = infer_update_topic(clean)
    terms = extract_common_term_hits(clean)
    metrics = extract_numeric_signals_for_translation(clean)

    sentence = f"{company} reported {topic}"
    if quarter_display:
        sentence += f" in {quarter_display}"
    if terms:
        sentence += f", covering {', '.join(terms)}"
    if metrics:
        sentence += f"; key figures include {', '.join(metrics)}"
    sentence += "."
    return polish_english_business_bullet(sentence, company_name)


def fallback_english_bullet_from_chinese(text: str, company_name: str) -> str:
    return compose_universal_english_business_bullet(text, company_name)


def translate_chinese_update_bullet_to_english(
    text: str,
    company_name: str,
    translation_memory: dict[str, str] | None = None,
    review_flags: list[dict[str, Any]] | None = None,
    use_model_translation: bool = False,
) -> str:
    text = strip_leading_list_marker(norm_space(text).rstrip("。"))
    if translation_memory and not use_model_translation:
        memory_hit = translation_memory_match(text, translation_memory)
        if memory_hit and should_use_translation_memory(text, memory_hit):
            return adapt_previous_english_translation(memory_hit, text, english_company_name(company_name))

    if use_model_translation:
        templated_project_amount = translate_project_amount_bullet_with_template(text, company_name)
        if templated_project_amount:
            return templated_project_amount
        model_translation, model_reason = translate_with_argos(text, company_name)
        if model_translation:
            return polish_english_business_bullet(model_translation, company_name)
        if review_flags is not None and model_reason and not ARGOS_TRANSLATION_STATUS.get("warned", False):
            review_flags.append(
                {
                    "id": "business-update-translation-model-unavailable",
                    "section": "Business Update",
                    "severity": "warning",
                    "message": f"Argos translation not used: {model_reason}",
                    "source": "system",
                }
            )
            ARGOS_TRANSLATION_STATUS["warned"] = True

    return compose_universal_english_business_bullet(text, company_name)


def build_english_business_update_from_chinese(
    paragraphs: list[str],
    company_name: str,
    max_bullets: int = 9,
    translation_memory: dict[str, str] | None = None,
    review_flags: list[dict[str, Any]] | None = None,
    use_model_translation: bool = False,
) -> list[str]:
    chinese_bullets = build_chinese_business_update_bullets(paragraphs)
    bullets: list[str] = []
    seen_keys: set[str] = set()
    for text in chinese_bullets:
        used_fallback = False
        bullet = translate_chinese_update_bullet_to_english(
            text,
            company_name,
            translation_memory=translation_memory,
            review_flags=review_flags,
            use_model_translation=use_model_translation,
        )
        if contains_chinese_chars(bullet):
            stripped = re.sub(r"[\u4e00-\u9fff]+", " ", bullet)
            stripped = re.sub(r"\s+", " ", stripped).strip(" ,.;:-")
            if len(re.findall(r"[A-Za-z]{2,}", stripped)) >= 8 and is_clean_english_bullet(stripped):
                bullet = stripped
            else:
                bullet = fallback_english_bullet_from_chinese(text, company_name)
                used_fallback = True
        bullet = polish_english_business_bullet(bullet, company_name)
        if contains_chinese_chars(bullet) or (not use_model_translation and not is_clean_english_bullet(bullet)):
            bullet = fallback_english_bullet_from_chinese(text, company_name)
            used_fallback = True
        if used_fallback and review_flags is not None:
            review_flags.append(
                {
                    "id": "business-update-translation-low-confidence",
                    "section": "Business Update",
                    "severity": "warning",
                    "message": "A Chinese update bullet required low-confidence fallback translation; review English wording.",
                    "source": "current_data_request",
                }
            )
        bullet = short_sentence(bullet, company_name, limit=600)
        key = re.sub(r"[^a-z0-9]", "", bullet.lower())
        if key in seen_keys:
            continue
        seen_keys.add(key)
        bullets.append(bullet)
        if len(bullets) >= max_bullets:
            break
    return bullets


def ensure_table_dimensions(table, required_rows: int, required_cols: int) -> bool:
    expanded = False
    if not table.rows:
        return False
    current_cols = len(table.rows[0].cells) if table.rows else 0
    grid_widths = get_table_grid_widths(table)
    last_width = grid_widths[-1] if grid_widths else 900

    while current_cols < required_cols:
        source_props = []
        for row in table.rows:
            source_cell = row.cells[-1]
            tc_pr = copy.deepcopy(source_cell._tc.tcPr) if source_cell._tc.tcPr is not None else None
            p_pr = None
            if source_cell.paragraphs:
                existing_ppr = source_cell.paragraphs[0]._element.find(qn("w:pPr"))
                if existing_ppr is not None:
                    p_pr = copy.deepcopy(existing_ppr)
            source_props.append((tc_pr, p_pr))
        table.add_column(Twips(last_width))
        if not source_props:
            current_cols += 1
            expanded = True
            continue
        for row_idx, row in enumerate(table.rows):
            new_cell = row.cells[-1]
            tc = new_cell._tc
            existing_tcpr = tc.tcPr
            if existing_tcpr is not None:
                tc.remove(existing_tcpr)
            source_tcpr, source_ppr = source_props[row_idx] if row_idx < len(source_props) else source_props[-1]
            if source_tcpr is not None:
                tc.insert(0, copy.deepcopy(source_tcpr))
            if new_cell.paragraphs:
                p = new_cell.paragraphs[0]
                existing_ppr = p._element.find(qn("w:pPr"))
                if existing_ppr is not None:
                    p._element.remove(existing_ppr)
                if source_ppr is not None:
                    p._element.insert(0, copy.deepcopy(source_ppr))
        current_cols += 1
        expanded = True

    while len(table.rows) < required_rows:
        source_row = table.rows[-1]
        source_props = []
        for source_cell in source_row.cells:
            tc_pr = copy.deepcopy(source_cell._tc.tcPr) if source_cell._tc.tcPr is not None else None
            p_pr = None
            if source_cell.paragraphs:
                existing_ppr = source_cell.paragraphs[0]._element.find(qn("w:pPr"))
                if existing_ppr is not None:
                    p_pr = copy.deepcopy(existing_ppr)
            source_props.append((tc_pr, p_pr))
        new_row = table.add_row()
        for cell_idx, new_cell in enumerate(new_row.cells[: len(source_props)]):
            tc = new_cell._tc
            existing_tcpr = tc.tcPr
            if existing_tcpr is not None:
                tc.remove(existing_tcpr)
            source_tcpr, source_ppr = source_props[cell_idx]
            if source_tcpr is not None:
                tc.insert(0, copy.deepcopy(source_tcpr))
            if new_cell.paragraphs:
                p = new_cell.paragraphs[0]
                existing_ppr = p._element.find(qn("w:pPr"))
                if existing_ppr is not None:
                    p._element.remove(existing_ppr)
                if source_ppr is not None:
                    p._element.insert(0, copy.deepcopy(source_ppr))
        expanded = True
    return expanded


def strip_response_prefix(text: str) -> str:
    """Strip Q&A response prefixes like 'CompanyName: ...' or 'Yes, CompanyName ...'."""
    # Strip "CompanyName: " or "CompanyName, " at start
    text = re.sub(r"^[A-Z][A-Za-z\s]{1,30}:\s*", "", text)
    # Strip "Yes, " or "Yes. " at start
    text = re.sub(r"^(?:Yes|No)[,.\s]+\s*", "", text, flags=re.I)
    return norm_space(text)


def should_merge_wrapped_update_lines(previous: str, current: str) -> bool:
    prev = norm_space(previous)
    curr = norm_space(current)
    if not prev or not curr:
        return False
    if is_current_business_update_heading(curr) or is_numbered_list_item_line(curr):
        return False
    # Wrapped PDF/DOCX lines often continue without terminal punctuation.
    if prev.endswith(("。", ".", "！", "!", "？", "?", "；", ";")):
        return False
    if prev.endswith(("，", ",", "：", ":")):
        return True
    if re.search(r"[A-Za-z]$", prev) and re.match(r"^[A-Za-z\u4e00-\u9fff]", curr):
        return True
    if re.search(r"[（(]$", prev) and re.match(r"^[A-Za-z0-9\u4e00-\u9fff]", curr):
        return True
    if re.search(r"(?:与|及|和|并|在|为|于|以及|多个|几家|若干|the|and|or|to|of|with|for|from|on|by|into)\s*$", prev, flags=re.I):
        return True
    return False


def clean_update_paragraphs(paragraphs: list[str]) -> list[str]:
    cleaned: list[str] = []
    for text in paragraphs:
        text = norm_space(text)
        if not text:
            continue
        if is_question_or_prompt(text):
            continue
        if is_standalone_label(text):
            continue
        text = strip_response_prefix(text)
        if not text:
            continue
        if cleaned and should_merge_wrapped_update_lines(cleaned[-1], text):
            cleaned[-1] = norm_space(f"{cleaned[-1]} {text}")
            continue
        if len(text) < 10:
            continue
        cleaned.append(text)
    return cleaned


def bullet_topic(text: str) -> str:
    lowered = norm_space(text).lower()
    if any(token in lowered for token in ["融资", "股改", "上市", "ipo", "equity", "fund", "loan", "facility", "capital", "cash flow", "现金流"]):
        return "finance"
    if any(token in lowered for token in ["项目", "交付", "签约", "回款", "合同", "project", "delivery", "contract", "customer", "operation", "运营", "milestone"]):
        return "operations"
    if any(token in lowered for token in ["计划", "扩张", "outlet", "growth", "expansion", "rollout", "trend", "progress", "channel", "product", "certification"]):
        return "growth"
    return "general"


def section_heading_preference(heading_text: str) -> list[str]:
    lowered = norm_space(heading_text).lower()
    if any(token in lowered for token in ["运营", "operation", "review", "业务回顾", "recent"]):
        return ["operations", "general", "growth", "finance"]
    if any(token in lowered for token in ["进展", "progress", "highlight", "trend", "业务进展", "business"]):
        return ["growth", "finance", "general", "operations"]
    return ["general", "operations", "growth", "finance"]


def allocate_middle_section_bullets(headings: list[str], bullets: list[str]) -> list[list[str]]:
    if not headings:
        return []
    if len(headings) == 1:
        return [bullets]

    remaining = [(bullet, bullet_topic(bullet)) for bullet in bullets]
    allocations: list[list[str]] = []
    for idx, heading in enumerate(headings):
        if idx == len(headings) - 1:
            allocations.append([bullet for bullet, _topic in remaining])
            break
        preferred = section_heading_preference(heading)
        chosen: list[str] = []
        still_remaining = []
        used_topics = set()
        for bullet, topic in remaining:
            if topic in preferred[:2] and topic not in used_topics:
                chosen.append(bullet)
                used_topics.add(topic)
            else:
                still_remaining.append((bullet, topic))
        if not chosen and remaining:
            chosen.append(remaining[0][0])
            still_remaining = remaining[1:]
        allocations.append(chosen)
        remaining = still_remaining
    while len(allocations) < len(headings):
        allocations.append([])
    return allocations


def normalize_grouped_bullets_for_template(grouped_bullets: list[list[str]], target_count: int) -> list[list[str]]:
    if target_count <= 0:
        return []
    normalized = [group for group in grouped_bullets if group]
    if not normalized:
        return [[] for _ in range(target_count)]
    if len(normalized) == target_count:
        return normalized
    if len(normalized) < target_count:
        return normalized + ([[]] * (target_count - len(normalized)))
    merged = normalized[: target_count - 1]
    tail: list[str] = []
    for group in normalized[target_count - 1 :]:
        tail.extend(group)
    merged.append(tail)
    return merged


def normalize_middle_chunks(chunk_list: list[list[str]], target_count: int) -> list[list[str]]:
    if target_count <= 0:
        return []
    normalized = chunk_list[:target_count] + ([[]] * max(0, target_count - len(chunk_list)))
    result: list[list[str]] = []
    for chunk in normalized:
        deduped: list[str] = []
        seen: set[str] = set()
        for bullet in chunk:
            cleaned = norm_space(bullet)
            key = normalized_middle_line_key(cleaned)
            if not key or key in seen:
                continue
            seen.add(key)
            deduped.append(cleaned)
        result.append(deduped)
    return result


def build_flat_middle_chunks_for_template(
    planned_occurrences: list[SectionOccurrence],
    grouped_bullets: list[list[str]] | None,
    flat_bullets: list[str] | None,
) -> list[list[str]]:
    if grouped_bullets:
        chunks = normalize_grouped_bullets_for_template(grouped_bullets, len(planned_occurrences))
    else:
        heading_texts = [planned.heading_text for planned in planned_occurrences]
        chunks = allocate_middle_section_bullets(heading_texts, flat_bullets or [])
    return normalize_middle_chunks(chunks, len(planned_occurrences))


def has_explicit_heading_number_prefix(text: str) -> bool:
    stripped = norm_space(text)
    if not stripped:
        return False
    if re.match(r"^\d+\s*[-–、.)]\s*\S", stripped):
        return True
    if re.match(r"^20\d{2}\s*Q[1-4]\b", stripped, flags=re.I):
        return True
    if re.match(r"^Q[1-4]\s*20\d{2}\b", stripped, flags=re.I):
        return True
    if re.match(r"^20\d{2}\s*年\s*Q?[1-4]\b", stripped):
        return True
    return False


def should_keep_middle_heading_flat_mode(planned: SectionOccurrence, position_in_language: int) -> bool:
    heading_text = norm_space(planned.heading_text)
    if position_in_language == 0:
        return True
    if canonical_heading(heading_text) == "Business Update":
        return True
    if has_explicit_heading_number_prefix(heading_text):
        return True
    return False


def build_previous_business_update_map(previous_occurrences: list[SectionOccurrence]) -> dict[str, list[str]]:
    previous_business_update_map: dict[str, list[str]] = {}
    for occurrence in previous_occurrences:
        if occurrence.canonical != "Business Update" or not occurrence.paragraphs:
            continue
        previous_business_update_map.setdefault(occurrence.language, [])
        previous_business_update_map[occurrence.language].extend(occurrence.paragraphs)
    return previous_business_update_map


def extract_middle_section_candidates(current_parsed: dict[str, Any]) -> MiddleSectionCandidates:
    raw_paragraphs = list(current_parsed.get("business_update_paragraphs") or [])
    source_groups = list(current_parsed.get("business_update_groups") or [])
    return MiddleSectionCandidates(
        raw_paragraphs=raw_paragraphs,
        source_groups=source_groups,
        chinese_lines=clean_update_paragraphs([p for p in raw_paragraphs if text_language(p) == "chinese"]),
        english_lines=clean_update_paragraphs([p for p in raw_paragraphs if text_language(p) == "english"]),
    )


def candidate_group_paragraphs(candidates: MiddleSectionCandidates) -> list[list[str]]:
    return [group["paragraphs"] for group in candidates.source_groups if group.get("paragraphs")]


def normalized_middle_line_key(text: str) -> str:
    key = strip_response_prefix(norm_space(text))
    key = re.sub(r"[。；;,.，:：!?？\s]+", "", key)
    return key.lower()


def is_admin_or_meta_chinese_line(text: str) -> bool:
    stripped = norm_space(text)
    if not stripped:
        return True
    if stripped in {"无", "暂无", "无重大变化", "无变化", "暂无变化", "不适用"}:
        return True
    admin_tokens = [
        "整理历史资料",
        "整理财务资料",
        "梳理关联方",
        "梳理业务与组织结构",
        "梳理业务",
        "梳理组织结构",
        "准备材料",
        "收集资料",
        "请提供",
        "请补充",
        "请说明",
        "见Q",
        "附录",
    ]
    if any(token in stripped for token in admin_tokens):
        return True
    return False


def chinese_middle_line_theme_key(text: str) -> str | None:
    stripped = norm_space(text)
    if not stripped:
        return None
    if "专利" in stripped:
        return "patent_status"
    if "产品情况" in stripped:
        return "product_summary"
    if "远红外超透镜" in stripped:
        return "product_far_ir"
    if "近红外超透镜" in stripped:
        return "product_near_ir"
    if "可见光超透镜" in stripped:
        return "product_visible"
    if "团队情况" in stripped or "总人数" in stripped or ("新增" in stripped and "离职" in stripped):
        return "team_update"
    if "人员总成本" in stripped:
        return "team_cost"
    if "融资情况" in stripped or "A+轮融资" in stripped or "工商变更" in stripped:
        return "financing_status"
    if "客户情况" in stripped or "客户已完成销售" in stripped:
        return "customer_status"
    if "到账" in stripped and "资金" in stripped:
        return "grant_income"
    if "入选" in stripped and not re.search(r"\d+(?:\.\d+)?\s*(万|亿|元)", stripped):
        return "award_or_pr"
    if any(token in stripped for token in ["融资额度", "估值", "投资方"]):
        return "fundraising_terms"
    if any(token in stripped for token in ["洽谈阶段", "财务尽调", "法务尽调", "尽调完成"]):
        return "fundraising_stage"
    if sum(token in stripped for token in ["森林消防", "智能巡检", "物流运输", "试飞", "巡检", "场景"]) >= 2:
        return "scenario_progress"
    if any(token in stripped for token in ["项目申报", "正式公示", "实验室项目", "国家实验室", "UUV"]):
        return "project_application"
    if any(token in stripped for token in ["水解制氢", "氢催化燃烧", "燃料电池", "电堆", "空冷", "液冷", "液氢"]):
        return "technical_product"
    return None


def summarize_selected_chinese_line(text: str, max_chars: int = 110) -> str:
    stripped = norm_space(text).rstrip("；;。")
    if len(stripped) <= max_chars:
        return stripped

    clauses = [norm_space(part).strip("，,") for part in re.split(r"[。；;]", stripped) if norm_space(part)]
    if len(clauses) <= 1:
        return stripped

    scored_clauses: list[tuple[float, int, str]] = []
    for idx, clause in enumerate(clauses):
        clause_score = score_useful_chinese_line(clause)
        if any(token in clause for token in ["正式公示", "项目申报", "国家实验室", "融资额度", "估值", "投资方", "试飞", "尽调"]):
            clause_score += 2.0
        if any(token in clause for token in ["目前开发了", "已开发", "已在", "计划在"]):
            clause_score += 1.5
        if any(token in clause for token in ["后续", "合作方式", "目标是", "知名度", "下一步"]):
            clause_score -= 3.0
        scored_clauses.append((clause_score, idx, clause))

    scored_clauses.sort(key=lambda item: (-item[0], item[1]))
    chosen = scored_clauses[:2]
    chosen.sort(key=lambda item: item[1])
    summary = "；".join(clause for _score, _idx, clause in chosen)
    if len(summary) > max_chars + 30:
        summary = "；".join(clause for _score, _idx, clause in chosen[:1])
    return summary.rstrip("；;。")


def is_company_amount_line_item(text: str) -> bool:
    stripped = norm_space(text)
    if not re.search(r"\d+(?:\.\d+)?\s*万", stripped):
        return False
    if not re.search(r"(有限公司|股份有限公司|科技有限公司|有限责任公司|公司)", stripped):
        return False
    project_context_tokens = [
        "铁路",
        "高铁",
        "隧道",
        "监测系统",
        "智能视频",
        "沿线",
        "技防监控",
        "围岩变形",
        "普铁",
    ]
    if any(token in stripped for token in project_context_tokens):
        return False
    strategic_tokens = [
        "签约",
        "签署",
        "交付",
        "导入",
        "合作",
        "联合开发",
        "量产",
        "测试",
        "验收",
        "试飞",
        "融资",
        "工商变更",
        "客户情况",
        "团队情况",
        "专利情况",
        "产品情况",
    ]
    if any(token in stripped for token in strategic_tokens):
        return False
    return True


def is_procurement_or_supplier_line(text: str) -> bool:
    stripped = norm_space(text)
    procurement_tokens = [
        "材料情况",
        "加工情况",
        "设备购买",
        "毛坯",
        "结构件",
        "模压小球",
        "毛坯镜片",
        "镀膜材料",
        "加工费用",
        "加工测试",
        "单点车尾款",
        "摇摆炉",
        "激光干涉仪",
        "精雕机",
    ]
    return any(token in stripped for token in procurement_tokens)


def is_generic_product_summary_line(text: str) -> bool:
    stripped = norm_space(text)
    if "产品情况" not in stripped:
        return False
    generic_tokens = ["本季度共", "项目", "新增项目", "未有新增", "暂无新增"]
    return sum(token in stripped for token in generic_tokens) >= 2


def has_concrete_product_detail_lines(paragraphs: list[str], exclude_text: str) -> bool:
    detail_tokens = [
        "超透镜",
        "量产",
        "交付",
        "导入",
        "联合开发",
        "测试验证",
        "小批量",
        "工艺开发",
        "样品",
        "客户",
    ]
    for paragraph in paragraphs:
        stripped = norm_space(paragraph)
        if not stripped or stripped == exclude_text:
            continue
        if "产品情况" in stripped:
            continue
        if sum(token in stripped for token in detail_tokens) >= 2:
            return True
    return False


def score_useful_chinese_line(text: str) -> float:
    stripped = norm_space(text)
    if not stripped:
        return -10.0
    score = 0.0
    technical_tokens = ["水解制氢", "氢催化燃烧", "燃料电池", "电堆", "空冷", "液冷", "液氢", "消氢", "储氢", "反应器", "控制策略"]
    technical_token_hits = sum(token in stripped for token in technical_tokens)
    if is_admin_or_meta_chinese_line(stripped):
        return -5.0
    if is_question_or_prompt(stripped):
        score -= 3.5
    if re.search(r"\d", stripped):
        score += 2.5
    if re.search(r"\d+(?:\.\d+)?\s*(?:万|亿|%|台|人|家|项|辆|月|年|天)", stripped):
        score += 3.0
    if any(token in stripped for token in ["签约", "签署", "交付", "回款", "到账", "收入", "量产", "开发", "合作", "采购", "报价", "交割"]):
        score += 2.5
    if any(token in stripped for token in ["不及预期", "低于预期", "弱于预期"]) and any(
        token in stripped for token in ["业务", "分润", "收入", "销售", "订单", "进展"]
    ):
        score += 4.0
    if any(token in stripped for token in ["融资", "估值", "投资方", "股改", "上市", "时间线", "进展", "计划"]):
        score += 2.0
    if (
        "融资" in stripped
        and any(token in stripped for token in ["计划", "进展", "情况"])
        and not re.search(r"\d", stripped)
        and not any(token in stripped for token in ["洽谈阶段", "财务尽调", "法务尽调", "投资方", "估值", "融资额度", "工商变更"])
    ):
        score -= 5.0
    if any(token in stripped for token in ["项目", "客户", "订单", "合同", "隧道", "监测系统", "NOA"]):
        score += 1.5
    if any(token in stripped for token in ["专利情况", "产品情况", "团队情况", "融资情况", "客户情况"]):
        score += 3.0
    if any(token in stripped for token in ["专利", "授权专利", "发明专利"]):
        score += 2.0
    if "已完成销售" in stripped:
        score += 1.5
    if any(token in stripped for token in ["洽谈阶段", "财务尽调", "法务尽调", "尽调完成"]):
        score += 3.0
    if any(token in stripped for token in ["项目申报", "正式公示", "国家实验室", "实验室项目", "UUV"]):
        score += 2.5
    if any(token in stripped for token in ["试飞", "巡检", "森林消防", "物流运输"]) or (
        "场景" in stripped and any(token in stripped for token in ["试飞", "巡检", "森林消防", "物流运输"])
    ):
        score += 1.5
    if len(stripped) <= 28 and "围绕" in stripped and "开展业务" in stripped:
        score += 1.5
    if re.search(r"[（(][^）)]*[，,][^）)]*[）)]", stripped) and not re.search(r"\d", stripped):
        score -= 3.0
    if is_company_amount_line_item(stripped):
        score -= 5.5
    if is_procurement_or_supplier_line(stripped):
        score -= 6.0
    if technical_token_hits >= 2 and not any(
        token in stripped for token in ["试飞", "场景", "融资", "估值", "投资方", "尽调", "项目申报", "正式公示", "国家实验室", "UUV"]
    ):
        score -= 4.0
    if not re.search(r"\d", stripped) and technical_token_hits >= 2:
        score -= 2.0
    if "技术方向" in stripped and "开发" in stripped and not any(
        token in stripped for token in ["验收", "交付", "签约", "销售", "公示", "试飞", "融资", "回款", "订单", "合同"]
    ):
        score -= 4.0
    if any(token in stripped for token in ["新拓展的一个方向", "进一步完善产品性能", "产品性能优势", "有利于下一步", "明确了方向"]):
        score -= 2.5
    if any(token in stripped for token in ["核名", "注册资料", "补充注册资料", "公司设立"]) and any(
        token in stripped for token in ["完成", "预计", "待完成", "进展"]
    ):
        score -= 4.5
    if "无重大人员变动" in stripped:
        score -= 4.0
    if any(token in stripped for token in ["工程化验证", "小试验证", "实施方案设计"]) and technical_token_hits >= 1:
        score -= 4.0
    if "完成" in stripped and "开发" in stripped and not re.search(r"\d", stripped) and any(
        token in stripped for token in ["无人机", "燃料电池", "产品"]
    ) and not any(token in stripped for token in ["试飞", "验收", "合同", "订单", "收入", "回款", "融资", "合作"]):
        score -= 3.0
    if "新拓展的一个方向" in stripped and "技术与产品开发" in stripped:
        score -= 4.0
    if any(token in stripped for token in ["基于过往", "市场调研与技术评估"]) and technical_token_hits >= 2 and not re.search(r"\d", stripped):
        score -= 2.5
    if any(token in stripped for token in ["采购上架", "采购", "上架"]) and any(
        token in stripped for token in ["以实现", "实现收入", "实现营收", "实现销售"]
    ):
        score -= 4.0
    if "1台套" in stripped and technical_token_hits >= 2:
        score -= 3.0
    if re.search(r"\d+\s*kw", stripped, re.IGNORECASE) and technical_token_hits >= 2:
        score -= 6.0
    if "技术与产品开发" in stripped and "应用场景" in stripped and not re.search(r"\d", stripped):
        score -= 4.0
    if "应用场景" in stripped and not any(token in stripped for token in ["森林消防", "智能巡检", "物流运输", "铁塔", "配送"]):
        score -= 2.0
    if "基于客户的需求" in stripped and technical_token_hits >= 2 and not re.search(r"\d", stripped):
        score -= 3.0
    if len(stripped) > 180 and not re.search(r"\d", stripped) and any(token in stripped for token in ["意向合作关系", "知名度", "推介公司产品"]):
        score -= 3.0
    if not re.search(r"\d", stripped):
        if any(token in stripped for token in ["现阶段", "继续找", "更多的"]):
            score -= 2.5
        if any(token in stripped for token in ["进展", "计划", "情况"]) and not any(
            token in stripped for token in ["签约", "交付", "交割", "回款", "合作", "采购", "报价", "收入"]
        ):
            score -= 2.0
    if len(stripped) < 6:
        score -= 2.0
    if len(stripped) > 120:
        score -= 1.0
    if stripped.endswith(("？", "?")):
        score -= 1.5
    # Prefer concrete progress over pure planning/timeline prompts.
    if re.search(r"^Q1\s*(?:重点目标|规划|计划)", stripped):
        score -= 6.0
    elif re.search(r"^Q[2-4](?:-Q[2-4])?\s*(?:规划|计划)", stripped):
        score -= 2.5
    if re.search(r"(?:业务规划|timeline|时间线)", stripped, flags=re.I):
        score -= 4.5
    if "转化率" in stripped and any(token in stripped for token in ["测试", "变现"]):
        score += 4.5
    if "Q2-Q4" in stripped and "转化率" in stripped:
        score += 5.5
    if any(token in stripped for token in ["将继续", "核心目标是", "计划于", "目标是", "规划"]) and not any(
        token in stripped
        for token in ["已完成", "完成了", "完成", "上线", "签约", "交付", "回款", "到账", "达成", "售出", "验收", "实现收入"]
    ):
        score -= 4.0
    if any(token in stripped for token in ["我们", "本季度我们"]) and any(
        token in stripped for token in ["重点目标", "规划", "计划", "目标"]
    ):
        score -= 3.0
    return score


def select_useful_chinese_groups(
    source_groups: list[dict[str, Any]],
    max_per_group: int = 7,
    max_total: int = 8,
) -> list[dict[str, Any]]:
    seen_candidate_keys: set[str] = set()
    # Keep selection stable by prioritizing earlier source groups first, then score within group.
    candidates_by_group: list[list[tuple[float, int, str, str, str | None]]] = []
    for group in source_groups:
        group_candidates: list[tuple[float, int, str, str, str | None]] = []
        group_paragraphs = group.get("paragraphs") or []
        for idx, text in enumerate(group_paragraphs):
            normalized = normalized_middle_line_key(text)
            if not normalized or normalized in seen_candidate_keys:
                continue
            seen_candidate_keys.add(normalized)
            if is_generic_product_summary_line(text) and has_concrete_product_detail_lines(group_paragraphs, norm_space(text)):
                continue
            score = score_useful_chinese_line(text)
            if score <= 2.0:
                continue
            group_candidates.append((score, idx, text, normalized, chinese_middle_line_theme_key(text)))
        group_candidates.sort(key=lambda item: (-item[0], item[1]))
        candidates_by_group.append(group_candidates)

    total_selected = 0
    seen_keys: set[str] = set()
    seen_themes: set[str] = set()
    per_group_counts: dict[int, int] = defaultdict(int)
    chosen_by_group: dict[int, list[tuple[int, str]]] = defaultdict(list)

    for group_idx, group_candidates in enumerate(candidates_by_group):
        if total_selected >= max_total:
            break
        for _score, idx, text, normalized, theme in group_candidates:
            if total_selected >= max_total:
                break
            if normalized in seen_keys:
                continue
            if theme and theme in seen_themes:
                continue
            if per_group_counts[group_idx] >= max_per_group:
                break
            summarized = summarize_selected_chinese_line(text)
            chosen_by_group[group_idx].append((idx, summarized))
            per_group_counts[group_idx] += 1
            total_selected += 1
            seen_keys.add(normalized)
            if theme:
                seen_themes.add(theme)

    selected_groups: list[dict[str, Any]] = []
    for group_idx, group in enumerate(source_groups):
        chosen = sorted(chosen_by_group.get(group_idx, []), key=lambda item: item[0])
        if chosen:
            selected_groups.append({"heading": group.get("heading"), "paragraphs": [text for _idx, text in chosen]})

    if selected_groups:
        return selected_groups

    # Safety fallback: keep the first non-admin unique line if scoring filtered everything out.
    for group in source_groups:
        for text in group.get("paragraphs") or []:
            normalized = normalized_middle_line_key(text)
            if not normalized or normalized in seen_keys or is_admin_or_meta_chinese_line(text):
                continue
            return [{"heading": group.get("heading"), "paragraphs": [text]}]
    return []


def is_admin_or_meta_english_line(text: str) -> bool:
    stripped = norm_space(text)
    if not stripped:
        return True
    if stripped.lower() in {"n/a", "none", "no material change", "equity financing"}:
        return True
    lowered = stripped.lower()
    words = re.findall(r"[A-Za-z]+", stripped)
    if words and len(words) <= 4 and not re.search(r"\d", stripped):
        lower_words = {word.lower() for word in words}
        common_verbs = {
            "is", "are", "was", "were", "be", "been", "being", "has", "have", "had",
            "will", "would", "can", "could", "should", "expected", "completed",
            "launched", "grew", "reached", "signed", "delivered", "focus", "focuses",
        }
        if lower_words.isdisjoint(common_verbs):
            return True
    if re.fullmatch(r"date as of \d{4}/\d{1,2}/\d{1,2}\.?", lowered):
        return True
    if lowered in {"appendix", "operation data", "financial statements as of 2025q4.", "financial statements as of 2025q4"}:
        return True
    if lowered.startswith("financial statements as of "):
        return True
    if lowered.startswith("operation data "):
        return True
    if lowered.startswith("please refer to q"):
        return True
    if lowered.startswith("date ") and re.search(r"\d{4}/\d{1,2}", lowered):
        return True
    if "see notes to investors" in lowered:
        cleaned = re.sub(r"(?:,\s*)?see notes to investors(?:\s+for [^.]+)?", "", stripped, flags=re.I).strip(" ,;:-")
        if len(cleaned) < 12:
            return True
    meta_tokens = [
        "please provide",
        "please elaborate",
        "please project",
        "please describe",
        "what is the strategic plan",
        "what is the company’s plan",
        "what is the company's plan",
        "what is the current timeline",
        "what is the latest",
        "what is the planned",
        "what is the company",
        "how many cars were delivered",
        "what is the current production capacity",
        "provide the sales pipeline",
        "brief intro of the business updates",
        "thanks to the relentless focus",
    ]
    return any(token in lowered for token in meta_tokens)


def strip_embedded_english_prompt(text: str) -> str:
    stripped = norm_space(text)
    if not stripped:
        return stripped
    if "?" not in stripped:
        return stripped
    prompt_tokens = [
        "please provide",
        "please elaborate",
        "please project",
        "please describe",
        "what is",
        "how many",
        "how was",
        "are there any",
        "if there is any",
        "provide the",
    ]
    head, tail = stripped.split("?", 1)
    if any(token in head.lower() for token in prompt_tokens):
        tail = norm_space(tail)
        if len(tail) >= 12:
            return tail
    return stripped


def english_middle_line_theme_key(text: str) -> str | None:
    stripped = norm_space(text).lower()
    if not stripped:
        return None
    if any(token in stripped for token in ["ebitda", "operating cash flow", "net burn"]):
        return "profitability_cashflow"
    if any(token in stripped for token in ["opening schedule", "outlets per month", "24 new outlets", "opening fees"]):
        return "outlet_schedule"
    if any(token in stripped for token in ["loan facility", "lender", "supplier financing", "hire purchase facility"]):
        return "debt_financing"
    if any(token in stripped for token in ["internal equity round", "equity round", "top-up"]):
        return "equity_financing"
    if any(token in stripped for token in ["halal certification", "sales uplift", "revenue growth potential"]):
        return "growth_initiative"
    return None


def score_useful_english_line(text: str) -> float:
    stripped = strip_embedded_english_prompt(text)
    if not stripped:
        return -10.0
    lowered = stripped.lower()
    if is_admin_or_meta_english_line(stripped):
        return -5.0

    score = 0.0
    if is_question_or_prompt(stripped):
        score -= 4.0
    if re.search(r"\$[\d,.]+(?:k|m|b)?|\brm[\d,.]+", stripped, re.I):
        score += 3.0
    if re.search(r"\b\d+(?:\.\d+)?%\b", stripped):
        score += 2.5
    if re.search(r"\b\d+(?:\.\d+)?%\s*(?:yoy|qoq|growth)\b", lowered):
        score += 2.5
    if re.search(r"\b(?:usd|rmb|rm)\s*[\d,.]+|\b[\d,.]+\s*(?:m|million|billion)\s*usd\b", lowered):
        score += 2.5
    if re.search(r"\b(?:q[1-4]\s*20\d{2}|20\d{2}q[1-4]|q[1-4])\b", lowered):
        score += 1.5
    if re.search(r"\b\d+\s+(?:outlets?|months?|locations?)\b", lowered):
        score += 2.0
    if re.search(r"\bheadcount\s+grew\b", lowered):
        score += 3.0
    if any(token in lowered for token in ["ebitda", "margin", "cash flow", "net burn", "gross profit", "revenue", "financing"]):
        score += 2.0
    if any(token in lowered for token in ["ipo", "audit", "going concern", "pre-ipo", "fundraising", "spa", "closing the deal"]):
        score += 2.5
    if any(token in lowered for token in ["capex", "factory", "final-assembly", "kd factory", "plant", "homologation", "production capacity"]):
        score += 2.0
    if any(token in lowered for token in ["achieved", "completed", "locked-in", "secured", "started", "expected", "positive", "milestone", "relaunch"]):
        score += 1.5
    if any(token in lowered for token in ["jv partner", "loan facility", "supplier financing", "hire purchase facility", "equity round", "top-up", "lender"]):
        score += 2.0
    if any(token in lowered for token in ["investor", "investors", "china mobile", "sfpim", "goodman", "hite"]):
        score += 2.0
    if any(token in lowered for token in ["outlet", "opening schedule", "sales uplift", "halal certification", "b2b", "event sales"]):
        score += 1.5
    if any(token in lowered for token in ["target to open", "at least 4 outlets per", "outlets per month starting in q2"]):
        score += 2.0
    if any(token in lowered for token in ["less than expected", "below expectations", "not as expected"]) and any(
        token in lowered for token in ["business", "performance", "progress", "revenue", "sales"]
    ):
        score += 3.5

    if any(token in lowered for token in ["cost cutting and optimization across all functions", "in-housing performance marketing and content marketing"]):
        score -= 3.0
    if any(token in lowered for token in ["our mission", "regional expansion", "500 outlets milestone"]) and not re.search(r"\$|\d+%", stripped):
        score -= 2.5
    if any(token in lowered for token in ["soft-launch", "planned roll-out in 2026", "growing interest from private investors"]) and not any(
        token in lowered for token in ["locked-in", "committed", "secured", "loan facility", "equity round"]
    ):
        score -= 1.5
    if any(token in lowered for token in ["in-house", "in-sourcing"]) and any(token in lowered for token in ["save", "uptime"]):
        score -= 4.0
    if any(token in lowered for token in ["other brands have seen", "estimated 6% revenue growth potential", "estimated 2% sales growth potential"]):
        score -= 4.0
    if any(token in lowered for token in ["secured 8 locations", "robust pipeline of interested investors"]):
        score -= 3.5
    if any(token in lowered for token in ["no plan for fundraising", "no fundraising at the moment", "no plan for fundraising at the moment"]):
        score -= 5.0
    if len(stripped) > 320:
        score -= 1.5
    elif len(stripped) > 220:
        score -= 0.5
    if len(stripped) < 20:
        score -= 2.0
    return score


def summarize_selected_english_line(text: str, max_chars: int = 260) -> str:
    stripped = strip_embedded_english_prompt(text)
    stripped = norm_space(stripped).rstrip(".")
    if len(stripped) <= max_chars:
        return stripped

    clauses = [norm_space(part).strip(" ,;:-") for part in re.split(r"(?<=[.;:])\s+|(?<=\.)\s+", stripped) if norm_space(part)]
    if len(clauses) <= 1:
        return stripped[:max_chars].rsplit(" ", 1)[0].rstrip(",;:-")

    scored_clauses: list[tuple[float, int, str]] = []
    for idx, clause in enumerate(clauses):
        clause_score = score_useful_english_line(clause)
        if any(token in clause.lower() for token in ["ebitda", "cash flow", "net burn", "opening schedule", "loan facility", "equity round", "top-up"]):
            clause_score += 2.0
        if any(token in clause.lower() for token in ["our mission", "foundation", "set the stage", "growing interest"]):
            clause_score -= 2.0
        scored_clauses.append((clause_score, idx, clause.rstrip(".;")))

    scored_clauses.sort(key=lambda item: (-item[0], item[1]))
    chosen = scored_clauses[:2]
    chosen.sort(key=lambda item: item[1])
    summary = ". ".join(clause for _score, _idx, clause in chosen)
    if len(summary) > max_chars + 30:
        summary = chosen[0][2]
    return summary.rstrip(".;")


def finalize_selected_english_line(text: str, company_name: str = "") -> str:
    text = strip_embedded_english_prompt(text)
    text = norm_space(text)
    text = re.sub(r"(?:,\s*)?see notes to investors(?:\s+for [^.]+)?", "", text, flags=re.I)
    text = re.sub(r"(?:,\s*)?see notes(?:\s+to investors)?(?:\s+for [^.]+)?", "", text, flags=re.I)
    text = re.sub(r"\bfocus is ([A-Za-z][A-Za-z\s&-]+),?\s*$", r"Focus is \1", text, flags=re.I)
    text = re.sub(r"\s+", " ", text).strip()
    text = text.strip(" ,;:-")
    if not text:
        return ""
    return polish_english_business_bullet(text, company_name)


def select_useful_english_groups(
    source_groups: list[dict[str, Any]],
    max_per_group: int = 6,
    max_total: int = 11,
) -> list[dict[str, Any]]:
    candidates: list[tuple[float, int, int, str, str, str | None]] = []
    seen_candidate_keys: set[str] = set()
    for group_idx, group in enumerate(source_groups):
        for idx, text in enumerate(group.get("paragraphs") or []):
            normalized = normalized_middle_line_key(text)
            if not normalized or normalized in seen_candidate_keys:
                continue
            seen_candidate_keys.add(normalized)
            score = score_useful_english_line(text)
            if score <= 2.0:
                continue
            candidates.append((score, group_idx, idx, text, normalized, english_middle_line_theme_key(text)))

    candidates.sort(key=lambda item: (-item[0], item[1], item[2]))

    seen_keys: set[str] = set()
    per_group_counts: dict[int, int] = defaultdict(int)
    chosen_by_group: dict[int, list[tuple[int, str]]] = defaultdict(list)

    for _score, group_idx, idx, text, normalized, theme in candidates:
        if len(seen_keys) >= max_total:
            break
        if normalized in seen_keys:
            continue
        if per_group_counts[group_idx] >= max_per_group:
            continue
        summarized = finalize_selected_english_line(summarize_selected_english_line(text))
        chosen_by_group[group_idx].append((idx, summarized))
        per_group_counts[group_idx] += 1
        seen_keys.add(normalized)

    selected_groups: list[dict[str, Any]] = []
    for group_idx, group in enumerate(source_groups):
        chosen = sorted(chosen_by_group.get(group_idx, []), key=lambda item: item[0])
        if chosen:
            selected_groups.append({"heading": group.get("heading"), "paragraphs": [text for _idx, text in chosen]})

    if selected_groups:
        return selected_groups

    for group in source_groups:
        for text in group.get("paragraphs") or []:
            normalized = normalized_middle_line_key(text)
            if not normalized or normalized in seen_keys or is_admin_or_meta_english_line(text):
                continue
            return [{"heading": group.get("heading"), "paragraphs": [finalize_selected_english_line(text)]}]
    return []


def build_middle_section_selection(source_language: str, candidates: MiddleSectionCandidates) -> MiddleSectionSelection:
    if source_language == "chinese":
        source_groups = candidates.source_groups or ([{"heading": None, "paragraphs": candidates.chinese_lines}] if candidates.chinese_lines else [])
        selected_groups = select_useful_chinese_groups(source_groups)
        chinese_lines = [paragraph for group in selected_groups for paragraph in group.get("paragraphs", [])]
        return MiddleSectionSelection(
            selected_groups=selected_groups,
            chinese_lines=chinese_lines,
            english_lines=[],
        )

    source_groups = candidates.source_groups or ([{"heading": None, "paragraphs": candidates.english_lines}] if candidates.english_lines else [])
    selected_groups = select_useful_english_groups(source_groups)
    english_lines = [paragraph for group in selected_groups for paragraph in group.get("paragraphs", [])]
    return MiddleSectionSelection(
        selected_groups=selected_groups,
        chinese_lines=[],
        english_lines=english_lines,
    )


def build_english_middle_section_output(
    source_language: str,
    company_name: str,
    template_languages: set[str],
    business_activity_map: dict[str, str],
    selection: MiddleSectionSelection,
    translation_memory: dict[str, str],
    review_flags: list[dict[str, Any]],
    use_model_translation_for_middle_english: bool,
) -> tuple[list[str], list[list[str]]]:
    if "english" not in template_languages:
        return [], []

    current_group_paragraphs = [group["paragraphs"] for group in selection.selected_groups if group.get("paragraphs")]
    if source_language == "chinese" and current_group_paragraphs:
        grouped = [
            build_english_business_update_from_chinese(
                group_paragraphs,
                company_name,
                translation_memory=translation_memory,
                review_flags=review_flags,
                use_model_translation=use_model_translation_for_middle_english,
            )
            for group_paragraphs in current_group_paragraphs
        ]
        return [bullet for group in grouped for bullet in group], grouped

    if source_language == "english" and current_group_paragraphs:
        grouped = [
            [
                finalize_selected_english_line(paragraph, company_name)
                for paragraph in group_paragraphs
                if finalize_selected_english_line(paragraph, company_name)
            ]
            for group_paragraphs in current_group_paragraphs
        ]
        grouped = [group for group in grouped if group]
        return [bullet for group in grouped for bullet in group], grouped

    if source_language == "chinese" and selection.chinese_lines:
        flat = build_english_business_update_from_chinese(
            selection.chinese_lines,
            company_name,
            translation_memory=translation_memory,
            review_flags=review_flags,
            use_model_translation=use_model_translation_for_middle_english,
        )
        return flat, []

    if selection.english_lines:
        flat = [
            finalize_selected_english_line(line, company_name)
            for line in selection.english_lines
            if finalize_selected_english_line(line, company_name)
        ]
        return flat, []

    return [], []


def build_chinese_middle_section_output(
    template_languages: set[str],
    business_activity_map: dict[str, str],
    selection: MiddleSectionSelection,
) -> tuple[list[str], list[list[str]]]:
    if "chinese" not in template_languages:
        return [], []

    current_group_paragraphs = [group["paragraphs"] for group in selection.selected_groups if group.get("paragraphs")]
    if current_group_paragraphs:
        grouped = [build_chinese_business_update_bullets(group_paragraphs) for group_paragraphs in current_group_paragraphs]
        return [bullet for group in grouped for bullet in group], grouped

    flat = build_chinese_business_update_bullets(selection.chinese_lines)
    return flat, []


def middle_section_translation_status(use_model_translation_for_middle_english: bool) -> tuple[str, str]:
    if use_model_translation_for_middle_english:
        argos_ready, argos_reason = ensure_argos_ready()
        return ("argos:zh-en", "") if argos_ready else ("heuristic-fallback", argos_reason)
    return (
        "disabled-for-non-bilingual-template",
        "Translation model is disabled because this template does not require an English middle section.",
    )


def resolve_middle_section_outputs(
    source_language: str,
    company_name: str,
    template_languages: set[str],
    business_activity_map: dict[str, str],
    previous_business_update_map: dict[str, list[str]],
    candidates: MiddleSectionCandidates,
    review_flags: list[dict[str, Any]],
    use_model_translation_for_middle_english: bool,
) -> MiddleSectionOutputs:
    translation_memory = build_business_update_translation_memory(previous_business_update_map)
    selection = build_middle_section_selection(source_language, candidates)

    flat_by_language: dict[str, list[str]] = {}
    grouped_by_language: dict[str, list[list[str]]] = {}

    english_flat, english_grouped = build_english_middle_section_output(
        source_language,
        company_name,
        template_languages,
        business_activity_map,
        selection,
        translation_memory,
        review_flags,
        use_model_translation_for_middle_english,
    )
    if english_flat:
        flat_by_language["english"] = english_flat
    if english_grouped:
        grouped_by_language["english"] = english_grouped

    chinese_flat, chinese_grouped = build_chinese_middle_section_output(
        template_languages,
        business_activity_map,
        selection,
    )
    if chinese_flat:
        flat_by_language["chinese"] = chinese_flat
    if chinese_grouped:
        grouped_by_language["chinese"] = chinese_grouped

    if not flat_by_language.get("english") and previous_business_update_map.get("english"):
        flat_by_language["english"] = previous_business_update_map["english"]
    if not flat_by_language.get("chinese") and previous_business_update_map.get("chinese"):
        flat_by_language["chinese"] = previous_business_update_map["chinese"]

    translation_engine, translation_engine_note = middle_section_translation_status(use_model_translation_for_middle_english)
    return MiddleSectionOutputs(
        flat_by_language=flat_by_language,
        grouped_by_language=grouped_by_language,
        translation_engine=translation_engine,
        translation_engine_note=translation_engine_note,
    )


def build_chinese_business_update_bullets(paragraphs: list[str]) -> list[str]:
    lines = [strip_leading_list_marker(text) for text in clean_update_paragraphs(paragraphs)]
    bullets: list[str] = []
    for text in lines:
        if len(text) < 4:
            continue
        if text.endswith("：") or text.endswith(":"):
            continue
        bullets.append(text.rstrip("；;。，,") + "。")
    deduped: list[str] = []
    seen = set()
    for bullet in bullets:
        key = bullet.lower()
        if key not in seen:
            seen.add(key)
            deduped.append(bullet)
    return deduped[:8]


def previous_review_has_logo(path: Path) -> bool:
    with zipfile.ZipFile(path) as zf:
        return any(name.startswith("word/media/") for name in zf.namelist())


def normalize_company_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", " ", name).strip()
    parts = [part for part in cleaned.split() if part]
    if not parts:
        return "PortfolioCompany"
    return "".join(parts)


def quarter_filename_label(quarter: Quarter) -> str:
    return f"{quarter.year}Q{quarter.quarter}"


def clean_main_output_folder(final_docx_path: Path) -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    for path in OUTPUT_DIR.iterdir():
        if path == DEBUG_OUTPUT_DIR:
            continue
        if path.is_file() and path.suffix.lower() != ".docx":
            path.unlink()


def clean_main_output_folder_for_batch() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    for path in OUTPUT_DIR.iterdir():
        if path == DEBUG_OUTPUT_DIR:
            continue
        if path.is_file():
            path.unlink()


def list_real_docx(folder: Path) -> list[Path]:
    return sorted(
        path
        for path in folder.glob("*.docx")
        if not path.name.startswith("~$")
    )


def list_current_inputs(folder: Path) -> list[Path]:
    return sorted(
        path
        for path in folder.iterdir()
        if path.is_file()
        and path.suffix.lower() in {".docx", ".pdf"}
        and not path.name.startswith("~$")
        and not path.name.startswith("_converted_")
    )


def previous_financial_lookup(parsed_prev: dict[str, Any]) -> dict[str, dict[str, str]]:
    return {row["label"]: row["values"] for row in parsed_prev["rows"]}


def latest_quarter_from_current(finance_source: dict[str, Any]) -> Quarter:
    for header in finance_source["header"]:
        quarter = parse_quarter_label(header)
        if quarter:
            return quarter
    raise ValueError("Could not identify current quarter from data request headers.")


def normalize_metric_label(text: str) -> str:
    text = text.lower()
    text = re.sub(r"\(=.*?\)", "", text)
    text = re.sub(r"\([^)]*\)", " ", text)
    text = text.replace("&", " and ")
    text = re.sub(r"[^a-z0-9\u4e00-\u9fff]+", " ", text)
    return norm_space(text)


def metric_tokens(text: str) -> set[str]:
    stopwords = {
        "of",
        "and",
        "the",
        "per",
        "company",
        "owned",
        "number",
        "accumulated",
    }
    normalized = normalize_metric_label(text)
    # For Chinese text, tokenize at the character level so partial matches work
    # e.g. "收入" tokens = {"收", "入", "收入"} which overlaps with "营业收入"
    chinese_chars = re.findall(r"[\u4e00-\u9fff]", normalized)
    if chinese_chars:
        tokens = set(chinese_chars)
        # Also add the full Chinese string and any contiguous Chinese substrings of length 2+
        chinese_only = "".join(chinese_chars)
        tokens.add(chinese_only)
        for length in range(2, len(chinese_chars)):
            for start in range(len(chinese_chars) - length + 1):
                tokens.add(chinese_only[start:start + length])
        # Also include any English tokens
        english_parts = re.sub(r"[\u4e00-\u9fff]+", " ", normalized)
        tokens |= {t for t in english_parts.split() if t not in stopwords}
        return tokens
    return {token for token in normalized.split() if token not in stopwords}


def find_current_quarter_total_header(finance_source: dict[str, Any], target_quarter: Quarter) -> str | None:
    for header in finance_source["header"]:
        if parse_quarter_label(header) == target_quarter:
            return header
    exact_subtotal = []
    yearly_total = []
    for header in finance_source["header"]:
        lowered = header.lower()
        if any(token in lowered for token in ["ytd", "accum", "exp", "qoq", "yoy", "预计", "预估"]):
            continue
        if parse_quarter_label(header) == target_quarter:
            exact_subtotal.append(header)
        # Match quarter number without year (e.g., "Q4", "第四季度")
        q_num = extract_quarter_number(header)
        if q_num == target_quarter.quarter and parse_quarter_label(header) is None:
            if not any(skip in lowered for skip in ["exp", "预计"]):
                exact_subtotal.append(header)
        # Match "总计/subtotal" with target quarter identifier
        if "总计" in header or "小计" in header or "subtotal" in lowered:
            parsed_q = parse_quarter_label(header)
            if parsed_q == target_quarter:
                exact_subtotal.append(header)
            elif parsed_q is None:
                q_in_header = extract_quarter_number(header)
                if q_in_header == target_quarter.quarter:
                    exact_subtotal.append(header)
                elif q_in_header is None:
                    # Generic yearly total (e.g. "2025总计") — use as last resort
                    if str(target_quarter.year) in header:
                        yearly_total.append(header)
    if exact_subtotal:
        return exact_subtotal[0]
    # Fall back to yearly total if no quarter-specific header found
    # (the actual Q4 value will come from month summation, but this header
    # is needed for choose_finance_row_label to filter rows that have data)
    if yearly_total:
        return yearly_total[0]
    # Last resort: any non-metadata column with numeric data
    skip_patterns = {"qoq", "yoy", "exp", "预计", "预估"}
    for header in finance_source["header"]:
        lowered = header.lower()
        if any(s in lowered for s in skip_patterns):
            continue
        if parse_month_header(header) is not None:
            continue
        # Check if at least one row has a parseable value
        for row_data in finance_source["rows"].values():
            if finance_parse_decimal(row_data.get(header, "")) is not None:
                return header
    return None


def find_header_value(values: dict[str, str], patterns: list[str]) -> str | None:
    for header, value in values.items():
        lowered = header.lower()
        if any(pattern in lowered for pattern in patterns):
            return value
    return None


def find_month_headers_for_quarter(finance_source: dict[str, Any], target_quarter: Quarter) -> list[str]:
    headers = []
    for header in finance_source["header"]:
        month_info = parse_month_header(header)
        if not month_info:
            continue
        year, month = month_info
        quarter = ((month - 1) // 3) + 1
        if year == target_quarter.year and quarter == target_quarter.quarter:
            headers.append(header)
    return headers


def sum_month_values(values: dict[str, str], month_headers: list[str]) -> Decimal | None:
    if not month_headers:
        return None
    month_values = [finance_parse_decimal(values.get(header, "")) for header in month_headers]
    if any(value is None for value in month_values):
        return None
    return sum(month_values)  # type: ignore[arg-type]


def find_fy_or_ytd_header(finance_source: dict[str, Any], target_quarter: Quarter) -> str | None:
    candidates = []
    year_text = str(target_quarter.year)
    for header in finance_source["header"]:
        lowered = header.lower()
        if year_text not in header:
            continue
        if any(token in lowered for token in ["exp", "qoq", "yoy"]):
            continue
        if parse_quarter_label(header) is not None:
            continue
        if parse_month_header(header) is not None:
            continue
        if any(token in lowered for token in ["ytd", "accum"]) or "总计" in header or "全年" in header:
            candidates.append(header)
    for header in candidates:
        if "总计" in header or "全年" in header:
            return header
    return candidates[0] if candidates else None


def normalize_previous_quarter_headers(raw_headers: list[str], previous_review_quarter: Quarter | None) -> list[tuple[str, str]]:
    quarter_headers = [header for header in raw_headers if parse_quarter_label(header)]
    if previous_review_quarter is None:
        return [(header, parse_quarter_label(header).display()) for header in quarter_headers if parse_quarter_label(header)]

    expected = []
    cursor = previous_review_quarter
    for _ in quarter_headers:
        expected.append(cursor)
        cursor = cursor.previous()

    normalized: list[tuple[str, str]] = []
    used_expected: set[Quarter] = set()
    for idx, header in enumerate(quarter_headers):
        parsed = parse_quarter_label(header)
        if parsed in expected and parsed not in used_expected:
            normalized.append((header, parsed.display()))
            used_expected.add(parsed)
            continue

        expected_q = expected[idx]
        parsed_qnum = extract_quarter_number(header)
        if expected_q not in used_expected and (parsed is None or parsed_qnum == expected_q.quarter or parsed not in expected):
            normalized.append((header, expected_q.display()))
            used_expected.add(expected_q)
            continue

        if parsed is not None and parsed not in used_expected:
            normalized.append((header, parsed.display()))
            used_expected.add(parsed)
            continue

        normalized.append((header, header))

    return normalized


def previous_row_value_for_normalized_header(
    prev_row: dict[str, Any],
    normalized_previous_headers: list[tuple[str, str]],
    target_header: str,
) -> str | None:
    compact_target = norm_space(target_header).replace(" ", "")
    for raw, normalized in normalized_previous_headers:
        compact_normalized = norm_space(normalized).replace(" ", "")
        compact_raw = norm_space(raw).replace(" ", "")
        if normalized == target_header or compact_normalized == compact_target or compact_raw == compact_target:
            return prev_row["values"].get(raw)
    for raw, value in prev_row["values"].items():
        if norm_space(raw).replace(" ", "") == compact_target:
            return value
    return None


def detect_unit_spec(text: str) -> UnitSpec | None:
    normalized = norm_space(text).lower()
    compact = normalized.replace(" ", "")
    # IMPORTANT: More specific (longer) patterns MUST come before shorter ones
    # e.g. "百万元人民币" before "万元人民币" since the latter is a substring of the former
    candidates: list[tuple[str, Decimal, str | None]] = [
        ("usdk", Decimal("1000"), "USD"),
        ("kusd", Decimal("1000"), "USD"),
        ("k usd", Decimal("1000"), "USD"),
        ("in thousand usd", Decimal("1000"), "USD"),
        ("amount in thousand usd", Decimal("1000"), "USD"),
        ("amount in k of us$", Decimal("1000"), "USD"),
        ("amount in k of usd", Decimal("1000"), "USD"),
        ("k of us$", Decimal("1000"), "USD"),
        ("k of usd", Decimal("1000"), "USD"),
        ("amount in usd million", Decimal("1000000"), "USD"),
        ("in millions usd", Decimal("1000000"), "USD"),
        ("in usd million", Decimal("1000000"), "USD"),
        ("usd million", Decimal("1000000"), "USD"),
        ("millions usd", Decimal("1000000"), "USD"),
        ("million usd", Decimal("1000000"), "USD"),
        ("financial data (unit: usd)", Decimal("1"), "USD"),
        ("finance data (usd)", Decimal("1"), "USD"),
        ("cash balance (usd)", Decimal("1"), "USD"),
        ("unit: usd", Decimal("1"), "USD"),
        # Chinese: 百万 (million) patterns BEFORE 万 (ten-thousand) patterns
        ("单位：百万元人民币", Decimal("1000000"), "RMB"),
        ("百万元人民币", Decimal("1000000"), "RMB"),
        ("单位：百万人民币", Decimal("1000000"), "RMB"),
        ("百万人民币", Decimal("1000000"), "RMB"),
        ("百万元", Decimal("1000000"), "RMB"),
        ("百万", Decimal("1000000"), "RMB"),
        ("单位：万元人民币", Decimal("10000"), "RMB"),
        ("万元人民币", Decimal("10000"), "RMB"),
        ("万元 人民币", Decimal("10000"), "RMB"),
        ("万元", Decimal("10000"), "RMB"),
        ("单位：人民币", Decimal("1"), "RMB"),
        ("元人民币", Decimal("1"), "RMB"),
        ("人民币", Decimal("1"), "RMB"),
        ("unit: rmb in millions", Decimal("1000000"), "RMB"),
        ("rmb in millions", Decimal("1000000"), "RMB"),
        ("amount in rmb mn", Decimal("1000000"), "RMB"),
        ("rmb mn", Decimal("1000000"), "RMB"),
        ("单位：美元", Decimal("1"), "USD"),
        ("美元", Decimal("1"), "USD"),
    ]
    for needle, factor, currency in candidates:
        if needle in normalized or needle.replace(" ", "") in compact:
            return UnitSpec(label=text, factor_to_base=factor, currency=currency)
    if re.search(r"\(\s*usd\s*\)", normalized):
        return UnitSpec(label=text, factor_to_base=Decimal("1"), currency="USD")
    if re.search(r"\b(?:k|thousand)\s+of\s+us\$?\b", normalized):
        return UnitSpec(label=text, factor_to_base=Decimal("1000"), currency="USD")
    if "unit: usd" in normalized or "(unit: usd)" in normalized:
        return UnitSpec(label=text, factor_to_base=Decimal("1"), currency="USD")
    return None


def extract_exchange_rates(blocks: list[dict[str, Any]]) -> ExchangeRates:
    texts: list[str] = []
    for block in blocks:
        if block["type"] == "paragraph":
            texts.append(block["text"])
        elif block["type"] == "table":
            for row in block["rows"]:
                texts.extend(row)

    patterns = [
        r"1\s*(?:美元|usd|us\$)\s*[=＝:：]\s*([0-9]+(?:\.[0-9]+)?)\s*(?:元\s*人民币|人民币|rmb|cny)",
        r"(?:人民币|rmb|cny)\s*([0-9]+(?:\.[0-9]+)?)\s*(?:元)?\s*[=＝:：]\s*1\s*(?:美元|usd|us\$)",
    ]
    for text in texts:
        normalized = norm_space(text)
        for pattern in patterns:
            match = re.search(pattern, normalized, flags=re.I)
            if match:
                return ExchangeRates(rmb_per_usd=Decimal(match.group(1)))
    return ExchangeRates()


def detect_current_unit_spec(blocks: list[dict[str, Any]]) -> UnitSpec | None:
    for block in blocks[:6]:
        if block["type"] == "paragraph":
            spec = detect_explicit_unit_spec(block["text"])
            if spec:
                return spec
        elif block["type"] == "table":
            for row in block.get("rows", [])[:2]:
                for cell in row[:3]:
                    spec = detect_explicit_unit_spec(cell)
                    if spec:
                        return spec
    for block in blocks:
        if block["type"] == "paragraph":
            spec = detect_explicit_unit_spec(block["text"])
            if spec:
                return spec
        elif block["type"] == "table":
            for row in block.get("rows", [])[:2]:
                for cell in row[:3]:
                    spec = detect_explicit_unit_spec(cell)
                    if spec:
                        return spec
    return None


def previous_quarter_value_from_prev_financial(
    prev_financial: dict[str, Any],
    label_aliases: list[str],
    target_quarter: Quarter,
    previous_review_quarter: Quarter | None,
) -> Decimal | None:
    prev_lookup = previous_financial_lookup(prev_financial)
    prev_label = find_first_finance_row_by_aliases(prev_lookup, label_aliases)
    if not prev_label:
        return None
    normalized_headers = normalize_previous_quarter_headers(prev_financial["headers"], previous_review_quarter)
    target_header = target_quarter.previous().display()
    for raw_header, normalized_header in normalized_headers:
        if normalized_header == target_header:
            return finance_parse_decimal(prev_lookup[prev_label].get(raw_header))
    return None


def infer_effective_current_unit_spec(
    current_finance: dict[str, Any],
    detected_spec: UnitSpec | None,
    reference_prev_financial: dict[str, Any] | None,
    target_quarter: Quarter,
    previous_review_quarter: Quarter | None,
    exchange_rates: ExchangeRates | None = None,
) -> UnitSpec | None:
    if detected_spec is None or reference_prev_financial is None:
        return detected_spec
    if detected_spec.currency not in {"RMB", "USD"} or detected_spec.factor_to_base == Decimal("1"):
        return detected_spec

    output_unit_spec = detect_unit_spec(reference_prev_financial.get("unit", ""))
    if output_unit_spec is None:
        return detected_spec

    current_total_header = find_current_quarter_total_header(current_finance, target_quarter)
    current_month_headers = find_month_headers_for_quarter(current_finance, target_quarter)
    if not current_total_header and not current_month_headers:
        return detected_spec

    candidate_specs = [
        detected_spec,
        UnitSpec(label=detected_spec.label, factor_to_base=Decimal("1"), currency=detected_spec.currency),
    ]
    metric_aliases = [
        ["收入", "营业收入", "revenue"],
        ["毛利", "毛利润", "gross profit"],
        ["净利", "净利润", "net profit"],
    ]

    def candidate_penalty(candidate_spec: UnitSpec) -> float:
        total = 0.0
        used = 0
        for aliases in metric_aliases:
            current_label = find_first_finance_row_by_aliases(current_finance["rows"], aliases)
            if not current_label:
                continue
            current_row = current_finance["rows"][current_label]
            current_value = sum_month_values(current_row, current_month_headers)
            if current_value is None and current_total_header:
                current_value = finance_parse_decimal(current_row.get(current_total_header, ""))
            previous_value = previous_quarter_value_from_prev_financial(
                reference_prev_financial,
                aliases,
                target_quarter,
                previous_review_quarter,
            )
            if current_value is None or previous_value is None:
                continue
            converted = convert_value_between_units(current_value, candidate_spec, output_unit_spec, exchange_rates)
            if converted is None:
                continue
            current_abs = max(abs(float(converted)), 0.01)
            previous_abs = max(abs(float(previous_value)), 0.01)
            ratio = current_abs / previous_abs
            if ratio < 1.0:
                ratio = 1.0 / ratio
            total += ratio
            used += 1
        return total / used if used else float("inf")

    detected_penalty = candidate_penalty(candidate_specs[0])
    raw_penalty = candidate_penalty(candidate_specs[1])

    if raw_penalty < detected_penalty / 100:
        return candidate_specs[1]
    return detected_spec


def detect_financial_section_unit(occurrence: SectionOccurrence) -> tuple[str | None, str]:
    outside_unit = None
    for paragraph in occurrence.paragraphs:
        if is_unit_text(paragraph):
            outside_unit = paragraph
            break
    inside_unit = None
    if occurrence.tables and occurrence.tables[0] and occurrence.tables[0][0]:
        inside_unit = norm_space(occurrence.tables[0][0][0])
        if not is_unit_text(inside_unit):
            inside_unit = None
    if outside_unit:
        return outside_unit, "outside"
    if inside_unit:
        return inside_unit, "inside"
    return inside_unit or outside_unit, "none"


def convert_value_between_units(
    value: Decimal | None,
    source_unit: UnitSpec | None,
    output_unit: UnitSpec | None,
    exchange_rates: ExchangeRates | None = None,
) -> Decimal | None:
    if value is None:
        return None
    if source_unit is None or output_unit is None:
        return value
    source_amount = value * source_unit.factor_to_base
    source_currency = source_unit.currency
    output_currency = output_unit.currency

    if source_currency and output_currency and source_currency != output_currency:
        rmb_per_usd = exchange_rates.rmb_per_usd if exchange_rates else None
        if rmb_per_usd:
            if source_currency == "RMB" and output_currency == "USD":
                source_amount = source_amount / rmb_per_usd
            elif source_currency == "USD" and output_currency == "RMB":
                source_amount = source_amount * rmb_per_usd
            else:
                return value
        else:
            return value

    return source_amount / output_unit.factor_to_base


def determine_target_quarter(current_path: Path, current_title: str, current_finance: dict[str, Any]) -> Quarter:
    monthly_headers = [h for h in current_finance["header"] if parse_month_header(h)]
    if monthly_headers:
        last_header = monthly_headers[-1]
        month_info = parse_month_header(last_header)
        if month_info:
            year, month = month_info
            quarter = ((month - 1) // 3) + 1
            return Quarter(year, quarter)
    for source in [current_title, current_path.stem]:
        quarter = parse_quarter_label(source)
        if quarter:
            return quarter
    return latest_quarter_from_current(current_finance)


def determine_target_quarter_with_debug(current_path: Path, current_title: str, current_finance: dict[str, Any]) -> tuple[Quarter, str]:
    monthly_headers = [h for h in current_finance["header"] if parse_month_header(h)]
    if monthly_headers:
        last_header = monthly_headers[-1]
        month_info = parse_month_header(last_header)
        if month_info:
            year, month = month_info
            quarter = Quarter(year, ((month - 1) // 3) + 1)
            return quarter, f"month-based inference from headers {monthly_headers}"
    for source_name, source in [("title", current_title), ("filename", current_path.stem)]:
        quarter = parse_quarter_label(source)
        if quarter:
            return quarter, f"fallback quarter-label inference from {source_name}: {source}"
    quarter = latest_quarter_from_current(current_finance)
    return quarter, f"fallback latest_quarter_from_current using finance headers {current_finance['header']}"


# Cross-language metric aliases: map English labels to their Chinese equivalents and vice versa
METRIC_ALIASES: dict[str, list[str]] = {
    "revenue": ["收入", "收入金额", "营业收入", "主营业务收入", "销售收入", "net revenue"],
    "net revenue": ["revenue", "收入", "收入金额", "营业收入", "主营业务收入", "销售收入"],
    "gross profit": ["毛利", "毛利润", "毛利额"],
    "net profit": ["净利", "净利润", "净利额", "纯利", "net income"],
    "net income": ["净利", "净利润", "净收入", "net profit"],
    "ebitda": ["ebitda", "息税折旧摊销前利润"],
    "cash inflow": ["现金流入", "经营性现金流入"],
    "cash outflow": ["现金支出", "经营性现金支出"],
    "burn rate": ["烧钱率", "月均净现金消耗"],
    "收入": ["revenue", "net revenue", "收入金额"],
    "毛利": ["gross profit", "毛利润"],
    "净利": ["net profit", "net income", "净利润"],
}


def is_margin_metric_label(label: str) -> bool:
    normalized = normalize_metric_label(label)
    return normalized in {"毛利率", "gross margin", "gross profit margin", "net margin", "净利率"} or normalized.endswith("率")


def is_amount_metric_label(label: str) -> bool:
    normalized = normalize_metric_label(label)
    return normalized in {
        "收入",
        "revenue",
        "gross profit",
        "毛利",
        "net profit",
        "net income",
        "净利",
        "净利润",
    }


def find_first_finance_row_by_aliases(finance_rows: dict[str, dict[str, str]], aliases: list[str]) -> str | None:
    normalized_aliases = [normalize_metric_label(alias) for alias in aliases]
    for candidate in finance_rows:
        norm_cand = normalize_metric_label(candidate)
        if any(norm_cand == alias or alias in norm_cand for alias in normalized_aliases):
            return candidate
    return None


def choose_finance_row_label(
    previous_label: str,
    finance_rows: dict[str, dict[str, str]],
    current_total_header: str,
    previous_current_value: str | None,
    flags: list[dict[str, Any]],
    used_source_labels: dict[str, str] | None = None,
    source_unit_spec: "UnitSpec | None" = None,
    output_unit_spec: "UnitSpec | None" = None,
    exchange_rates: "ExchangeRates | None" = None,
) -> str | None:
    prev_tokens = metric_tokens(previous_label)
    used_source_labels = used_source_labels or {}

    exact_matches = []
    for candidate, values in finance_rows.items():
        if candidate in used_source_labels and used_source_labels[candidate] != previous_label:
            continue
        norm_cand = normalize_metric_label(candidate)
        norm_prev = normalize_metric_label(previous_label)
        if norm_cand == norm_prev:
            exact_matches.append(candidate)
    if len(exact_matches) == 1:
        return exact_matches[0]
    if len(exact_matches) > 1:
        flags.append(
            {
                "id": f"mapping-{normalize_metric_label(previous_label).replace(' ', '-')}",
                "section": "Financial Update",
                "severity": "warning",
                "message": f"Multiple exact metric matches found for '{previous_label}': {', '.join(exact_matches)}.",
                "source": "current_data_request",
            }
        )
        return exact_matches[0]

    # Try cross-language alias matching before fuzzy
    norm_prev = normalize_metric_label(previous_label)
    aliases = METRIC_ALIASES.get(norm_prev, [])
    alias_matches = []
    for candidate, values in finance_rows.items():
        if candidate in used_source_labels and used_source_labels[candidate] != previous_label:
            continue
        norm_cand = normalize_metric_label(candidate)
        if is_amount_metric_label(previous_label) and is_margin_metric_label(candidate):
            continue
        for alias in aliases:
            if norm_cand == normalize_metric_label(alias) or normalize_metric_label(alias) in norm_cand:
                alias_matches.append(candidate)
                break
    if len(alias_matches) == 1:
        return alias_matches[0]
    if len(alias_matches) > 1:
        return alias_matches[0]

    best_label: str | None = None
    best_score: tuple[int, Decimal] | None = None
    ambiguous = False

    # Also check for Chinese substring containment: if one label contains the other
    # (e.g. "收入" in "收入金额"), give bonus token score
    def containment_bonus(prev_label: str, candidate_label: str) -> int:
        prev_cn = "".join(re.findall(r"[\u4e00-\u9fff]", prev_label))
        cand_cn = "".join(re.findall(r"[\u4e00-\u9fff]", candidate_label))
        if prev_cn and cand_cn:
            if prev_cn in cand_cn or cand_cn in prev_cn:
                return 5
        return 0

    for candidate, values in finance_rows.items():
        if candidate in used_source_labels and used_source_labels[candidate] != previous_label:
            continue
        if is_amount_metric_label(previous_label) and is_margin_metric_label(candidate):
            continue
        candidate_value = finance_parse_decimal(values.get(current_total_header, ""))
        if candidate_value is None:
            continue
        cand_tokens = metric_tokens(candidate)
        raw_overlap = len(prev_tokens & cand_tokens)
        bonus = containment_bonus(previous_label, candidate)
        if bonus == 0 and raw_overlap < 2:
            continue
        token_score = raw_overlap + bonus

        previous_value = finance_parse_decimal(previous_current_value or "")
        continuity_penalty = Decimal("999999")
        if previous_value is not None:
            # Convert candidate value to output units using actual unit specs
            if source_unit_spec is not None and output_unit_spec is not None:
                candidate_converted = convert_value_between_units(candidate_value, source_unit_spec, output_unit_spec, exchange_rates) or candidate_value
            else:
                candidate_converted = candidate_value / Decimal("1000")
            continuity_penalty = abs(candidate_converted - previous_value)

        # Token overlap is PRIMARY (higher = better); continuity penalty is tiebreaker (lower = better)
        score = (token_score, -continuity_penalty)
        if best_score is None or score > best_score:
            best_label = candidate
            best_score = score
            ambiguous = False
        elif score == best_score:
            ambiguous = True

    if best_label is None or best_score is None:
        if any(
            normalize_metric_label(candidate) == normalize_metric_label(previous_label)
            for candidate in finance_rows
        ):
            flags.append(
                {
                    "id": f"mapping-{normalize_metric_label(previous_label).replace(' ', '-')}",
                    "section": "Financial Update",
                    "severity": "warning",
                    "message": f"Metric '{previous_label}' had a potential source-row collision and was left unmapped.",
                    "source": "current_data_request",
                }
            )
            return None
        flags.append(
            {
                "id": f"mapping-{normalize_metric_label(previous_label).replace(' ', '-')}",
                "section": "Financial Update",
                "severity": "warning",
                "message": f"Could not map previous-review metric '{previous_label}' to a current data request row.",
                "source": "current_data_request",
            }
        )
        return None

    if best_score[0] <= 0:
        flags.append(
            {
                "id": f"mapping-{normalize_metric_label(previous_label).replace(' ', '-')}",
                "section": "Financial Update",
                "severity": "warning",
                "message": f"No reliable metric-token match found for '{previous_label}'; leaving it unmapped instead of forcing a fuzzy row match.",
                "source": "current_data_request",
            }
        )
        return None

    if ambiguous:
        flags.append(
            {
                "id": f"mapping-{normalize_metric_label(previous_label).replace(' ', '-')}",
                "section": "Financial Update",
                "severity": "warning",
                "message": f"Metric mapping for '{previous_label}' is ambiguous; selected '{best_label}' by token overlap.",
                "source": "current_data_request",
            }
        )
    return best_label


def build_financial_update(
    prev_financial: dict[str, Any],
    current_finance: dict[str, Any],
    current_operation_rows: list[list[str]] | None,
    flags: list[dict[str, Any]],
    target_quarter: Quarter | None = None,
    source_unit_spec: UnitSpec | None = None,
    output_unit_spec: UnitSpec | None = None,
    previous_review_quarter: Quarter | None = None,
    exchange_rates: ExchangeRates | None = None,
) -> dict[str, Any]:
    previous_lookup = previous_financial_lookup(prev_financial)
    finance_rows = current_finance["rows"]
    target_quarter = target_quarter or latest_quarter_from_current(current_finance)

    normalized_previous_headers = normalize_previous_quarter_headers(prev_financial["headers"], previous_review_quarter)
    normalized_previous_quarter_labels = [normalized for _, normalized in normalized_previous_headers]

    rolling_headers = [target_quarter.display()]
    quarter_cursor = target_quarter.previous()
    while len(rolling_headers) < 5:
        rolling_headers.append(quarter_cursor.display())
        quarter_cursor = quarter_cursor.previous()

    final_total_header = f"{target_quarter.year} FY" if target_quarter.quarter == 4 else f"{target_quarter.year} YTD"
    estimate_header = f"{target_quarter.year + 1}E" if target_quarter.quarter == 4 else f"{target_quarter.year}E"
    output_headers = rolling_headers + ["QoQ", "YoY", final_total_header, estimate_header]
    current_total_header = find_current_quarter_total_header(current_finance, target_quarter)
    current_month_headers = find_month_headers_for_quarter(current_finance, target_quarter)
    fy_or_ytd_header = find_fy_or_ytd_header(current_finance, target_quarter)
    if current_total_header is None:
        current_total_header = ""
    operation_rows = current_operation_rows or [["", ""]]
    company_outlets, jv_outlets = extract_outlet_counts(operation_rows)
    company_outlets_total, jv_outlets_total = extract_outlet_counts(operation_rows)

    rows_output: list[dict[str, Any]] = []
    used_source_labels: dict[str, str] = {}
    revenue_row_label = find_first_finance_row_by_aliases(finance_rows, ["营业收入", "收入", "revenue", "sales revenue"])
    gross_margin_row_label = find_first_finance_row_by_aliases(finance_rows, ["毛利率", "gross margin", "gross profit margin"])
    for prev_row in prev_financial["rows"]:
        label = prev_row["label"]
        if label == "#of Company-Owned Outlets":
            source_info = {"type": "derived", "label": "company_outlets"}
        elif label == "#of JV Outlets":
            source_info = {"type": "derived", "label": "jv_outlets"}
        else:
            previous_current_value = None
            if previous_review_quarter is not None:
                for raw, normalized in normalized_previous_headers:
                    if normalized == previous_review_quarter.display():
                        previous_current_value = prev_row["values"].get(raw)
                        break
            mapped_label = choose_finance_row_label(
                label,
                finance_rows,
                current_total_header,
                previous_current_value,
                flags,
                used_source_labels,
                source_unit_spec,
                output_unit_spec,
                exchange_rates,
            )
            source_info = {"type": "finance_row", "label": mapped_label} if mapped_label else None
            if mapped_label:
                used_source_labels[mapped_label] = label
            elif normalize_metric_label(label) in {"毛利", "毛利润", "gross profit"} and revenue_row_label and gross_margin_row_label:
                source_info = {
                    "type": "derived_margin_amount",
                    "label": gross_margin_row_label,
                    "revenue_label": revenue_row_label,
                }
        row_flags: list[str] = []
        values: dict[str, str | None] = {}
        source_trace: dict[str, str] = {}
        backfilled_current_year_history = False

        source_row = finance_rows.get(source_info["label"]) if source_info and source_info["type"] == "finance_row" else None
        prev_quarters = {normalized: prev_row["values"].get(raw) for raw, normalized in normalized_previous_headers}

        current_quarter_value: str | None = None
        if source_info:
            if source_info["type"] == "finance_row" and source_row:
                decimal_value = sum_month_values(source_row, current_month_headers)
                if decimal_value is None and current_total_header:
                    raw_current = source_row.get(current_total_header)
                    decimal_value = finance_parse_decimal(raw_current)
                if decimal_value is None and finance_row_current_period_is_blank(
                    source_row,
                    current_month_headers,
                    current_total_header,
                    fy_or_ytd_header,
                ):
                    decimal_value = Decimal("0")
                if decimal_value is not None:
                    converted_value = convert_value_between_units(decimal_value, source_unit_spec, output_unit_spec, exchange_rates)
                    current_quarter_value = format_decimal(converted_value)
                    source_trace[target_quarter.display()] = "current_data_request"
                else:
                    current_quarter_value = default_missing_financial_value(label)
                    row_flags.append("Missing current-quarter value in the company data request.")
            elif source_info["type"] == "derived_margin_amount":
                margin_row = finance_rows.get(source_info["label"], {})
                revenue_row = finance_rows.get(source_info["revenue_label"], {})
                revenue_value = sum_month_values(revenue_row, current_month_headers)
                if revenue_value is None and current_total_header:
                    revenue_value = finance_parse_decimal(revenue_row.get(current_total_header, ""))
                margin_percent = parse_percent(margin_row.get(current_total_header, "")) if current_total_header else None
                if revenue_value is not None and margin_percent is not None:
                    decimal_value = revenue_value * (margin_percent / Decimal("100"))
                    converted_value = convert_value_between_units(decimal_value, source_unit_spec, output_unit_spec, exchange_rates)
                    current_quarter_value = format_decimal(converted_value)
                    source_trace[target_quarter.display()] = "derived_from_margin_and_revenue"
                    row_flags.append(f"Derived from {source_info['revenue_label']} and {source_info['label']}.")
                else:
                    current_quarter_value = default_missing_financial_value(label)
                    row_flags.append("Could not derive gross profit from revenue and gross margin.")
            elif source_info["type"] == "finance_row" and source_row is None:
                current_quarter_value = default_missing_financial_value(label)
                row_flags.append("No current data request row could be mapped to this metric.")
            elif source_info["label"] == "company_outlets":
                current_quarter_value = company_outlets
                if company_outlets is None:
                    current_quarter_value = default_missing_financial_value(label)
                    row_flags.append("Could not parse company-owned outlet count from operation data.")
                else:
                    source_trace[target_quarter.display()] = "current_data_request"
            elif source_info["label"] == "jv_outlets":
                current_quarter_value = jv_outlets
                if jv_outlets is None:
                    current_quarter_value = default_missing_financial_value(label)
                    row_flags.append("Could not parse JV outlet count from operation data.")
                else:
                    source_trace[target_quarter.display()] = "current_data_request"
        if current_quarter_value is None:
            current_quarter_value = default_missing_financial_value(label)
            if not any("Missing current-quarter value" in flag or "No current data request row could be mapped" in flag for flag in row_flags):
                row_flags.append("Missing current-quarter value in the company data request.")

        for header in rolling_headers:
            if header == target_quarter.display():
                values[header] = current_quarter_value
                continue
            values[header] = prev_quarters.get(header) or "0"
            if prev_quarters.get(header):
                source_trace[header] = "previous_review"
            elif parse_quarter_label(header) and parse_quarter_label(header).year == target_quarter.year:
                backfilled_current_year_history = True

        qoq_value = None
        yoy_value = None

        if source_info and source_info["type"] == "finance_row" and source_row:
            row_current_dec = finance_parse_decimal(values[target_quarter.display()] or "")
            if row_current_dec is not None and label in {"Revenue", "Gross profit", "EBITDA"}:
                row_current_dec = Decimal(values[target_quarter.display()])  # same unit in output

            previous_q_dec = finance_parse_decimal(values[rolling_headers[1]] or "")
            yoy_dec = finance_parse_decimal(values[rolling_headers[4]] or "")
            qoq_value = display_percent_text(values[target_quarter.display()], values[rolling_headers[1]] if len(rolling_headers) > 1 else None)
            yoy_value = display_percent_text(values[target_quarter.display()], values[rolling_headers[4]] if len(rolling_headers) > 4 else None)
            if qoq_value is None:
                row_flags.append("QoQ could not be determined from the rolled quarter values.")
            if yoy_value is None:
                row_flags.append("YoY could not be determined from the rolled quarter values.")

            current_fy_source = finance_parse_decimal(source_row.get(fy_or_ytd_header or "", ""))
            estimate_source = finance_parse_decimal(find_header_value(source_row, ["exp", "预计", "全年预计"]) or "")
            if current_fy_source is not None and not backfilled_current_year_history:
                values[final_total_header] = format_decimal(convert_value_between_units(current_fy_source, source_unit_spec, output_unit_spec, exchange_rates))
                source_trace[final_total_header] = "current_data_request"
            else:
                q_values = [finance_parse_decimal(values[h] or "") for h in rolling_headers if parse_quarter_label(h) and parse_quarter_label(h).year == target_quarter.year]
                if all(v is not None for v in q_values):
                    values[final_total_header] = format_decimal(sum(q_values))  # type: ignore[arg-type]
                    source_trace[final_total_header] = "derived_from_quarters"
                else:
                    values[final_total_header] = "0.00"
                    row_flags.append(f"{final_total_header} is missing and was defaulted to 0.00.")
            if estimate_source is not None:
                values[estimate_header] = format_decimal(convert_value_between_units(estimate_source, source_unit_spec, output_unit_spec, exchange_rates))
                source_trace[estimate_header] = "current_data_request"
            else:
                carried_estimate = previous_row_value_for_normalized_header(prev_row, normalized_previous_headers, estimate_header)
                if carried_estimate is not None and not is_empty_numeric_cell(carried_estimate):
                    values[estimate_header] = carried_estimate
                    source_trace[estimate_header] = "previous_review"
                    row_flags.append(f"{estimate_header} is missing in the current data request and was carried forward from the previous review.")
                else:
                    values[estimate_header] = default_missing_financial_value(label)
                    row_flags.append(f"{estimate_header} is missing in the current data request and was defaulted to {values[estimate_header]}.")

        else:
            row_current_dec = finance_parse_decimal(values[target_quarter.display()] or "")
            previous_q_dec = finance_parse_decimal(values[rolling_headers[1]] or "")
            yoy_dec = finance_parse_decimal(values[rolling_headers[4]] or "")
            qoq_value = display_percent_text(values[target_quarter.display()], values[rolling_headers[1]] if len(rolling_headers) > 1 else None)
            yoy_value = display_percent_text(values[target_quarter.display()], values[rolling_headers[4]] if len(rolling_headers) > 4 else None)
            if qoq_value is None:
                row_flags.append("QoQ could not be determined from the rolled quarter values.")
            if yoy_value is None:
                row_flags.append("YoY could not be determined from the rolled quarter values.")

            if source_info and source_info["label"] == "company_outlets":
                values[final_total_header] = company_outlets_total
                if company_outlets_total is not None:
                    source_trace[final_total_header] = "current_data_request"
                else:
                    values[final_total_header] = "0"
                    row_flags.append(f"{final_total_header} is missing for company-owned outlets and was defaulted to 0.")
            elif source_info and source_info["label"] == "jv_outlets":
                values[final_total_header] = jv_outlets_total
                if jv_outlets_total is not None:
                    source_trace[final_total_header] = "current_data_request"
                else:
                    values[final_total_header] = "0"
                    row_flags.append(f"{final_total_header} is missing for JV outlets and was defaulted to 0.")
            else:
                q_values = [
                    finance_parse_decimal(values[h] or "")
                    for h in rolling_headers
                    if parse_quarter_label(h) and parse_quarter_label(h).year == target_quarter.year
                ]
                if q_values and all(v is not None for v in q_values):
                    values[final_total_header] = format_decimal(sum(q_values))  # type: ignore[arg-type]
                    source_trace[final_total_header] = "derived_from_quarters"
                else:
                    values[final_total_header] = default_missing_financial_value(label)
                    row_flags.append(f"{final_total_header} is missing and was defaulted to {values[final_total_header]}.")
            carried_estimate = previous_row_value_for_normalized_header(prev_row, normalized_previous_headers, estimate_header)
            if carried_estimate is not None and not is_empty_numeric_cell(carried_estimate):
                values[estimate_header] = carried_estimate
                source_trace[estimate_header] = "previous_review"
                row_flags.append(f"{estimate_header} is missing in the current data request and was carried forward from the previous review.")
            else:
                values[estimate_header] = default_missing_financial_value(label)
                row_flags.append(f"{estimate_header} is missing in the current data request and was defaulted to {values[estimate_header]}.")

        values["QoQ"] = qoq_value
        values["YoY"] = yoy_value

        rows_output.append(
            {
                "label": label,
                "mapped_source_label": source_info["label"] if source_info else None,
                "values": {header: values.get(header) for header in output_headers},
                "source_trace": source_trace,
                "flags": row_flags,
            }
        )

    return {
        "status": "updated_with_flags",
        "unit": prev_financial["unit"] or "in millions USD",
        "columns": output_headers,
        "rows": rows_output,
    }


def build_markdown(
    business_activities: str,
    financial_update: dict[str, Any],
    business_update_bullets: list[str],
    risk_exit_bullets: list[str],
) -> str:
    lines = [
        "# Portfolio Review Draft",
        "",
        "## 1. Business Activities",
        business_activities,
        "",
        "## 2. Financial Update",
        f"({financial_update['unit']})",
        "",
    ]

    headers = [financial_update["unit"]] + financial_update["columns"]
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in financial_update["rows"]:
        row_values = [row["label"]] + [row["values"].get(header) or "" for header in financial_update["columns"]]
        lines.append("| " + " | ".join(row_values) + " |")

    lines.extend(["", "## 3. Business Update"])
    for bullet in business_update_bullets:
        lines.append(f"- {bullet}")

    lines.extend(["", "## 4. Risk & Exit"])
    for bullet in risk_exit_bullets:
        lines.append(f"- {bullet}")

    return "\n".join(lines) + "\n"


def split_evenly(items: list[str], count: int) -> list[list[str]]:
    if count <= 0:
        return []
    if not items:
        return [[] for _ in range(count)]
    base, remainder = divmod(len(items), count)
    chunks: list[list[str]] = []
    cursor = 0
    for idx in range(count):
        size = base + (1 if idx < remainder else 0)
        chunks.append(items[cursor : cursor + size])
        cursor += size
    return chunks


def detect_doc_section_occurrences(doc: Document) -> list[DocSectionOccurrence]:
    occurrences: list[DocSectionOccurrence] = []
    after_financial: dict[str, bool] = {"english": False, "chinese": False}
    after_risk: dict[str, bool] = {"english": False, "chinese": False}
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        language = text_language(text)
        canonical = fixed_anchor_heading(text)
        if canonical:
            occurrences.append(
                DocSectionOccurrence(
                    canonical=canonical,
                    language=language,
                    heading_index=idx,
                )
            )
            if canonical == "Financial Update":
                after_financial[language] = True
                after_risk[language] = False
            elif canonical == "Risk & Exit":
                after_risk[language] = True
            continue

        canonical = canonical_heading(text)
        if canonical == "Business Update":
            occurrences.append(
                DocSectionOccurrence(
                    canonical=canonical,
                    language=language,
                    heading_index=idx,
                )
            )
            continue

        if after_financial[language] and not after_risk[language] and is_middle_section_heading_candidate(paragraph):
            occurrences.append(
                DocSectionOccurrence(
                    canonical="Business Update",
                    language=language,
                    heading_index=idx,
                )
            )
    return occurrences


def resolve_section_plan_to_doc(doc: Document, section_plan: list[SectionOccurrence]) -> list[DocSectionOccurrence]:
    occurrences: list[DocSectionOccurrence] = []
    paragraph_index = 0
    for planned in section_plan:
        target = norm_space(planned.heading_text)
        for idx in range(paragraph_index, len(doc.paragraphs)):
            if norm_space(doc.paragraphs[idx].text) == target:
                occurrences.append(
                    DocSectionOccurrence(
                        canonical=planned.canonical,
                        language=planned.language,
                        heading_index=idx,
                    )
                )
                paragraph_index = idx + 1
                break
    return occurrences


def find_resolved_occurrence_index(
    doc: Document,
    occurrences: list[DocSectionOccurrence],
    planned: SectionOccurrence,
) -> int | None:
    target = norm_space(planned.heading_text)
    for idx, occurrence in enumerate(occurrences):
        if occurrence.canonical != planned.canonical or occurrence.language != planned.language:
            continue
        current = norm_space(doc.paragraphs[occurrence.heading_index].text)
        if current == target:
            return idx
    fallback_matches = [
        idx
        for idx, occurrence in enumerate(occurrences)
        if occurrence.canonical == planned.canonical and occurrence.language == planned.language
    ]
    if len(fallback_matches) == 1:
        return fallback_matches[0]
    return None


def template_language_mode(occurrences: list[DocSectionOccurrence]) -> str:
    languages = {occ.language for occ in occurrences}
    if "english" in languages and "chinese" in languages:
        return "bilingual"
    if "english" in languages:
        return "English-only"
    return "Chinese-only"


def middle_section_indices_by_language(occurrences: list[DocSectionOccurrence]) -> dict[str, list[int]]:
    result: dict[str, list[int]] = {"english": [], "chinese": []}
    for language in ["english", "chinese"]:
        language_occurrences = [(idx, occ) for idx, occ in enumerate(occurrences) if occ.language == language]
        financial_idx = next((idx for idx, occ in language_occurrences if occ.canonical == "Financial Update"), None)
        risk_idx = next((idx for idx, occ in language_occurrences if occ.canonical == "Risk & Exit" and financial_idx is not None and idx > financial_idx), None)
        if financial_idx is None:
            continue
        upper = risk_idx if risk_idx is not None else len(occurrences)
        for idx, occ in enumerate(occurrences):
            if occ.language != language:
                continue
            if idx <= financial_idx or idx >= upper:
                continue
            if occ.canonical in {"Business Activities", "Financial Update", "Risk & Exit"}:
                continue
            result[language].append(idx)
    return result


def normalize_middle_section_occurrence_indices(
    doc: Document,
    occurrences: list[DocSectionOccurrence],
    occurrence_indices: list[int],
) -> list[int]:
    normalized: list[int] = []
    for position, occ_idx in enumerate(occurrence_indices):
        occ = occurrences[occ_idx]
        body = get_body_paragraphs_for_occurrence(doc, occurrences, occ_idx)
        has_body_text = any(paragraph.text.strip() for paragraph in body)
        next_idx = occurrence_indices[position + 1] if position + 1 < len(occurrence_indices) else None
        if (
            not has_body_text
            and next_idx is not None
            and occurrences[next_idx].language == occ.language
            and occurrences[next_idx].canonical == "Business Update"
            and occurrences[next_idx].heading_index == occ.heading_index + 1
        ):
            continue
        normalized.append(occ_idx)
    return normalized


def get_body_paragraphs_for_occurrence(doc: Document, occurrences: list[DocSectionOccurrence], occurrence_index: int):
    start = occurrences[occurrence_index].heading_index + 1
    end = occurrences[occurrence_index + 1].heading_index if occurrence_index + 1 < len(occurrences) else len(doc.paragraphs)
    return doc.paragraphs[start:end]


def set_paragraph_text_preserve(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text_preserve(cell, text: str) -> None:
    """Set cell text while preserving paragraph and run formatting.
    Unlike cell.text = value, this doesn't destroy the cell's formatting."""
    # Remove ALL paragraphs except the first
    while len(cell.paragraphs) > 1:
        last_p = cell.paragraphs[-1]
        last_p._element.getparent().remove(last_p._element)
    p = cell.paragraphs[0]
    # Remove ALL child elements except pPr (paragraph properties)
    for child in list(p._element):
        if child.tag != qn("w:pPr"):
            p._element.remove(child)
    # Add a single clean run with the text
    run_elem = OxmlElement("w:r")
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")
    run_elem.append(t_elem)
    p._element.append(run_elem)
    # Force compact spacing by creating/replacing <w:spacing> in <w:pPr>
    pPr = p._element.get_or_add_pPr()
    old_spacing = pPr.find(qn("w:spacing"))
    if old_spacing is not None:
        pPr.remove(old_spacing)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def table_display_text(text: str, language: str) -> str:
    text = text.replace("\xa0", " ")
    numeric = text.strip()
    numeric_match = re.fullmatch(r"([+\-\u2212]?)(\d+(?:,\d{3})*)(\.\d+)?", numeric)
    if numeric_match:
        sign, integer_part, decimal_part = numeric_match.groups()
        integer_digits = integer_part.replace(",", "")
        grouped = f"{int(integer_digits):,}"
        text = f"{sign}{grouped}{decimal_part or ''}"
    if re.match(r"^-\d", text):
        text = "\u2212" + text[1:]
    if text.endswith("%"):
        text = text[:-1] + "\u2060%"
    # Keep structural financial headers on one line regardless of language.
    # (Bilingual Chinese pages often still use English quarter/header tokens.)
    if re.match(r"^Q[1-4] 20\d{2}$", text):
        return text.replace(" ", "\xa0")
    if re.match(r"^20\d{2} (?:FY|YTD)$", text):
        return text.replace(" ", "\xa0")
    if re.fullmatch(r"20\d{2}E", text):
        return text
    if text in {"QoQ", "YoY"}:
        return text
    if language != "english":
        return text
    normalized = text
    if is_unit_text(normalized) and " " in normalized:
        return normalized.replace(" ", "\xa0")
    if normalized.startswith("Q") and re.match(r"^Q[1-4] 20\d{2}$", normalized):
        return normalized.replace(" ", "\xa0")
    protected = {
        "Gross Profit",
        "Net Profit",
        "Cash Flow",
        "Business Update",
        "Risk & Exit",
        "Financial Update",
        "Financial Updates",
        "2025 FY",
        "2026E",
    }
    if normalized in protected:
        return normalized.replace(" ", "\xa0")
    if " " in normalized and re.fullmatch(r"[A-Za-z#()&/\- ]+", normalized):
        return normalized.replace(" ", "\xa0")
    return normalized


def set_cell_no_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is None:
        no_wrap = OxmlElement("w:noWrap")
        tc_pr.append(no_wrap)


def clear_cell_no_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is not None:
        tc_pr.remove(no_wrap)


def get_table_layout_type(table) -> str | None:
    tbl_layout = table._tbl.tblPr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        return None
    return tbl_layout.get(qn("w:type"))


def set_table_layout_fixed(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    try:
        table.autofit = False
    except Exception:
        pass


def looks_like_financial_table(table) -> bool:
    if not table.rows:
        return False
    header_tokens = [norm_space(cell.text).replace("\xa0", " ") for cell in table.rows[0].cells]
    joined = " ".join(header_tokens).lower()
    if "qoq" in joined or "yoy" in joined:
        return True
    if re.search(r"q[1-4]\s*20\d{2}", joined):
        return True
    if re.search(r"20\d{2}\s*(fy|ytd|e)\b", joined):
        return True
    return False


def get_table_grid_widths(table) -> list[int]:
    grid = table._tbl.tblGrid
    if grid is None:
        return []
    widths: list[int] = []
    for col in grid.iterchildren():
        width = col.get(qn("w:w"))
        widths.append(int(width) if width else 0)
    return widths


def estimate_table_text_width(text: str, language: str) -> int:
    total = 180
    for char in text:
        if char == "\u2060":
            continue
        if char in {" ", "\xa0"}:
            total += 35
        elif char.isdigit():
            total += 80
        elif char in {".", ",", ":", ";", "!", "|"}:
            total += 40
        elif char in {"-", "\u2212", "%", "/", "&", "#", "(", ")"}:
            total += 55
        elif ord(char) > 127 or unicodedata.east_asian_width(char) in {"W", "F"}:
            total += 120
        elif char.isupper():
            total += 95
        else:
            total += 85
    if language != "english":
        total += 40
    return total


def financial_column_min_width(header_text: str, language: str, col_idx: int) -> int:
    normalized = header_text.replace("\xa0", " ").strip()
    if col_idx == 0:
        return 980
    if re.fullmatch(r"Q[1-4] 20\d{2}", normalized):
        return 760
    if normalized in {"QoQ", "YoY"}:
        return 820
    if normalized.endswith("FY") or normalized.endswith("YTD"):
        return 820
    if re.fullmatch(r"20\d{2}E", normalized):
        return 680
    return 0


def financial_column_weight(header_text: str, col_idx: int) -> float:
    normalized = header_text.replace("\xa0", " ").strip()
    if col_idx == 0:
        return 1.65
    if re.fullmatch(r"Q[1-4] 20\d{2}", normalized):
        return 1.02
    if normalized == "QoQ":
        return 1.00
    if normalized == "YoY":
        return 1.10
    if normalized.endswith("FY") or normalized.endswith("YTD"):
        return 1.10
    if re.fullmatch(r"20\d{2}E", normalized):
        return 0.85
    return 1.00


def distribute_total_width(total_width: int, weights: list[float]) -> list[int]:
    if total_width <= 0 or not weights:
        return []
    weight_sum = sum(weights)
    if weight_sum <= 0:
        return []
    raw = [total_width * weight / weight_sum for weight in weights]
    widths = [int(value) for value in raw]
    remainder = total_width - sum(widths)
    order = sorted(range(len(raw)), key=lambda idx: raw[idx] - widths[idx], reverse=True)
    for idx in order[:remainder]:
        widths[idx] += 1
    return widths


def page_usable_width_twips(doc: Document) -> int | None:
    try:
        section = doc.sections[0]
    except IndexError:
        return None
    usable_emu = section.page_width - section.left_margin - section.right_margin
    if usable_emu is None:
        return None
    return max(0, int(round(usable_emu / 635)))


def rebalance_table_widths(widths: list[int], required: list[int], minimums: list[int], max_total: int | None = None) -> list[int]:
    if not widths or len(widths) != len(required) or len(widths) != len(minimums):
        return widths
    adjusted = [max(widths[i], minimums[i], required[i]) for i in range(len(widths))]
    if max_total is not None and sum(adjusted) > max_total:
        excess = sum(adjusted) - max_total
        shrinkable = [max(0, adjusted[i] - minimums[i]) for i in range(len(widths))]
        total_shrinkable = sum(shrinkable)
        if total_shrinkable > 0:
            for i, available in enumerate(shrinkable):
                if excess <= 0:
                    break
                if available <= 0:
                    continue
                shrink = min(available, round((sum(adjusted) - max_total) * available / total_shrinkable))
                adjusted[i] -= shrink
                excess -= shrink
            if excess > 0:
                for i, available in enumerate(shrinkable):
                    if excess <= 0:
                        break
                    extra = min(max(0, adjusted[i] - minimums[i]), excess)
                    adjusted[i] -= extra
                    excess -= extra
    return adjusted


def set_table_grid_widths(table, widths: list[int], sync_cell_widths: bool = True) -> None:
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(1, tbl_grid)
    grid_cols = list(tbl_grid.iterchildren())
    while len(grid_cols) < len(widths):
        new_col = OxmlElement("w:gridCol")
        tbl_grid.append(new_col)
        grid_cols = list(tbl_grid.iterchildren())
    for idx, width in enumerate(widths):
        grid_cols[idx].set(qn("w:w"), str(int(width)))
    for extra in grid_cols[len(widths):]:
        tbl_grid.remove(extra)
    if sync_cell_widths:
        for row in table.rows:
            for idx, cell in enumerate(row.cells[: len(widths)]):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:type"), "dxa")
                tc_w.set(qn("w:w"), str(int(widths[idx])))


def set_table_preferred_width(table, width_twips: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_twips)))


def set_table_indent_zero(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")


def set_table_cell_margins(table, top: int = 20, left: int = 30, bottom: int = 20, right: int = 30) -> None:
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for side, value in {"top": top, "left": left, "bottom": bottom, "right": right}.items():
        side_el = cell_mar.find(qn(f"w:{side}"))
        if side_el is None:
            side_el = OxmlElement(f"w:{side}")
            cell_mar.append(side_el)
        side_el.set(qn("w:w"), str(int(value)))
        side_el.set(qn("w:type"), "dxa")


def fit_financial_table_columns(doc: Document, table, language: str) -> None:
    financial_like = looks_like_financial_table(table)
    if financial_like:
        # Prevent Word auto-shrinking from forcing wrapped quarter headers.
        set_table_layout_fixed(table)
        usable_width = page_usable_width_twips(doc)
        if usable_width:
            set_table_preferred_width(table, usable_width)
        set_table_indent_zero(table)
        set_table_cell_margins(table)
    widths = get_table_grid_widths(table)
    layout_type = get_table_layout_type(table)
    if widths:
        required: list[int] = []
        minimums: list[int] = []
        num_cols = min(len(widths), max(len(row.cells) for row in table.rows))
        usable_width = page_usable_width_twips(doc)
        header_texts = [
            table.rows[0].cells[col_idx].text.replace("\n", " ").strip() if table.rows and col_idx < len(table.rows[0].cells) else ""
            for col_idx in range(num_cols)
        ]
        if financial_like and usable_width and num_cols >= 8:
            weighted_widths = distribute_total_width(
                usable_width,
                [financial_column_weight(header_texts[col_idx], col_idx) for col_idx in range(num_cols)],
            )
            if len(weighted_widths) == num_cols:
                set_table_grid_widths(table, weighted_widths, sync_cell_widths=True)
                layout_type = get_table_layout_type(table)
                for row in table.rows:
                    for cell in row.cells:
                        if layout_type == "fixed":
                            set_cell_no_wrap(cell)
                return
        for col_idx in range(num_cols):
            max_width = widths[col_idx]
            header_text = header_texts[col_idx]
            for row in table.rows:
                if col_idx >= len(row.cells):
                    continue
                cell_text = row.cells[col_idx].text.replace("\n", " ")
                if not cell_text:
                    continue
                max_width = max(max_width, estimate_table_text_width(cell_text, language))
            minimum_width = financial_column_min_width(header_text, language, col_idx)
            max_width = max(max_width, minimum_width)
            required.append(max_width)
            minimums.append(minimum_width)
        adjusted = rebalance_table_widths(
            widths[: len(required)],
            required,
            minimums,
            usable_width,
        )
        set_table_grid_widths(table, adjusted, sync_cell_widths=(layout_type == "fixed"))
    for row in table.rows:
        for cell in row.cells:
            if layout_type == "fixed":
                set_cell_no_wrap(cell)
            else:
                # Keep explicitly-set no-wrap flags for auto-layout tables too.
                # Clearing here causes header wrapping regressions in some templates.
                pass


def capture_run_format(paragraph) -> tuple[str | None, str | None, float | None, bool | None]:
    """Returns (font_name, east_asia_font, size_pt, bold)."""
    for run in paragraph.runs:
        font_name = run.font.name
        east_asia_font = None
        rpr = run._element.find(qn("w:rPr"))
        if rpr is not None:
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                east_asia_font = rfonts.get(qn("w:eastAsia"))
        size_pt = float(run.font.size.pt) if run.font.size is not None else None
        bold = run.font.bold
        return font_name, east_asia_font, size_pt, bold
    return None, None, None, None


def apply_paragraph_run_format(paragraph, font_name: str | None, east_asia_font: str | None, size_pt: float | None, bold: bool | None) -> None:
    for run in paragraph.runs:
        if font_name:
            run.font.name = font_name
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:ascii"), font_name)
            rfonts.set(qn("w:hAnsi"), font_name)
            rfonts.set(qn("w:cs"), font_name)
            # Use the East Asian font if captured, otherwise fall back to font_name
            rfonts.set(qn("w:eastAsia"), east_asia_font or font_name)
        elif east_asia_font:
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:eastAsia"), east_asia_font)
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        run.font.bold = bold


def apply_run_font_format(run, font_name: str | None, east_asia_font: str | None, size_pt: float | None, bold: bool | None) -> None:
    if font_name:
        run.font.name = font_name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        rfonts.set(qn("w:cs"), font_name)
        rfonts.set(qn("w:eastAsia"), east_asia_font or font_name)
    elif east_asia_font:
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), east_asia_font)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.font.bold = bold


def apply_subheader_prefix_format(paragraph, text: str, font_name: str | None, east_asia_font: str | None, size_pt: float | None) -> None:
    stripped = strip_leading_list_marker(norm_space(text))
    match = re.match(r"^(.*?[：:])(\s*)(.*)$", stripped)
    if not match:
        return
    prefix, spacing, suffix = match.groups()
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        first_run = paragraph.runs[0]
    else:
        first_run = paragraph.add_run()
    first_run.text = prefix
    apply_run_font_format(first_run, font_name, east_asia_font, size_pt, True)
    first_run.font.underline = False
    if spacing:
        space_run = paragraph.add_run(spacing)
        apply_run_font_format(space_run, font_name, east_asia_font, size_pt, False)
        space_run.font.underline = False
    if suffix:
        suffix_run = paragraph.add_run(suffix)
        apply_run_font_format(suffix_run, font_name, east_asia_font, size_pt, False)
        suffix_run.font.underline = False


def clear_paragraph_preserve(paragraph, normal_style=None) -> None:
    set_paragraph_text_preserve(paragraph, "")
    if normal_style is not None:
        try:
            paragraph.style = normal_style
        except Exception:
            pass


def copy_paragraph_layout(target, source) -> None:
    try:
        target.style = source.style
    except Exception:
        pass
    target.alignment = source.alignment
    src_fmt = source.paragraph_format
    dst_fmt = target.paragraph_format
    dst_fmt.left_indent = src_fmt.left_indent
    dst_fmt.right_indent = src_fmt.right_indent
    dst_fmt.first_line_indent = src_fmt.first_line_indent
    dst_fmt.space_before = src_fmt.space_before
    dst_fmt.space_after = src_fmt.space_after
    dst_fmt.line_spacing = src_fmt.line_spacing
    dst_fmt.line_spacing_rule = src_fmt.line_spacing_rule
    dst_fmt.keep_together = src_fmt.keep_together
    dst_fmt.keep_with_next = src_fmt.keep_with_next
    dst_fmt.page_break_before = src_fmt.page_break_before
    dst_fmt.widow_control = src_fmt.widow_control
    # Copy bullet/numbering properties (numPr) from source to target
    src_pPr = source._element.find(qn("w:pPr"))
    if src_pPr is not None:
        src_numPr = src_pPr.find(qn("w:numPr"))
        if src_numPr is not None:
            tgt_pPr = target._element.get_or_add_pPr()
            old_numPr = tgt_pPr.find(qn("w:numPr"))
            if old_numPr is not None:
                tgt_pPr.remove(old_numPr)
            tgt_pPr.append(copy.deepcopy(src_numPr))


def paragraph_is_list_like(paragraph) -> bool:
    style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
    if "list" in style_name or "列表" in style_name:
        return True
    ppr = getattr(paragraph._element, "pPr", None)
    if ppr is not None and getattr(ppr, "numPr", None) is not None:
        return True
    stripped = paragraph.text.strip()
    return stripped.startswith(("-", "•", "*"))


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def is_business_update_root_heading_text(text: str) -> bool:
    stripped = norm_space(text)
    if not stripped:
        return False
    if canonical_heading(stripped) == "Business Update":
        return True
    normalized = stripped.lower()
    return any(
        token in normalized
        for token in [
            "business operations",
            "业务情况",
        ]
    )


def paragraph_has_emphasis(paragraph) -> bool:
    for run in paragraph.runs:
        if not run.text.strip():
            continue
        if bool(run.bold) or bool(run.underline):
            return True
    return False


def looks_like_middle_subheader(text: str) -> bool:
    stripped = strip_leading_list_marker(norm_space(text))
    if not stripped:
        return False
    if len(stripped) <= 72 and (":" in stripped or "：" in stripped):
        return True
    if len(stripped) <= 56 and not re.search(r"[。.!?]$", stripped):
        if not re.search(r"\b(?:is|are|was|were|has|have|had|will|can|could|should|would)\b", stripped, flags=re.I):
            return True
    return False


def is_standalone_labeled_middle_bullet(text: str) -> bool:
    stripped = strip_leading_list_marker(norm_space(text))
    if not stripped:
        return False
    match = re.match(r"^([^：:]{1,40}[：:])\s*(.+)$", stripped)
    if not match:
        return False
    body = norm_space(match.group(2))
    if len(body) < 12:
        return False
    if re.search(r"[。.!?；;]$", body):
        return True
    if len(body) >= 28:
        return True
    return False


def should_preserve_middle_subheader_hierarchy(texts: list[str]) -> bool:
    normalized = [norm_space(text) for text in texts if norm_space(text)]
    if not normalized:
        return False
    subheader_like = [text for text in normalized if looks_like_middle_subheader(text)]
    if not subheader_like:
        return False
    standalone_labeled = [text for text in normalized if is_standalone_labeled_middle_bullet(text)]
    # If every generated line is already a complete labeled bullet like
    # "Patents: ..." or "Team: ...", keep the output flat instead of
    # inheriting the template's nested subheader/body list structure.
    if standalone_labeled and len(standalone_labeled) == len(normalized):
        return False
    return True


def redistribute_flat_middle_chunks(
    planned_occurrences: list[SectionOccurrence],
    chunks: list[list[str]],
) -> list[list[str]]:
    if len(planned_occurrences) <= 1:
        return chunks
    flat_bullets = [bullet for chunk in chunks for bullet in chunk if norm_space(bullet)]
    if not flat_bullets:
        return chunks
    if should_preserve_middle_subheader_hierarchy(flat_bullets):
        return chunks

    root_positions = [
        idx
        for idx, planned in enumerate(planned_occurrences)
        if is_business_update_root_heading_text(planned.heading_text)
    ]
    if not root_positions:
        target_idx = 0
    elif len(root_positions) == 1:
        target_idx = root_positions[0]
    else:
        # Templates like Yiqing intentionally use a numbered root heading
        # followed by a second root heading ("Business Progress"). In flat mode,
        # keep both headings but place the bullet body under the last root.
        target_idx = root_positions[-1]

    redistributed = [[] for _ in planned_occurrences]
    redistributed[target_idx] = flat_bullets
    return redistributed


def ensure_middle_occurrence_body_slots(doc: Document, occurrences: list[DocSectionOccurrence], occ_idx: int, texts: list[str]):
    body = list(get_body_paragraphs_for_occurrence(doc, occurrences, occ_idx))
    if body or not texts:
        return body
    next_heading_index = occurrences[occ_idx + 1].heading_index if occ_idx + 1 < len(occurrences) else None
    if next_heading_index is not None and next_heading_index < len(doc.paragraphs):
        placeholder = doc.paragraphs[next_heading_index].insert_paragraph_before("")
        return [placeholder]
    return body


def fill_section_paragraphs(
    doc: Document,
    paragraphs,
    texts: list[str],
    force_bullets: bool = False,
    preserve_subheader_hierarchy: bool = False,
    force_dash_bullets: bool = False,
) -> None:
    def dash_bullet_text(text: str) -> str:
        stripped = strip_leading_list_marker(norm_space(text))
        stripped = re.sub(r"^(?:[-•*]\s*)+", "", stripped).strip()
        if stripped.startswith("- "):
            return stripped
        return f"- {stripped}" if stripped else "-"

    slots = list(paragraphs)
    normal_style = None
    try:
        normal_style = doc.styles["Normal"]
    except Exception:
        pass

    template_style = None
    template_run_format = (None, None, None, None)
    for paragraph in slots:
        if paragraph.text.strip():
            template_style = paragraph.style
            template_run_format = capture_run_format(paragraph)
            break
    if template_style is None and slots:
        template_style = slots[0].style
        template_run_format = capture_run_format(slots[0])

    working_slots = slots
    bullet_reference = None
    uniform_bullet_mode = False
    if force_bullets:
        list_like = [
            paragraph
            for paragraph in slots
            if paragraph.text.strip() and paragraph_is_list_like(paragraph)
        ]
        if list_like:
            working_slots = list_like
            bullet_reference = list_like[0]
        elif slots:
            working_slots = slots
            bullet_reference = slots[0]
        if bullet_reference is not None and not paragraph_is_list_like(bullet_reference):
            fallback_bullet = next((paragraph for paragraph in doc.paragraphs if paragraph_is_list_like(paragraph)), None)
            if fallback_bullet is not None:
                bullet_reference = fallback_bullet
        structured_mode = preserve_subheader_hierarchy and any(looks_like_middle_subheader(text) for text in texts)
        if not structured_mode:
            plain_bullet = next((paragraph for paragraph in list_like if not paragraph_has_emphasis(paragraph)), None)
            if plain_bullet is not None:
                bullet_reference = plain_bullet
            uniform_bullet_mode = bullet_reference is not None

    remove_later = [paragraph for paragraph in slots if paragraph not in working_slots]
    slot_formats = [capture_run_format(paragraph) for paragraph in working_slots]
    for idx, paragraph in enumerate(working_slots):
        if idx < len(texts):
            output_text = dash_bullet_text(texts[idx]) if force_dash_bullets else texts[idx]
            set_paragraph_text_preserve(paragraph, output_text)
            if force_bullets and uniform_bullet_mode and bullet_reference is not None:
                copy_paragraph_layout(paragraph, bullet_reference)
                ref_fmt = capture_run_format(bullet_reference)
                apply_paragraph_run_format(paragraph, ref_fmt[0], ref_fmt[1], ref_fmt[2], False)
                for run in paragraph.runs:
                    run.font.bold = False
                    run.font.underline = False
            else:
                # Keep each template slot's own paragraph style (important for mixed bullet templates).
                fmt = slot_formats[idx] if idx < len(slot_formats) else template_run_format
                apply_paragraph_run_format(paragraph, *fmt)
                if preserve_subheader_hierarchy and looks_like_middle_subheader(texts[idx]):
                    apply_subheader_prefix_format(paragraph, texts[idx], fmt[0], fmt[1], fmt[2])
                else:
                    # For plain bullet bodies, avoid inheriting full-line emphasis from the template.
                    for run in paragraph.runs:
                        run.font.bold = False
                        run.font.underline = False
            if force_dash_bullets:
                remove_paragraph_numbering(paragraph)
        else:
            remove_later.append(paragraph)

    if len(texts) <= len(working_slots):
        for paragraph in reversed(remove_later):
            element = paragraph._element
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
        return

    anchor = working_slots[-1] if working_slots else None
    for text in texts[len(working_slots) :]:
        if anchor is None:
            break
        output_text = dash_bullet_text(text) if force_dash_bullets else text
        new_paragraph = anchor.insert_paragraph_before(output_text)
        if force_bullets and bullet_reference is not None:
            copy_paragraph_layout(new_paragraph, bullet_reference)
            if uniform_bullet_mode:
                ref_fmt = capture_run_format(bullet_reference)
                apply_paragraph_run_format(new_paragraph, ref_fmt[0], ref_fmt[1], ref_fmt[2], False)
                for run in new_paragraph.runs:
                    run.font.bold = False
                    run.font.underline = False
            else:
                apply_paragraph_run_format(new_paragraph, *template_run_format)
                if preserve_subheader_hierarchy and looks_like_middle_subheader(text):
                    apply_subheader_prefix_format(new_paragraph, text, template_run_format[0], template_run_format[1], template_run_format[2])
                else:
                    for run in new_paragraph.runs:
                        run.font.bold = False
                        run.font.underline = False
            if force_dash_bullets:
                remove_paragraph_numbering(new_paragraph)
        elif template_style is not None:
            try:
                new_paragraph.style = template_style
            except Exception:
                pass
            apply_paragraph_run_format(new_paragraph, *template_run_format)
            if preserve_subheader_hierarchy and looks_like_middle_subheader(text):
                apply_subheader_prefix_format(new_paragraph, text, template_run_format[0], template_run_format[1], template_run_format[2])
            else:
                for run in new_paragraph.runs:
                    run.font.bold = False
                    run.font.underline = False
            if force_dash_bullets:
                remove_paragraph_numbering(new_paragraph)
        anchor = new_paragraph

    for paragraph in reversed(remove_later):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//*[local-name()='drawing' or local-name()='pict']"))


def paragraph_has_page_break(paragraph) -> bool:
    for element in paragraph._element.iter():
        if element.tag.endswith("}br") and element.get(qn("w:type")) == "page":
            return True
    return False


def is_plain_blank_paragraph(paragraph) -> bool:
    return paragraph.text.strip() == "" and not paragraph_has_drawing(paragraph) and not paragraph_has_page_break(paragraph)


def is_cover_title_paragraph(paragraph) -> bool:
    text = norm_space(paragraph.text)
    if not text:
        return False
    if is_numbered_section_line(text) or is_top_level_numbered_section_line(text) or is_decimal_section_heading_line(text):
        return False
    if re.search(r"[。.!?；;]$", text):
        return False
    if len(text) > 80:
        return False
    if len(re.findall(r"[A-Za-z]+", text)) > 10:
        return False
    return True


def remove_page_break_runs(paragraph) -> None:
    to_remove = []
    for element in paragraph._element.iter():
        if element.tag.endswith("}br") and element.get(qn("w:type")) == "page":
            to_remove.append(element)
    for element in to_remove:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def remove_paragraph_numbering(paragraph) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is not None:
        ppr.remove(num_pr)


def find_logo_info(doc: Document, template_path: Path) -> tuple[Path | None, Any | None, Any | None, Any | None]:
    logo_paragraph = None
    for paragraph in doc.paragraphs:
        if paragraph_has_drawing(paragraph):
            logo_paragraph = paragraph
            break
    if logo_paragraph is None:
        return None, None, None, None

    alignment = logo_paragraph.alignment
    style = logo_paragraph.style
    width = None
    height = None
    if doc.inline_shapes:
        shape = doc.inline_shapes[0]
        width = shape.width
        height = shape.height

    with zipfile.ZipFile(template_path) as zf:
        media_names = sorted(name for name in zf.namelist() if name.startswith("word/media/"))
        if not media_names:
            return None, alignment, style, None
        media_name = media_names[0]
        suffix = Path(media_name).suffix or ".png"
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        temp_file.write(zf.read(media_name))
        temp_file.flush()
        temp_file.close()
    return Path(temp_file.name), alignment, style, (width, height)


def insert_page_break_and_repeated_logo(doc: Document, template_path: Path, occurrences: list[DocSectionOccurrence] | None = None) -> None:
    occurrences = occurrences or detect_doc_section_occurrences(doc)
    english_occurrences = [occ for occ in occurrences if occ.language == "english"]
    chinese_occurrences = [occ for occ in occurrences if occ.language == "chinese"]
    if not english_occurrences or not chinese_occurrences:
        return

    english_start = doc.paragraphs[english_occurrences[0].heading_index]
    logo_path, alignment, style, size = find_logo_info(doc, template_path)

    # Insert logo first, then page break before it, so final order is:
    # [page break] [logo] [english heading]
    logo_ref = english_start  # reference for page break insertion
    if logo_path is not None:
        logo_paragraph = english_start.insert_paragraph_before("")
        if style is not None:
            try:
                logo_paragraph.style = style
            except Exception:
                pass
        logo_paragraph.alignment = alignment
        run = logo_paragraph.add_run()
        width, height = size if size is not None else (None, None)
        if width is not None:
            run.add_picture(str(logo_path), width=width, height=height)
        else:
            run.add_picture(str(logo_path))
        logo_ref = logo_paragraph  # insert page break before the logo
        try:
            logo_path.unlink()
        except Exception:
            pass

    break_paragraph = logo_ref.insert_paragraph_before("")
    break_paragraph.add_run().add_break(WD_BREAK.PAGE)


def remove_blank_paragraphs_inside_sections(doc: Document, occurrences: list[DocSectionOccurrence] | None = None) -> None:
    occurrences = occurrences or detect_doc_section_occurrences(doc)
    to_remove = []
    for occ_idx, _occurrence in enumerate(occurrences):
        body = list(get_body_paragraphs_for_occurrence(doc, occurrences, occ_idx))
        for paragraph in body:
            if is_plain_blank_paragraph(paragraph):
                to_remove.append(paragraph)
    for paragraph in reversed(to_remove):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def normalize_bullet_sections(doc: Document, occurrences: list[DocSectionOccurrence] | None = None) -> None:
    occurrences = occurrences or detect_doc_section_occurrences(doc)
    for occ_idx, occurrence in enumerate(occurrences):
        # Keep Business Update's original mixed bullet hierarchy from the template.
        # Only normalize Risk & Exit where homogeneous bullets are expected.
        if occurrence.canonical not in {"Risk & Exit"}:
            continue
        raw_body = list(get_body_paragraphs_for_occurrence(doc, occurrences, occ_idx))
        body: list[Any] = []
        for paragraph in raw_body:
            if paragraph_has_drawing(paragraph) or paragraph_has_page_break(paragraph):
                break
            body.append(paragraph)
        content_paragraphs = [p for p in body if p.text.strip()]
        if not content_paragraphs:
            continue
        bullet_reference = next((p for p in content_paragraphs if paragraph_is_list_like(p)), content_paragraphs[0])
        font_name, east_asia_font, size_pt, _bold = capture_run_format(bullet_reference)
        for paragraph in content_paragraphs:
            copy_paragraph_layout(paragraph, bullet_reference)
            # Always set bold=False for content paragraphs; headings get bold separately
            apply_paragraph_run_format(paragraph, font_name, east_asia_font, size_pt, False)


def normalize_cover_title_paragraphs(doc: Document, occurrences: list[DocSectionOccurrence] | None = None) -> None:
    occurrences = occurrences or detect_doc_section_occurrences(doc)
    first_by_language: dict[str, DocSectionOccurrence] = {}
    for occurrence in sorted(occurrences, key=lambda occ: occ.heading_index):
        if occurrence.language not in first_by_language:
            first_by_language[occurrence.language] = occurrence

    normal_style = None
    with contextlib.suppress(Exception):
        normal_style = doc.styles["Normal"]

    for occurrence in first_by_language.values():
        heading_idx = occurrence.heading_index
        logo_idx = None
        idx = heading_idx - 1
        while idx >= 0:
            paragraph = doc.paragraphs[idx]
            if paragraph_has_drawing(paragraph):
                logo_idx = idx
                break
            if is_plain_blank_paragraph(paragraph) or paragraph_has_page_break(paragraph) or is_cover_title_paragraph(paragraph):
                idx -= 1
                continue
            break
        if logo_idx is None:
            continue

        for idx in range(logo_idx + 1, heading_idx):
            paragraph = doc.paragraphs[idx]
            if is_plain_blank_paragraph(paragraph) or paragraph_has_page_break(paragraph):
                continue
            # Never treat actual section headings as cover titles, even if
            # occurrence resolution missed them earlier.
            if fixed_anchor_heading(paragraph.text) or canonical_heading(paragraph.text):
                break
            if not is_cover_title_paragraph(paragraph):
                break
            remove_paragraph_numbering(paragraph)
            if normal_style is not None:
                with contextlib.suppress(Exception):
                    paragraph.style = normal_style
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.bold = False
                run.font.underline = False


def reapply_english_template_format_for_bilingual(doc: Document, template_path: Path) -> None:
    current_occurrences = detect_doc_section_occurrences(doc)
    if template_language_mode(current_occurrences) != "bilingual":
        return
    try:
        template_doc = Document(str(template_path))
    except Exception:
        return
    template_occurrences = detect_doc_section_occurrences(template_doc)

    current_en_entries = [(idx, occ) for idx, occ in enumerate(current_occurrences) if occ.language == "english"]
    template_en_entries = [(idx, occ) for idx, occ in enumerate(template_occurrences) if occ.language == "english"]
    if not current_en_entries or not template_en_entries:
        return

    current_en = [occ for _idx, occ in current_en_entries]
    template_en = [occ for _idx, occ in template_en_entries]

    def collect_cover_title_indices(local_doc: Document, first_heading_idx: int) -> list[int]:
        logo_idx = None
        idx = first_heading_idx - 1
        while idx >= 0:
            paragraph = local_doc.paragraphs[idx]
            if paragraph_has_drawing(paragraph):
                logo_idx = idx
                break
            if is_plain_blank_paragraph(paragraph) or paragraph_has_page_break(paragraph) or is_cover_title_paragraph(paragraph):
                idx -= 1
                continue
            break
        if logo_idx is None:
            # Some outputs have only a page break before the English title line
            # (no explicit logo drawing paragraph). Use that boundary too.
            idx = first_heading_idx - 1
            while idx >= 0:
                paragraph = local_doc.paragraphs[idx]
                if paragraph_has_page_break(paragraph):
                    logo_idx = idx
                    break
                if is_plain_blank_paragraph(paragraph) or is_cover_title_paragraph(paragraph):
                    idx -= 1
                    continue
                break
        indices: list[int] = []
        if logo_idx is not None:
            for p_idx in range(logo_idx + 1, first_heading_idx):
                paragraph = local_doc.paragraphs[p_idx]
                if is_plain_blank_paragraph(paragraph) or paragraph_has_page_break(paragraph):
                    continue
                if not is_cover_title_paragraph(paragraph):
                    break
                indices.append(p_idx)
            if indices:
                return indices

        # Final fallback: grab contiguous cover-title lines immediately before
        # the first English heading (works for templates without logo/page marker).
        idx = first_heading_idx - 1
        while idx >= 0 and is_plain_blank_paragraph(local_doc.paragraphs[idx]):
            idx -= 1
        reversed_indices: list[int] = []
        while idx >= 0:
            paragraph = local_doc.paragraphs[idx]
            if paragraph_has_drawing(paragraph) or paragraph_has_page_break(paragraph):
                break
            if is_cover_title_paragraph(paragraph):
                reversed_indices.append(idx)
                idx -= 1
                continue
            break
        return list(reversed(reversed_indices))

    def template_body_reference(template_occ_global_idx: int, template_en_pos: int, canonical: str):
        template_body = [
            p
            for p in get_body_paragraphs_for_occurrence(template_doc, template_occurrences, template_occ_global_idx)
            if p.text.strip() and not paragraph_has_drawing(p) and not paragraph_has_page_break(p)
        ]
        if template_body:
            template_list_ref = next((p for p in template_body if paragraph_is_list_like(p)), None)
            if canonical in {"Business Update", "Risk & Exit"} and template_list_ref is not None:
                return template_list_ref
            return template_body[0]

        # Some templates (e.g., Honghu root Business Update heading) keep no body
        # under the root heading itself. Borrow style from the next Business Update
        # sub-block that has real bullet content.
        if canonical == "Business Update":
            for next_pos in range(template_en_pos + 1, len(template_en_entries)):
                next_global_idx, next_occ = template_en_entries[next_pos]
                if next_occ.canonical != "Business Update":
                    break
                next_body = [
                    p
                    for p in get_body_paragraphs_for_occurrence(template_doc, template_occurrences, next_global_idx)
                    if p.text.strip() and not paragraph_has_drawing(p) and not paragraph_has_page_break(p)
                ]
                if not next_body:
                    continue
                next_list_ref = next((p for p in next_body if paragraph_is_list_like(p)), None)
                return next_list_ref or next_body[0]
        return None

    # Re-apply cover title font/style from template (e.g., English company name line).
    current_cover = collect_cover_title_indices(doc, current_en_entries[0][1].heading_index)
    template_cover = collect_cover_title_indices(template_doc, template_en_entries[0][1].heading_index)
    for curr_idx, tpl_idx in zip(current_cover, template_cover):
        paragraph_out = doc.paragraphs[curr_idx]
        paragraph_tpl = template_doc.paragraphs[tpl_idx]
        original_text = paragraph_out.text
        copy_paragraph_layout(paragraph_out, paragraph_tpl)
        set_paragraph_text_preserve(paragraph_out, original_text)
        apply_paragraph_run_format(paragraph_out, *capture_run_format(paragraph_tpl))
        remove_paragraph_numbering(paragraph_out)

    def matching_template_occurrence(target: DocSectionOccurrence, start_pos: int, target_heading_text: str) -> tuple[int | None, int]:
        target_is_root_middle = target.canonical == "Business Update" and is_business_update_root_heading_text(target_heading_text)
        fallback_pos = start_pos
        for pos in range(start_pos, len(template_en_entries)):
            _global_idx, candidate = template_en_entries[pos]
            if candidate.canonical == target.canonical:
                if target.canonical != "Business Update":
                    return pos, pos + 1
                candidate_text = template_doc.paragraphs[candidate.heading_index].text
                candidate_is_root_middle = is_business_update_root_heading_text(candidate_text)
                if candidate_is_root_middle == target_is_root_middle:
                    return pos, pos + 1
                if fallback_pos == start_pos:
                    fallback_pos = pos + 1
        if fallback_pos > start_pos:
            return fallback_pos - 1, fallback_pos
        return None, start_pos

    template_cursor = 0
    for current_global_idx, current in current_en_entries:
        current_heading_text = doc.paragraphs[current.heading_index].text
        template_match_pos, template_cursor = matching_template_occurrence(current, template_cursor, current_heading_text)
        if template_match_pos is None:
            continue
        template_global_idx, template_match = template_en_entries[template_match_pos]

        # Keep output text, but force heading layout/font from template.
        heading_out = doc.paragraphs[current.heading_index]
        heading_tpl = template_doc.paragraphs[template_match.heading_index]
        heading_text = heading_out.text
        heading_fmt = capture_run_format(heading_tpl)
        copy_paragraph_layout(heading_out, heading_tpl)
        set_paragraph_text_preserve(heading_out, heading_text)
        apply_paragraph_run_format(heading_out, *heading_fmt)

        current_body = [
            p
            for p in get_body_paragraphs_for_occurrence(doc, current_occurrences, current_global_idx)
            if p.text.strip() and not paragraph_has_drawing(p) and not paragraph_has_page_break(p)
        ]
        body_ref = template_body_reference(template_global_idx, template_match_pos, current.canonical)
        if not current_body or body_ref is None:
            continue

        body_fmt = capture_run_format(body_ref)
        for paragraph in current_body:
            text = paragraph.text
            copy_paragraph_layout(paragraph, body_ref)
            set_paragraph_text_preserve(paragraph, text)
            apply_paragraph_run_format(paragraph, *body_fmt)
            if current.canonical in {"Business Update", "Risk & Exit"}:
                for run in paragraph.runs:
                    run.font.bold = False
            if current.canonical == "Business Update" and norm_space(paragraph.text).startswith("-"):
                remove_paragraph_numbering(paragraph)


def remove_empty_business_update_subheadings(doc: Document, occurrences: list[DocSectionOccurrence] | None = None) -> None:
    occurrences = occurrences or detect_doc_section_occurrences(doc)
    for occ_idx in reversed(range(len(occurrences))):
        occurrence = occurrences[occ_idx]
        if occurrence.canonical != "Business Update":
            continue
        heading = doc.paragraphs[occurrence.heading_index]
        if is_business_update_root_heading_text(heading.text):
            continue
        body = list(get_body_paragraphs_for_occurrence(doc, occurrences, occ_idx))
        has_body_text = any(paragraph.text.strip() for paragraph in body)
        if not has_body_text:
            remove_paragraph(heading)


def normalize_exact_single_blank_between_sections(doc: Document, occurrences: list[DocSectionOccurrence] | None = None) -> None:
    """Insert exactly one blank paragraph before each section heading (except the very first)."""
    occurrences = occurrences or detect_doc_section_occurrences(doc)
    if not occurrences:
        return

    heading_paragraphs = []
    for occ in occurrences:
        if occ.heading_index >= len(doc.paragraphs):
            continue
        heading = doc.paragraphs[occ.heading_index]
        if occ.canonical in {"Business Activities", "Financial Update", "Risk & Exit"}:
            heading_paragraphs.append(heading)
            continue
        if occ.canonical == "Business Update" and is_business_update_root_heading_text(heading.text):
            heading_paragraphs.append(heading)

    # Deduplicate while preserving order.
    deduped: list[Any] = []
    seen_ids: set[int] = set()
    for paragraph in heading_paragraphs:
        pid = id(paragraph._element)
        if pid in seen_ids:
            continue
        seen_ids.add(pid)
        deduped.append(paragraph)
    heading_paragraphs = deduped

    for idx, paragraph in enumerate(heading_paragraphs):
        # Skip the very first heading in the document (no blank needed above it)
        if idx == 0:
            continue

        # Check if there's already a blank paragraph immediately before this heading
        previous = paragraph._element.getprevious()
        if previous is not None:
            prev_p = next((p for p in doc.paragraphs if p._element is previous), None)
            if prev_p is not None and is_plain_blank_paragraph(prev_p):
                continue  # already has a blank

        # Insert a blank paragraph before this heading
        blank = paragraph.insert_paragraph_before("")
        try:
            blank.style = doc.styles["Normal"]
        except Exception:
            pass


def center_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")


def financial_table_backbone_score(prev_financial: dict[str, Any], previous_review_quarter: Quarter | None) -> tuple[int, int]:
    normalized_headers = normalize_previous_quarter_headers(prev_financial["headers"], previous_review_quarter)
    labels = [normalized for _raw, normalized in normalized_headers]
    expected: list[str] = []
    if previous_review_quarter is not None:
        cursor = previous_review_quarter
        for _ in range(5):
            expected.append(cursor.display())
            cursor = cursor.previous()
    coverage = len([label for label in labels if label in expected]) if expected else len(labels)
    nonempty = 0
    for row in prev_financial["rows"]:
        for raw, _normalized in normalized_headers:
            if norm_space(row["values"].get(raw, "")):
                nonempty += 1
    return coverage, nonempty


def localize_financial_update(
    source_update: dict[str, Any],
    target_prev_financial: dict[str, Any],
    source_unit_spec: UnitSpec | None = None,
    target_unit_spec: UnitSpec | None = None,
    exchange_rates: ExchangeRates | None = None,
) -> dict[str, Any]:
    localized_rows: list[dict[str, Any]] = []
    for idx, source_row in enumerate(source_update["rows"]):
        target_label = target_prev_financial["rows"][idx]["label"] if idx < len(target_prev_financial["rows"]) else source_row["label"]
        localized_values: dict[str, str | None] = {}
        for header, raw_value in source_row["values"].items():
            if raw_value is None or header in {"QoQ", "YoY"}:
                localized_values[header] = raw_value
                continue
            if isinstance(raw_value, str) and "%" in raw_value:
                localized_values[header] = raw_value
                continue
            decimal_value = finance_parse_decimal(str(raw_value))
            if decimal_value is None:
                localized_values[header] = raw_value
                continue
            converted = convert_value_between_units(decimal_value, source_unit_spec, target_unit_spec, exchange_rates)
            localized_values[header] = format_decimal(converted) if converted is not None else raw_value
        localized_rows.append(
            {
                "label": target_label,
                "values": localized_values,
                "source_trace": dict(source_row.get("source_trace", {})),
                "flags": list(source_row.get("flags", [])),
            }
        )
    return {
        "status": source_update.get("status"),
        "unit": target_prev_financial.get("unit") or source_update.get("unit"),
        "columns": list(source_update["columns"]),
        "rows": localized_rows,
    }


def section_table_entries(section_plan: list[SectionOccurrence], canonical_filter: set[str] | None = None) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    global_table_index = 0
    for occurrence in section_plan:
        for local_table_index, _rows in enumerate(occurrence.tables):
            if canonical_filter is None or occurrence.canonical in canonical_filter:
                entries.append(
                    {
                        "occurrence": occurrence,
                        "local_table_index": local_table_index,
                        "global_table_index": global_table_index,
                    }
                )
            global_table_index += 1
    return entries


def group_middle_table_occurrences(previous_occurrences: list[SectionOccurrence]) -> list[list[SectionOccurrence]]:
    groups: dict[int, list[SectionOccurrence]] = {}
    counters: dict[str, int] = {"chinese": 0, "english": 0}
    for occurrence in previous_occurrences:
        if occurrence.canonical != "Business Update" or not occurrence.tables:
            continue
        group_index = counters.get(occurrence.language, 0)
        counters[occurrence.language] = group_index + 1
        groups.setdefault(group_index, []).append(occurrence)
    return [groups[idx] for idx in sorted(groups)]


def build_section_table_update(
    prev_table: dict[str, Any],
    current_table: dict[str, Any],
    flags: list[dict[str, Any]],
    target_quarter: Quarter | None = None,
    source_unit_spec: UnitSpec | None = None,
    output_unit_spec: UnitSpec | None = None,
    previous_review_quarter: Quarter | None = None,
    exchange_rates: ExchangeRates | None = None,
) -> dict[str, Any]:
    target_quarter = target_quarter or latest_quarter_from_current(current_table)
    normalized_previous_headers = normalize_previous_quarter_headers(prev_table["headers"], previous_review_quarter)
    previous_quarter_labels = [normalized for _raw, normalized in normalized_previous_headers if normalized]
    quarter_count = len(previous_quarter_labels)
    rolling_headers: list[str] = []
    quarter_cursor = target_quarter
    while len(rolling_headers) < quarter_count:
        rolling_headers.append(quarter_cursor.display())
        quarter_cursor = quarter_cursor.previous()

    trailing_headers = [header for header in prev_table["headers"] if not parse_quarter_label(header)]
    output_headers = rolling_headers + trailing_headers
    current_total_header = find_current_quarter_total_header(current_table, target_quarter) or ""
    current_month_headers = find_month_headers_for_quarter(current_table, target_quarter)
    fy_or_ytd_header = find_fy_or_ytd_header(current_table, target_quarter)

    rows_output: list[dict[str, Any]] = []
    used_source_labels: dict[str, str] = {}
    for prev_row in prev_table["rows"]:
        label = prev_row["label"]
        previous_current_value = None
        if previous_review_quarter is not None:
            for raw, normalized in normalized_previous_headers:
                if normalized == previous_review_quarter.display():
                    previous_current_value = prev_row["values"].get(raw)
                    break
        mapped_label = choose_finance_row_label(
            label,
            current_table["rows"],
            current_total_header,
            previous_current_value,
            flags,
            used_source_labels,
            source_unit_spec,
            output_unit_spec,
            exchange_rates,
        )
        if mapped_label:
            used_source_labels[mapped_label] = label
        source_row = current_table["rows"].get(mapped_label) if mapped_label else None

        values: dict[str, str | None] = {}
        source_trace: dict[str, str] = {}
        row_flags: list[str] = []
        prev_quarters = {normalized: prev_row["values"].get(raw) for raw, normalized in normalized_previous_headers}

        current_quarter_value: str | None = None
        if source_row:
            decimal_value = sum_month_values(source_row, current_month_headers)
            if decimal_value is None and current_total_header:
                decimal_value = finance_parse_decimal(source_row.get(current_total_header, ""))
            if decimal_value is None and finance_row_current_period_is_blank(source_row, current_month_headers, current_total_header, fy_or_ytd_header):
                decimal_value = Decimal("0")
            if decimal_value is not None:
                converted_value = convert_value_between_units(decimal_value, source_unit_spec, output_unit_spec, exchange_rates)
                current_quarter_value = format_decimal(converted_value)
                source_trace[target_quarter.display()] = "current_data_request"
            else:
                current_quarter_value = default_missing_financial_value(label)
                row_flags.append("Missing current-quarter value in the company data request.")
        else:
            current_quarter_value = default_missing_financial_value(label)
            row_flags.append("No current data request row could be mapped to this metric.")

        for header in rolling_headers:
            if header == target_quarter.display():
                values[header] = current_quarter_value
            else:
                values[header] = prev_quarters.get(header)
                if prev_quarters.get(header):
                    source_trace[header] = "previous_review"

        for header in trailing_headers:
            normalized_header = header.replace("\xa0", " ").strip()
            if normalized_header in {"QoQ", "YoY"}:
                comparison_header = rolling_headers[1] if normalized_header == "QoQ" and len(rolling_headers) > 1 else rolling_headers[4] if normalized_header == "YoY" and len(rolling_headers) > 4 else None
                values[header] = display_percent_text(
                    values.get(target_quarter.display()),
                    values.get(comparison_header) if comparison_header else None,
                )
            elif normalized_header.endswith("FY") or normalized_header.endswith("YTD"):
                source_value = finance_parse_decimal(source_row.get(fy_or_ytd_header or "", "")) if source_row and fy_or_ytd_header else None
                if source_value is not None:
                    values[header] = format_decimal(convert_value_between_units(source_value, source_unit_spec, output_unit_spec, exchange_rates))
                    source_trace[header] = "current_data_request"
                else:
                    quarter_values = [finance_parse_decimal(values.get(qh) or "") for qh in rolling_headers if qh.startswith(f"Q") and qh.endswith(str(target_quarter.year))]
                    if quarter_values and all(v is not None for v in quarter_values):
                        values[header] = format_decimal(sum(quarter_values))  # type: ignore[arg-type]
                        source_trace[header] = "derived_from_quarters"
                    else:
                        values[header] = default_missing_financial_value(label)
                        row_flags.append(f"{header} is missing in the current data request and was defaulted to {values[header]}.")
            elif re.fullmatch(r"20\d{2}E", normalized_header):
                estimate_source = None
                if source_row:
                    estimate_source = finance_parse_decimal(source_row.get(header, ""))
                    if estimate_source is None:
                        estimate_source = finance_parse_decimal(find_header_value(source_row, ["exp", "预计", "全年预计"]) or "")
                if estimate_source is not None:
                    values[header] = format_decimal(convert_value_between_units(estimate_source, source_unit_spec, output_unit_spec, exchange_rates))
                    source_trace[header] = "current_data_request"
                else:
                    carried_estimate = previous_row_value_for_normalized_header(prev_row, normalized_previous_headers, normalized_header)
                    if carried_estimate is not None and not is_empty_numeric_cell(carried_estimate):
                        values[header] = carried_estimate
                        source_trace[header] = "previous_review"
                        row_flags.append(f"{header} is missing in the current data request and was carried forward from the previous review.")
                    else:
                        values[header] = default_missing_financial_value(label)
                        row_flags.append(f"{header} is missing in the current data request and was defaulted to {values[header]}.")
            else:
                if source_row:
                    raw_source_value = source_row.get(header, "")
                    decimal_source = finance_parse_decimal(raw_source_value)
                    if decimal_source is not None:
                        values[header] = format_decimal(convert_value_between_units(decimal_source, source_unit_spec, output_unit_spec, exchange_rates))
                        source_trace[header] = "current_data_request"
                    else:
                        values[header] = norm_space(raw_source_value) or default_missing_financial_value(label)
                        if not norm_space(raw_source_value):
                            row_flags.append(f"{header} is missing in the current data request and was defaulted to {values[header]}.")
                else:
                    values[header] = default_missing_financial_value(label)
                    row_flags.append(f"{header} is missing in the current data request and was defaulted to {values[header]}.")

        rows_output.append(
            {
                "label": label,
                "mapped_source_label": mapped_label,
                "values": {header: values.get(header) for header in output_headers},
                "source_trace": source_trace,
                "flags": row_flags,
            }
        )

    return {
        "status": "updated_with_flags",
        "unit": prev_table.get("unit") or "",
        "columns": output_headers,
        "rows": rows_output,
    }


def enforce_bilingual_page_break(doc: Document, occurrences: list[DocSectionOccurrence]) -> None:
    english_occurrences = [occ for occ in occurrences if occ.language == "english"]
    chinese_occurrences = [occ for occ in occurrences if occ.language == "chinese"]
    if not english_occurrences or not chinese_occurrences:
        return

    first_english_heading = english_occurrences[0].heading_index
    start_idx = first_english_heading
    while start_idx > 0:
        previous = doc.paragraphs[start_idx - 1]
        if (
            is_plain_blank_paragraph(previous)
            or paragraph_has_drawing(previous)
            or paragraph_has_page_break(previous)
            or is_cover_title_paragraph(previous)
        ):
            start_idx -= 1
            continue
        break

    # Remove stale page breaks between the cover block start and first English heading,
    # then add exactly one page break before the cover block.
    for idx in range(start_idx, first_english_heading):
        remove_page_break_runs(doc.paragraphs[idx])

    target = doc.paragraphs[start_idx]
    previous = target._element.getprevious()
    previous_para = next((p for p in doc.paragraphs if p._element is previous), None) if previous is not None else None
    if paragraph_has_page_break(target):
        return
    if previous_para is not None and paragraph_has_page_break(previous_para):
        return

    break_paragraph = target.insert_paragraph_before("")
    break_paragraph.add_run().add_break(WD_BREAK.PAGE)


def normalize_logo_alignment(doc: Document, occurrences: list[DocSectionOccurrence]) -> None:
    # Logos are drawing paragraphs immediately before a language block heading.
    # Center them consistently for all companies/languages.
    first_by_language: dict[str, DocSectionOccurrence] = {}
    for occurrence in sorted(occurrences, key=lambda occ: occ.heading_index):
        if occurrence.language not in first_by_language:
            first_by_language[occurrence.language] = occurrence

    centered_indices: set[int] = set()
    for occurrence in first_by_language.values():
        idx = occurrence.heading_index - 1
        while idx >= 0:
            paragraph = doc.paragraphs[idx]
            if paragraph_has_drawing(paragraph):
                if idx not in centered_indices:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    centered_indices.add(idx)
                idx -= 1
                continue
            if is_plain_blank_paragraph(paragraph) or paragraph_has_page_break(paragraph):
                idx -= 1
                continue
            break


def write_docx(
    output_path: Path,
    template_path: Path,
    section_plan: list[SectionOccurrence],
    business_activities_map: dict[str, str],
    financial_updates: list[dict[str, Any]],
    middle_table_updates: dict[tuple[str, str, int], dict[str, Any]] | None,
    business_update_map: dict[str, list[str]],
    business_update_group_map: dict[str, list[list[str]]] | None,
    risk_exit_map: dict[str, list[str]],
    financial_note_map: dict[str, str] | None = None,
) -> None:
    doc = Document(str(template_path))
    financial_note_map = financial_note_map or {}
    middle_table_updates = middle_table_updates or {}
    target_quarter = current_quarter_from_financial_updates(financial_updates)

    def apply_table_run_format(run, font_name: str | None, east_asia_font: str | None, size_pt: float | None, bold: bool | None) -> None:
        if font_name:
            run.font.name = font_name
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:ascii"), font_name)
            rfonts.set(qn("w:hAnsi"), font_name)
            rfonts.set(qn("w:cs"), font_name)
            rfonts.set(qn("w:eastAsia"), east_asia_font or font_name)
        elif east_asia_font:
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:eastAsia"), east_asia_font)
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        run.font.bold = bold

    # --- Fill Business Activities (re-resolve before each fill so later headings don't drift) ---
    for planned in section_plan:
        if planned.canonical != "Business Activities":
            continue
        occurrences = resolve_section_plan_to_doc(doc, section_plan)
        occ_idx = find_resolved_occurrence_index(doc, occurrences, planned)
        if occ_idx is None:
            continue
        occurrence = occurrences[occ_idx]
        body = get_body_paragraphs_for_occurrence(doc, occurrences, occ_idx)
        fill_section_paragraphs(doc, body, [business_activities_map.get(occurrence.language, business_activities_map.get("english", ""))])

    # --- Fill Business Update middle sections using the original template structure ---
    planned_middle_by_language: dict[str, list[SectionOccurrence]] = {"english": [], "chinese": []}
    for planned in section_plan:
        if planned.canonical == "Business Update":
            planned_middle_by_language.setdefault(planned.language, []).append(planned)

    for language, planned_occurrences in planned_middle_by_language.items():
        if not planned_occurrences:
            continue
        occurrences = resolve_section_plan_to_doc(doc, section_plan)
        resolved_pairs: list[tuple[SectionOccurrence, int]] = []
        for planned in planned_occurrences:
            occ_idx = find_resolved_occurrence_index(doc, occurrences, planned)
            if occ_idx is not None:
                resolved_pairs.append((planned, occ_idx))
        if not resolved_pairs:
            continue
        resolved_planned = [planned for planned, _occ_idx in resolved_pairs]
        keep_positions = [
            idx
            for idx, planned in enumerate(resolved_planned)
            if should_keep_middle_heading_flat_mode(planned, idx)
        ]
        if not keep_positions and resolved_planned:
            keep_positions = [0]
        kept_planned = [resolved_planned[idx] for idx in keep_positions]
        grouped_bullets = (business_update_group_map or {}).get(language, [])
        kept_chunks = build_flat_middle_chunks_for_template(
            kept_planned,
            grouped_bullets,
            business_update_map.get(language, []),
        )
        chunks = [[] for _ in resolved_planned]
        for pos, chunk in zip(keep_positions, kept_chunks):
            chunks[pos] = chunk

        # Re-resolve from the template plan before each fill so paragraph deletions/additions
        # do not change which original middle-section slot we are editing.
        for pos, texts in enumerate(chunks):
            if pos >= len(resolved_planned):
                break
            planned = resolved_planned[pos]
            occurrences = resolve_section_plan_to_doc(doc, section_plan)
            occ_idx = find_resolved_occurrence_index(doc, occurrences, planned)
            if occ_idx is None:
                continue
            occurrence = occurrences[occ_idx]
            body = ensure_middle_occurrence_body_slots(doc, occurrences, occ_idx, texts)
            fill_section_paragraphs(
                doc,
                body,
                texts,
                force_bullets=True,
                preserve_subheader_hierarchy=False,
                force_dash_bullets=True,
            )

    # Risk & Exit: DO NOT rewrite. The template (previous quarter) already has
    # the correct content. Rewriting via XML extraction and fill_section_paragraphs
    # is lossy (merges paragraphs, creates duplicates, loses formatting).

    occurrences = resolve_section_plan_to_doc(doc, section_plan)
    financial_entries = section_table_entries(section_plan, {"Financial Update"})
    for update_idx, entry in enumerate(financial_entries):
        if update_idx >= len(financial_updates):
            break
        table_idx = entry["global_table_index"]
        if table_idx >= len(doc.tables):
            continue
        table = doc.tables[table_idx]
        plan_occurrence = entry["occurrence"]
        occurrence_language = plan_occurrence.language
        center_table(table)
        financial_update = financial_updates[update_idx]
        ensure_table_dimensions(table, len(financial_update["rows"]) + 1, len(financial_update["columns"]) + 1)
        formatting: dict[tuple[int, int], tuple[str | None, str | None, float | None, bool | None]] = {}
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                font_name = None
                east_asia_font = None
                size_pt = None
                bold = None
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font_name = run.font.name or font_name
                        rpr = run._element.find(qn("w:rPr"))
                        if rpr is not None:
                            rfonts = rpr.find(qn("w:rFonts"))
                            if rfonts is not None:
                                east_asia_font = rfonts.get(qn("w:eastAsia")) or east_asia_font
                        if run.font.size is not None:
                            size_pt = float(run.font.size.pt)
                        if run.font.bold is not None:
                            bold = run.font.bold
                        break
                    if font_name or size_pt is not None or bold is not None:
                        break
                formatting[(r_idx, c_idx)] = (font_name, east_asia_font, size_pt, bold)

        unit_label = financial_update["unit"]
        top_left_original = norm_space(table.cell(0, 0).text)
        inside_unit = is_unit_text(top_left_original)
        outside_unit_paragraph = None
        occ_index = None
        for i, occ in enumerate(occurrences):
            if occ.canonical != "Financial Update" or occ.language != plan_occurrence.language:
                continue
            heading_text = norm_space(doc.paragraphs[occ.heading_index].text)
            if heading_text == norm_space(plan_occurrence.heading_text):
                occ_index = i
                break
        if occ_index is None:
            occ_index = next(
                (i for i, occ in enumerate(occurrences) if occ.canonical == "Financial Update" and occ.language == plan_occurrence.language),
                None,
            )
        if occ_index is not None:
            for paragraph in get_body_paragraphs_for_occurrence(doc, occurrences, occ_index):
                if paragraph.text.strip() and is_unit_text(paragraph.text):
                    outside_unit_paragraph = paragraph
                    break

        if outside_unit_paragraph is not None:
            # Only rewrite if the text actually changed (rewriting can lose font formatting)
            if norm_space(outside_unit_paragraph.text) != norm_space(unit_label):
                unit_run_format = capture_run_format(outside_unit_paragraph)
                set_paragraph_text_preserve(outside_unit_paragraph, unit_label)
                apply_paragraph_run_format(outside_unit_paragraph, *unit_run_format)
            set_cell_text_preserve(table.cell(0, 0), "")
        elif inside_unit:
            set_cell_text_preserve(table.cell(0, 0), unit_label)

        if occ_index is not None:
            note_text = financial_note_map.get(plan_occurrence.language)
            if note_text:
                for paragraph in get_body_paragraphs_for_occurrence(doc, occurrences, occ_index):
                    existing = norm_space(paragraph.text)
                    if not existing or is_unit_text(existing):
                        continue
                    if (
                        "银行账户现金余额" in existing
                        or "cash balance" in existing.lower()
                        or "bank account cash balance" in existing.lower()
                    ):
                        para_format = capture_run_format(paragraph)
                        set_paragraph_text_preserve(paragraph, note_text)
                        apply_paragraph_run_format(paragraph, para_format[0], para_format[1], para_format[2], False)
                        break

        for idx, header in enumerate(financial_update["columns"], start=1):
            if idx < len(table.rows[0].cells):
                cell = table.cell(0, idx)
                set_cell_text_preserve(cell, table_display_text(header, occurrence_language))
                set_cell_no_wrap(cell)
        for row_idx, row_data in enumerate(financial_update["rows"], start=1):
            if row_idx >= len(table.rows):
                break
            cell_language = occurrence_language
            label_cell = table.cell(row_idx, 0)
            set_cell_text_preserve(label_cell, table_display_text(row_data["label"], cell_language))
            set_cell_no_wrap(label_cell)
            center_cell_paragraphs(label_cell)
            for col_idx, header in enumerate(financial_update["columns"], start=1):
                if col_idx >= len(table.rows[row_idx].cells):
                    break
                value_cell = table.cell(row_idx, col_idx)
                set_cell_text_preserve(value_cell, table_display_text(row_data["values"].get(header) or "", cell_language))
                if header in {"QoQ", "YoY", "2025 FY", "2026E"} or header.endswith("FY") or header.endswith("YTD") or header.endswith("E"):
                    set_cell_no_wrap(value_cell)

        fit_financial_table_columns(doc, table, occurrence_language)

        # Build numeric-cell fallback format per row so added QoQ/YoY/FY/Estimate cells
        # inherit a consistent visual style across the entire row.
        row_defaults: dict[int, tuple] = {}
        row_numeric_defaults: dict[int, tuple] = {}
        drop_table_font = table_requires_font_drop(table, occurrence_language)
        for r_idx in range(len(table.rows)):
            preferred_indices = list(range(1, len(table.rows[r_idx].cells))) + [0]
            for c_idx in preferred_indices:
                fmt = formatting.get((r_idx, c_idx), (None, None, None, None))
                if any(v is not None for v in fmt):
                    row_defaults[r_idx] = fmt
                    break
            # Prefer a historical quarter style (col>=2) as the row numeric reference
            # so Q4/percent/FY cells don't look visually different.
            numeric_preferred = [idx for idx in range(2, len(table.rows[r_idx].cells))]
            numeric_preferred += [1]
            for c_idx in numeric_preferred:
                fmt = formatting.get((r_idx, c_idx), (None, None, None, None))
                if any(v is not None for v in fmt):
                    row_numeric_defaults[r_idx] = fmt
                    break

        for r_idx, row in enumerate(table.rows):
            fallback = row_defaults.get(r_idx, (None, None, None, None))
            numeric_fallback = row_numeric_defaults.get(r_idx, fallback)
            for c_idx, cell in enumerate(row.cells):
                fmt = formatting.get((r_idx, c_idx), (None, None, None, None))
                # If this cell had no formatting, use the row default
                if not any(v is not None for v in fmt):
                    fmt = fallback
                if r_idx >= 1 and c_idx >= 1:
                    fmt = numeric_fallback
                font_name, east_asia_font, size_pt, bold = fmt
                column_header = financial_update["columns"][c_idx - 1] if c_idx >= 1 and c_idx - 1 < len(financial_update["columns"]) else ""
                cell_text = cell.text
                for paragraph in cell.paragraphs:
                    if c_idx == 0:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        apply_table_run_format(run, font_name, east_asia_font, size_pt, bold)
                        target_size = None
                        if r_idx == 0:
                            target_size = 9.0
                        elif column_header in {"QoQ", "YoY"}:
                            target_size = percent_cell_font_size(cell_text)
                        elif size_pt is not None:
                            target_size = size_pt
                        if target_size is not None:
                            if drop_table_font:
                                target_size = max(7.0, target_size - 1.0)
                            run.font.size = Pt(target_size)
                        # Keep all numeric data cells non-bold; only header row/label column may retain bold.
                        if r_idx >= 1 and c_idx >= 1:
                            run.font.bold = False
        fit_financial_table_columns(doc, table, occurrence_language)

    additional_table_entries = section_table_entries(section_plan, {"Business Update"})
    for entry in additional_table_entries:
        key = (entry["occurrence"].language, entry["occurrence"].heading_text, entry["local_table_index"])
        table_update = middle_table_updates.get(key)
        if table_update is None:
            continue
        table_idx = entry["global_table_index"]
        if table_idx >= len(doc.tables):
            continue
        table = doc.tables[table_idx]
        center_table(table)
        formatting: dict[tuple[int, int], tuple[str | None, str | None, float | None, bool | None]] = {}
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                font_name = None
                east_asia_font = None
                size_pt = None
                bold = None
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font_name = run.font.name or font_name
                        rpr = run._element.find(qn("w:rPr"))
                        if rpr is not None:
                            rfonts = rpr.find(qn("w:rFonts"))
                            if rfonts is not None:
                                east_asia_font = rfonts.get(qn("w:eastAsia")) or east_asia_font
                        if run.font.size is not None:
                            size_pt = float(run.font.size.pt)
                        if run.font.bold is not None:
                            bold = run.font.bold
                        break
                    if font_name or size_pt is not None or bold is not None:
                        break
                formatting[(r_idx, c_idx)] = (font_name, east_asia_font, size_pt, bold)

        occurrence = entry["occurrence"]
        unit_label = table_update.get("unit") or ""
        top_left_original = norm_space(table.cell(0, 0).text)
        inside_unit = is_unit_text(top_left_original)
        if inside_unit:
            set_cell_text_preserve(table.cell(0, 0), unit_label if unit_label else "")

        for idx, header in enumerate(table_update["columns"], start=1):
            if idx < len(table.rows[0].cells):
                cell = table.cell(0, idx)
                set_cell_text_preserve(cell, table_display_text(header, occurrence.language))
        for row_idx, row_data in enumerate(table_update["rows"], start=1):
            if row_idx >= len(table.rows):
                break
            label_cell = table.cell(row_idx, 0)
            set_cell_text_preserve(label_cell, table_display_text(row_data["label"], occurrence.language))
            center_cell_paragraphs(label_cell)
            for col_idx, header in enumerate(table_update["columns"], start=1):
                if col_idx >= len(table.rows[row_idx].cells):
                    break
                value_cell = table.cell(row_idx, col_idx)
                set_cell_text_preserve(value_cell, table_display_text(row_data["values"].get(header) or "", occurrence.language))

        fit_financial_table_columns(doc, table, occurrence.language)

        row_defaults: dict[int, tuple] = {}
        row_numeric_defaults: dict[int, tuple] = {}
        drop_table_font = table_requires_font_drop(table, occurrence.language)
        for r_idx in range(len(table.rows)):
            for c_idx in range(len(table.rows[r_idx].cells)):
                fmt = formatting.get((r_idx, c_idx), (None, None, None, None))
                if any(v is not None for v in fmt):
                    row_defaults[r_idx] = fmt
                    break
            numeric_preferred = [idx for idx in range(2, len(table.rows[r_idx].cells))]
            numeric_preferred += [1]
            for c_idx in numeric_preferred:
                fmt = formatting.get((r_idx, c_idx), (None, None, None, None))
                if any(v is not None for v in fmt):
                    row_numeric_defaults[r_idx] = fmt
                    break

        for r_idx, row in enumerate(table.rows):
            fallback = row_defaults.get(r_idx, (None, None, None, None))
            numeric_fallback = row_numeric_defaults.get(r_idx, fallback)
            for c_idx, cell in enumerate(row.cells):
                fmt = formatting.get((r_idx, c_idx), (None, None, None, None))
                if not any(v is not None for v in fmt):
                    fmt = fallback
                if r_idx >= 1 and c_idx >= 1:
                    fmt = numeric_fallback
                font_name, east_asia_font, size_pt, bold = fmt
                column_header = table_update["columns"][c_idx - 1] if c_idx >= 1 and c_idx - 1 < len(table_update["columns"]) else ""
                cell_text = cell.text
                for paragraph in cell.paragraphs:
                    if c_idx == 0:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        apply_table_run_format(run, font_name, east_asia_font, size_pt, bold)
                        target_size = None
                        if r_idx >= 1 and column_header in {"QoQ", "YoY"}:
                            target_size = percent_cell_font_size(cell_text)
                        elif size_pt is not None:
                            target_size = size_pt
                        if target_size is not None:
                            if drop_table_font:
                                target_size = max(7.0, target_size - 1.0)
                            run.font.size = Pt(target_size)
        fit_financial_table_columns(doc, table, occurrence.language)

    def safe_occurrences() -> list[DocSectionOccurrence]:
        detected = [
            occ
            for occ in detect_doc_section_occurrences(doc)
            if occ.heading_index < len(doc.paragraphs)
        ]
        resolved = [
            occ
            for occ in resolve_section_plan_to_doc(doc, section_plan)
            if occ.heading_index < len(doc.paragraphs)
        ]
        if resolved and len(resolved) >= len(detected):
            return resolved
        return detected

    occurrences = safe_occurrences()
    remove_blank_paragraphs_inside_sections(doc, occurrences)
    occurrences = safe_occurrences()
    remove_blank_paragraphs_inside_sections(doc, occurrences)
    normalize_middle_section_heading_quarters(doc, occurrences, target_quarter)
    remove_ignorable_middle_headings(doc, safe_occurrences())
    normalize_bullet_sections(doc, occurrences)
    remove_empty_business_update_subheadings(doc, safe_occurrences())
    reapply_english_template_format_for_bilingual(doc, template_path)

    # Ensure all section headings are bold (after all paragraph shifts are done)
    for occ in safe_occurrences():
        heading_paragraph = doc.paragraphs[occ.heading_index]
        for run in heading_paragraph.runs:
            run.font.bold = True

    # Remove ALL blank paragraphs first, then re-insert exactly one before each heading
    to_remove = []
    for idx, paragraph in enumerate(doc.paragraphs):
        if is_plain_blank_paragraph(paragraph):
            to_remove.append(paragraph)
    for paragraph in reversed(to_remove):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    # Now insert exactly one blank paragraph before each section heading
    normalize_exact_single_blank_between_sections(doc, safe_occurrences())
    if FORCE_BILINGUAL_PAGE_BREAK:
        enforce_bilingual_page_break(doc, safe_occurrences())
    normalize_logo_alignment(doc, safe_occurrences())
    normalize_cover_title_paragraphs(doc, safe_occurrences())

    doc.save(str(output_path))


def extract_company_name(title: str, business_activities: str) -> str:
    match = re.search(r',\s*([^,]+?) \(the "Company"\)', business_activities)
    if match:
        return norm_space(match.group(1))
    cleaned = re.sub(r"\b20\d{2}\b", "", title)
    cleaned = re.sub(r"\bQ[1-4]\b", "", cleaned, flags=re.I)
    cleaned = re.sub(r"[_-]+", " ", cleaned)
    cleaned = norm_space(cleaned)
    if cleaned:
        return re.sub(r"(?<=[a-z])(?=[A-Z])", " ", cleaned)
    return "The Company"


def normalize_name_tokens(text: str) -> list[str]:
    text = re.sub(r"(?<=[A-Za-z])(?=[\u4e00-\u9fff])", " ", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])(?=[A-Za-z])", " ", text)
    text = re.sub(r"20\d{2}Q[1-4]", " ", text, flags=re.I)
    text = re.sub(r"20\d{2}", " ", text)
    text = re.sub(r"q[1-4]", " ", text, flags=re.I)
    text = re.sub(r"data request|datarequest|previous quarter template|template|samples|review|portfolio|quarter", " ", text, flags=re.I)
    base_text = re.sub(r"[^A-Za-z\u4e00-\u9fff]+", " ", text)

    tokens: list[str] = []
    stop = {"docx", "pdf", "ventech", "项目", "报告", "终稿"}
    for part in base_text.split():
        lowered = part.lower()
        if len(lowered) > 1 and lowered not in stop:
            tokens.append(lowered)

        camel_split = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", part)
        if camel_split != part:
            for sub in camel_split.split():
                lowered_sub = sub.lower()
                if len(lowered_sub) > 1 and lowered_sub not in stop:
                    tokens.append(lowered_sub)

        for suffix in ["项目", "公司"]:
            if part.endswith(suffix):
                stem = part[: -len(suffix)]
                lowered_stem = stem.lower()
                if len(lowered_stem) > 1 and lowered_stem not in stop:
                    tokens.append(lowered_stem)

    deduped: list[str] = []
    seen: set[str] = set()
    for token in tokens:
        if token not in seen:
            seen.add(token)
            deduped.append(token)
    return deduped


def jaccard_score(a: set[str], b: set[str]) -> float:
    if not a or not b:
        return 0.0
    return len(a & b) / len(a | b)


def infer_current_company_name(path: Path) -> str:
    blocks = load_blocks(path)
    parsed = parse_current_blocks(blocks)
    title = parsed.get("title", "")
    if title:
        cleaned = re.sub(r"\b20\d{2}\b", "", title)
        cleaned = re.sub(r"\bQ[1-4]\b", "", cleaned, flags=re.I)
        cleaned = re.sub(r"第[一二三四1234]季度.*", "", cleaned)
        cleaned = re.sub(r"\(.*?view.*?\)", "", cleaned, flags=re.I)
        cleaned = re.sub(r"业务概览.*", "", cleaned)
        cleaned = norm_space(cleaned)
        if cleaned:
            return re.sub(r"(?<=[a-z])(?=[A-Z])", " ", cleaned)
    return path.stem


def infer_previous_company_name(path: Path) -> str:
    blocks = load_docx_blocks(path)
    sections = extract_previous_sections(blocks)
    business = " ".join(sections.get("Business Activities", {}).get("paragraphs", []))
    if business:
        name = extract_company_name("", business)
        if name != "The Company":
            return name
    cleaned = re.sub(r"^\d+\s+", "", path.stem)
    cleaned = re.sub(r"\b20\d{2}Q[1-4]\b", "", cleaned, flags=re.I)
    cleaned = re.sub(r"\bdocx\b", "", cleaned, flags=re.I)
    cleaned = norm_space(cleaned)
    return cleaned


def extract_identity_text(path: Path, is_previous: bool) -> str:
    blocks = load_blocks(path) if not is_previous else load_docx_blocks(path)
    if is_previous:
        sections = extract_previous_sections(blocks)
        return " ".join(sections.get("Business Activities", {}).get("paragraphs", []))
    parsed = parse_current_blocks(blocks)
    return " ".join(parsed.get("business_update_paragraphs", [])[:3] + [parsed.get("title", "")])


def match_current_to_previous(current_files: list[Path], previous_files: list[Path]) -> list[dict[str, Any]]:
    previous_meta = []
    for prev in previous_files:
        prev_name = infer_previous_company_name(prev)
        prev_tokens = set(normalize_name_tokens(prev.stem + " " + prev_name))
        prev_identity = set(normalize_name_tokens(extract_identity_text(prev, is_previous=True)))
        prev_quarter = parse_quarter_label(prev.stem) or parse_quarter_label(docx_to_text(prev)[:50])
        previous_meta.append(
            {
                "path": prev,
                "company_name": prev_name,
                "name_tokens": prev_tokens,
                "identity_tokens": prev_identity,
                "quarter": prev_quarter,
            }
        )

    reports = []
    for current in current_files:
        current_name = infer_current_company_name(current)
        current_tokens = set(normalize_name_tokens(current.stem + " " + current_name))
        current_identity = set(normalize_name_tokens(extract_identity_text(current, is_previous=False)))
        current_quarter = parse_quarter_label(current.stem) or parse_quarter_label(docx_to_text(current)[:50])

        scored = []
        for prev in previous_meta:
            score = 0.0
            reasons = []

            name_score = jaccard_score(current_tokens, prev["name_tokens"])
            score += name_score * 0.55
            if name_score > 0:
                reasons.append(f"filename/company tokens overlap ({name_score:.2f})")

            current_name_tokens = set(normalize_name_tokens(current_name))
            prev_name_tokens = set(normalize_name_tokens(prev["company_name"]))
            direct_name_overlap = current_name_tokens & prev_name_tokens
            if direct_name_overlap:
                score += 0.20
                reasons.append(f"direct company-name overlap ({', '.join(sorted(direct_name_overlap))})")

            identity_score = jaccard_score(current_identity, prev["identity_tokens"])
            score += identity_score * 0.30
            if identity_score > 0:
                reasons.append(f"content identity overlap ({identity_score:.2f})")

            if current_quarter and prev["quarter"]:
                if current_quarter.year == prev["quarter"].year and current_quarter.quarter == prev["quarter"].quarter + 1:
                    score += 0.15
                    reasons.append("quarter progression aligns (current is next quarter)")
                elif current_quarter.year == prev["quarter"].year and current_quarter.quarter != prev["quarter"].quarter:
                    score += 0.05
                    reasons.append("same year quarter reference present")

            scored.append((score, prev, reasons))

        scored.sort(key=lambda x: x[0], reverse=True)
        best_score, best_prev, best_reasons = scored[0]
        second_score = scored[1][0] if len(scored) > 1 else 0.0
        gap = best_score - second_score

        if best_score >= 0.65 and gap >= 0.15:
            confidence = "high"
        elif best_score >= 0.45 and gap >= 0.08:
            confidence = "medium"
        else:
            confidence = "ambiguous / review needed"

        reports.append(
            {
                "current_file": str(current.relative_to(ROOT)),
                "matched_previous_file": str(best_prev["path"].relative_to(ROOT)),
                "confidence": confidence,
                "score": round(best_score, 3),
                "explanation": "; ".join(best_reasons) if best_reasons else "weak textual similarity only",
            }
        )
    return reports


def run_matching_report() -> None:
    current_files = list_current_inputs(DATA_REQUEST_DIR)
    previous_files = list_real_docx(PREVIOUS_REVIEW_DIR)
    reports = match_current_to_previous(current_files, previous_files)
    for item in reports:
        pass


def generate_review_for_pair(current_path: Path, previous_path: Path) -> dict[str, Any]:
    OUTPUT_DIR.mkdir(exist_ok=True)
    DEBUG_OUTPUT_DIR.mkdir(exist_ok=True)
    OLLAMA_TRANSLATION_STATUS["checked"] = False
    OLLAMA_TRANSLATION_STATUS["available"] = False
    OLLAMA_TRANSLATION_STATUS["reason"] = ""
    OLLAMA_TRANSLATION_STATUS["model"] = ""
    OLLAMA_TRANSLATION_STATUS["warned"] = False
    ARGOS_TRANSLATION_STATUS["checked"] = False
    ARGOS_TRANSLATION_STATUS["available"] = False
    ARGOS_TRANSLATION_STATUS["reason"] = ""
    ARGOS_TRANSLATION_STATUS["warned"] = False
    ARGOS_RESOURCES["translation"] = None

    spec_text = docx_to_text(UNIVERSAL_SPEC)
    prev_blocks = load_docx_blocks(previous_path)
    current_blocks = load_blocks(current_path)
    template_occurrences = detect_doc_section_occurrences(Document(str(previous_path)))
    template_languages = {occ.language for occ in template_occurrences}
    language_mode = template_language_mode(template_occurrences)
    previous_sections = extract_previous_sections(prev_blocks)
    previous_occurrences = extract_previous_section_occurrences(prev_blocks)
    previous_financial_occurrences = [occ for occ in previous_occurrences if occ.canonical == "Financial Update" and occ.tables]
    previous_tables = []
    for occurrence in previous_financial_occurrences:
        unit_label, _ = detect_financial_section_unit(occurrence)
        previous_tables.append(parse_previous_financial_table_with_unit(occurrence.tables[0], unit_label))
    current_parsed = parse_current_blocks(current_blocks)
    current_finance = normalize_finance_source(current_parsed["finance_table"])
    current_unit_spec = detect_current_unit_spec(current_blocks)
    current_exchange_rates = extract_exchange_rates(current_blocks)
    target_quarter = determine_target_quarter(current_path, current_parsed["title"], current_finance)
    previous_review_quarter = parse_quarter_label(previous_path.stem)
    financial_note_map = {
        language: note
        for language, note in {
            "chinese": extract_cash_balance_line(current_parsed.get("balance_table"), "chinese"),
            "english": extract_cash_balance_line(current_parsed.get("balance_table"), "english"),
        }.items()
        if note
    }

    review_flags: list[dict[str, Any]] = []
    source_language, language_evidence = detect_language(blocks_to_text(current_blocks))
    use_model_translation_for_middle_english = source_language == "chinese" and "english" in template_languages
    company_name = infer_previous_company_name(previous_path)
    final_company_name = normalize_company_filename(company_name)
    final_quarter_label = quarter_filename_label(target_quarter)
    final_docx_path = OUTPUT_DIR / f"{final_company_name}_{final_quarter_label}.docx"
    json_output = DEBUG_OUTPUT_DIR / f"{final_company_name}_{final_quarter_label}.json"
    markdown_output = DEBUG_OUTPUT_DIR / f"{final_company_name}_{final_quarter_label}.md"

    strongest_financial_idx = 0
    if previous_tables:
        strongest_financial_idx = max(
            range(len(previous_tables)),
            key=lambda idx: financial_table_backbone_score(previous_tables[idx], previous_review_quarter),
        )
    reference_prev_financial = previous_tables[strongest_financial_idx] if previous_tables else None
    current_unit_spec = infer_effective_current_unit_spec(
        current_finance,
        current_unit_spec,
        reference_prev_financial,
        target_quarter,
        previous_review_quarter,
        current_exchange_rates,
    )

    source_financial_update = None
    source_financial_unit_spec = None
    if previous_tables:
        source_prev_financial = previous_tables[strongest_financial_idx]
        source_financial_unit_spec = detect_unit_spec(source_prev_financial.get("unit", ""))
        source_financial_update = build_financial_update(
            source_prev_financial,
            current_finance,
            current_parsed.get("operation_table"),
            review_flags,
            target_quarter,
            current_unit_spec,
            source_financial_unit_spec,
            previous_review_quarter,
            current_exchange_rates,
        )

    financial_updates = []
    for idx, prev_financial in enumerate(previous_tables):
        target_unit_spec = detect_unit_spec(prev_financial.get("unit", ""))
        if source_financial_update is not None and len(prev_financial.get("rows", [])) == len(source_financial_update.get("rows", [])):
            financial_updates.append(
                localize_financial_update(
                    source_financial_update,
                    prev_financial,
                    source_financial_unit_spec,
                    target_unit_spec,
                    current_exchange_rates,
                )
            )
        else:
            financial_updates.append(
                build_financial_update(
                    prev_financial,
                    current_finance,
                    current_parsed.get("operation_table"),
                    review_flags,
                    target_quarter,
                    current_unit_spec,
                    target_unit_spec,
                    previous_review_quarter,
                    current_exchange_rates,
                )
            )

    middle_table_updates: dict[tuple[str, str, int], dict[str, Any]] = {}
    current_extra_tables: list[list[list[str]]] = list(current_parsed.get("extra_tables") or [])
    middle_table_groups = group_middle_table_occurrences(previous_occurrences)
    for group_idx, group in enumerate(middle_table_groups):
        if group_idx >= len(current_extra_tables):
            continue
        current_source_table = normalize_finance_source(current_extra_tables[group_idx])
        parsed_group_tables = []
        for occurrence in group:
            unit_label, _placement = detect_financial_section_unit(occurrence)
            parsed_group_tables.append((occurrence, parse_previous_financial_table_with_unit(occurrence.tables[0], unit_label)))
        strongest_idx = max(
            range(len(parsed_group_tables)),
            key=lambda idx: financial_table_backbone_score(parsed_group_tables[idx][1], previous_review_quarter),
        )
        source_occurrence, source_prev_table = parsed_group_tables[strongest_idx]
        source_middle_update = build_section_table_update(
            source_prev_table,
            current_source_table,
            review_flags,
            target_quarter,
            current_unit_spec,
            detect_unit_spec(source_prev_table.get("unit", "")),
            previous_review_quarter,
            current_exchange_rates,
        )
        for occurrence, prev_table in parsed_group_tables:
            if len(prev_table.get("rows", [])) == len(source_middle_update.get("rows", [])):
                middle_table_updates[(occurrence.language, occurrence.heading_text, 0)] = localize_financial_update(source_middle_update, prev_table)
            else:
                middle_table_updates[(occurrence.language, occurrence.heading_text, 0)] = build_section_table_update(
                    prev_table,
                    current_source_table,
                    review_flags,
                    target_quarter,
                    current_unit_spec,
                    detect_unit_spec(prev_table.get("unit", "")),
                    previous_review_quarter,
                    current_exchange_rates,
                )

    business_activity_map: dict[str, str] = {}
    risk_exit_map: dict[str, list[str]] = {}
    for occurrence in previous_occurrences:
        if occurrence.canonical == "Business Activities" and occurrence.paragraphs:
            business_activity_map.setdefault(occurrence.language, " ".join(occurrence.paragraphs).strip())
        elif occurrence.canonical == "Risk & Exit" and occurrence.paragraphs:
            if occurrence.language not in risk_exit_map:
                seen = set()
                deduped = []
                for p in occurrence.paragraphs:
                    key = norm_space(p).lower()
                    if key not in seen:
                        seen.add(key)
                        deduped.append(p)
                risk_exit_map[occurrence.language] = deduped
    previous_business_update_map = build_previous_business_update_map(previous_occurrences)
    middle_section_source_parsed = current_parsed
    if current_path.suffix.lower() == ".pdf" and not current_parsed.get("business_update_paragraphs"):
        converted_path = converted_pdf_docx_path(current_path)
        if not converted_path.exists():
            maybe_converted = convert_pdf_via_word(current_path)
            if maybe_converted is not None:
                converted_path = maybe_converted
        if converted_path.exists():
            try:
                converted_parsed = parse_current_blocks(load_docx_blocks(converted_path))
                if converted_parsed.get("business_update_paragraphs"):
                    middle_section_source_parsed = converted_parsed
                    review_flags.append(
                        {
                            "id": "middle-source-converted-docx",
                            "section": "Business Update",
                            "severity": "info",
                            "message": "Business Update source was extracted from auto-converted DOCX because raw PDF extraction had no usable update paragraphs.",
                            "source": "current_data_request",
                        }
                    )
            except Exception:
                pass
    middle_section_candidates = extract_middle_section_candidates(middle_section_source_parsed)
    middle_section_outputs = resolve_middle_section_outputs(
        source_language,
        company_name,
        template_languages,
        business_activity_map,
        previous_business_update_map,
        middle_section_candidates,
        review_flags,
        use_model_translation_for_middle_english,
    )
    business_update_map = middle_section_outputs.flat_by_language
    business_update_group_map = middle_section_outputs.grouped_by_language

    if not previous_review_has_logo(previous_path):
        review_flags.append(
            {
                "id": "logo-missing",
                "section": "Business Activities",
                "severity": "warning",
                "message": "No reusable company logo was found in the previous review.",
                "source": "previous_review",
            }
        )

    translation_engine = middle_section_outputs.translation_engine
    translation_engine_note = middle_section_outputs.translation_engine_note

    for table_update in financial_updates:
        for row in table_update["rows"]:
            for idx, message in enumerate(row["flags"], start=1):
                review_flags.append(
                    {
                        "id": f"{row['label'].lower().replace('#', 'num').replace(' ', '-')}-flag-{idx}",
                        "section": "Financial Update",
                        "severity": "warning" if "missing" in message.lower() or "could not" in message.lower() else "info",
                        "message": f"{row['label']}: {message}",
                        "source": "current_data_request" if "current data request" in message.lower() else "previous_review",
                    }
                )

    primary_financial_update = financial_updates[-1] if financial_updates else {}
    primary_business_activities = business_activity_map.get("english") or next(iter(business_activity_map.values()), "")
    primary_risk = risk_exit_map.get("english") or next(iter(risk_exit_map.values()), [])
    primary_business_update = business_update_map.get("english") or next(iter(business_update_map.values()), [])

    output_payload = {
        "prototype_version": "v3",
        "test_case": {
            "company": company_name,
            "target_quarter": target_quarter.display(),
            "files": {
                "universal_spec": {"path": str(UNIVERSAL_SPEC)},
                "current_data_request": {"path": str(current_path)},
                "previous_review": {"path": str(previous_path)},
            },
        },
        "language_detection": {
            "source_language": source_language,
            "template_language_mode": language_mode,
            "evidence": language_evidence,
            "translation_engine": translation_engine,
            "translation_engine_note": translation_engine_note,
        },
        "extracted_inputs": {
            "universal_spec": {"found": True, "raw_text": spec_text},
            "previous_review": {
                "business_activities": {"found": True, "raw_text": primary_business_activities},
                "financial_update": {"found": True, "parsed_table": primary_financial_update},
                "business_update": {"found": True, "raw_bullets": previous_business_update_map.get("english") or previous_business_update_map.get("chinese", [])},
                "risk_exit": {"found": True, "raw_bullets": primary_risk},
            },
            "current_data_request": {
                "financial_source": {"found": True, "raw_table": current_parsed["finance_table"], "parsed_metrics": current_finance},
                "operation_source": {"found": current_parsed.get("operation_table") is not None, "raw_table": current_parsed.get("operation_table")},
                "business_update_source": {
                    "found": True,
                    "source_paragraphs": middle_section_candidates.raw_paragraphs,
                    "source_groups": middle_section_candidates.source_groups,
                    "candidate_lines": {
                        "chinese": middle_section_candidates.chinese_lines,
                        "english": middle_section_candidates.english_lines,
                    },
                },
            },
        },
        "proposed_outputs": {
            "business_activities": {"status": "carried_forward", "text": primary_business_activities},
            "financial_update": primary_financial_update,
            "business_update": {
                "status": "drafted_from_current_data",
                "bullets": primary_business_update,
                "grouped_bullets": business_update_group_map,
                "source_trace": middle_section_source_parsed["business_update_paragraphs"],
            },
            "risk_exit": {"status": "carried_forward", "bullets": primary_risk},
        },
        "review_flags": review_flags,
    }

    json_output.write_text(json.dumps(output_payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    markdown_output.write_text(
        build_markdown(primary_business_activities, primary_financial_update, primary_business_update, primary_risk),
        encoding="utf-8",
    )
    write_docx(
        final_docx_path,
        previous_path,
        previous_occurrences,
        business_activity_map,
        financial_updates,
        middle_table_updates,
        business_update_map,
        business_update_group_map,
        risk_exit_map,
        financial_note_map,
    )

    return {
        "current_file": str(current_path.relative_to(ROOT)),
        "matched_previous_file": str(previous_path.relative_to(ROOT)),
        "output_file": str(final_docx_path.relative_to(ROOT)),
        "language_mode": language_mode,
        "company_naming_preserved": "yes",
    }


def debug_generation_for_pair(current_path: Path, previous_path: Path) -> None:
    output_path = OUTPUT_DIR / f"{normalize_company_filename(infer_previous_company_name(previous_path))}_{quarter_filename_label(determine_target_quarter(current_path, parse_current_blocks(load_blocks(current_path))['title'], normalize_finance_source(parse_current_blocks(load_blocks(current_path))['finance_table'])))}.docx"
    existed_before = output_path.exists()
    before_mtime = datetime.fromtimestamp(output_path.stat().st_mtime).astimezone().isoformat(timespec="seconds") if existed_before else None

    prev_blocks = load_docx_blocks(previous_path)
    current_blocks = load_blocks(current_path)
    previous_occurrences = extract_previous_section_occurrences(prev_blocks)
    previous_financial_occurrences = [occ for occ in previous_occurrences if occ.canonical == "Financial Update" and occ.tables]
    current_parsed = parse_current_blocks(current_blocks)
    current_finance = normalize_finance_source(current_parsed["finance_table"])
    target_quarter, quarter_debug = determine_target_quarter_with_debug(current_path, current_parsed["title"], current_finance)
    current_unit_spec = detect_current_unit_spec(current_blocks)
    current_exchange_rates = extract_exchange_rates(current_blocks)
    previous_review_quarter = parse_quarter_label(previous_path.stem)


    for idx, occurrence in enumerate(previous_financial_occurrences):
        unit_label, placement = detect_financial_section_unit(occurrence)
        prev_financial = parse_previous_financial_table_with_unit(occurrence.tables[0], unit_label)
        output_unit_spec = detect_unit_spec(prev_financial.get("unit", ""))
        conversion_factor = None
        if current_unit_spec and output_unit_spec:
            conversion_factor = current_unit_spec.factor_to_base / output_unit_spec.factor_to_base
        historical_headers = [header for header in prev_financial["headers"] if parse_quarter_label(header)]
        normalized_historical = normalize_previous_quarter_headers(prev_financial["headers"], previous_review_quarter)
        flags: list[dict[str, Any]] = []
        financial_update = build_financial_update(
            prev_financial,
            current_finance,
            current_parsed.get("operation_table"),
            flags,
            target_quarter,
            current_unit_spec,
            output_unit_spec,
            previous_review_quarter,
            current_exchange_rates,
        )

    result = generate_review_for_pair(current_path, previous_path)
    regenerated_output = ROOT / result["output_file"]
    existed_after = regenerated_output.exists()
    after_mtime = datetime.fromtimestamp(regenerated_output.stat().st_mtime).astimezone().isoformat(timespec="seconds") if existed_after else None

def run_batch_generation() -> None:
    current_files = list_current_inputs(DATA_REQUEST_DIR)
    previous_files = list_real_docx(PREVIOUS_REVIEW_DIR)
    reports = match_current_to_previous(current_files, previous_files)
    clean_main_output_folder_for_batch()

    for item in reports:
        current_path = ROOT / item["current_file"]
        previous_path = ROOT / item["matched_previous_file"]
        try:
            generate_review_for_pair(current_path, previous_path)
        except Exception as exc:
            print(f"  -> skipped {current_path.name} ({exc})")
            continue


def main() -> None:
    current_files = list_current_inputs(DATA_REQUEST_DIR)
    previous_files = list_real_docx(PREVIOUS_REVIEW_DIR)
    if not current_files:
        print(f"ERROR: No .docx or .pdf files found in {DATA_REQUEST_DIR}")
        sys.exit(1)
    if not previous_files:
        print(f"ERROR: No .docx files found in {PREVIOUS_REVIEW_DIR}")
        sys.exit(1)
    if not UNIVERSAL_SPEC.exists():
        print(f"WARNING: Universal Spec not found at {UNIVERSAL_SPEC}")
    reports = match_current_to_previous(current_files, previous_files)
    OUTPUT_DIR.mkdir(exist_ok=True)
    for item in reports:
        current_path = ROOT / item["current_file"]
        previous_path = ROOT / item["matched_previous_file"]
        print(f"Processing: {current_path.name} <-> {previous_path.name} (confidence: {item['confidence']})")
        try:
            result = generate_review_for_pair(current_path, previous_path)
        except Exception as exc:
            print(f"  -> skipped ({exc})")
            continue
        print(f"  -> {result['output_file']}")


if __name__ == "__main__":
    if "--match-report" in sys.argv:
        run_matching_report()
    elif "--generate-matched" in sys.argv:
        run_batch_generation()
    elif "--debug-company" in sys.argv:
        idx = sys.argv.index("--debug-company")
        keyword = sys.argv[idx + 1]
        reports = match_current_to_previous(list_current_inputs(DATA_REQUEST_DIR), list_real_docx(PREVIOUS_REVIEW_DIR))
        keyword_lower = keyword.lower()
        candidates = [
            item
            for item in reports
            if keyword_lower in item["current_file"].lower() or keyword_lower in item["matched_previous_file"].lower()
        ]
        matched = None
        if candidates:
            def candidate_key(item: dict[str, Any]) -> tuple[int, float]:
                current_name = str(item.get("current_file", ""))
                # Prefer original inputs over helper artifacts when scores are close.
                converted_penalty = 1 if Path(current_name).name.startswith("_converted_") else 0
                return (-converted_penalty, float(item.get("score", 0.0)))

            matched = max(candidates, key=candidate_key)
        if matched is None:
            raise SystemExit(f"No matched pair found for keyword: {keyword}")
        debug_generation_for_pair(ROOT / matched["current_file"], ROOT / matched["matched_previous_file"])
    else:
        main()
